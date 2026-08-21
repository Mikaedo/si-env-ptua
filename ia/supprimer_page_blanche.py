# -*- coding: utf-8 -*-
"""
Supprime la page blanche coincee entre les deux pages de garde.

La garde en double est une norme de l'ecole et se conserve. La feuille
entierement vide qui les separait, non : elle venait de trois paragraphes vides
laisses apres la premiere garde, qui debordaient sur une page a eux seuls.

La correction ne se contente pas de les effacer. Sans eux, la seconde garde
remonterait au bas de la premiere page. Un saut de page explicite prend donc
leur place : la mise en page ne depend plus alors du volume de la premiere
garde, alors qu'avec des lignes vides le moindre ajustement de police la
deplacerait a nouveau.
"""
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_page_blanche.docx")


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    paras = doc.paragraphs

    # La seconde garde commence a la deuxieme occurrence du titre du memoire.
    reperes = [i for i, p in enumerate(paras)
               if p.text.strip() == "MÉMOIRE DE FIN DE CYCLE"]
    if len(reperes) < 2:
        raise SystemExit("les deux pages de garde n'ont pas ete retrouvees")
    seconde = reperes[1]
    print(f"seconde page de garde au paragraphe {seconde}")

    # Les paragraphes vides qui la precedent immediatement sont retires.
    retires = 0
    i = seconde - 1
    while i > 0 and not paras[i].text.strip():
        paras[i]._element.getparent().remove(paras[i]._element)
        retires += 1
        i -= 1
    print(f"{retires} paragraphe(s) vide(s) retire(s)")

    # Le saut remplace ces lignes vides : il ne se decale pas. La liste des
    # paragraphes est relue, les suppressions ayant decale les rangs.
    reperes = [p for p in doc.paragraphs
               if p.text.strip() == "MÉMOIRE DE FIN DE CYCLE"]
    reperes[1].paragraph_format.page_break_before = True
    print("saut de page pose sur la seconde garde")

    doc.save(SOURCE)
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
