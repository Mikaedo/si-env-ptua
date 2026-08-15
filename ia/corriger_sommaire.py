# -*- coding: utf-8 -*-
"""
Remet le sommaire dans l'ordre reel du document.

La table des matieres est un champ Word : elle s'est reordonnee toute seule au
deplacement du resume et de l'abstract. Le sommaire, lui, est saisi a la main
et annoncait encore « Resume, Abstract, Annexes » alors que le document se
termine desormais par « Annexes, Resume, Abstract ». Un sommaire qui contredit
la pagination est une faute que le jury releve avant meme d'ouvrir le corps.

Seul le texte des paragraphes est permute, pas les paragraphes eux-memes : ils
partagent la meme mise en forme, et reecrire leur contenu evite de manipuler
l'arbre XML pour un resultat identique.
"""
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_sommaire.docx")

ATTENDU = ["Résumé", "Abstract", "Annexes"]
CORRIGE = ["Annexes", "Résumé", "Abstract"]


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    paras = doc.paragraphs

    debut = next(i for i, p in enumerate(paras) if p.text.strip() == "SOMMAIRE")
    fin = next(i for i, p in enumerate(paras)
               if i > debut and p.text.strip() == "LISTE DES FIGURES")

    # Reperer les trois lignes consecutives a permuter.
    lignes = [i for i in range(debut, fin) if paras[i].text.strip()]
    rang = None
    for k in range(len(lignes) - 2):
        trio = [paras[j].text.strip() for j in lignes[k:k + 3]]
        if trio == ATTENDU:
            rang = lignes[k:k + 3]
            break

    if rang is None:
        raise SystemExit("les trois lignes attendues n'ont pas ete trouvees")

    for indice, libelle in zip(rang, CORRIGE):
        p = paras[indice]
        for seg in list(p.runs)[1:]:
            seg._element.getparent().remove(seg._element)
        if p.runs:
            p.runs[0].text = libelle
        else:
            p.add_run(libelle)
        print(f"  ligne {indice} : {libelle}")

    doc.save(SOURCE)
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
