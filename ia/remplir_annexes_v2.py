"""
remplir_annexes_v2.py
=====================
Remplace le placeholder des annexes par le contenu reel.
Approche : vider le placeholder, puis inserer le contenu dans l'ordre.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
from lxml import etree

CHEMIN_ENTREE = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v48.docx"
CHEMIN_SORTIE = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v49.docx"

doc = Document(CHEMIN_ENTREE)

# Trouver le placeholder des annexes
placeholder = None
for i, p in enumerate(doc.paragraphs):
    if "contenu des annexes" in p.text:
        placeholder = p
        print(f"Placeholder trouve: paragraphe {i}")
        break

if placeholder is None:
    print("ERREUR: Placeholder non trouve")
    exit(1)

# Vider le placeholder
for run in placeholder.runs:
    run.text = ""

# Contenu des annexes dans l'ordre correct
# Format: (texte, bold, size, is_code)
contenu = [
    # Annexe A
    ("Annexe A : Extraits de code : hyperparamètres et paramètres d'entraînement", True, 13, False),
    ("", False, 11, False),
    ("Les extraits ci-dessous présentent la configuration des hyperparamètres pour les deux modèles d'intelligence artificielle.", False, 11, False),
    ("", False, 11, False),
    ("A.1 – Entraînement YOLOv8n (détection) :", True, 11, False),
    ("", False, 11, False),
    ("from ultralytics import YOLO", False, 10, True),
    ("", False, 10, True),
    ("CHEMIN_DATA_YAML = 'dataset/data.yaml'", False, 10, True),
    ("NB_EPOCHS = 10", False, 10, True),
    ("TAILLE_IMAGE = 320", False, 10, True),
    ("BATCH = 8", False, 10, True),
    ("", False, 10, True),
    ("model = YOLO('yolov8n.pt')", False, 10, True),
    ("resultats = model.train(", False, 10, True),
    ("    data=CHEMIN_DATA_YAML,", False, 10, True),
    ("    epochs=NB_EPOCHS,", False, 10, True),
    ("    imgsz=TAILLE_IMAGE,", False, 10, True),
    ("    batch=BATCH,", False, 10, True),
    ("    lr0=0.0001,", False, 10, True),
    ("    optimizer='AdamW',", False, 10, True),
    ("    patience=20,", False, 10, True),
    ("    augment=True,", False, 10, True),
    ("    project='runs_detection',", False, 10, True),
    ("    name='dechets_yolov8n',", False, 10, True),
    (")", False, 10, True),
    ("", False, 10, True),
    ("# Export ONNX pour déploiement mobile", False, 10, True),
    ("model.export(format='onnx')", False, 10, True),
    ("", False, 11, False),
    ("A.2 – Entraînement MobileNetV2 (classification) :", True, 11, False),
    ("", False, 11, False),
    ("import torch, torch.nn as nn", False, 10, True),
    ("from torchvision import datasets, transforms, models", False, 10, True),
    ("", False, 10, True),
    ("DOSSIER = 'dataset_criticite'", False, 10, True),
    ("NB_EPOCHS = 10", False, 10, True),
    ("BATCH = 8", False, 10, True),
    ("LR = 0.0001", False, 10, True),
    ("", False, 10, True),
    ("# Augmentation de données", False, 10, True),
    ("transform = transforms.Compose([", False, 10, True),
    ("    transforms.Resize((224, 224)),", False, 10, True),
    ("    transforms.RandomHorizontalFlip(),", False, 10, True),
    ("    transforms.RandomRotation(15),", False, 10, True),
    ("    transforms.ColorJitter(brightness=0.2),", False, 10, True),
    ("    transforms.ToTensor(),", False, 10, True),
    ("    transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225]),", False, 10, True),
    ("])", False, 10, True),
    ("", False, 10, True),
    ("# Transfer learning : MobileNetV2 pré-entraîné sur ImageNet", False, 10, True),
    ("model = models.mobilenet_v2(weights=models.MobileNet_V2_Weights.DEFAULT)", False, 10, True),
    ("model.classifier[1] = nn.Linear(model.last_channel, 3)  # 3 classes", False, 10, True),
    ("", False, 10, True),
    ("critere = nn.CrossEntropyLoss()", False, 10, True),
    ("optim = torch.optim.AdamW(model.parameters(), lr=LR)", False, 10, True),
    ("", False, 10, True),
    ("# Boucle d'entraînement", False, 10, True),
    ("for epoch in range(NB_EPOCHS):", False, 10, True),
    ("    model.train()", False, 10, True),
    ("    for images, labels in train_dl:", False, 10, True),
    ("        optim.zero_grad()", False, 10, True),
    ("        sorties = model(images)", False, 10, True),
    ("        perte = critere(sorties, labels)", False, 10, True),
    ("        perte.backward()", False, 10, True),
    ("        optim.step()", False, 10, True),
    ("", False, 10, True),
    ("# Export ONNX", False, 10, True),
    ("exemple = torch.randn(1, 3, 224, 224)", False, 10, True),
    ("torch.onnx.export(model, exemple, 'mobilenetv2_criticite.onnx')", False, 10, True),
    ("", False, 11, False),
    ("─" * 60, False, 11, False),
    ("", False, 11, False),
    # Annexe B
    ("Annexe B : Extraits de code : calcul des métriques d'évaluation", True, 13, False),
    ("", False, 11, False),
    ("Les métriques d'évaluation (précision, rappel, F1-score, mAP) sont calculées à l'aide des bibliothèques Ultralytics (détection) et scikit-learn (classification).", False, 11, False),
    ("", False, 11, False),
    ("B.1 – Évaluation YOLOv8n (détection) :", True, 11, False),
    ("", False, 11, False),
    ("from ultralytics import YOLO", False, 10, True),
    ("model = YOLO('runs_detection/dechets_yolov8n/weights/best.pt')", False, 10, True),
    ("metriques = model.val(data='dataset/data.yaml')", False, 10, True),
    ("", False, 10, True),
    ("# Métriques calculées automatiquement :", False, 10, True),
    ("# mAP@0.5      = metriques.box.map50   # 0.365", False, 10, True),
    ("# mAP@0.5:0.95 = metriques.box.map      # 0.217", False, 10, True),
    ("# Précision    = metriques.box.mp       # 0.522", False, 10, True),
    ("# Rappel       = metriques.box.mr       # 0.370", False, 10, True),
    ("# F1 = 2 * (P * R) / (P + R)            # 0.434", False, 10, True),
    ("", False, 11, False),
    ("B.2 – Évaluation MobileNetV2 (classification) :", True, 11, False),
    ("", False, 11, False),
    ("from sklearn.metrics import classification_report, confusion_matrix", False, 10, True),
    ("import seaborn as sns, matplotlib.pyplot as plt", False, 10, True),
    ("", False, 10, True),
    ("report = classification_report(vrais, predits, target_names=classes)", False, 10, True),
    ("print(report)", False, 10, True),
    ("#               precision  recall  f1-score  support", False, 10, True),
    ("#       faible       0.77     0.85      0.81       115", False, 10, True),
    ("#      modere       0.42     0.35      0.38        45", False, 10, True),
    ("#    important       0.12     0.08      0.10        14", False, 10, True),
    ("#    accuracy                           0.61       174", False, 10, True),
    ("#   macro avg       0.44     0.43      0.43       174", False, 10, True),
    ("# weighted avg       0.56     0.61      0.57       174", False, 10, True),
    ("", False, 11, False),
    ("B.3 – Matrice de confusion :", True, 11, False),
    ("", False, 11, False),
    ("cm = confusion_matrix(vrais, predits)", False, 10, True),
    ("sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',", False, 10, True),
    ("            xticklabels=classes, yticklabels=classes)", False, 10, True),
    ("plt.xlabel('Prédit'); plt.ylabel('Réel')", False, 10, True),
    ("plt.title('Matrice de confusion - MobileNetV2')", False, 10, True),
    ("plt.savefig('matrice_confusion.png', dpi=200)", False, 10, True),
    ("", False, 11, False),
    ("─" * 60, False, 11, False),
    ("", False, 11, False),
    # Annexe C
    ("Annexe C : Cartes satellite : risque pluie/relief et NO2 avant/après travaux", True, 13, False),
    ("", False, 11, False),
    ("Les cartes ci-dessous sont produites avec Google Earth Engine (GEE) en utilisant les données CHIRPS (précipitations 48h), SRTM (relief) et Sentinel-5P TROPOMI (NO2).", False, 11, False),
    ("", False, 11, False),
    ("Figure C.1 : Carte de risque pluie/relief – Zone du PTUA (CHIRPS 48h + SRTM). Les zones en rouge indiquent un fort cumul pluviométrique sur terrain plat (pente < 3°), propice à la formation d'eaux stagnantes.", False, 11, False),
    ("", False, 11, False),
    ("Figure C.2 : Carte de concentration NO2 – Avant travaux (Sentinel-5P TROPOMI, moyenne mensuelle). Les valeurs sont exprimées en µmol/m². Les zones de chantier apparaissent en jaune-orange.", False, 11, False),
    ("", False, 11, False),
    ("Figure C.3 : Carte de concentration NO2 – Pendant travaux (Sentinel-5P TROPOMI, moyenne mensuelle). Une augmentation du NO2 est visible sur les zones de construction (engins, circulation).", False, 11, False),
    ("", False, 11, False),
    ("Figure C.4 : Carte NDWI – Détection des surfaces d'eau (Sentinel-2, composite mensuel). Les pixels bleus indiquent les plans d'eau persistants, susceptibles de constituer des gîtes larvaires.", False, 11, False),
    ("", False, 11, False),
    ("Méthodologie :", True, 11, False),
    ("• CHIRPS : cumul des précipitations sur 48 heures, seuil > 50 mm", False, 11, False),
    ("• SRTM : pente calculée par dérivée du MNT, seuil < 3° (terrain plat)", False, 11, False),
    ("• Sentinel-5P : NO2 troposphérique, filtrage qualité > 0.5, agrégation mensuelle", False, 11, False),
    ("• Sentinel-2 : NDWI = (Green - NIR) / (Green + NIR), seuil > 0.3", False, 11, False),
    ("• Superposition : intersection des zones à forte pluie + faible pente = risque élevé", False, 11, False),
]

# Inserer le contenu apres le placeholder, dans l'ordre
# On utilise addnext mais on insere en partant de la fin pour que l'ordre soit correct
ref_element = placeholder._element

# Creer tous les paragraphes et les inserer en ordre inverse avec addnext
# pour qu'ils apparaissent dans l'ordre correct
paragraphs_to_insert = []

for texte, bold, size, is_code in contenu:
    new_p = doc.add_paragraph()
    run = new_p.add_run(texte)
    if is_code:
        run.font.name = 'Consolas'
    else:
        run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if texte.startswith("─"):
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    paragraphs_to_insert.append(new_p)

# Inserer en ordre inverse apres le placeholder
for p in reversed(paragraphs_to_insert):
    ref_element.addnext(p._element)

doc.save(CHEMIN_SORTIE)
print(f">> Document sauvegarde : {CHEMIN_SORTIE}")

# Verifier
doc2 = Document(CHEMIN_SORTIE)
for i in range(700, min(730, len(doc2.paragraphs))):
    p = doc2.paragraphs[i]
    t = p.text.strip()
    label = t[:100] if t else '(vide)'
    print(f'{i}: {label}')
