# -*- coding: utf-8 -*-
"""
Remonte la capture du tableau de bord dans le corps du memoire.

Une soutenance montre l'interface : laisser toutes les captures du web en
annexe oblige le lecteur a quitter le fil pour voir ce que le chapitre decrit.
Celle du tableau de bord rejoint donc le paragraphe 5.4, qui le presente, et
reprend le numero 5.3, libre depuis que les trois captures etaient descendues
en annexe.

L'annexe G est renumerotee dans la foulee : sans cela elle sauterait de G.1 a
G.3, ce qui se remarque plus qu'un numero manquant au milieu d'un chapitre.

Un defaut de placement est corrige au passage. Le titre de l'annexe avait ete
insere devant la premiere legende, mais l'image qu'elle commente la precede :
le titre se retrouvait donc apres cette image, qui flottait au-dessus de son
annexe. L'ordre attendu est titre, image, legende.

Chaque numero vit a trois endroits : la legende, l'entree de la liste des
figures et les renvois du texte. En oublier un laisserait le corps appeler une
figure qui n'existe plus.
"""
import shutil
from pathlib import Path

from docx import Document

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_remontee.docx")

LEGENDE_5_3 = ("Figure 5.3 : Tableau de bord, vue d'ensemble avec statistiques "
               "et carte des chantiers.")
LISTE_5_3 = "Figure 5.3 : Tableau de bord, vue d'ensemble et carte"


def porte_image(p):
    return bool(p._element.findall(f".//{A}blip"))


def reecrire(p, texte):
    for fragment in list(p.runs)[1:]:
        fragment._element.getparent().remove(fragment._element)
    if p.runs:
        p.runs[0].text = texte
    else:
        p.add_run(texte)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    paras = doc.paragraphs

    # 1. Reperer le trio de l'annexe G.
    # La liste des figures, en tete de document, porte les memes libelles que
    # les legendes du corps. Chercher sans borne y trouverait l'entree de liste
    # au lieu de la legende, et l'image attendue juste avant n'existerait pas.
    i_titre = next(i for i, p in enumerate(paras)
                   if p.text.strip().startswith("Annexe G") and i > 300)
    i_g2 = next(i for i, p in enumerate(paras)
                if p.text.strip().startswith("Figure G.2 :") and i > 300)
    i_g3 = next(i for i, p in enumerate(paras)
                if p.text.strip().startswith("Figure G.3 :") and i > 300)
    img_g2 = paras[i_g2 - 1]
    if not porte_image(img_g2):
        raise SystemExit("l'image de G.2 n'est pas la ou on l'attend")

    # 2. Le titre de l'annexe passe avant l'image de G.1, qui le precedait.
    avant_titre = paras[i_titre - 1]
    if porte_image(avant_titre):
        el = paras[i_titre]._element
        el.getparent().remove(el)
        avant_titre._element.addprevious(el)
        print("  titre de l'annexe G replace avant sa premiere image")

    # 3. Deplacer l'image et sa legende vers le paragraphe 5.4.
    paras = doc.paragraphs
    ancre = next(p for p in paras
                 if p.text.strip().startswith("Le module de rapports produit"))
    # python-docx reconstruit ses objets Paragraph a chaque acces : chercher
    # ensuite l'objet par identite echoue. Le rang est donc retenu au moment ou
    # on le trouve, et l'image lue dans la meme liste.
    courants = doc.paragraphs
    rang = next(i for i, p in enumerate(courants)
                if p.text.strip().startswith("Figure G.2 :") and i > 300)
    legende_g2, image = courants[rang], courants[rang - 1]
    if not porte_image(image):
        raise SystemExit("l'image attendue avant la legende G.2 est absente")

    for element in (legende_g2._element, image._element):
        element.getparent().remove(element)
    ancre._element.addnext(legende_g2._element)
    ancre._element.addnext(image._element)
    reecrire(legende_g2, LEGENDE_5_3)
    print("  capture remontee au paragraphe 5.4, devient Figure 5.3")

    # 4. L'annexe se resserre : G.3 devient G.2.
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Figure G.3 :") and i > 300:
            reecrire(p, p.text.replace("Figure G.3", "Figure G.2", 1))
            print("  Figure G.3 devient Figure G.2")

    # 5. La liste des figures suit.
    paras = doc.paragraphs
    debut = next(i for i, p in enumerate(paras)
                 if p.text.strip() == "LISTE DES FIGURES")
    fin = next(i for i, p in enumerate(paras)
               if i > debut and p.text.strip() == "LISTE DES TABLEAUX")
    entree_g2 = entree_g3 = ancre_liste = None
    for p in paras[debut + 1:fin]:
        t = p.text.strip()
        if t.startswith("Figure G.2 :"):
            entree_g2 = p
        elif t.startswith("Figure G.3 :"):
            entree_g3 = p
        elif t.startswith("Figure 5.2 bis"):
            ancre_liste = p
    if entree_g2 is not None and ancre_liste is not None:
        el = entree_g2._element
        el.getparent().remove(el)
        ancre_liste._element.addnext(el)
        reecrire(entree_g2, LISTE_5_3)
        print("  liste : entree replacee apres 5.2 bis")
    if entree_g3 is not None:
        reecrire(entree_g3, entree_g3.text.replace("Figure G.3", "Figure G.2", 1))

    # 6. Le texte doit appeler la figure, sinon elle flotte.
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Le tableau de bord sert au Spécialiste") \
                and "figure 5.3" not in t:
            reecrire(p, t.rstrip() + " La figure 5.3 en donne la vue "
                                     "d'ensemble.")
            print("  renvoi ajoute dans le texte de 5.4")
            break

    doc.save(SOURCE)

    controle = Document(SOURCE)
    import re
    txt = "\n".join(p.text for p in controle.paragraphs)
    leg = sorted({m.group(1) for m in
                  re.finditer(r"^Figure (G\.\d|5\.3)\s*:", txt, re.M)})
    print(f"\nfigures G et 5.3 presentes : {leg}")
    print(f"renvoi « figure 5.3 » dans le texte : "
          f"{'figure 5.3' in txt.lower()}")
    print(f"Sauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
