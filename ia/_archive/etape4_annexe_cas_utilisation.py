# -*- coding: utf-8 -*-
"""
Etape 4 : deplacer les descriptions textuelles des cas d'utilisation vers une
annexe, et recuperer les pages blanches.

Justification du choix : les six tableaux 6.6 a 6.11 detaillent chacun un cas
d'utilisation sur huit lignes. Ils occupent quatre a cinq pages du chapitre
Conception sans etre indispensables au fil du raisonnement, le diagramme de cas
d'utilisation et sa lecture suffisant a la demonstration. Les placer en annexe
respecte le guide de forme - numerotation conservee, titre au-dessus du tableau,
renvoi depuis le texte - et les sort du decompte des 50 pages, les annexes n'y
etant pas comptees.

Sont egalement retirees les trois pages blanches produites par des paires de
sauts de page consecutifs.

L'espacement des paragraphes n'est PAS modifie : l'essai precedent a montre que
le forcer degradait la pagination au lieu de l'ameliorer.

Sortie : MEMOIRE_ETAPE4.docx
"""
import os
import re
import shutil

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_ETAPE2.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_ETAPE4.docx")

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

#: Anciens numeros des tableaux a deplacer -> nouveau numero d'annexe.
DEPLACES = {
    '6.6': 'E.1', '6.7': 'E.2', '6.8': 'E.3',
    '6.9': 'E.4', '6.10': 'E.5', '6.11': 'E.6',
}


def texte(el):
    return ''.join(t.text or '' for t in el.iter('{%s}t' % W)).strip()


def sauts_page(el):
    return [br for br in el.findall('.//{%s}br' % W)
            if br.get('{%s}type' % W) == 'page']


def supprimer_pages_blanches(doc):
    """Retire le second saut de chaque paire consecutive sans contenu."""
    body = doc.element.body
    retires = 0
    saut_en_attente = False
    for el in list(body):
        if not el.tag.endswith('}p'):
            saut_en_attente = False
            continue
        vide = not texte(el) and not el.findall('.//{%s}blip' % A_NS)
        sauts = sauts_page(el)
        if sauts and vide:
            if saut_en_attente:
                for br in sauts:
                    br.getparent().remove(br)
                retires += 1
                saut_en_attente = False
            else:
                saut_en_attente = True
        elif not vide:
            saut_en_attente = False
    return retires


def main():
    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)
    body = doc.element.body

    # ── Reperage des couples (titre, tableau) a deplacer ─────────────────────
    a_deplacer = []
    elements = list(body)
    for i, el in enumerate(elements):
        if not el.tag.endswith('}p'):
            continue
        t = texte(el)
        m = re.match(r'^Tableau\s+(6\.\d+)\s*[:.]', t)
        if not m or m.group(1) not in DEPLACES:
            continue
        # Le tableau suit son titre, eventuellement apres un paragraphe vide
        tbl = None
        for j in range(i + 1, min(i + 4, len(elements))):
            if elements[j].tag.endswith('}tbl'):
                tbl = elements[j]
                break
        if tbl is None:
            print("  ! tableau introuvable pour %s" % m.group(1))
            continue
        a_deplacer.append((m.group(1), el, tbl))

    print("  couples titre+tableau reperes : %d" % len(a_deplacer))
    if len(a_deplacer) != len(DEPLACES):
        raise SystemExit("Reperage incomplet, on n'applique rien")

    # ── Renvoi laisse dans le corps, a la place du premier tableau ───────────
    premier_titre = a_deplacer[0][1]
    renvoi = premier_titre.makeelement('{%s}p' % W, {})
    premier_titre.addprevious(renvoi)
    from docx.text.paragraph import Paragraph
    par = Paragraph(renvoi, doc.paragraphs[0]._parent)
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.add_run(
        "Les descriptions textuelles detaillees des six cas d'utilisation "
        "principaux (acteurs, preconditions, scenario nominal, scenarios "
        "alternatifs et postconditions) sont presentees en annexe E, aux "
        "tableaux E.1 a E.6.")
    print("  renvoi insere dans le corps")

    # ── Creation de l'annexe E en fin de document ───────────────────────────
    doc.add_page_break()
    titre_annexe = doc.add_paragraph()
    titre_annexe.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    titre_annexe.add_run(
        "Annexe E : Descriptions textuelles des cas d'utilisation").bold = True

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "Les tableaux E.1 a E.6 detaillent les six cas d'utilisation "
        "principaux du SI-ENV, selon le formalisme employe au chapitre "
        "consacre a la conception : acteurs concernes, preconditions, "
        "scenario nominal, scenarios alternatifs et postconditions.")

    ancre = doc.paragraphs[-1]._element

    # ── Deplacement effectif ────────────────────────────────────────────────
    for ancien, titre_el, tbl_el in a_deplacer:
        nouveau = DEPLACES[ancien]
        # Renumerotation du titre
        for t in titre_el.iter('{%s}t' % W):
            if t.text and 'Tableau' in t.text:
                t.text = t.text.replace('Tableau %s' % ancien,
                                        'Tableau %s' % nouveau)
        # Deplacement : le titre puis son tableau, a la suite de l'ancre
        ancre.addnext(tbl_el)
        ancre.addnext(titre_el)
        ancre = tbl_el
        print("    Tableau %-5s -> Tableau %s" % (ancien, nouveau))

    # ── Mise a jour des renvois et de la liste des tableaux ─────────────────
    remplaces = 0
    for p in doc.paragraphs:
        for run in p.runs:
            for ancien, nouveau in DEPLACES.items():
                for forme in ('tableau %s' % ancien, 'Tableau %s' % ancien,
                              'tableaux %s' % ancien):
                    if forme in run.text:
                        run.text = run.text.replace(
                            forme, forme.split()[0] + ' ' + nouveau)
                        remplaces += 1
    print("  renvois et entrees de liste mis a jour : %d" % remplaces)

    blanches = supprimer_pages_blanches(doc)
    print("  pages blanches supprimees : %d" % blanches)

    doc.save(SORTIE)
    print("\nEnregistre : %s" % SORTIE)


if __name__ == "__main__":
    main()
