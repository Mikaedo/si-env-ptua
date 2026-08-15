# -*- coding: utf-8 -*-
"""
Ramene la troisieme conclusion partielle a la longueur des deux autres.

Les trois conclusions de partie faisaient 508, 803 et 1549 caracteres. La
derniere pesait donc le triple de la premiere, ce qui se voit a l'oeil nu et
donne l'impression d'une partie plus importante que les autres alors qu'elle
ne l'est pas. Une conclusion de partie annonce ce qui suit, elle ne resume pas
le chapitre une seconde fois.

Le texte resserre conserve les trois enseignements et la limite finale, en
supprimant les redites : les trois lecons etaient chacune introduites par une
formule d'annonce, et l'ouverture repetait ce que le chapitre venait de dire.
Le fond est intact, seule la charpente rhetorique disparait.

Le texte est lu depuis un fichier en UTF-8. Compose sur une ligne de commande,
il perdrait ses accents.
"""
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_cp3.docx")
TEXTE = Path(r"C:\Users\DELL\AppData\Local\Temp\claude"
             r"\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25"
             r"\scratchpad\cp3.txt")


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    paras = doc.paragraphs

    titres = [i for i, p in enumerate(paras)
              if p.style.name.startswith("Heading")
              and p.text.strip() == "Conclusion partielle"]
    if len(titres) != 3:
        raise SystemExit(f"attendu 3 conclusions, trouve {len(titres)}")

    cible = paras[titres[2] + 1]
    nouveau = TEXTE.read_text(encoding="utf-8").strip()

    # Le premier run est reecrit et les suivants supprimes : la mise en forme
    # du paragraphe et de son texte est ainsi conservee telle quelle.
    for seg in list(cible.runs)[1:]:
        seg._element.getparent().remove(seg._element)
    if cible.runs:
        cible.runs[0].text = nouveau
    else:
        cible.add_run(nouveau)

    doc.save(SOURCE)

    controle = Document(SOURCE)
    paras = controle.paragraphs
    titres = [i for i, p in enumerate(paras)
              if p.style.name.startswith("Heading")
              and p.text.strip() == "Conclusion partielle"]
    for n, i in enumerate(titres, 1):
        texte = paras[i + 1].text.strip()
        accents = sum(texte.count(c) for c in "éèêàçôûîœ")
        print(f"  conclusion {n} : {len(texte)} caracteres, {accents} accents")

    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
