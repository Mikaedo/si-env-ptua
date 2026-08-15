# -*- coding: utf-8 -*-
"""
Etape 13 : accentuer le sommaire reecrit, et rendre aux figures une taille
lisible.

  a) Le sommaire reecrit a l'etape 11 avait ete saisi sans accents. Il est
     reecrit ici avec l'orthographe correcte, et confronte aux titres reellement
     presents dans le corps pour verifier qu'aucune entree ne divergeait.

  b) Le corps etant descendu a 47 pages, soit trois pages sous le plafond de 50,
     cette marge est rendue aux figures : elles passent de 45 % a 60 % de leur
     taille d'origine, ce qui les rend a nouveau lisibles a l'impression. Les
     bornes du corps sont reperees explicitement, la detection automatique
     ayant attrape les entrees de la table des matieres generee.

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os
import re

from docx import Document

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

#: Sommaire accentue, dans l'ordre, tel qu'il doit apparaitre.
SOMMAIRE = [
    "Introduction générale",
    "Première partie : Présentation générale",
    "Chapitre 1 : Cadre de l'étude et problématique",
    "Chapitre 2 : État de l'art et analyse de l'existant",
    "Conclusion partielle",
    "Deuxième partie : Analyse et conception",
    "Chapitre 3 : Analyse des besoins",
    "Chapitre 4 : Conception",
    "Conclusion partielle",
    "Troisième partie : Réalisation et résultats",
    "Chapitre 5 : Réalisation : plateforme, intelligence artificielle "
    "et télédétection",
    "Chapitre 6 : Déploiement, tests et discussion",
    "Conclusion partielle",
    "Conclusion générale",
]

#: Bornes du corps, en index de doc.paragraphs (releve par diagnostic).
CORPS_DEBUT, CORPS_FIN = 249, 548

TAILLE_VOULUE = 0.60
TAILLE_ACTUELLE = 0.45
FACTEUR = TAILLE_VOULUE / TAILLE_ACTUELLE

SEUIL_LOGO_CM = 3.0
LARGEUR_UTILE_CM = 16.0


def ecrire(paragraphe, contenu):
    noeuds = list(paragraphe._element.iter('{%s}t' % W))
    if not noeuds:
        paragraphe.add_run(contenu)
        return
    noeuds[0].text = contenu
    for n in noeuds[1:]:
        n.text = ''


def sans_accents(s):
    table = str.maketrans("àâäéèêëîïôöùûüçÀÂÄÉÈÊËÎÏÔÖÙÛÜÇ",
                          "aaaeeeeiioouuucAAAEEEEIIOOUUUC")
    return s.translate(table)


def main():
    doc = Document(CIBLE)
    paras = doc.paragraphs

    # ── a) Accentuation du sommaire ─────────────────────────────────────────
    # On repere les lignes par leur equivalent non accentue, ce qui evite de
    # dependre d'index absolus.
    attendu = {sans_accents(x): x for x in SOMMAIRE}
    corriges = 0
    for i, p in enumerate(paras):
        if i > 130:
            break
        t = p.text.strip()
        if not t:
            continue
        bon = attendu.get(t)
        if bon and bon != t:
            ecrire(p, bon)
            corriges += 1
    print("  entrees du sommaire accentuees : %d / %d" % (corriges, len(SOMMAIRE)))

    # ── Verification : titres reellement presents dans le corps ─────────────
    titres_corps = []
    for i in range(CORPS_DEBUT, min(CORPS_FIN, len(paras))):
        t = paras[i].text.strip()
        m = re.match(r'^Chapitre\s+([1-6])\s*$', t)
        if m:
            # le titre suit sur le paragraphe non vide suivant
            for j in range(i + 1, min(i + 4, len(paras))):
                suite = paras[j].text.strip()
                if suite:
                    titres_corps.append((m.group(1), suite))
                    break
    print("  chapitres trouves dans le corps : %d" % len(titres_corps))
    for num, titre in titres_corps:
        entree = next((x for x in SOMMAIRE
                       if x.startswith('Chapitre %s :' % num)), None)
        attendu_titre = entree.split(' : ', 1)[1] if entree else ''
        etat = 'ok' if titre.lower()[:22] == attendu_titre.lower()[:22] else 'DIVERGE'
        print("    ch.%s %-8s corps=%s" % (num, etat, titre[:52]))

    # ── b) Figures du corps ─────────────────────────────────────────────────
    agrandies = plafonnees = 0
    for i in range(CORPS_DEBUT, min(CORPS_FIN, len(paras))):
        for ext in paras[i]._element.iter('{%s}ext' % A_NS):
            cx, cy = ext.get('cx'), ext.get('cy')
            if not (cx and cy):
                continue
            l_cm, h_cm = int(cx) / 360000.0, int(cy) / 360000.0
            if l_cm < SEUIL_LOGO_CM:
                continue
            f = FACTEUR
            if l_cm * f > LARGEUR_UTILE_CM:
                f = LARGEUR_UTILE_CM / l_cm
                plafonnees += 1
            ext.set('cx', str(int(round(l_cm * f * 360000))))
            ext.set('cy', str(int(round(h_cm * f * 360000))))
            agrandies += 1
    print("  figures portees a %d %% : %d (dont %d plafonnees a la largeur "
          "de page)" % (TAILLE_VOULUE * 100, agrandies, plafonnees))

    doc.save(CIBLE)
    print("\nEnregistre : %s" % CIBLE)


if __name__ == "__main__":
    main()
