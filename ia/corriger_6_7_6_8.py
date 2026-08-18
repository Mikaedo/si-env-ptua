# -*- coding: utf-8 -*-
"""
Corrige deux erreurs de fond dans les paragraphes 6.7 et 6.8.

La premiere est arithmetique. Le tableau 6.5 annoncait une licence « a partir
de 500 USD par utilisateur et par mois », puis un cout annuel de vingt mille
dollars par utilisateur. Douze mois a cinq cents dollars font six mille, non
vingt mille. La conversion en francs, elle, etait juste pour vingt mille, si
bien que l'erreur se propageait : le tableau affichait douze millions la ou le
calcul donne trois millions six cent mille. C'est le genre d'ecart qu'un jury
verifie de tete.

La seconde tient a la conclusion tiree de ce tableau. Le texte situait le
SI-ENV « trois a quatre ordres de grandeur » sous une solution commerciale,
soit mille a dix mille fois moins cher. Les chiffres du tableau ne le
soutiennent pas, meme apres correction. Plutot que d'avancer un multiplicateur
qui se refute par une division, le texte compare desormais les echelles et
souligne ce qui creuse reellement l'ecart : les licences commerciales se
facturent au poste, pas le SI-ENV.

S'y ajoute une contradiction entre deux paragraphes voisins. L'un presentait le
serveur prive virtuel comme la cible du pilote AGEROUTE, l'autre la pile Docker
locale pour le meme pilote. Le premier visait en realite la mise en
exploitation, et le dit maintenant.

L'en-tete de colonne « Montant estime » perd enfin son adjectif : depuis que la
prime de stage y figure pour son montant reel, toutes les lignes ne sont plus
des estimations.
"""
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_6768.docx")

CONCLUSION = (
    "Ces montants, présentés à titre de comparaison d'ordre de grandeur et non "
    "comme un coût effectivement évité par AGEROUTE, situent le SI-ENV très "
    "en dessous d'une solution commerciale équivalente : le premier se chiffre "
    "en centaines de milliers de francs par an, les secondes en millions. "
    "L'écart se creuse en outre avec le nombre d'utilisateurs, les licences "
    "commerciales étant facturées au poste alors que le coût du SI-ENV ne "
    "dépend que de l'hébergement. Le système reste ainsi reproductible pour "
    "d'autres projets sans surcoût de licence."
)

REMPLACEMENTS = [
    # Contradiction sur la destination du serveur prive virtuel.
    ("une projection de production sur serveur privé virtuel, envisagée pour "
     "un déploiement pilote AGEROUTE",
     "une projection de mise en exploitation sur serveur privé virtuel"),
]

CELLULES = {
    # (entete de la colonne, ancien contenu) -> nouveau contenu
    "Montant estimé": "Montant",
    "≈ 20 000 USD/utilisateur/an, soit ≈ 12 millions FCFA":
        "≈ 6 000 USD/utilisateur/an, soit ≈ 3,6 millions FCFA",
}


def reecrire(paragraphe, texte):
    for fragment in list(paragraphe.runs)[1:]:
        fragment._element.getparent().remove(fragment._element)
    if paragraphe.runs:
        paragraphe.runs[0].text = texte
    else:
        paragraphe.add_run(texte)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    faits = []

    for p in doc.paragraphs:
        for avant, apres in REMPLACEMENTS:
            if avant in p.text:
                reecrire(p, p.text.replace(avant, apres))
                faits.append("contradiction pilote / exploitation")

        if p.text.strip().startswith("Ces montants, présentés à titre"):
            reecrire(p, CONCLUSION)
            faits.append("conclusion du tableau 6.5")

        # Une double espace subsistait avant la reprise sur Hugging Face.
        if "cette pratique.  Un premier essai" in p.text:
            reecrire(p, p.text.replace("cette pratique.  Un premier essai",
                                       "cette pratique. Un premier essai"))
            faits.append("double espace, paragraphe 6.8")

    for table in doc.tables:
        for ligne in table.rows:
            for cellule in ligne.cells:
                contenu = cellule.text.strip()
                if contenu in CELLULES:
                    reecrire(cellule.paragraphs[0], CELLULES[contenu])
                    for p in list(cellule.paragraphs)[1:]:
                        p._element.getparent().remove(p._element)
                    faits.append(f"cellule « {contenu[:44]} »")

    doc.save(SOURCE)
    for f in faits:
        print(f"  corrige : {f}")

    # Controle des totaux du tableau 6.4, recalcules a la main.
    print("\ncontrole arithmetique")
    print("  investissement initial : 0 + (3 500 a 80 000) + 9 500 "
          "= 13 000 a 89 500 FCFA")
    print("  fonctionnement annuel  : (3 500 a 80 000) x 12 + 9 500 "
          "= 51 500 a 969 500 FCFA")
    print("  licence commerciale    : 500 USD x 12 = 6 000 USD/an, "
          "soit 3 600 000 FCFA au taux de 600")
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
