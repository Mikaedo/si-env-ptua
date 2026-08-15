# -*- coding: utf-8 -*-
"""Restaure les accents manquants dans les paragraphes et cellules de tableau
que j'ai ajoutes/modifies au cours de la session (chapitres 7, 8, 9, 10,
legendes de figures, tableau des sprints, benchmarks IA). Les runs sont
fusionnes en un seul run par paragraphe corrige (la mise en forme etait
uniforme - aucun gras/italique perdu)."""
from docx import Document

SRC = r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_N'GUESSAN_TDM_corrigee.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_N'GUESSAN_v2.docx"

doc = Document(SRC)

def fix_para(p, new_text):
    old = p.text
    runs = p.runs
    assert runs, f"paragraphe sans run : {old!r}"
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''
    return old

PARA_FIXES = {
    105: "Figure 8.1 : Métriques par classe du modèle YOLOv8n (Recycle Trash, GPU T4, 100 epochs).",
    106: "Figure 8.2 : Matrice de confusion normalisée du modèle YOLOv8n (détection de déchets, 6 classes).",
    107: "Figure 8.3 : Courbe précision-rappel du modèle YOLOv8n (mAP@0.5 global = 0,798).",
    108: "Figure 8.4 : Matrice de confusion du modèle MobileNetV2 (classification de criticité, accuracy = 86,6 %).",
    444: ("L'application fonctionne hors connexion d'abord : SQLite stocke les signalements en attente, "
          "synchronisés au retour du réseau. L'interface est optimisée terrain (boutons larges, contraste "
          "élevé, saisie < 1 minute). L'architecture suit le pattern BLoC ; onnxruntime exécute YOLOv8 "
          "localement (13-27 ms sur CPU desktop selon le modèle, cf. tableau 10.3). Le tableau 7.3 liste "
          "les packages principaux. La figure 7.2 présente la liste des signalements, la carte des "
          "chantiers et le détail d'un signalement."),
    450: ("Le module de rapports produit un PGES PDF en un clic : le serveur agrège les données par "
          "chantier, applique le modèle BAD et renvoie le fichier. Cette opération, autrefois manuelle et "
          "sur plusieurs jours, prend quelques secondes. La figure 7.3 présente le tableau de bord, avec "
          "la carte des tracés PTUA et la répartition des signalements."),
    469: ("L'entraînement d'un réseau de neurones convolutifs exige un volume massif d'images annotées. "
          "N'ayant pas eu l'autorisation de photographier extensivement les chantiers du PTUA durant notre "
          "stage, nous nous sommes appuyés sur Recycle Trash, un jeu de données public de détection de "
          "déchets recyclables, disponible sur GitHub et Roboflow Universe. Ce dataset comporte 2 462 "
          "images annotées au format YOLO, réparties en six catégories (métal, plastique, papier, carton, "
          "verre, organique), avec des contextes variés (extérieur, intérieur, surfaces mixtes) proches "
          "des conditions d'un chantier de construction."),
    476: "Figure 8.1 : Métriques par classe du modèle YOLOv8n (Recycle Trash, GPU T4, 100 epochs).",
    478: "Figure 8.2 : Matrice de confusion normalisée du modèle YOLOv8n (détection de déchets, 6 classes).",
    480: "Figure 8.3 : Courbe précision-rappel du modèle YOLOv8n (mAP@0.5 global = 0,798).",
    485: "Figure 8.4 : Matrice de confusion du modèle MobileNetV2 (classification de criticité, accuracy = 86,6 %).",
    486: ("Les objectifs de performance ont été calibrés par comparaison avec des travaux publiés sur des "
          "tâches de vision par ordinateur comparables, plutôt que fixés arbitrairement. Pour la "
          "détection, une étude récente sur la détection de déchets de chantier par YOLOv8n rapporte un "
          "mAP@0.5 de 89,8 % [20] ; notre modèle atteint un mAP@0.5 de 0,798, soit 79,8 %, ce qui se "
          "situe dans un ordre de grandeur cohérent avec cette référence. L'écart résiduel s'explique par "
          "la différence de dataset (6 classes vs 1 classe spécialisée) et l'absence d'optimisation "
          "architecturale (FE-YOLO). Pour la classification, une étude comparative reporte un F1-score de "
          "0,93 pour MobileNetV2 [21] ; notre modèle atteint un F1-score pondéré de 0,86, cohérent avec un "
          "dataset de criticité plus petit et des classes déséquilibrées."),
    511: ("Le Spec. Env. sélectionne une zone d'intérêt sur le dashboard. FastAPI transmet la requête à "
          "GEE (source, dates, indicateur). Pour le NO2, GEE calcule la concentration pixel par pixel et "
          "applique un masque de nuages avant de renvoyer une heatmap ; pour le risque pluie/relief, GEE "
          "croise directement le cumul CHIRPS et la pente SRTM, sans masque de nuages nécessaire. Le "
          "résultat est stocké dans PostGIS. Le dashboard permet de comparer deux périodes (avant/après "
          "travaux) pour alimenter les rapports PGES. La figure 9.1 illustre ces indices, calculés à "
          "partir de données Google Earth Engine réelles, pour les six chantiers PTUA."),
    513: ("Figure 9.1 : Indices environnementaux calculés via Google Earth Engine (NO2, NDVI, NDWI, risque "
          "pluie/relief) pour les six chantiers PTUA — données réelles Sentinel-5P/Sentinel-2, capture du "
          "31 juillet 2026."),
    526: ("Les tests unitaires (pytest, PostGIS, Flutter) vérifient les composants isolés. Les tests "
          "d'intégration valident la chaîne mobile -> backend -> dashboard. La suite pytest (32 tests "
          "automatisés, fichier tests/test_functional.py) a été exécutée le 31 juillet 2026 sur la stack "
          "Docker réelle (base PostGIS, backend FastAPI, nginx) : 32/32 tests passés. Le tableau 10.2 "
          "synthétise les douze scénarios fonctionnels couverts."),
    531: ("Les indicateurs ci-dessous ont été mesurés le 31 juillet 2026 sur un déploiement Docker local "
          "(3 conteneurs : PostGIS, FastAPI, nginx), et non sur le VPS cible de production (2 vCPU / 4 Go "
          "RAM) : les temps réseau/API pourront différer légèrement en production. Les temps d'inférence "
          "ONNX ont été mesurés sur processeur d'ordinateur (CPU desktop) et restent à confirmer sur un "
          "téléphone réel. Le modèle de détection embarqué dans l'application mobile est actuellement une "
          "version d'entraînement rapide mono-classe ; le modèle Recycle Trash à six classes validé au "
          "chapitre 8 (GPU, mAP@0.5 = 0,798) n'a pas encore été réexporté et intégré à l'application (cf. "
          "discussion 10.5)."),
    534: ("Le SI-ENV répond aux six lacunes du chapitre 4 : signalement instantané avec diagnostic IA, "
          "subjectivité encadrée, géolocalisation automatique, données centralisées, rapports en secondes, "
          "alertes quasi temps réel (rafraîchissement automatique du tableau de bord toutes les 10 à 15 "
          "secondes ; la synchronisation mobile reste déclenchée manuellement par l'utilisateur). Les "
          "limites : tests fonctionnels exécutés en environnement de développement local (pas encore sur "
          "le VPS de production), dataset Recycle Trash non encore validé sur des photographies réelles "
          "des chantiers du PTUA, modèle de détection six classes validé au chapitre 8 pas encore "
          "réexporté vers l'application mobile (version embarquée : modèle mono-classe d'entraînement "
          "rapide), effet de propagation d'erreur inhérent au pipeline en cascade (une détection manquée "
          "ou imprécise par YOLOv8 dégrade la classification de criticité par MobileNetV2 en aval), "
          "dépendance Internet pour le satellite, résolution Sentinel-5P trop grossière pour le niveau "
          "chantier."),
    536: ("Perspectives : au-delà du déploiement pilote et de l'enrichissement du dataset (5 000 images), "
          "trois axes prolongeraient naturellement le SI-ENV. Premièrement, l'extension des capteurs IoT "
          "à d'autres nuisances que celles couvertes par la vision par ordinateur : sondes de turbidité et "
          "de niveau d'eau pour les eaux stagnantes difficiles d'accès, sonomètres connectés pour le "
          "bruit, capteurs PM2.5/PM10 pour les poussières ; couplées à des drones pour cartographier "
          "régulièrement les zones inaccessibles au sol (bassins de rétention, points bas de chantier), en "
          "complément des indices Sentinel dont la résolution reste trop grossière à l'échelle d'un "
          "chantier. Deuxièmement, le remplacement des seuils empiriques actuels (alertes GEE, criticité) "
          "par un modèle prédictif entraîné sur l'historique des signalements, des indices satellitaires "
          "et des données météorologiques, afin d'anticiper un risque environnemental avant qu'il ne se "
          "matérialise plutôt que de le constater après coup. Troisièmement, l'extension du système à "
          "d'autres projets AGEROUTE au-delà du PTUA. Le SI-ENV reste par ailleurs open source et "
          "fonctionne hors ligne, contrairement aux solutions commerciales (Enablon, Cority) qui "
          "nécessitent licences et connectivité permanente."),
}

paras = doc.paragraphs
for idx, new_text in PARA_FIXES.items():
    old = fix_para(paras[idx], new_text)
    assert paras[idx].text == new_text, f"echec paragraphe {idx}"
print(f"[OK] {len(PARA_FIXES)} paragraphes corriges (accents)")

TABLE_FIXES = {
    (9, 6, 3): ("Déploiement Docker complet (3 conteneurs), câblage du déclenchement automatique des "
                "alertes (seuil -> alerte -> email), correction du pipeline de détection mobile (NMS), "
                "fiabilisation de la synchronisation satellite GEE."),
    (9, 7, 3): ("Exécution de la suite de tests fonctionnels (32 scénarios), mesure des performances "
                "réelles (API, synchronisation, inférence), correction des écarts entre résultats "
                "attendus et résultats mesurés, rédaction finale du mémoire."),
    (24, 1, 0): "Recycle Trash (entraînement)",
    (24, 4, 0): "Augmentation de données",
    (24, 4, 2): "0 (intégrée via Ultralytics)",
    (25, 0, 0): "Modèle",
    (25, 0, 2): "Précision",
    (26, 0, 0): "Modèle",
    (26, 0, 1): "Précision",
    (30, 4, 2): "Score en < 200 ms (mesure CPU desktop, modèle embarqué actuel)",
    (31, 0, 2): "Seuil ou référence",
    (31, 5, 0): "Récupération des indices satellitaires GEE (6 chantiers PTUA)",
}

for (ti, ri, ci), new_text in TABLE_FIXES.items():
    cell = doc.tables[ti].rows[ri].cells[ci]
    p = cell.paragraphs[0]
    fix_para(p, new_text)
    assert cell.text == new_text, f"echec table{ti} r{ri}c{ci} : {cell.text!r}"
print(f"[OK] {len(TABLE_FIXES)} cellules de tableau corrigees (accents)")

doc.save(DST)
print(f"\n=== SAUVEGARDE : {DST} ===")
