# -*- coding: utf-8 -*-
"""
Met a jour le memoire apres la refonte du tableau de bord et de la page
d'analyse satellitaire :

  - remplace la capture de la Figure 7.3 (tableau de bord) ;
  - remplace la capture de la Figure 9.1 (indices satellitaires) ;
  - corrige la duree de validite du jeton JWT, portee de 1 h a 12 h dans la
    configuration du backend : le texte du chapitre 7 annoncait encore 1 h.

Sortie : MEMOIRE_N'GUESSAN_v9.docx (l'original n'est jamais modifie).
"""
import io
import os
import shutil

from docx import Document
from docx.shared import Inches

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v7.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v9.docx")
CAPTURES = r"D:\etude_soutenance\SI-ENV\ia\captures_annexes"

# Paragraphe (index) -> capture a y placer
REMPLACEMENTS_IMAGES = {
    454: "web_dashboard_v2.png",   # Figure 7.3 : tableau de bord
    515: "web_satellite_v2.png",   # Figure 9.1 : indices satellitaires
}


def remplacer_image(doc, index_par, fichier, largeur_pouces=6.3):
    """Vide le paragraphe puis y insere la nouvelle capture."""
    par = doc.paragraphs[index_par]
    for run in list(par.runs):
        run._element.getparent().remove(run._element)
    par.add_run().add_picture(os.path.join(CAPTURES, fichier),
                             width=Inches(largeur_pouces))


def compter_images(par):
    return len(par._element.findall(
        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'))


def main():
    for f in REMPLACEMENTS_IMAGES.values():
        chemin = os.path.join(CAPTURES, f)
        if not os.path.exists(chemin):
            raise SystemExit("Capture manquante : %s" % chemin)

    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)

    # ── Images ───────────────────────────────────────────────────────────────
    for index_par, fichier in REMPLACEMENTS_IMAGES.items():
        par = doc.paragraphs[index_par]
        avant = compter_images(par)
        if avant == 0:
            # Le paragraphe attendu ne porte pas d'image : on cherche le
            # paragraphe illustre le plus proche pour ne pas ecrire au hasard.
            trouve = None
            for delta in (1, -1, 2, -2, 3, -3):
                voisin = index_par + delta
                if 0 <= voisin < len(doc.paragraphs) and compter_images(doc.paragraphs[voisin]):
                    trouve = voisin
                    break
            if trouve is None:
                print("  ! aucune image trouvee autour du paragraphe %d, ignore" % index_par)
                continue
            print("  paragraphe %d sans image -> utilise %d" % (index_par, trouve))
            index_par = trouve
        remplacer_image(doc, index_par, fichier)
        print("  image remplacee au paragraphe %d par %s" % (index_par, fichier))

    # ── Duree du jeton JWT ───────────────────────────────────────────────────
    corriges = 0
    for par in doc.paragraphs:
        if "HS256" in par.text and "1h" in par.text:
            for run in par.runs:
                if "1h" in run.text:
                    run.text = run.text.replace("1h", "12h")
                    corriges += 1
    print("  mentions de duree du jeton corrigees : %d" % corriges)

    doc.save(SORTIE)
    print("\nEnregistre : %s" % SORTIE)
    print("Paragraphes : %d | Tableaux : %d"
          % (len(doc.paragraphs), len(doc.tables)))


if __name__ == "__main__":
    main()
