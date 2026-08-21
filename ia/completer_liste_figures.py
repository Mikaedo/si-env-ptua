# -*- coding: utf-8 -*-
"""
Remet la liste des figures d'aplomb.

Trois defauts s'y trouvaient. La figure 1.1 y occupait cent dix-huit
caracteres, seule entree a deborder. Les figures 4.10 et 5.2 bis, inserees
apres coup dans le corps, n'y figuraient pas du tout : une figure absente de la
liste est introuvable pour qui s'y reporte. Et deux entrees annoncaient un
libelle que le corps ne portait plus, depuis que ses legendes ont ete
resserrees.

Les entrees sont inserees a leur rang de numerotation et non en fin de liste :
c'est l'ordre des numeros qui fait la fonction d'une telle liste.

Le libelle de la liste reste plus court que la legende du corps la ou celle-ci
precise quelque chose que la liste n'a pas a porter. Les deux se rejoignent
partout ailleurs.
"""
import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_liste_figures.docx")

# Entrees a corriger : ancien libelle -> nouveau.
CORRECTIONS = [
    ("Figure 1.1 : Organigramme de la Direction des Affaires Juridiques, "
     "Moyens Généraux et Système d'Information (AGEROUTE)",
     "Figure 1.1 : Organigramme de la direction d'accueil à l'AGEROUTE"),
    ("Figure 5.2 : Captures d'écran de l'application mobile",
     "Figure 5.2 : Application mobile agent : liste, carte et détail"),
    ("Figure 5.3 : Captures d'écran du tableau de bord",
     "Figure 5.3 : Tableau de bord : vue d'ensemble, statistiques et carte"),
]

# Entrees manquantes : (libelle, entree apres laquelle l'inserer).
AJOUTS = [
    ("Figure 4.10 : Tableau de bord vu par l'ANDE, en consultation seule",
     "Figure 4.9"),
    ("Figure 5.2 bis : Application citoyenne : accueil et dépôt d'une doléance",
     "Figure 5.2"),
]


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    paras = doc.paragraphs

    debut = next(i for i, p in enumerate(paras)
                 if p.text.strip() == "LISTE DES FIGURES")
    fin = next(i for i, p in enumerate(paras)
               if i > debut and p.text.strip() == "LISTE DES TABLEAUX")

    for avant, apres in CORRECTIONS:
        pose = False
        for p in paras[debut + 1:fin]:
            if avant in p.text:
                entier = p.text.replace(avant, apres)
                for fragment in list(p.runs)[1:]:
                    fragment._element.getparent().remove(fragment._element)
                p.runs[0].text = entier
                pose = True
                break
        print(f"  {'corrigee' if pose else 'INTROUVABLE'} : {apres[:56]}")

    for libelle, apres_quoi in AJOUTS:
        paras = doc.paragraphs
        debut = next(i for i, p in enumerate(paras)
                     if p.text.strip() == "LISTE DES FIGURES")
        fin = next(i for i, p in enumerate(paras)
                   if i > debut and p.text.strip() == "LISTE DES TABLEAUX")
        ancre = None
        for p in paras[debut + 1:fin]:
            if p.text.strip().startswith(apres_quoi):
                ancre = p
        if ancre is None:
            print(f"  ancre introuvable : {apres_quoi}")
            continue
        element = copy.deepcopy(ancre._element)
        ancre._element.addnext(element)
        ligne = Paragraph(element, ancre._parent)
        for fragment in list(ligne.runs)[1:]:
            fragment._element.getparent().remove(fragment._element)
        ligne.runs[0].text = libelle
        print(f"  ajoutee apres {apres_quoi} : {libelle[:52]}")

    doc.save(SOURCE)

    controle = Document(SOURCE)
    paras = controle.paragraphs
    debut = next(i for i, p in enumerate(paras)
                 if p.text.strip() == "LISTE DES FIGURES")
    fin = next(i for i, p in enumerate(paras)
               if i > debut and p.text.strip() == "LISTE DES TABLEAUX")
    entrees = [p.text.strip() for p in paras[debut + 1:fin] if p.text.strip()]
    longues = [e for e in entrees if len(e) > 110]
    print(f"\n{len(entrees)} entrees dans la liste, "
          f"{len(longues)} au-dela de 110 caracteres")

    # Toute figure legendee dans le corps doit figurer a la liste.
    corps = {m.group(1) for p in paras[fin:]
             for m in [re.match(r"^(Figure [\d.]+(?: bis)?)\s*:", p.text.strip())]
             if m}
    listees = {m.group(1) for e in entrees
               for m in [re.match(r"^(Figure [\d.]+(?: bis)?)\s*:", e)] if m}
    manquantes = sorted(corps - listees)
    print(f"figures du corps absentes de la liste : "
          f"{manquantes if manquantes else 'aucune'}")
    print(f"Sauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
