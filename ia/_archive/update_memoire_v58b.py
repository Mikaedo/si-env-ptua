# -*- coding: utf-8 -*-
"""
v58 -> v59 : enrichissement de la section 10.6 Perspectives (IoT pour
d'autres nuisances, drones, modele predictif de risque), sans toucher au
reste du document. Verification passee : aucun autre fichier .onnx que les
deux deja identifies n'existe dans le depot (recherche find -iname "*.onnx"
sur tout le projet), donc la mention du modele mono-classe embarque au
chapitre 10 reste un fait verifiable, pas une supposition.
"""
from docx import Document

SRC = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v58.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v59.docx"

doc = Document(SRC)

def replace_paragraph_text(idx, new_text):
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)
    print(f"[OK] Paragraphe {idx} remplace")

replace_paragraph_text(550,
    "Perspectives : au-dela du deploiement pilote et de l'enrichissement du dataset (5 000 images), "
    "trois axes prolongeraient naturellement le SI-ENV. Premierement, l'extension des capteurs IoT a "
    "d'autres nuisances que celles couvertes par la vision par ordinateur : sondes de turbidite et de "
    "niveau d'eau pour les eaux stagnantes difficiles d'acces, sonometres connectes pour le bruit, "
    "capteurs PM2.5/PM10 pour les poussieres ; couplees a des drones pour cartographier regulierement "
    "les zones inaccessibles au sol (bassins de retention, points bas de chantier), en complement des "
    "indices Sentinel dont la resolution reste trop grossiere a l'echelle d'un chantier. Deuxiemement, "
    "le remplacement des seuils empiriques actuels (alertes GEE, criticite) par un modele predictif "
    "entraine sur l'historique des signalements, des indices satellitaires et des donnees "
    "meteorologiques, afin d'anticiper un risque environnemental avant qu'il ne se materialise plutot "
    "que de le constater apres coup. Troisiemement, l'extension du systeme a d'autres projets AGEROUTE "
    "au-dela du PTUA. Le SI-ENV reste par ailleurs open source et fonctionne hors ligne, contrairement "
    "aux solutions commerciales (Enablon, Cority) qui necessitent licences et connectivite permanente."
)

doc.save(DST)
print(f"\n=== MEMOIRE SAUVEGARDE : {DST} ===")
