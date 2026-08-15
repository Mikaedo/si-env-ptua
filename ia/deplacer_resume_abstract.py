# -*- coding: utf-8 -*-
"""
Place le resume et l'abstract sur la derniere page, partagee entre les deux.

L'ecole demande que ces deux textes ferment le document, apres les annexes, et
qu'ils tiennent sur une seule et meme page : le resume en haut, l'abstract en
bas. Ils se trouvaient jusqu'ici avant les annexes, sur deux pages distinctes.

Le deplacement se fait au niveau XML plutot que sur la liste des paragraphes.
La raison est la meme qui avait deja fausse une insertion precedente : la
propriete paragraphs de python-docx ignore les tableaux, si bien qu'un
decoupage fonde sur elle laisserait derriere lui tout element non textuel
intercale. En travaillant sur les enfants du corps, on emporte le bloc entier
quel qu'en soit le contenu.

Les elements sont reinseres avant le sectPr final, qui porte la mise en page de
la derniere section et doit rester le dernier enfant du corps.

Reste a repartir les deux textes sur la hauteur de la page. Plutot que de
deviner un nombre de lignes vides, ce qui se decale au moindre changement de
police, un paragraphe unique separe les deux blocs et recoit un espacement
calcule apres coup, une fois la page mesuree.
"""
import shutil
from pathlib import Path

from docx import Document

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_deplacement.docx")


def rang_dans_corps(body, paragraphe):
    for i, enfant in enumerate(body):
        if enfant is paragraphe._element:
            return i
    raise ValueError("paragraphe introuvable dans le corps")


def titre(doc, libelle):
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip() == libelle:
            return p
    raise SystemExit(f"titre absent : {libelle}")


def purger_sauts(p):
    """Retire les sauts de page manuels portes par un paragraphe."""
    p.paragraph_format.page_break_before = False
    for br in p._element.findall(f".//{W}br[@{W}type='page']"):
        br.getparent().remove(br)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    body = doc.element.body

    resume = titre(doc, "RÉSUMÉ")
    abstract = titre(doc, "ABSTRACT")
    annexes = titre(doc, "ANNEXES")

    debut = rang_dans_corps(body, resume)
    fin = rang_dans_corps(body, annexes)          # exclu
    bloc = list(body[debut:fin])
    print(f"{len(bloc)} elements deplaces, de RÉSUMÉ jusqu'avant ANNEXES")

    # Les paragraphes vides qui precedaient le resume portaient le saut de page.
    # Ils restent en place et serviront desormais aux annexes.
    for enfant in bloc:
        body.remove(enfant)

    sectPr = body.find(f"{W}sectPr")
    for enfant in bloc:
        sectPr.addprevious(enfant)

    # Le resume ouvre une page neuve, l'abstract enchaine sur la meme.
    resume.paragraph_format.page_break_before = True
    purger_sauts(abstract)
    annexes.paragraph_format.page_break_before = True

    # Aucun saut ne doit subsister a l'interieur du bloc deplace, ni apres lui :
    # un saut final laisserait une page blanche en fin de document.
    #
    # La comparaison porte sur l'element XML et non sur l'objet Paragraph.
    # python-docx reconstruit ces objets a chaque acces a doc.paragraphs, si
    # bien qu'un test d'identite entre un paragraphe garde en variable et celui
    # rencontre dans la boucle est toujours faux. Une premiere version faisait
    # cette erreur et ne purgeait rien.
    dedans = False
    for p in doc.paragraphs:
        if p._element is resume._element:
            dedans = True
            continue
        if dedans:
            purger_sauts(p)

    doc.save(SOURCE)

    ordre = [p.text.strip() for p in Document(SOURCE).paragraphs
             if p.style.name.startswith("Heading 1") and p.text.strip()]
    print("fin du document :", " puis ".join(ordre[-4:]))
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
