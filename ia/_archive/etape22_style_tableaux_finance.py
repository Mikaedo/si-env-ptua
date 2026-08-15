# -*- coding: utf-8 -*-
"""
Etape 22 : redessiner les tableaux 6.4 et 6.5 a l'identique du reste du
memoire.

Le tableau ajoute a l'etape 21 utilisait le style Word par defaut (« Normal
Table », sans bordure) et une legende alignee a gauche en texte courant. Tous
les autres tableaux du memoire suivent une autre recette, extraite ici
directement du tableau 6.3 :

  Legende  : centree, gras, italique, noir, 10 pt, keepNext, espacement
             40/160 twips avant/apres.
  Tableau  : style « Grilledutableau » (Table Grid), centre sur la page,
             largeur totale ~9000 twips.
  Entete   : fond gris clair (E8E8E8), texte 9 pt, sans gras (le contraste
             vient du fond, pas de la graisse).
  Corps    : fond blanc, texte 9 pt, interligne simple, espacement reduit
             (20 twips apres chaque paragraphe de cellule).

Cette etape supprime les tableaux 6.4 et 6.5 actuels (mal habilles) et les
reconstruit avec exactement cette recette, sans toucher a leur contenu.

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table

sys.stdout.reconfigure(encoding='utf-8')

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WQ = '{%s}' % W

LARGEURS = [2200, 4300, 2500]  # twips, total 9000, memes 3 colonnes pour 6.4 et 6.5


def qw(tag):
    return qn('w:' + tag)


def construire_legende(ancre_el, parent, texte):
    """Reproduit exactement la legende du tableau 6.3 : centree, gras,
    italique, noir, 10 pt, keepNext."""
    p_el = OxmlElement('w:p')
    ancre_el.addprevious(p_el)
    ppr = OxmlElement('w:pPr')
    p_el.append(ppr)
    ppr.append(OxmlElement('w:keepNext'))
    spacing = OxmlElement('w:spacing')
    spacing.set(qw('before'), '40')
    spacing.set(qw('after'), '160')
    ppr.append(spacing)
    jc = OxmlElement('w:jc')
    jc.set(qw('val'), 'center')
    ppr.append(jc)

    r = OxmlElement('w:r')
    p_el.append(r)
    rpr = OxmlElement('w:rPr')
    r.append(rpr)
    rpr.append(OxmlElement('w:b'))
    rpr.append(OxmlElement('w:i'))
    color = OxmlElement('w:color')
    color.set(qw('val'), '000000')
    rpr.append(color)
    sz = OxmlElement('w:sz')
    sz.set(qw('val'), '20')
    rpr.append(sz)
    t = OxmlElement('w:t')
    t.text = texte
    r.append(t)
    return Paragraph(p_el, parent)


def cellule(tc_el, texte, entete=False):
    tcpr = OxmlElement('w:tcPr')
    tc_el.append(tcpr)
    tcw = OxmlElement('w:tcW')
    tcw.set(qw('type'), 'dxa')
    tcpr.append(tcw)  # largeur posee par l'appelant via tcw.set('w', ...)
    if entete:
        shd = OxmlElement('w:shd')
        shd.set(qw('val'), 'clear')
        shd.set(qw('color'), 'auto')
        shd.set(qw('fill'), 'E8E8E8')
        tcpr.append(shd)

    p_el = OxmlElement('w:p')
    tc_el.append(p_el)
    ppr = OxmlElement('w:pPr')
    p_el.append(ppr)
    spacing = OxmlElement('w:spacing')
    spacing.set(qw('after'), '20')
    spacing.set(qw('line'), '240')
    spacing.set(qw('lineRule'), 'auto')
    ppr.append(spacing)

    r = OxmlElement('w:r')
    p_el.append(r)
    rpr = OxmlElement('w:rPr')
    r.append(rpr)
    sz = OxmlElement('w:sz')
    sz.set(qw('val'), '18')
    rpr.append(sz)
    lang = OxmlElement('w:lang')
    lang.set(qw('val'), 'fr-FR')
    rpr.append(lang)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = texte
    r.append(t)
    return tcw


def construire_tableau(ancre_el, doc, lignes, largeurs=LARGEURS):
    tbl = OxmlElement('w:tbl')
    ancre_el.addprevious(tbl)

    tblpr = OxmlElement('w:tblPr')
    tbl.append(tblpr)
    style = OxmlElement('w:tblStyle')
    style.set(qw('val'), 'Grilledutableau')
    tblpr.append(style)
    tblw = OxmlElement('w:tblW')
    tblw.set(qw('w'), '0')
    tblw.set(qw('type'), 'auto')
    tblpr.append(tblw)
    jc = OxmlElement('w:jc')
    jc.set(qw('val'), 'center')
    tblpr.append(jc)
    look = OxmlElement('w:tblLook')
    look.set(qw('val'), '04A0')
    look.set(qw('firstRow'), '1')
    look.set(qw('lastRow'), '0')
    look.set(qw('firstColumn'), '1')
    look.set(qw('lastColumn'), '0')
    look.set(qw('noHBand'), '0')
    look.set(qw('noVBand'), '1')
    tblpr.append(look)

    grid = OxmlElement('w:tblGrid')
    tbl.append(grid)
    for w in largeurs:
        gc = OxmlElement('w:gridCol')
        gc.set(qw('w'), str(w))
        grid.append(gc)

    for i, ligne in enumerate(lignes):
        tr = OxmlElement('w:tr')
        tbl.append(tr)
        for j, valeur in enumerate(ligne):
            tc = OxmlElement('w:tc')
            tr.append(tc)
            tcw_el = cellule(tc, valeur, entete=(i == 0))
            tcw_el.set(qw('w'), str(largeurs[j]))

    return Table(tbl, doc)


def main():
    doc = Document(CIBLE)
    paras = doc.paragraphs
    body = doc.element.body

    # ── Reperage des elements a remplacer : legende 6.4, tableau, legende 6.5,
    #    tableau, jusqu'au paragraphe qui suit immediatement (deja rempli) ──
    els = list(body)
    i_leg64 = i_leg65 = None
    for i, el in enumerate(els):
        if not el.tag.endswith('}p'):
            continue
        txt = ''.join(t.text or '' for t in el.iter(WQ + 't')).strip()
        if txt.startswith('Tableau 6.4'):
            i_leg64 = i
        elif txt.startswith('Tableau 6.5'):
            i_leg65 = i
    assert i_leg64 is not None and i_leg65 is not None, "légendes introuvables"

    def tbl_apres(i):
        for j in range(i + 1, min(i + 3, len(els))):
            if els[j].tag.endswith('}tbl'):
                return j
        return None
    i_tbl64 = tbl_apres(i_leg64)
    i_tbl65 = tbl_apres(i_leg65)
    assert i_tbl64 is not None and i_tbl65 is not None, "tableaux introuvables"
    print("  ancien : légende 6.4=el%d tableau=el%d | légende 6.5=el%d "
          "tableau=el%d" % (i_leg64, i_tbl64, i_leg65, i_tbl65))

    donnees_64 = [
        ("Poste", "Hypothèse retenue et source", "Montant estimé"),
        ("Ressources humaines (stage)",
         "3 mois (03 mai – 03 août 2026), valorisés au tarif d'entrée d'un "
         "développeur web junior à Abidjan, 200 000 FCFA/mois [24]",
         "600 000 FCFA"),
        ("Encadrement académique et technique",
         "Mission institutionnelle de l'encadrant UPB et du maître de "
         "stage AGEROUTE",
         "Non facturé"),
        ("Licences logicielles",
         "Stack entièrement open source : FastAPI, Angular, Flutter, "
         "PostgreSQL/PostGIS, Docker, ONNX Runtime",
         "0 FCFA"),
        ("Hébergement (VPS Systalink)",
         "Entrée de gamme à partir de 3 500 FCFA/mois ; un plan avec "
         "≥ 2 Go de RAM, Docker et PostGIS se situe au-dessus, jusqu'à "
         "40 000-80 000 FCFA/mois pour les offres hautes performances [25], "
         "palier exact à confirmer par devis",
         "3 500 à 80 000 FCFA / mois"),
        ("Nom de domaine .ci",
         "Enregistrement NIC-CI via bureau accrédité, 9 000 à 9 500 "
         "FCFA/an [26]",
         "≈ 9 500 FCFA / an"),
        ("Total investissement initial (hors valorisation du stage)",
         "Licences + premier mois d'hébergement + domaine",
         "13 000 à 89 500 FCFA"),
        ("Coût de fonctionnement annuel récurrent",
         "Hébergement (12 mois) + domaine",
         "51 500 à 969 500 FCFA / an"),
    ]
    donnees_65 = [
        ("Solution", "Modèle de licence publié", "Coût annuel indicatif"),
        ("SI-ENV",
         "Open source, aucune licence récurrente",
         "≈ 0,05 à 1 million FCFA/an (hébergement seul)"),
        ("Enablon CSR Suite",
         "À partir de 500 USD/utilisateur/mois [27]",
         "≈ 20 000 USD/utilisateur/an, soit ≈ 12 millions FCFA"),
        ("Enablon Risk Management",
         "Déploiement entreprise [27]",
         "À partir de 50 000 USD/an, soit ≈ 30 millions FCFA"),
        ("Cority et plateformes EHS comparables",
         "40 à 50 USD/utilisateur/mois + 15 000 à 200 000 USD de mise en "
         "œuvre [27]",
         "Plusieurs dizaines de millions de FCFA la première année"),
    ]

    # ── Suppression des quatre elements (dans l'ordre inverse pour ne pas
    #    invalider les indices) ────────────────────────────────────────────
    for idx in sorted([i_leg64, i_tbl64, i_leg65, i_tbl65], reverse=True):
        el = els[idx]
        el.getparent().remove(el)
    print("  anciens éléments retirés")

    # ── Reconstruction, ancree juste avant le paragraphe qui suivait 6.5 ────
    # (paragraphe "L'écart entre..." ou celui d'après, deja bien forme,
    # retrouve par une nouvelle lecture du corps)
    doc2 = Document(CIBLE)  # relecture impossible ici (pas encore sauve) ;
    # on continue avec le meme objet doc, en reperant l'ancre par contenu.
    ancre_texte = "L'écart entre les deux bornes"
    ancre_el = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(ancre_texte):
            ancre_el = p._element
            break
    assert ancre_el is not None, "paragraphe d'ancrage introuvable"

    parent = doc.paragraphs[0]._parent
    construire_legende(ancre_el, parent,
                       "Tableau 6.4 : Coûts de développement et de "
                       "fonctionnement du SI-ENV, sources publiques (août "
                       "2026).")
    construire_tableau(ancre_el, doc, donnees_64)

    # Ancre pour 6.5 : paragraphe "Ces montants, cités"
    ancre2_el = None
    for p in doc.paragraphs:
        if p.text.strip().startswith('Ces montants, cités'):
            ancre2_el = p._element
            break
    assert ancre2_el is not None, "paragraphe d'ancrage 6.5 introuvable"
    construire_legende(ancre2_el, parent,
                       "Tableau 6.5 : Comparaison indicative avec des "
                       "plateformes commerciales de gestion HSE (août "
                       "2026).")
    construire_tableau(ancre2_el, doc, donnees_65)

    print("  tableaux 6.4 et 6.5 reconstruits avec le style Grilledutableau")

    doc.save(CIBLE)
    print("\nEnregistré : %s" % CIBLE)


if __name__ == "__main__":
    main()
