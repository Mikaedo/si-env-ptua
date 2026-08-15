# -*- coding: utf-8 -*-
"""
Relecture du corps : numerotation, repetitions, passages trop maigres.

Trois natures de defaut ont ete relevees a la lecture suivie des chapitres.

La premiere est une erreur de numerotation : les tableaux du chapitre 3 se
suivent dans l'ordre 3.1, 3.3, 3.2. Un lecteur qui remonte du texte vers le
tableau annonce tombe sur le mauvais, et la liste liminaire aggrave la
confusion puisqu'elle les range dans l'ordre des numeros.

La deuxieme tient a des repetitions. La liste des six dysfonctionnements
figure trois fois presque a l'identique, la problematique deux fois, le cout
du projet deux fois. Repeter n'est pas insister : au troisieme passage, le
lecteur cesse de lire.

La troisieme concerne deux paragraphes reduits a une enumeration seche, qui
ressemblent a des notes que l'on aurait oublie de redirger. Le paragraphe 2.4
en particulier annonce une figure puis se contente de lister six mots.
"""
import re
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_relecture.docx")

journal = []


def reecrire(p, texte):
    for seg in list(p.runs)[1:]:
        seg._element.getparent().remove(seg._element)
    if p.runs:
        p.runs[0].text = texte
    else:
        p.add_run(texte)


# ── Textes de remplacement ────────────────────────────────────────────────

# 1.6 : la problematique reste, la liste des six part. Elle sera dressee au
# chapitre 2, a sa place, apres la figure qui montre le processus actuel.
PROBLEMATIQUE = (
    "Le suivi environnemental du PTUA s'appuie sur un socle réglementaire "
    "solide, rappelé au paragraphe précédent. Le dispositif qui doit le mettre "
    "en œuvre, lui, repose sur des fiches papier et des tableurs. C'est de cet "
    "écart que naît la question à laquelle ce mémoire tente de répondre.\n"
    "Comment l'AGEROUTE peut-elle transformer ce suivi artisanal en un "
    "dispositif fiable, géolocalisé et réactif, conforme aux exigences de la "
    "BAD, sans budget dédié et sans connexion Internet garantie sur les "
    "chantiers ? Le chapitre suivant examine les limites précises de "
    "l'existant, dont dépend la forme que prendra la réponse."
)

# 2.4 : le paragraphe se contentait d'une enumeration de six mots.
EXISTANT = (
    "La figure 2.1 illustre le processus de suivi environnemental actuellement "
    "en vigueur, fondé sur des outils bureautiques disjoints. L'agent constate "
    "une nuisance, la consigne sur une fiche, la transmet en fin de semaine ; "
    "le spécialiste ressaisit ces fiches dans un tableur, puis compose son "
    "rapport trimestriel à partir de fichiers qu'aucun lien ne rattache les uns "
    "aux autres.\n"
    "Six limites découlent de ce fonctionnement. La détection est tardive, un "
    "incident pouvant rester ignoré plusieurs semaines. L'évaluation de la "
    "gravité dépend de l'appréciation de celui qui constate, sans référentiel "
    "commun. Aucune position n'est enregistrée, si bien qu'un signalement ne "
    "peut être rattaché avec certitude à un ouvrage. Les données se dispersent "
    "entre messageries, classeurs et postes de travail. Le rapport réglementaire "
    "se compose à la main, au prix de plusieurs jours de ressaisie. Aucune "
    "alerte, enfin, ne se déclenche lorsqu'un seuil est franchi : il faut que "
    "quelqu'un s'en aperçoive.\n"
    "Ces six limites servent de fil conducteur au reste du mémoire. Le "
    "paragraphe suivant indique ce que le SI-ENV leur oppose, et le tableau de "
    "traçabilité du chapitre 3 rattache chacune d'elles à un besoin fonctionnel "
    "précis."
)

# L'introduction generale et le paragraphe 1.3 citaient tous deux le cout du
# projet avec sa reference. La seconde occurrence s'allege.
COUT_13 = (
    "L'envergure du PTUA a nécessité la mise en place d'une organisation projet "
    "spécifique et rigoureuse, rattachée à la Direction Générale de l'AGEROUTE. "
    "L'organigramme du projet se décline en trois grands niveaux de "
    "responsabilité (figure 1.2)."
)

# Le paludisme est cite dans l'introduction avec ses chiffres ; en 2.6 on
# garde le raisonnement et non la statistique deja donnee.
SANITAIRE = (
    "Les eaux stagnantes sur les chantiers constituent des gîtes larvaires pour "
    "les moustiques vecteurs du paludisme, dont l'introduction a rappelé le "
    "poids en Afrique. C'est cet enjeu qui justifie la possibilité de signaler "
    "manuellement toute zone d'eau stagnante observée sur le terrain (tableau "
    "3.1) : le phénomène est suffisamment visible pour que le Responsable "
    "Environnement ou l'Expert HSE le repère sans recourir à un module dédié "
    "d'intelligence artificielle. Les émissions atmosphériques du chantier "
    "affectent par ailleurs la qualité de l'air ambiant et augmentent les "
    "risques respiratoires pour les riverains ; c'est cet enjeu qui justifie le "
    "suivi du NO2 par télédétection satellitaire (chapitre 5)."
)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    # ── Numerotation des tableaux 3.2 et 3.3 ──────────────────────────────
    # Le tableau des profils apparait avant celui de la tracabilite : il doit
    # donc porter le numero inferieur. On echange les deux, legendes, renvois
    # et liste liminaire compris.
    def echanger(texte):
        texte = texte.replace("Tableau 3.3", "@@TMP@@").replace("tableau 3.3", "@@tmp@@")
        texte = texte.replace("Tableau 3.2", "Tableau 3.3").replace("tableau 3.2", "tableau 3.3")
        return texte.replace("@@TMP@@", "Tableau 3.2").replace("@@tmp@@", "tableau 3.2")

    n = 0
    for p in doc.paragraphs:
        if re.search(r"[Tt]ableau 3\.[23]", p.text):
            reecrire(p, echanger(p.text))
            n += 1
    journal.append(f"numerotation des tableaux 3.2 et 3.3 : {n} mention(s) echangee(s)")

    # ── Repetitions et paragraphes trop maigres ───────────────────────────
    reprises = [
        ("Le suivi environnemental repose sur quatre textes", PROBLEMATIQUE,
         "1.6 : la liste des six sort, la problematique reste"),
        ("La figure 2.1 illustre le processus", EXISTANT,
         "2.4 : enumeration seche remplacee par un texte suivi"),
        ("L'envergure du PTUA, dont le coût global", COUT_13,
         "1.3 : cout du projet, deuxieme occurrence allegee"),
        ("Les eaux stagnantes sur les chantiers constituent", SANITAIRE,
         "2.6 : statistique du paludisme, deuxieme occurrence allegee"),
    ]
    for debut, remplacement, libelle in reprises:
        for p in doc.paragraphs:
            if p.text.strip().startswith(debut):
                reecrire(p, remplacement.replace("\n", "  "))
                journal.append(libelle)
                break
        else:
            journal.append(f"MANQUE : {libelle}")

    # Le paragraphe 2.4 comportait une seconde ligne reduite a six mots, qui
    # fait desormais double emploi avec le texte ci-dessus.
    for p in doc.paragraphs:
        if p.text.strip().startswith("Six limites : détection tardive"):
            reecrire(p, "")
            journal.append("2.4 : ligne residuelle supprimee")
            break

    doc.save(SOURCE)
    for l in journal:
        print("  " + l)
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
