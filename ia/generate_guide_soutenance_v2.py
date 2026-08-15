# -*- coding: utf-8 -*-
"""
Genere la version actualisee du guide de preparation a la soutenance SI-ENV.

Contenu enrichi par rapport a la version precedente :
  - questions/reponses classees du plus simple au plus difficile, avec les
    questions pieges ;
  - toutes les evolutions techniques recentes (invitation par jeton, jeton JWT
    de 12 h, letterbox et TTA, detection temps reel, cache satellitaire
    persistant, rapport PGES aux normes officielles, multi-typage) ;
  - les limites a assumer honnetement, avec la formulation a employer ;
  - le deroule minute par minute de la demonstration et la liste de controle
    technique.

Ecrase le fichier existant dans le dossier MEMOIRE.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SORTIE = r"C:\Users\DELL\Downloads\MEMOIRE\GUIDE_PREPARATION_SOUTENANCE_SI-ENV.docx"

ORANGE = RGBColor(0xE8, 0x6C, 0x00)
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GRAY = RGBColor(0x47, 0x55, 0x69)
GREEN = RGBColor(0x16, 0x7C, 0x3C)
RED = RGBColor(0xB9, 0x1C, 0x1C)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin = sec.bottom_margin = Cm(2.5)


def _police(style, nom='Times New Roman'):
    rpr = style.element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(a), nom)
    rpr.append(rf)


normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
_police(normal)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(8)

for nom, taille, couleur in [('Heading 1', 19, ORANGE),
                             ('Heading 2', 14.5, NAVY),
                             ('Heading 3', 12.5, GRAY)]:
    st = doc.styles[nom]
    st.font.name = 'Times New Roman'
    st.font.size = Pt(taille)
    st.font.bold = True
    st.font.color.rgb = couleur
    _police(st)

st_q = doc.styles.add_style('Question', 1)
st_q.base_style = doc.styles['Normal']
st_q.font.bold = True
st_q.font.color.rgb = NAVY
st_q.font.size = Pt(12)
st_q.paragraph_format.space_before = Pt(10)
st_q.paragraph_format.space_after = Pt(3)

st_c = doc.styles.add_style('PointCle', 1)
st_c.base_style = doc.styles['Normal']
st_c.font.italic = True
st_c.font.color.rgb = ORANGE
st_c.font.size = Pt(11.5)
st_c.paragraph_format.left_indent = Cm(0.5)
st_c.paragraph_format.space_before = Pt(4)
st_c.paragraph_format.space_after = Pt(10)


def h1(t):
    doc.add_heading(level=1).add_run(t)


def h2(t):
    doc.add_heading(level=2).add_run(t)


def h3(t):
    doc.add_heading(level=3).add_run(t)


def para(t, gras=False, italique=False, taille=None, couleur=None, centre=False):
    p = doc.add_paragraph()
    if centre:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t)
    r.bold, r.italic = gras, italique
    if taille:
        r.font.size = Pt(taille)
    if couleur:
        r.font.color.rgb = couleur
    return p


def puce(t):
    doc.add_paragraph(t, style='List Bullet')


def num(t):
    doc.add_paragraph(t, style='List Number')


def cle(t):
    doc.add_paragraph(style='PointCle').add_run(t)


def qr(question, reponse, complement=None):
    """Question du jury suivie de la reponse a donner."""
    doc.add_paragraph(style='Question').add_run("Q. " + question)
    para(reponse)
    if complement:
        cle(complement)


def tableau(entetes, lignes, largeurs=None):
    t = doc.add_table(rows=1, cols=len(entetes))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, e in enumerate(entetes):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(e)
        r.bold = True
        r.font.size = Pt(10.5)
    for ligne in lignes:
        cells = t.add_row().cells
        for i, v in enumerate(ligne):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(10.5)
    if largeurs:
        for ligne in t.rows:
            for i, l in enumerate(largeurs):
                ligne.cells[i].width = Cm(l)
    doc.add_paragraph()
    return t


# ═══════════════════════════════════════════════════════════════════════════
para("GUIDE DE PRÉPARATION À LA SOUTENANCE", gras=True, taille=20,
     couleur=ORANGE, centre=True)
para("SI-ENV — Système d'Information Environnemental du PTUA", gras=True,
     taille=13, couleur=NAVY, centre=True)
para("Questions et réponses, du plus basique au plus difficile — "
     "version actualisée et complète", italique=True, taille=11,
     couleur=GRAY, centre=True)

doc.add_paragraph()
h2("Comment utiliser ce guide")
para("Ce guide n'est pas à apprendre par cœur. Il sert à trois choses : "
     "vérifier que vous savez répondre à chaque question sans hésiter, "
     "connaître les formulations qui évitent les pièges, et savoir quoi "
     "répondre quand vous ne savez pas.")

cle("La règle la plus importante : un jury ne cherche pas à vous prendre en "
    "défaut, il cherche à savoir si vous maîtrisez ce que vous avez produit. "
    "Reconnaître une limite en l'expliquant vaut toujours mieux que de bluffer. "
    "Un « je ne l'ai pas traité, et voici pourquoi » est une réponse "
    "recevable ; une affirmation fausse ne l'est pas.")

h2("Sommaire")
for i, s in enumerate([
    "Le projet en deux minutes (à savoir réciter)",
    "Questions de niveau 1 — comprendre le projet",
    "Questions de niveau 2 — les choix techniques",
    "Questions de niveau 3 — l'intelligence artificielle",
    "Questions de niveau 4 — sécurité, données, architecture",
    "Questions pièges et comment y répondre",
    "Les limites à assumer (et la formulation à employer)",
    "Déroulé de la démonstration, minute par minute",
    "Liste de contrôle technique avant de présenter",
    "Fiche de secours — les chiffres et comptes",
], 1):
    para("%d. %s" % (i, s))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("1. Le projet en deux minutes")
# ═══════════════════════════════════════════════════════════════════════════

para("À préparer mot pour mot, car c'est ce que l'on vous demandera en premier.")

para("« Le PTUA, Projet de Transport Urbain d'Abidjan, est soumis à un Plan de "
     "Gestion Environnementale et Sociale. Ce plan impose un suivi documenté "
     "des nuisances de chantier et un mécanisme de gestion des plaintes. Dans "
     "les faits, ce suivi reposait sur des relevés papier et des tableurs, avec "
     "trois conséquences : des délais de remontée longs, des données non "
     "géolocalisées, et une traçabilité difficile à produire devant "
     "l'ANDE.", italique=True)

para("« SI-ENV répond à ce besoin par trois composants : une application "
     "mobile pour les agents de terrain, qui fonctionne hors connexion et "
     "embarque une détection automatique des déchets ; un tableau de bord web "
     "pour les spécialistes du bureau, avec cartographie et analyse "
     "satellitaire ; et une interface de programmation qui centralise les "
     "données et génère le rapport PGES réglementaire.", italique=True)

para("« L'apport principal n'est pas l'informatisation en elle-même, mais la "
     "traçabilité : chaque signalement est horodaté, géolocalisé, photographié "
     "et associé à un diagnostic automatique. »", italique=True)

cle("Trois mots à placer absolument : traçabilité, hors connexion, "
    "réglementaire. Ce sont eux qui montrent que vous avez compris le besoin "
    "métier et pas seulement la technique.")

h2("Les rôles et qui utilise quoi")

tableau(
    ["Profil", "Outil", "Ce qu'il fait"],
    [
        ["Responsable Environnement", "Mobile",
         "Saisit les signalements sur le chantier (entreprise de travaux)"],
        ["Expert HSE", "Mobile",
         "Contrôle externe quotidien (bureau de contrôle)"],
        ["Spécialiste Suivi Env.", "Web",
         "Instruit les signalements, analyse satellitaire, produit le PGES"],
        ["Spécialiste P.A.R", "Web",
         "Traite les plaintes communautaires (mécanisme MGP)"],
        ["Administrateur", "Web",
         "Gère les comptes, les référentiels et la journalisation"],
    ],
    largeurs=[4.2, 2.2, 9.6],
)

cle("Précision qui fait bonne impression : les deux applications sont "
    "mutuellement exclusives. Un agent de terrain est refusé sur le web, un "
    "profil bureau est refusé sur le mobile. Ce n'est pas une limitation, c'est "
    "une conséquence du modèle de rôles : l'outil suit la fonction.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("2. Questions de niveau 1 — comprendre le projet")
# ═══════════════════════════════════════════════════════════════════════════

qr("Qu'est-ce qu'un PGES ?",
   "Un Plan de Gestion Environnementale et Sociale. C'est le document qui, à "
   "la suite d'une étude d'impact, recense les mesures d'atténuation que le "
   "projet s'engage à appliquer, désigne les responsables et fixe les "
   "indicateurs de suivi. Sa mise en œuvre doit être documentée et rendue "
   "compte à l'autorité environnementale.")

qr("Quel problème concret votre travail résout-il ?",
   "Le suivi environnemental reposait sur des fiches papier ressaisies dans "
   "des tableurs. Trois problèmes en découlaient : le délai entre le constat "
   "et sa remontée, l'absence de localisation fiable des nuisances, et la "
   "difficulté à produire une preuve datée en cas de contrôle. SI-ENV traite "
   "les trois.")

qr("Pourquoi une application mobile plutôt qu'un simple formulaire web ?",
   "Parce que le constat se fait sur le chantier, où la couverture réseau "
   "n'est pas garantie. L'application fonctionne hors connexion : les "
   "signalements sont stockés localement puis synchronisés dès qu'une "
   "connexion est disponible. Un formulaire web exigerait du réseau au moment "
   "précis où il en manque le plus.",
   "C'est l'argument « offline-first ». Il justifie à lui seul le choix du "
   "développement mobile natif plutôt qu'un site adaptatif.")

qr("Quelles technologies avez-vous utilisées, et pourquoi ?",
   "Flutter pour le mobile, car un seul code source couvre Android et iOS et "
   "l'exécution reste native. Angular pour le tableau de bord, pour sa "
   "structure adaptée aux applications de gestion. FastAPI en Python côté "
   "serveur, pour sa rapidité de développement et sa documentation "
   "automatique. PostgreSQL avec l'extension PostGIS, indispensable dès que "
   "l'on manipule des coordonnées et des emprises géographiques.")

qr("Combien de chantiers et de types de nuisance gérez-vous ?",
   "Six axes du programme PTUA sont suivis : le 4e Pont, la rocade Y4, le "
   "boulevard Latrille, les sorties Est et Ouest et les échangeurs. Cinq types "
   "de nuisance sont prévus : déchets de chantier, eaux stagnantes, "
   "poussières, bruit et dégradation de la végétation. Un même signalement "
   "peut cumuler plusieurs types, car une situation réelle est rarement pure.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("3. Questions de niveau 2 — les choix techniques")
# ═══════════════════════════════════════════════════════════════════════════

qr("Comment fonctionne le mode hors connexion ?",
   "Le signalement est d'abord écrit dans une base SQLite locale sur le "
   "téléphone, avec un identifiant unique généré côté mobile. Il est ensuite "
   "transmis au serveur dès que le réseau le permet. Cet identifiant unique "
   "est ce qui garantit qu'un signalement synchronisé deux fois ne crée pas de "
   "doublon.",
   "Le terme à employer est « idempotence » : rejouer la même synchronisation "
   "produit le même résultat, sans duplication.")

qr("Comment les photos remontent-elles du téléphone vers le web ?",
   "Elles sont envoyées au serveur par une requête multipart, stockées sur le "
   "disque du serveur, et leur chemin est enregistré en base. Le tableau de "
   "bord les affiche ensuite dans la fiche du signalement, et elles sont "
   "reprises en annexe du rapport PGES comme preuves visuelles.")

qr("Comment gérez-vous les rôles et les permissions ?",
   "Par un contrôle d'accès fondé sur les rôles. Le jeton d'authentification "
   "porte le rôle de l'utilisateur, et chaque point d'entrée protégé du serveur "
   "vérifie que ce rôle est autorisé. La vérification est faite côté serveur, "
   "pas seulement dans l'interface : masquer un bouton ne protège rien.")

qr("Comment un nouvel utilisateur obtient-il son compte ?",
   "L'administrateur crée le compte en saisissant seulement le nom, "
   "l'adresse et le rôle : il ne choisit aucun mot de passe. Le serveur génère "
   "un jeton d'invitation à usage unique, valable 72 heures, et envoie par "
   "courriel un lien d'activation. L'agent clique, définit lui-même son mot de "
   "passe, et le jeton est immédiatement invalidé.",
   "Point important à souligner : l'administrateur ne connaît jamais le mot de "
   "passe des utilisateurs. C'est un principe de sécurité élémentaire mais "
   "qu'il faut savoir énoncer.")

qr("Que se passe-t-il si la session expire pendant l'utilisation ?",
   "Le jeton a une durée de validité de 12 heures. Lorsque le serveur refuse "
   "une requête pour cause de session expirée, l'application purge la session "
   "locale et ramène l'utilisateur à l'écran de connexion avec un message "
   "explicite, au lieu de laisser des écrans en erreur permanente.")

qr("Comment le rapport PGES est-il produit ?",
   "Le serveur agrège les signalements, alertes et plaintes de la période et "
   "des chantiers demandés, puis compose un PDF. Le document suit les codes "
   "d'un rapport officiel : page de garde, fiche signalétique avec référence "
   "unique, sommaire paginé, sigles, résumé exécutif, cadre juridique, matrice "
   "de suivi, conclusion et annexes.",
   "Détail technique valorisant : le sommaire ne peut pas être paginé en une "
   "seule passe, puisqu'il faut connaître la page de chaque titre avant de "
   "l'écrire. Le document est donc composé deux fois. C'est le genre de "
   "précision qui montre que vous avez compris ce que vous avez utilisé.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("4. Questions de niveau 3 — l'intelligence artificielle")
# ═══════════════════════════════════════════════════════════════════════════

qr("Que fait exactement votre module d'intelligence artificielle ?",
   "Deux choses, dans cet ordre. YOLOv8n détecte où sont les déchets sur la "
   "photo et de quelle matière ils sont, parmi six classes : carton, verre, "
   "métal, organique, papier, plastique. Puis MobileNetV2 reçoit chaque zone "
   "détectée et lui attribue une criticité environnementale : faible, modérée "
   "ou élevée.")

qr("Quels résultats obtenez-vous ?",
   "Une mAP@50 globale de 0,807, soit 80,7 %, avec une précision de 0,797 et "
   "un rappel de 0,717, mesurés sur 247 images de validation représentant 568 "
   "objets annotés. Le métal est la classe la mieux reconnue avec 0,970, le "
   "plastique la plus difficile avec 0,690.")

qr("Pourquoi le plastique est-il moins bien détecté alors qu'il est la classe "
   "la plus fréquente ?",
   "Parce que la difficulté ne vient pas du volume de données mais de la "
   "nature de l'objet. Un sachet plastique est transparent, déformable et sans "
   "forme stable : il épouse son support et se confond avec l'arrière-plan. Le "
   "métal, à l'inverse, a des reflets et des formes régulières très "
   "caractéristiques.",
   "Cette réponse est importante : elle montre que vous interprétez vos "
   "résultats physiquement, et pas seulement statistiquement.")

qr("Le modèle tourne-t-il sur le téléphone ou sur le serveur ?",
   "Sur le téléphone. Les deux modèles sont exportés au format ONNX et "
   "exécutés localement, sans appel réseau. C'est cohérent avec le mode hors "
   "connexion : un agent sans réseau doit tout de même obtenir son diagnostic.")

qr("Comment évitez-vous que l'analyse gèle l'application ?",
   "L'exécution du modèle se fait sur un isolate séparé, c'est-à-dire un fil "
   "d'exécution indépendant avec sa propre mémoire. La bibliothèque d'inférence "
   "expose une fonction synchrone : l'appeler sur le fil de l'interface "
   "bloquerait l'affichage pendant toute la durée du calcul.")

qr("Qu'est-ce que le letterbox et pourquoi l'avoir ajouté ?",
   "Le modèle attend une image carrée de 320 pixels de côté. Or une photo de "
   "téléphone est au format 3:4. La redimensionner directement en carré "
   "l'écrase et déforme les objets. Le letterbox conserve les proportions et "
   "complète les bords par une teinte neutre, exactement comme pendant "
   "l'entraînement. Mesuré sur un échantillon, la confiance sur les vraies "
   "détections passe par exemple de 32 % à 77 % sur une même photo.",
   "Preuve à citer : sur une image déjà carrée, le résultat est rigoureusement "
   "identique avec ou sans letterbox. C'est ce qui démontre que le gain vient "
   "de la correction de la déformation et non du hasard.")

qr("Pourquoi un seuil de confiance de 0,25 et non 0,5 ?",
   "C'est une décision métier, pas un réglage arbitraire. Le rappel du "
   "plastique est de 0,559 : un seuil à 0,5 écarterait de nombreux déchets "
   "réellement présents. En suivi environnemental, manquer une pollution est "
   "plus grave que demander une vérification inutile à l'agent.")

qr("L'utilisateur peut-il corriger le diagnostic de l'IA ?",
   "Oui, et c'est volontaire. La criticité proposée par le modèle est "
   "pré-sélectionnée mais reste modifiable. Les deux valeurs sont enregistrées "
   "séparément : la criticité retenue par l'agent et celle suggérée par l'IA. "
   "Cela permet de mesurer plus tard le taux d'accord entre l'humain et le "
   "modèle.",
   "Argument fort : l'IA assiste la décision, elle ne la remplace pas. C'est "
   "aussi ce qui rend le système acceptable pour les agents.")

qr("D'où viennent vos étiquettes de criticité ? Le corpus ne les contient pas.",
   "Exact, et c'est un point que j'explicite dans le mémoire. Elles ont été "
   "construites par une règle déterministe appliquée aux annotations : le "
   "nombre d'objets sur l'image. Un ou deux objets valent une accumulation "
   "faible, trois à cinq une accumulation modérée, six ou plus une accumulation "
   "importante. L'étiquetage est donc reproductible et vérifiable par "
   "quiconque.",
   "C'est LA question à laquelle il faut être prêt. Répondez-la avant qu'on "
   "vous la pose, en enchaînant sur la question suivante.")

qr("Alors votre classifieur ne fait que compter des objets ?",
   "C'est la conséquence lucide de cette définition, et je l'assume. Séparer "
   "cinq objets de six n'a aucune signature visuelle globale : la tâche relève "
   "du dénombrement, ce pour quoi un classifieur d'image entière n'est pas "
   "conçu. Cela borne par construction la performance sur les classes "
   "intermédiaires, et c'est précisément ce que montrent les F1 de 0,59 sur "
   "modérée et 0,67 sur importante.",
   "En reconnaissant la cause structurelle, vous transformez un mauvais chiffre "
   "en analyse. C'est très différent de le subir.")

qr("Que faudrait-il faire alors ?",
   "Trois choses, que j'ai portées en perspectives. D'abord dériver la "
   "criticité du comptage des objets retournés par le détecteur plutôt que "
   "d'entraîner un second réseau : c'est plus économe et exactement aussi "
   "fidèle à la définition retenue. Ensuite ramener la décision à sa forme "
   "utile, qui est binaire : faut-il intervenir ou non. Enfin ajouter une "
   "option de rejet, le modèle s'abstenant sous un seuil de confiance pour "
   "laisser l'agent trancher.",
   "L'option de rejet est l'argument le plus fort : elle convertit une "
   "performance moyenne en engagement exploitable, une justesse élevée sur une "
   "fraction assumée des cas.")

qr("Votre module de classification sert-il vraiment à quelque chose ?",
   "Il démontre que l'inférence embarquée fonctionne hors connexion sur un "
   "téléphone, ce qui était l'enjeu technique. Mais je ne le présente pas comme "
   "un instrument de mesure de la gravité : la criticité enregistrée reste "
   "celle que l'agent déclare, celle de l'IA n'étant conservée qu'à côté, pour "
   "pouvoir mesurer plus tard le taux d'accord. C'est une démonstration de "
   "faisabilité, et l'IA est portée en perspectives comme un levier à "
   "développer.",
   "Ne défendez pas la performance, défendez le positionnement. C'est la "
   "réponse qui vous fait gagner des points.")

qr("Pourquoi ne pas dépasser 90 % ?",
   "Il faudrait réentraîner avec une résolution de 640 pixels au lieu de 320, "
   "une architecture plus grande que la version nano, et davantage d'époques. "
   "Chacun de ces leviers alourdit l'inférence sur téléphone. Le choix "
   "assumé est de privilégier un modèle utilisable sur le terrain à un score "
   "maximal en laboratoire.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("5. Questions de niveau 4 — sécurité, données, architecture")
# ═══════════════════════════════════════════════════════════════════════════

qr("Comment les mots de passe sont-ils stockés ?",
   "Ils ne sont jamais stockés. Seule une empreinte calculée avec bcrypt est "
   "enregistrée. Bcrypt est conçu pour être lent et intègre un sel, ce qui "
   "rend les attaques par dictionnaire coûteuses. Vérifier un mot de passe "
   "consiste à recalculer l'empreinte et à comparer.")

qr("Pourquoi PostGIS et pas simplement deux colonnes latitude et longitude ?",
   "Parce que PostGIS permet des requêtes spatiales : savoir si un point est "
   "dans l'emprise d'un chantier, calculer des distances, agréger par zone. "
   "Avec deux colonnes numériques, il faudrait reprogrammer cette géométrie à "
   "la main, avec des erreurs assurées sur les projections.")

qr("Comment garantissez-vous que la donnée de terrain est fiable ?",
   "Par trois éléments joints automatiquement au signalement : l'horodatage "
   "serveur, les coordonnées relevées par le téléphone, et la photo. Le mode "
   "d'acquisition de la position est également enregistré, automatique ou "
   "manuel, ce qui permet de distinguer une position mesurée d'une position "
   "saisie.")

qr("Comment fonctionne l'analyse satellitaire ?",
   "Elle s'appuie sur Google Earth Engine et calcule quatre indices sur "
   "l'emprise des chantiers : le NO₂ pour la qualité de l'air à partir de "
   "Sentinel-5P, le NDVI pour le couvert végétal et le NDWI pour l'humidité "
   "des sols à partir de Sentinel-2, et un indice de risque d'érosion "
   "combinant pluviométrie CHIRPS et relief SRTM.")

qr("Vos données satellitaires sont-elles en temps réel ?",
   "Non, et il serait faux de le prétendre. Sentinel-2 repasse au même endroit "
   "tous les cinq jours environ, et la couverture nuageuse tropicale impose de "
   "travailler sur des composites de plusieurs jours. Un cache de sept jours "
   "est donc utilisé, ce qui est cohérent avec la nature même de la donnée.",
   "Cette réponse est un test d'honnêteté scientifique. Répondre « oui, c'est "
   "en temps réel » serait une faute que le jury relèverait immédiatement.")

qr("Qu'est-ce que le NDWI, et quel est le piège ?",
   "C'est un indice d'humidité de surface. Le piège est que, contrairement aux "
   "trois autres, un niveau BAS constitue l'alerte : il signale une sécheresse "
   "et un stress hydrique, pas un excès d'eau.")

qr("Comment votre architecture pourrait-elle passer à l'échelle ?",
   "Trois axes. D'abord la base : les requêtes spatiales sont déjà indexées "
   "par PostGIS. Ensuite le serveur : l'interface de programmation est sans "
   "état, l'authentification reposant sur un jeton, donc plusieurs instances "
   "peuvent être placées derrière un répartiteur de charge. Enfin les fichiers : "
   "le stockage des photos sur disque local devrait être remplacé par un "
   "service de stockage objet.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("6. Questions pièges et comment y répondre")
# ═══════════════════════════════════════════════════════════════════════════

qr("Votre IA est-elle vraiment utile, ou juste un effet de mode ?",
   "Elle répond à un besoin précis : homogénéiser l'appréciation de la "
   "criticité. Deux agents différents face au même tas de déchets ne "
   "l'évaluent pas identiquement. Le modèle propose une référence commune, que "
   "l'agent peut corriger. Le gain n'est pas la vitesse mais la cohérence des "
   "données dans le temps.")

qr("Que se passe-t-il si l'IA se trompe ?",
   "Rien d'irréversible, car elle ne décide pas. Elle pré-remplit une "
   "proposition que l'agent valide ou corrige, et les deux valeurs sont "
   "conservées. Une erreur du modèle est donc visible et rattrapable, jamais "
   "propagée silencieusement.")

qr("Avez-vous testé votre application avec de vrais utilisateurs ?",
   "À dire honnêtement selon votre situation réelle. Si ce n'est pas le cas : "
   "« Non, la validation a porté sur le fonctionnement technique et sur des "
   "jeux de données réels, pas sur une expérimentation terrain avec les agents. "
   "C'est la suite logique du travail, et cela suppose un accord "
   "institutionnel. »",
   "Ne jamais inventer une expérimentation qui n'a pas eu lieu. Une limite "
   "assumée est recevable, une affirmation invérifiable ne l'est pas.")

qr("Pourquoi ne pas avoir utilisé un service d'IA existant, dans le nuage ?",
   "Parce que le besoin est justement de fonctionner sans réseau. Un service "
   "distant exigerait une connexion au moment précis où elle manque. "
   "L'embarquement du modèle sur le téléphone est une conséquence directe de la "
   "contrainte de terrain, pas une préférence technique.")

qr("Votre système ne fait-il pas doublon avec les rapports du bureau de "
   "contrôle ?",
   "Non, il les alimente. Le bureau de contrôle produit un avis d'expert ; "
   "SI-ENV fournit la donnée horodatée et géolocalisée sur laquelle cet avis "
   "s'appuie. Le rapport PGES généré est un document de synthèse, pas un "
   "substitut au contrôle externe.")

qr("Qu'est-ce qui, dans ce travail, est vraiment de vous ?",
   "Question fréquente et légitime. Répondez par la conception : le modèle de "
   "données, l'architecture en trois composants, le choix de l'offline-first, "
   "l'intégration de la chaîne de détection sur téléphone, la structure du "
   "rapport réglementaire. Les briques technologiques sont des outils "
   "existants ; l'assemblage au service d'un besoin PGES est le travail "
   "propre.")

qr("Si je vous demande de créer un utilisateur maintenant, cela fonctionne-t-il ?",
   "Oui. Vous devez pouvoir le faire en direct : création du compte par "
   "l'administrateur, réception du lien d'invitation, définition du mot de "
   "passe, puis connexion. Entraînez-vous à cet enchaînement, c'est une "
   "demande de démonstration très probable.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("7. Les limites à assumer")
# ═══════════════════════════════════════════════════════════════════════════

para("Annoncer soi-même ses limites désarme la critique et démontre du recul. "
     "Voici celles à connaître, avec la formulation à employer.")

tableau(
    ["Limite", "Formulation à employer"],
    [
        ["Performance du modèle plafonnée à 80,7 %",
         "« C'est le compromis d'un modèle nano à 320 pixels, choisi pour rester "
         "utilisable sur un téléphone de terrain. »"],
        ["Plastique souple mal détecté (0,690)",
         "« Cas difficile reconnu : matière transparente et déformable. "
         "L'amélioration passerait par un enrichissement ciblé du jeu de "
         "données sur cette classe. »"],
        ["Données satellitaires non temps réel",
         "« Contrainte physique des satellites et de la couverture nuageuse : "
         "on travaille sur des composites de plusieurs jours. »"],
        ["Photos stockées sur le disque du serveur",
         "« Suffisant pour un déploiement unique, à remplacer par un stockage "
         "objet en production pour la résistance aux pannes. »"],
        ["Pas d'expérimentation terrain formalisée",
         "« La validation a porté sur la technique et sur des données réelles ; "
         "l'expérimentation avec les agents est l'étape suivante. »"],
        ["Validation en environnement local, pas en production",
         "« La suite de 32 tests automatisés passe sur la pile Docker réelle, "
         "mais le déploiement sur le serveur cible de l'AGEROUTE reste à "
         "faire. »"],
        ["Criticité IA peu fiable sur les classes intermédiaires",
         "« Conséquence assumée de la définition de l'étiquette par comptage "
         "d'objets : je l'explique et je propose la décision binaire et "
         "l'option de rejet en perspectives. »"],
    ],
    largeurs=[6.0, 10.0],
)

cle("Si l'on vous demande une limite et que vous répondez « je n'en vois pas », "
    "vous perdez immédiatement en crédibilité. Ayez toujours deux ou trois "
    "limites prêtes, formulées comme des perspectives et non comme des échecs.")

h2("Ce que vous pouvez présenter comme démarche de qualité")

para("Plusieurs défauts ont été détectés puis corrigés au cours du "
     "développement. Savoir les raconter est un atout, car cela prouve une "
     "démarche de vérification et non une confiance aveugle.")

puce("Le prétraitement des images déformait les photos avant de les donner au "
     "modèle : corrigé par le letterbox, avec un gain mesuré.")
puce("La suppression des détections redondantes était globale au lieu d'être "
     "faite matière par matière, ce qui faisait disparaître un déchet lorsqu'un "
     "autre le recouvrait.")
puce("Deux statuts de signalement n'étaient comptés dans aucun indicateur du "
     "tableau de bord : la somme des indicateurs ne retombait pas sur le total.")
puce("Le calcul de l'évolution mensuelle reculait par tranches de 30 jours, ce "
     "qui pouvait faire disparaître un mois de la courbe.")
puce("Les coordonnées géographiques n'étaient pas transmises à l'interface web "
     "faute d'une dépendance optionnelle absente, l'échec étant masqué "
     "silencieusement.")

cle("Formulation utile : « j'ai systématiquement vérifié les sorties plutôt "
    "que de faire confiance au code, et cela m'a permis de détecter plusieurs "
    "défauts que les tests fonctionnels n'auraient pas révélés. »")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("8. Déroulé de la démonstration")
# ═══════════════════════════════════════════════════════════════════════════

para("Une démonstration réussie raconte une histoire : un agent constate une "
     "nuisance, elle remonte, elle est instruite, elle devient un rapport "
     "réglementaire. Suivez cet ordre, il est logique et facile à commenter.")

tableau(
    ["Temps", "Étape", "À dire pendant que ça tourne"],
    [
        ["0:00", "Connexion mobile en tant que Responsable Environnement",
         "« Voici l'outil de l'agent de terrain »"],
        ["0:30", "Nouveau signalement, ouverture de la caméra",
         "« Les cadres verts apparaissent en direct, avant même la prise de "
         "vue : le modèle tourne sur le téléphone »"],
        ["1:15", "Prise de vue, diagnostic automatique",
         "« La matière et la criticité sont proposées ; l'agent peut corriger »"],
        ["1:45", "Ajout d'une seconde photo, sélection de plusieurs types",
         "« Le diagnostic est recalculé sur l'ensemble des photos »"],
        ["2:15", "Enregistrement du signalement",
         "« Il est d'abord écrit localement, donc cela fonctionne sans réseau »"],
        ["2:45", "Passage au web, connexion Spécialiste Suivi Env.",
         "« Changement de profil et donc d'outil »"],
        ["3:00", "Tableau de bord : indicateurs, carte, tracés PTUA",
         "« Le signalement vient d'apparaître, géolocalisé »"],
        ["3:45", "Clic sur un indicateur",
         "« Chaque chiffre est cliquable et ouvre les dossiers concernés »"],
        ["4:15", "Fiche du signalement avec la photo",
         "« La photo prise sur le terrain est consultable ici »"],
        ["4:45", "Analyse satellitaire",
         "« Quatre indices, avec une synthèse en clair et l'action attendue »"],
        ["5:30", "Génération du rapport PGES",
         "« Sommaire, sigles, cadre juridique, matrice, annexes photos »"],
        ["6:15", "Création d'un utilisateur par l'administrateur",
         "« Invitation par lien, l'administrateur ne connaît pas le mot de "
         "passe »"],
    ],
    largeurs=[1.6, 6.4, 8.0],
)

cle("Ne commentez jamais un écran vide en attendant qu'il charge. Préparez une "
    "phrase pour chaque temps de latence, ou ouvrez la page à l'avance dans un "
    "autre onglet.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("9. Liste de contrôle technique avant de présenter")
# ═══════════════════════════════════════════════════════════════════════════

para("À dérouler dans cet ordre, une heure avant la soutenance.")

num("Démarrer les conteneurs du serveur et attendre que la documentation de "
    "l'interface de programmation réponde.")
num("Vérifier la connexion avec un compte web et un compte mobile : les deux "
    "doivent aboutir.")
num("Lancer le tableau de bord et ouvrir chaque page une fois, pour amorcer "
    "les caches — en particulier la page satellitaire, dont le premier appel "
    "peut demander une trentaine de secondes.")
num("Brancher le téléphone, vérifier qu'il est bien reconnu, puis établir la "
    "redirection de port. C'est le point le plus fragile de l'installation.")
num("Ouvrir l'application mobile et se connecter une fois.")
num("Faire un signalement de test complet, avec photo, et vérifier qu'il "
    "apparaît côté web.")
num("Générer un rapport PGES et l'ouvrir, pour vérifier qu'il contient bien le "
    "sommaire et les annexes.")
num("Supprimer les données de test créées pendant cette vérification.")

cle("Le point de fragilité identifié est le câble. Si le téléphone est "
    "débranché ou déplacé, la redirection de port est perdue et l'application "
    "affiche une erreur de chargement. Le réflexe à avoir : rétablir la "
    "redirection, sans paniquer ni relancer toute l'application.")

h2("En cas de problème pendant la démonstration")

tableau(
    ["Symptôme", "Cause la plus probable", "Réflexe"],
    [
        ["Mobile : erreur de chargement",
         "Redirection de port perdue (câble déplacé)",
         "Rétablir la redirection"],
        ["Web : session expirée",
         "Jeton arrivé à échéance",
         "Se reconnecter, la durée est de 12 heures"],
        ["Page satellitaire en chargement",
         "Cache vide après un redémarrage",
         "Patienter, ou l'avoir ouverte avant"],
        ["Aucune détection sur une photo",
         "L'objet n'appartient pas aux six matières",
         "Viser un déchet réel : bouteille, carton, canette"],
    ],
    largeurs=[4.5, 5.5, 6.0],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("10. Fiche de secours")
# ═══════════════════════════════════════════════════════════════════════════

h2("Les chiffres à connaître par cœur")

tableau(
    ["Grandeur", "Valeur"],
    [
        ["mAP@50 du modèle de détection", "0,807 (80,7 %)"],
        ["Précision / rappel globaux", "0,797 / 0,717"],
        ["Meilleure classe / la plus difficile", "metal 0,970 / plastic 0,690"],
        ["Jeu de validation", "247 images, 568 objets annotés"],
        ["Nombre de classes de matières", "6"],
        ["Résolution d'inférence", "320 × 320 pixels"],
        ["Seuil de confiance retenu", "0,25"],
        ["Durée de validité du jeton d'accès", "12 heures"],
        ["Validité du lien d'invitation", "72 heures, usage unique"],
        ["Indices satellitaires suivis", "4 (NO₂, NDVI, NDWI, risque pluie)"],
        ["Profils utilisateurs", "5"],
        ["Types de nuisance", "5, cumulables"],
        ["Chantiers PTUA suivis", "6"],
        ["Tables de la base", "11"],
        ["Tests fonctionnels automatisés", "32, tous passants"],
        ["Latence classification (médiane / 95e centile)", "8,6 ms / 17,5 ms"],
        ["Latence détection (médiane / 95e centile)", "23,8 ms / 48,9 ms"],
        ["Limitation de débit de l'API", "100 requêtes/minute/IP"],
        ["Rétention du journal d'audit", "30 jours"],
    ],
    largeurs=[8.5, 7.5],
)

h2("Comptes de démonstration")

tableau(
    ["Application", "Identifiant", "Mot de passe", "Profil"],
    [
        ["Web", "admin@sienv.ci", "admin123", "Administrateur"],
        ["Web", "spec.env@ageroute.ci", "spec123", "Spécialiste Suivi Env."],
        ["Web", "spec.par@ageroute.ci", "spec123", "Spécialiste P.A.R"],
        ["Mobile", "resp.env@ageroute.ci", "env123", "Resp. Environnement"],
        ["Mobile", "expert.hse@ageroute.ci", "expert123", "Expert HSE"],
    ],
    largeurs=[2.6, 5.4, 3.0, 5.0],
)

cle("Rappel : un compte web est refusé sur le mobile et inversement. Si une "
    "connexion échoue en démonstration, vérifiez d'abord que vous utilisez le "
    "compte correspondant à l'application.")

h2("Les cinq phrases qui sauvent")

puce("« C'est une limite que j'assume, et voici pourquoi ce choix a été fait. »")
puce("« Je n'ai pas traité ce point, mais voici comment je m'y prendrais. »")
puce("« Cette valeur provient d'une mesure sur l'application, je peux vous "
     "montrer comment elle est obtenue. »")
puce("« L'IA propose, l'agent décide : les deux valeurs sont conservées. »")
puce("« Le besoin métier était la traçabilité ; la technique en découle. »")
puce("« Ce module démontre la faisabilité ; je porte son industrialisation en "
     "perspectives, pas comme un résultat acquis. »")

doc.add_paragraph()
para("Bonne soutenance.", gras=True, couleur=NAVY, centre=True)
para("Fin du document — SI-ENV / PTUA", italique=True, taille=10,
     couleur=GRAY, centre=True)

doc.save(SORTIE)
print("Enregistre :", SORTIE)
print("Paragraphes :", len(doc.paragraphs), "| Tableaux :", len(doc.tables))
