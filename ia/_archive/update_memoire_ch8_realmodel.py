# -*- coding: utf-8 -*-
"""Met a jour le chapitre 8 et la discussion 10.5 avec les VRAIS resultats
du reentrainement Colab du 3 aout 2026 (dataset Recycle Trash reel, 6
classes, GPU T4, 100 epochs, imgsz=320), et supprime la reserve
"modele mono-classe pas encore integre" puisque le modele est desormais
reellement embarque dans l'application mobile (assets/models/detection_yolov8n.onnx)."""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

SRC = r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_N'GUESSAN_v6.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_N'GUESSAN_v7.docx"
FIG_DIR = r"D:\etude_soutenance\SI-ENV\ia\captures_annexes"

doc = Document(SRC)
paras = doc.paragraphs

def fix_para(p, new_text):
    runs = p.runs
    assert runs, f"paragraphe sans run : {p.text!r}"
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''

def replace_image(pidx, image_path, width_inches=6.1):
    p = paras[pidx]
    for r in p.runs:
        drawings = r.element.findall(qn('w:drawing'))
        if drawings:
            for d in drawings:
                d.getparent().remove(d)
            r.add_picture(image_path, width=Inches(width_inches))
            print(f"[OK] Image remplacee au paragraphe {pidx}")
            return
    raise RuntimeError(f"Aucune image trouvee au paragraphe {pidx}")

# --- 1) Legendes Figure 8.3 (mAP global 0,798 -> 0,807), 2 occurrences ---
for i in (107, 483):
    old = paras[i].text
    assert '0,798' in old, f"p{i}: {old!r}"
    fix_para(paras[i], old.replace('0,798', '0,807'))
print("[OK] Legendes Figure 8.3 mises a jour (2x)")

# --- 2) Paragraphe discussion mAP (0,798 -> 0,807, 79,8% -> 80,7%) ---
p489 = paras[489].text
assert '0,798' in p489
new489 = p489.replace('0,798, soit 79,8 %', '0,807, soit 80,7 %')
fix_para(paras[489], new489)
print("[OK] p489 (comparaison mAP) mis a jour")

# --- 3) Rappel plastique 0,504 -> 0,559 ---
p494 = paras[494].text
assert '0,504' in p494
fix_para(paras[494], p494.replace('Rappel = 0,504', 'Rappel = 0,559'))
print("[OK] p494 (faux negatifs plastique) mis a jour")

# --- 4) Paragraphe 534 : suppression de la reserve mono-classe ---
p534 = paras[534].text
assert 'mono-classe' in p534
old_tail = ("Le mod\xe8le de d\xe9tection embarqu\xe9 dans l'application mobile est actuellement une "
            "version d'entra\xeenement rapide mono-classe ; le mod\xe8le Recycle Trash \xe0 six classes "
            "valid\xe9 au chapitre 8 (GPU, mAP@0.5 = 0,798) n'a pas encore \xe9t\xe9 r\xe9export\xe9 et "
            "int\xe9gr\xe9 \xe0 l'application (cf. discussion 10.5).")
new_tail = ("Le mod\xe8le de d\xe9tection embarqu\xe9 dans l'application mobile est d\xe9sormais le mod\xe8le "
            "Recycle Trash \xe0 six classes valid\xe9 au chapitre 8, r\xe9entra\xeen\xe9 le 3 ao\xfbt 2026 "
            "(GPU T4, 100 epochs, mAP@0.5 = 0,807), export\xe9 au format ONNX et int\xe9gr\xe9 \xe0 "
            "l'application : le pipeline de d\xe9tection d\xe9ploy\xe9 correspond donc d\xe9sormais "
            "exactement au mod\xe8le valid\xe9 au chapitre 8.")
assert old_tail in p534, "texte mono-classe non trouve tel quel dans p534"
fix_para(paras[534], p534.replace(old_tail, new_tail))
print("[OK] p534 (reserve performance) mise a jour")

# --- 5) Paragraphe 537 : retirer la limite "mono-classe" de la liste ---
p537 = paras[537].text
old_clause = ("mod\xe8le de d\xe9tection six classes valid\xe9 au chapitre 8 pas encore r\xe9export\xe9 "
              "vers l'application mobile (version embarqu\xe9e : mod\xe8le mono-classe d'entra\xeenement "
              "rapide), ")
assert old_clause in p537, "clause mono-classe non trouvee dans p537"
fix_para(paras[537], p537.replace(old_clause, ""))
print("[OK] p537 (liste des limites) mise a jour")

# --- 6) Table 8.2 : ligne YOLOv8n ---
found_table = False
for tbl in doc.tables:
    for row in tbl.rows:
        cells = row.cells
        if cells and cells[0].text.strip() == 'YOLOv8n':
            cells[1].text = '0,807'
            cells[2].text = '0,797'
            cells[3].text = '0,717'
            cells[4].text = '0,755'
            found_table = True
            print("[OK] Tableau 8.2 (ligne YOLOv8n) mis a jour")
assert found_table

# --- 7) Remplacement des figures 8.1, 8.2, 8.3 ---
replace_image(478, FIG_DIR + r"\yolo_results_v2.png", width_inches=6.3)
replace_image(480, FIG_DIR + r"\yolo_confusion_matrix_v2.png", width_inches=5.6)
replace_image(482, FIG_DIR + r"\yolo_PR_curve_v2.png", width_inches=5.8)

doc.save(DST)
print(f"\n=== SAUVEGARDE : {DST} ===")
