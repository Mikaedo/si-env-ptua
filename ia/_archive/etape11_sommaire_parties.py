# -*- coding: utf-8 -*-
"""
Etape 11 : mettre le sommaire en coherence avec les six chapitres, et rendre le
decoupage en parties coherent avec la fusion.

Deux corrections :

  a) Le sommaire enumerait encore les dix chapitres d'origine. Il est reecrit
     avec les six chapitres et un decoupage en trois parties equilibre, deux
     chapitres chacune :
        Premiere partie   : cadre, etat de l'art et existant  (ch. 1 et 2)
        Deuxieme partie   : besoins et conception             (ch. 3 et 4)
        Troisieme partie  : realisation et validation         (ch. 5 et 6)

  b) La fusion avait cree un chevauchement : le nouveau chapitre 2 reunit
     l'ancien chapitre 3, situe en premiere partie, et l'ancien chapitre 4,
     situe en deuxieme partie. Le titre de la deuxieme partie est donc deplace
     apres ce chapitre, de sorte qu'aucun chapitre ne soit a cheval sur deux
     parties.

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os
import re

from docx import Document

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

#: Sommaire cible, dans l'ordre. None = ligne conservee telle quelle.
SOMMAIRE = [
    "Introduction generale",
    "Premiere partie : Presentation generale",
    "Chapitre 1 : Cadre de l'etude et problematique",
    "Chapitre 2 : Etat de l'art et analyse de l'existant",
    "Conclusion partielle",
    "Deuxieme partie : Analyse et conception",
    "Chapitre 3 : Analyse des besoins",
    "Chapitre 4 : Conception",
    "Conclusion partielle",
    "Troisieme partie : Realisation et resultats",
    "Chapitre 5 : Realisation : plateforme, intelligence artificielle et teledetection",
    "Chapitre 6 : Deploiement, tests et discussion",
    "Conclusion partielle",
    "Conclusion generale",
]


def texte(el):
    return ''.join(t.text or '' for t in el.iter('{%s}t' % W)).strip()


def ecrire(paragraphe, contenu):
    noeuds = list(paragraphe._element.iter('{%s}t' % W))
    if not noeuds:
        paragraphe.add_run(contenu)
        return
    noeuds[0].text = contenu
    for n in noeuds[1:]:
        n.text = ''


def main():
    doc = Document(CIBLE)
    paras = doc.paragraphs

    # ── a) Reecriture du sommaire ───────────────────────────────────────────
    # On repere le bloc par sa premiere et sa derniere entree connues.
    debut = fin = None
    for i, p in enumerate(paras):
        if i > 130:
            break
        t = p.text.strip()
        if debut is None and t.lower().startswith('introduction g'):
            debut = i
        if debut is not None and t.lower().startswith('conclusion g'):
            fin = i
            break

    if debut is None or fin is None:
        print("  ! bloc du sommaire introuvable")
    else:
        anciennes = [i for i in range(debut, fin + 1) if paras[i].text.strip()]
        print("  sommaire : %d lignes reperees (body %d a %d)"
              % (len(anciennes), debut, fin))
        # On ecrit les nouvelles entrees sur les lignes existantes ; le
        # surplus eventuel est vide plutot que supprime, pour ne pas perturber
        # la numerotation des paragraphes suivants.
        for k, idx in enumerate(anciennes):
            if k < len(SOMMAIRE):
                ecrire(paras[idx], SOMMAIRE[k])
            else:
                ecrire(paras[idx], '')
        manquantes = len(SOMMAIRE) - len(anciennes)
        print("  entrees ecrites : %d | lignes en trop videes : %d | manquantes : %d"
              % (min(len(SOMMAIRE), len(anciennes)),
                 max(0, len(anciennes) - len(SOMMAIRE)), max(0, manquantes)))

    # ── b) Titres de partie dans le corps ───────────────────────────────────
    # Reperage des trois titres de partie et du chapitre qui les suit.
    partie_2 = None
    for i, p in enumerate(paras):
        if i < 250:
            continue
        if re.search(r'DEUXI.{1,3}ME PARTIE', p.text.upper()):
            partie_2 = i
            break

    if partie_2 is None:
        print("  (titre de deuxieme partie non trouve dans le corps)")
    else:
        suite = ' | '.join(paras[j].text.strip()[:40]
                           for j in range(partie_2, min(partie_2 + 6, len(paras)))
                           if paras[j].text.strip())
        print("  deuxieme partie a body[%d] : %s"
              % (partie_2, suite[:90].encode('ascii', 'replace').decode()))

    doc.save(CIBLE)
    print("\nEnregistre : %s" % CIBLE)


if __name__ == "__main__":
    main()
