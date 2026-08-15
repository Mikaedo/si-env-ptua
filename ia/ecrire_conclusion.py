# -*- coding: utf-8 -*-
"""
Ecrit la conclusion generale depuis un fichier texte.

Le passage precedent avait tente de la composer dans une ligne de commande, ou
les accents ne survivent pas a l'encodage du shell : le memoire s'etait
retrouve avec « memoire », « reglementaire » et « horodatees ». Le texte est
donc lu depuis un fichier en UTF-8, seule facon fiable de faire transiter du
francais accentue.

La conclusion est resserree pour tenir sur une page. C'est souvent la seule
que le jury relit avant sa premiere question : elle gagne a se lire d'un trait,
sans que le regard ait a tourner.
"""
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_conclusion_accents.docx")
TEXTE = Path(r"C:\Users\DELL\AppData\Local\Temp\claude"
             r"\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25"
             r"\scratchpad\conclusion.txt")


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    paragraphes = [l.strip() for l in
                   TEXTE.read_text(encoding="utf-8").split("\n") if l.strip()]
    print(f"{len(paragraphes)} paragraphes, "
          f"{sum(len(p) for p in paragraphes)} caracteres")

    # Les paragraphes existants de la conclusion sont recuperes puis reecrits :
    # cela preserve leur style, et les surnumeraires sont vides plutot que
    # supprimes, pour ne pas toucher a la structure du document.
    cibles = []
    dans = False
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading 1"):
            dans = p.text.strip() == "Conclusion générale"
            continue
        if dans:
            cibles.append(p)

    if len(cibles) < len(paragraphes):
        print(f"  ATTENTION : {len(cibles)} paragraphes disponibles pour "
              f"{len(paragraphes)} a ecrire")

    for i, cible in enumerate(cibles):
        texte = paragraphes[i] if i < len(paragraphes) else ""
        for seg in list(cible.runs)[1:]:
            seg._element.getparent().remove(seg._element)
        if cible.runs:
            cible.runs[0].text = texte
        elif texte:
            cible.add_run(texte)

    doc.save(SOURCE)

    # Verification : les accents ont-ils bien traverse ?
    controle = Document(SOURCE)
    dans, echantillon = False, ""
    for p in controle.paragraphs:
        if p.style.name.startswith("Heading 1"):
            dans = p.text.strip() == "Conclusion générale"
            continue
        if dans and p.text.strip():
            echantillon = p.text.strip()[:90]
            break
    print(f"\npremier paragraphe : {echantillon}")
    accents = sum(echantillon.count(c) for c in "éèêàçôûîœ")
    print(f"caracteres accentues dans l'echantillon : {accents}")


if __name__ == "__main__":
    main()
