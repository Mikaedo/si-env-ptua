# -*- coding: utf-8 -*-
"""
Corrige le tableau 6.4 et remet la bibliographie d'aplomb.

Trois corrections, dont deux touchent a l'exactitude.

La ligne des ressources humaines valorisait le stage au tarif du marche, deux
cent mille francs par mois sur trois mois. C'est une convention comptable
defendable, mais elle ne dit pas ce qui a ete percu : la prime de stage s'est
elevee a cent cinquante mille francs pour la periode. Le chiffre reel remplace
l'estimation, et la source salariale qui la justifiait n'a plus d'objet.

L'hebergement designait un prestataire nomme comme « cible retenue » alors
qu'aucun devis n'a ete demande ni aucun engagement pris. La formule laissait
croire a un choix arrete, voire a un service deja souscrit. La ligne devient
une projection explicite, le prestataire n'y figurant plus que comme reference
de prix, ce qu'il a toujours ete.

Retirer une entree de bibliographie oblige enfin a renumeroter celles qui la
suivent, faute de quoi les appels du corps pointeraient sur la mauvaise
source. Le decalage est applique en ordre croissant : commencer par la fin
ecraserait un numero encore utilise.
"""
import re
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_finances.docx")

# Cellules du tableau 6.4 : (ligne, colonne) -> nouveau texte.
CELLULES = {
    (1, 1): "Prime de stage effectivement perçue sur les trois mois "
            "(03 mai – 03 août 2026)",
    (1, 2): "150 000 FCFA",
    (4, 0): "Hébergement de production (projection)",
    (4, 1): "Aucun engagement pris à ce jour : les montants sont relevés chez "
            "un hébergeur VPS ivoirien pris comme référence de prix [24]. "
            "L'entrée de gamme démarre à 3 500 FCFA/mois ; un plan disposant "
            "d'au moins 2 Go de mémoire, de Docker et de PostGIS se situe "
            "au-dessus, jusqu'à 40 000 à 80 000 FCFA/mois pour les offres "
            "hautes performances. Le palier exact reste à établir par devis.",
    (7, 0): "Total investissement initial (hors prime de stage)",
    (8, 1): "Environnement de validation (Render, Supabase, Cloudflare) : "
            "gratuit. Projection de production : hébergement sur douze mois "
            "et nom de domaine.",
}

# Ancien numero -> nouveau, applique en ordre croissant.
RENUMEROTATION = [("[25]", "[24]"), ("[26]", "[25]"), ("[27]", "[26]")]


def remplacer(paragraphe, avant, apres):
    """Substitue dans un paragraphe, y compris a cheval sur deux fragments."""
    if avant not in paragraphe.text:
        return False
    for fragment in paragraphe.runs:
        if avant in fragment.text:
            fragment.text = fragment.text.replace(avant, apres)
    if avant in paragraphe.text:
        entier = paragraphe.text.replace(avant, apres)
        for fragment in list(paragraphe.runs)[1:]:
            fragment._element.getparent().remove(fragment._element)
        if paragraphe.runs:
            paragraphe.runs[0].text = entier
        else:
            paragraphe.add_run(entier)
    return True


def ecrire_cellule(cellule, texte):
    """Reecrit une cellule en conservant la mise en forme du premier fragment."""
    premier = cellule.paragraphs[0]
    for p in list(cellule.paragraphs)[1:]:
        p._element.getparent().remove(p._element)
    for fragment in list(premier.runs)[1:]:
        fragment._element.getparent().remove(fragment._element)
    if premier.runs:
        premier.runs[0].text = texte
    else:
        premier.add_run(texte)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    # 1. Le tableau des couts.
    cible = None
    for t in doc.tables:
        if [c.text.strip() for c in t.rows[0].cells][:1] == ["Poste"]:
            cible = t
            break
    if cible is None:
        raise SystemExit("tableau 6.4 introuvable")

    for (ligne, colonne), texte in CELLULES.items():
        ecrire_cellule(cible.rows[ligne].cells[colonne], texte)
        print(f"  cellule ({ligne},{colonne}) : {texte[:56]}")

    # 2. La prose qui commente le tableau.
    for p in doc.paragraphs:
        remplacer(p, "la cible de production sur serveur privé virtuel Systalink, "
                     "retenue pour un déploiement pilote AGEROUTE",
                  "une projection de production sur serveur privé virtuel, "
                  "envisagée pour un déploiement pilote AGEROUTE")
        remplacer(p, "l'absence de grille tarifaire publique pour le palier "
                     "intermédiaire chez Systalink : seuls le tarif d'entrée "
                     "et le tarif haute performance sont publiés [25]",
                  "l'absence de grille tarifaire publique pour le palier "
                  "intermédiaire chez l'hébergeur pris comme référence : seuls "
                  "le tarif d'entrée et le tarif haute performance sont "
                  "publiés [25]")
        remplacer(p, "le serveur privé virtuel du tableau 6.4",
                  "le serveur privé virtuel projeté au tableau 6.4")

    # 3. La source salariale, devenue sans objet.
    bibliographie = next(i for i, p in enumerate(doc.paragraphs)
                         if p.text.strip() == "BIBLIOGRAPHIE")
    for p in list(doc.paragraphs[bibliographie + 1:]):
        if p.text.strip().startswith("[24]Simoon"):
            p._element.getparent().remove(p._element)
            print("  reference [24] retiree (Simoon CV)")
            break

    # 4. Renumerotation des suivantes, corps et bibliographie.
    for avant, apres in RENUMEROTATION:
        touches = 0
        for p in doc.paragraphs:
            if remplacer(p, avant, apres):
                touches += 1
        for t in doc.tables:
            for ligne in t.rows:
                for cellule in ligne.cells:
                    for p in cellule.paragraphs:
                        if remplacer(p, avant, apres):
                            touches += 1
        print(f"  {avant} devient {apres} : {touches} emplacement(s)")

    doc.save(SOURCE)

    # Controle : plus aucun appel orphelin.
    controle = Document(SOURCE)
    appels, sources = set(), set()
    for p in controle.paragraphs:
        for m in re.finditer(r"\[(\d+)\]", p.text):
            (sources if p.text.strip().startswith(f"[{m.group(1)}]")
             else appels).add(int(m.group(1)))
    for t in controle.tables:
        for ligne in t.rows:
            for cellule in ligne.cells:
                for m in re.finditer(r"\[(\d+)\]", cellule.text):
                    appels.add(int(m.group(1)))
    orphelins = sorted(appels - sources)
    print(f"\nsources : {len(sources)}, appels orphelins : "
          f"{orphelins if orphelins else 'aucun'}")
    print(f"Sauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
