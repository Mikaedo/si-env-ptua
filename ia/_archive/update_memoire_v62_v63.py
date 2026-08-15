# -*- coding: utf-8 -*-
"""
v62 -> v63 :
1. Renomme la Figure 1.1 (liste des figures + legende) pour refleter ce qui
   est reellement represente : la Direction des Affaires Juridiques, Moyens
   Generaux et Systeme d'Information, pas "l'AGEROUTE" dans son ensemble.
   Corrige le paragraphe 1.2 qui presentait la DSI comme une Direction a
   part entiere alors qu'elle est un departement au sein de cette direction
   (coherence avec l'organigramme reel).
2. Redesign des 3 pages "PARTIE" (Premiere/Deuxieme/Troisieme) : suppression
   des filets orange, centrage vertical reel sur la page (table invisible
   pleine hauteur, technique standard Word - pas de rupture de section pour
   ne pas fragiliser la pagination existante), police plus soignee pour les
   titres de partie.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v62.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v63.docx"

doc = Document(SRC)

# ============================================================
# 1. Renommage Figure 1.1 + correction du paragraphe 1.2
# ============================================================
OLD_CAPTION = "Organigramme de l'AGEROUTE et de la DSI"
NEW_CAPTION = "Organigramme de la Direction des Affaires Juridiques, Moyens Généraux et Système d'Information (AGEROUTE)"

def replace_in_paragraph(idx, old, new):
    p = doc.paragraphs[idx]
    full = p.text
    if old not in full:
        print(f"[ATTENTION] texte attendu absent au paragraphe {idx}")
        return
    new_full = full.replace(old, new)
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_full
    else:
        p.add_run(new_full)
    print(f"[OK] Paragraphe {idx} renomme")

replace_in_paragraph(86, OLD_CAPTION, NEW_CAPTION)
replace_in_paragraph(259, OLD_CAPTION, NEW_CAPTION)

def replace_paragraph_text(idx, new_text):
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)
    print(f"[OK] Paragraphe {idx} remplace")

replace_paragraph_text(256,
    "Au sein de l'AGEROUTE, la fonction Système d'Information — désignée dans ce mémoire par l'acronyme "
    "DSI — relève du Département Système d'Information et Reprographie, l'un des trois départements de "
    "la Direction des Affaires Juridiques, Moyens Généraux et Système d'Information (figure 1.1). Elle "
    "est structurée autour de deux services distincts : le Service Études et Développement Applicatif "
    "(SEDA) et le Service Informatique, Sécurité et Technologie (SIST)."
)

# ============================================================
# 2. Redesign des 3 pages "PARTIE"
# ============================================================

def remove_border(paragraph):
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is not None:
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            pPr.remove(pBdr)

def set_run_font(run, name, size_pt=None, spacing_pt=None):
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), name)
    if size_pt:
        run.font.size = Pt(size_pt)
    if spacing_pt:
        rPr_spacing = OxmlElement('w:spacing')
        rPr_spacing.set(qn('w:val'), str(int(spacing_pt * 20)))
        rPr.append(rPr_spacing)

def set_table_no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)

def build_centered_divider(title_idx, subtitle_idx, desc_idx, blanks_before, blanks_after):
    title_p = doc.paragraphs[title_idx]
    subtitle_p = doc.paragraphs[subtitle_idx]
    desc_p = doc.paragraphs[desc_idx]

    remove_border(title_p)
    remove_border(desc_p)

    for r in title_p.runs:
        set_run_font(r, 'Constantia', size_pt=34, spacing_pt=1.5)
    for r in subtitle_p.runs:
        set_run_font(r, 'Constantia', size_pt=19, spacing_pt=1)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_no_borders(table)
    cell = table.cell(0, 0)
    cell.width = Inches(6.1)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '13000')
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    default_p = cell.paragraphs[0]._p
    for p in (title_p, subtitle_p, desc_p):
        default_p.addprevious(p._p)
    default_p.getparent().remove(default_p)

    anchor = doc.paragraphs[blanks_before[0]]._p
    anchor.addprevious(table._tbl)

    for idx in list(blanks_before) + list(blanks_after):
        p_el = doc.paragraphs[idx]._p
        p_el.getparent().remove(p_el)

    print(f"[OK] Page partie centree (titre ex-paragraphe {title_idx})")

# Traite en ordre INVERSE (paragraphe le plus loin d'abord) pour ne pas
# decaler les index des blocs traites ensuite.
build_centered_divider(441, 442, 443, blanks_before=[439, 440], blanks_after=[444])
build_centered_divider(306, 307, 308, blanks_before=[303, 304, 305], blanks_after=[309])
build_centered_divider(243, 244, 245, blanks_before=[241, 242], blanks_after=[246])

doc.save(DST)
print(f"\n=== MEMOIRE SAUVEGARDE : {DST} ===")
