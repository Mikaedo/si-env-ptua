# -*- coding: utf-8 -*-
"""
Retire du memoire le dispositif de supervision, hors sujet pour le lecteur.

Le paragraphe 6.8 detaillait un mecanisme de surveillance en trois couches :
interrogation periodique, rejeu de la chaine complete, cascade de reparation
automatique. Cet outillage sert l'exploitation du service, non la demonstration
academique, et n'a pas sa place dans le document. Ne subsistent donc que les
technologies d'hebergement et la facon dont le code y est publie.

La limite des offres gratuites, elle, reste ecrite. La passer sous silence
apres avoir supprime la parade laisserait croire qu'elle n'existe pas, alors
qu'un evaluateur qui ouvrirait l'adresse apres plusieurs heures constaterait le
delai de reveil. Mieux vaut l'annoncer et en tirer l'argument qui va avec :
c'est l'une des raisons pour lesquelles une mise en exploitation demanderait un
serveur dedie.

Le titre perd « haute disponibilite », promesse que le texte ne soutient plus
et que rien ne mesurait sur une periode significative.
"""
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_retrait_supervision.docx")

TITRE_AVANT = "6.8  Déploiement continu et haute disponibilité"
TITRE_APRES = "6.8  Déploiement continu et hébergement"

LIMITE = (
    "L'hébergement gratuit a une limite structurelle qu'il faut assumer : les "
    "plateformes suspendent les services après un délai d'inactivité, quinze "
    "minutes pour Render et sept jours pour Supabase. La première requête qui "
    "suit une suspension paie donc un délai de réveil de plusieurs dizaines de "
    "secondes. La contrainte est acceptable pour une validation académique, où "
    "les consultations sont ponctuelles ; elle ne le serait pas en "
    "exploitation, et c'est l'une des raisons pour lesquelles la cible de "
    "production repose sur un serveur privé virtuel."
)

ADRESSES = (
    "Le tableau de bord est publié à l'adresse si-env-ptua.pages.dev et le "
    "backend à si-env-ptua.onrender.com, l'un et l'autre accessibles sans "
    "installation préalable. Le passage à la cible de production, le serveur "
    "privé virtuel projeté au tableau 6.4, se fait par simple changement de la "
    "variable DATABASE_URL et de la commande de lancement, sans modification du "
    "code : c'est une conséquence directe de la conteneurisation."
)

CELLULE = (
    "Pile réellement déployée pour la validation du mémoire : Render (backend "
    "FastAPI en Docker), Supabase (PostgreSQL 16 + PostGIS + stockage des "
    "photos), Cloudflare Pages (tableau de bord Angular), GitHub Actions "
    "(intégration et livraison continue). Aucune carte bancaire requise, "
    "aucune limite pratique de bande passante ; endormissement des services "
    "après un délai d'inactivité, propre aux offres gratuites."
)


def reecrire(paragraphe, texte):
    """Remplace le contenu en conservant la mise en forme du premier fragment."""
    for fragment in list(paragraphe.runs)[1:]:
        fragment._element.getparent().remove(fragment._element)
    if paragraphe.runs:
        paragraphe.runs[0].text = texte
    else:
        paragraphe.add_run(texte)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    faits = []
    for p in doc.paragraphs:
        depart = p.text.strip()

        if depart == TITRE_AVANT and p.style.name.startswith("Heading"):
            reecrire(p, TITRE_APRES)
            faits.append("titre 6.8")

        elif depart.startswith("L'hébergement gratuit a une limite"):
            reecrire(p, LIMITE)
            faits.append("limite des offres gratuites")

        elif depart.startswith("Ce dispositif ne se substitue pas"):
            reecrire(p, ADRESSES)
            faits.append("adresses de publication")

    for table in doc.tables:
        for ligne in table.rows:
            for cellule in ligne.cells:
                if "watchdog" in cellule.text:
                    reecrire(cellule.paragraphs[0], CELLULE)
                    for p in list(cellule.paragraphs)[1:]:
                        p._element.getparent().remove(p._element)
                    faits.append("tableau 6.4, ligne hébergement gratuit")

    doc.save(SOURCE)
    for f in faits:
        print(f"  corrige : {f}")

    # Controle : plus aucune trace du dispositif.
    interdits = ("watchdog", "self-heal", "dispositif de surveillance",
                 "health check", "redémarrage complet", "republication",
                 "haute disponibilité", "trois couches")
    controle = Document(SOURCE)
    restes = []
    for p in controle.paragraphs:
        for mot in interdits:
            if mot.lower() in p.text.lower():
                restes.append(f"{mot} : {p.text.strip()[:70]}")
    for t in controle.tables:
        for l in t.rows:
            for c in l.cells:
                for mot in interdits:
                    if mot.lower() in c.text.lower():
                        restes.append(f"{mot} : {c.text.strip()[:70]}")

    print(f"\nmentions restantes : {len(restes)}")
    for r in restes:
        print(f"  {r}")
    print(f"Sauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
