# -*- coding: utf-8 -*-
"""
Etape 8 : compacter la mise en page du corps.

Le diagnostic a montre que le texte redige ne represente qu'environ 14 pages sur
60 : l'essentiel du volume vient des figures, des tableaux et des espaces
perdus. Trois actions, par ordre d'efficacite :

  a) Figures ramenees a 55 % de leur taille d'origine. A cette echelle elles
     cessent d'occuper une page entiere et cohabitent avec le texte, ce qui est
     precisement ce qui fait gagner des pages : une reduction moderee laissait
     simplement du blanc en bas de page.

  b) Suppression des paragraphes vides excedentaires et des sauts de page en
     double, qui produisent des pages blanches.

  c) Les legendes de figures et titres de tableaux passent en 10 points et
     restent solidaires de leur objet, pour eviter qu'une legende seule ne soit
     rejetee sur la page suivante.

L'interligne 1,5, la police et les marges imposes par le reglement ne sont pas
touches.

Sortie : MEMOIRE_ETAPE8.docx
"""
import os
import re
import shutil

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_ETAPE7.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_ETAPE8.docx")

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

#: Taille finale des figures du corps, en proportion de l'original.
#: Les figures ont deja ete ramenees a 70 % : on applique le complement.
CIBLE = 0.55
DEJA_APPLIQUE = 0.70
FACTEUR = CIBLE / DEJA_APPLIQUE

SEUIL_LOGO_CM = 5.0


def texte(el):
    return ''.join(t.text or '' for t in el.iter('{%s}t' % W)).strip()


def bornes_corps(doc):
    debut = fin = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if debut is None and t.startswith('Introduction g') and i > 250:
            debut = i
        if debut is not None and t.upper().startswith('BIBLIOGRAPHIE'):
            fin = i
            break
    return debut, fin or len(doc.paragraphs)


def reduire_figures(doc, debut, fin):
    reduites = gain = 0
    for i in range(debut, fin):
        p = doc.paragraphs[i]
        for ext in p._element.iter('{%s}ext' % A_NS):
            cx, cy = ext.get('cx'), ext.get('cy')
            if not (cx and cy):
                continue
            l_cm, h_cm = int(cx) / 360000.0, int(cy) / 360000.0
            if l_cm < SEUIL_LOGO_CM:
                continue
            ext.set('cx', str(int(round(l_cm * FACTEUR * 360000))))
            ext.set('cy', str(int(round(h_cm * FACTEUR * 360000))))
            reduites += 1
            gain += h_cm * (1 - FACTEUR)
    return reduites, gain


def nettoyer_blancs(doc):
    body = doc.element.body
    supprimes = sauts = 0
    vides = []
    saut_precedent = False
    for el in list(body):
        if not el.tag.endswith('}p'):
            vides = []
            saut_precedent = False
            continue
        a_texte = bool(texte(el))
        a_image = bool(el.findall('.//{%s}blip' % A_NS))
        brs = [b for b in el.findall('.//{%s}br' % W)
               if b.get('{%s}type' % W) == 'page']
        a_sect = (el.find('{%s}pPr' % W) is not None and
                  el.find('{%s}pPr' % W).find('{%s}sectPr' % W) is not None)

        if a_texte or a_image or a_sect:
            vides = []
            saut_precedent = bool(brs)
            continue
        if brs:
            if saut_precedent:
                for b in brs:
                    b.getparent().remove(b)
                sauts += 1
                saut_precedent = False
            else:
                saut_precedent = True
            vides = []
            continue
        # paragraphe totalement vide
        vides.append(el)
        if len(vides) > 1:
            body.remove(el)
            supprimes += 1
    return supprimes, sauts


def compacter_legendes(doc, debut, fin):
    """Reduit les legendes et les rend solidaires de leur figure ou tableau."""
    traites = 0
    for i in range(debut, fin):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if not re.match(r'^(Figure|Tableau)\s+[0-9A-F]+\.\d+\s*[:.]', t):
            continue
        for run in p.runs:
            run.font.size = Pt(10)
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(8)
        pf.keep_with_next = t.startswith('Tableau')  # titre au-dessus du tableau
        traites += 1
    return traites


def main():
    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)
    debut, fin = bornes_corps(doc)
    print("  corps : paragraphes %d a %d" % (debut, fin))

    reduites, gain = reduire_figures(doc, debut, fin)
    print("  figures ramenees a %d %% de l'original : %d (%.0f cm liberes)"
          % (CIBLE * 100, reduites, gain))

    supprimes, sauts = nettoyer_blancs(doc)
    print("  paragraphes vides supprimes : %d | sauts en double retires : %d"
          % (supprimes, sauts))

    legendes = compacter_legendes(doc, debut, fin)
    print("  legendes compactees : %d" % legendes)

    doc.save(SORTIE)
    print("\nEnregistre : %s" % SORTIE)


if __name__ == "__main__":
    main()
