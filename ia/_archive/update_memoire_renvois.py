# -*- coding: utf-8 -*-
"""Etape 3/3 : ajoute un renvoi explicite dans le texte pour les 7 figures
et 13 tableaux qui n'etaient reference nulle part en prose (juste
legendes), comme l'exige le guide UPB ("Toujours referencer dans le
texte")."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

SRC = r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_N'GUESSAN_v5.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_N'GUESSAN_v6.docx"

doc = Document(SRC)
paras = doc.paragraphs

def fix_para(p, new_text):
    runs = p.runs
    assert runs, f"paragraphe sans run : {p.text!r}"
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''

def insert_new_paragraph_before(anchor_para, text):
    new_p = anchor_para.insert_paragraph_before(text)
    new_p.style = doc.styles['Normal']
    return new_p

EDITS = {
    260: ("L'envergure du PTUA, dont le coût global s'élève à 657,8 milliards de FCFA [1], a nécessité "
          "la mise en place d'une organisation projet spécifique et rigoureuse, rattachée à la Direction "
          "Générale de l'AGEROUTE. L'organigramme du projet se décline en trois grands niveaux de "
          "responsabilité (figure 1.2)."),
    356: ("Le SI-ENV comprend cinq modules : application mobile Flutter, backend FastAPI, base "
          "PostgreSQL/PostGIS, tableau de bord Angular et module IA exécuté localement avec ONNX Runtime. "
          "Le mobile stocke les signalements dans SQLite hors connexion ; le backend expose les services "
          "REST et centralise les données ; le dashboard assure la supervision et les rapports ; le "
          "module IA réalise le diagnostic des images. Google Earth Engine complète l'architecture pour "
          "les analyses satellitaires (figure 6.1)."),
    364: ("Pour la réalisation du projet, sept (7) acteurs ont été identifiés : cinq acteurs principaux "
          "interagissant directement avec le système, et deux acteurs secondaires, à savoir des systèmes "
          "externes sollicités par le SI-ENV (tableaux 6.4 et 6.5)."),
    373: ("Le diagramme de cas d'utilisation ne décrit pas de manière détaillée le dialogue entre les "
          "acteurs et les cas d'utilisation. Il est recommandé de rédiger une description textuelle "
          "appelée scénario du cas d'utilisation. Nous présentons ci-dessous les scénarios des six (6) "
          "cas les plus pertinents (tableaux 6.6 à 6.11)."),
    439: ("Le backend FastAPI (Python 3.12) expose les services REST : authentification, signalements, "
          "alertes, rapports et analyse satellitaire. La validation des données est assurée par "
          "Pydantic, l'accès à la base par SQLAlchemy, et les requêtes spatiales par PostGIS. La "
          "documentation Swagger UI est auto-générée (figure 7.1), dont les principaux endpoints sont "
          "détaillés au tableau 7.2. L'architecture suit le pattern Repository (routes → services → "
          "repositories). L'authentification utilise JWT (1h, renouvelable) et le RBAC garantit le "
          "cloisonnement des profils."),
    474: ("Le tableau 8.2 compare les modèles de détection testés selon ces critères. Le Rappel (Recall) "
          "mérite une attention particulière : un faux négatif (déchets non détectés) a des conséquences "
          "pratiques directes sur la fiabilité du diagnostic automatique. Ce critère, au même titre que "
          "le mAP et les autres métriques retenues [18], justifie le choix du modèle retenu pour le "
          "SI-ENV ; la figure 8.1 détaille ces métriques par classe."),
    483: ("Le tableau 8.3 compare les modèles de classification testés. Le modèle retenu offre le "
          "meilleur compromis performance/taille pour une exécution locale sur les téléphones des agents "
          "terrain (figure 8.4)."),
    486: ("Les objectifs de performance ont été calibrés par comparaison avec des travaux publiés sur des "
          "tâches de vision par ordinateur comparables, plutôt que fixés arbitrairement. Pour la "
          "détection, une étude récente sur la détection de déchets de chantier par YOLOv8n rapporte un "
          "mAP@0.5 de 89,8 % [22] ; notre modèle atteint un mAP@0.5 de 0,798, soit 79,8 % (figure 8.3), "
          "ce qui se situe dans un ordre de grandeur cohérent avec cette référence. L'écart résiduel "
          "s'explique par la différence de dataset (6 classes vs 1 classe spécialisée) et l'absence "
          "d'optimisation architecturale (FE-YOLO). Pour la classification, une étude comparative "
          "reporte un F1-score de 0,93 pour MobileNetV2 [23] ; notre modèle atteint un F1-score pondéré "
          "de 0,86, cohérent avec un dataset de criticité plus petit et des classes déséquilibrées."),
    491: ("Les faux négatifs observés pour YOLOv8 (figure 8.2) concernent principalement la classe "
          "plastique (Rappel = 0,504), due à la transparence et au faible contraste des objets en "
          "plastique. Les faux positifs les plus probables portent sur des matériaux de chantier "
          "visuellement proches de déchets (gravats, débris de coffrage). Pour MobileNetV2, la "
          "classification est plus fiable sur la classe faible (F1 = 0,93) que sur les classes modérée "
          "(F1 = 0,59) et importante (F1 = 0,67), un biais courant en classification multi-classe avec "
          "déséquilibre. L'utilisation de WeightedRandomSampler et de class weights a permis d'atténuer "
          "ce déséquilibre sans toutefois le résoudre entièrement."),
    524: ("Le tableau 10.1 détaille la configuration des conteneurs Docker Compose. Nginx termine TLS "
          "(Let's Encrypt) et route vers FastAPI ou Angular. Les secrets sont dans un fichier .env. "
          "Déploiement : docker compose up -d."),
}

for idx, new_text in EDITS.items():
    fix_para(paras[idx], new_text)
    assert paras[idx].text == new_text, f"echec p{idx}"
print(f"[OK] {len(EDITS)} paragraphes existants etendus avec un renvoi explicite")

# Figure 4.1 : aucun paragraphe de prose avant l'image -> on insere une phrase d'intro
insert_new_paragraph_before(
    paras[308],
    "La figure 4.1 illustre le processus de suivi environnemental actuellement en vigueur, fondé sur "
    "des outils bureautiques disjoints."
)
print("[OK] Phrase d'introduction inseree avant la figure 4.1")

# Tableau 8.4 (hyperparametres) : aucune phrase ne le mentionne -> ajout d'une phrase apres la legende
paras = doc.paragraphs  # recharger apres insertion (les indices ont decale de 1 apres p308)
target = None
for p in paras:
    if p.text.strip() == "Tableau 8.4 : Hyperparamètres optimisés pour YOLOv8 et MobileNetV2.":
        target = p
        break
assert target is not None, "legende Tableau 8.4 introuvable"
# python-docx n'a pas d'insert_paragraph_after natif : on cree l'element XML directement
from docx.text.paragraph import Paragraph
p_after = OxmlElement('w:p')
target._p.addnext(p_after)
para_after = Paragraph(p_after, target._parent)
para_after.style = doc.styles['Normal']
para_after.add_run(
    "Le tableau 8.4 récapitule les hyperparamètres retenus pour les deux modèles après optimisation."
)
print("[OK] Phrase ajoutee apres le tableau 8.4 (hyperparametres)")

doc.save(DST)
print(f"\n=== SAUVEGARDE : {DST} ===")
