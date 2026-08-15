# -*- coding: utf-8 -*-
"""v61 -> v62 : corrige les organigrammes 1.1/1.2 (traits de liaison manquants
pour les enfants uniques, suppression du titre duplique dans l'image)."""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

SRC = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v61.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v62.docx"
SCRATCH = r"C:\Users\DELL\AppData\Local\Temp\claude\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25\scratchpad"

doc = Document(SRC)

def replace_image(pidx, image_path, width_inches=6.1):
    p = doc.paragraphs[pidx]
    for r in p.runs:
        drawings = r.element.findall(qn('w:drawing'))
        if drawings:
            for d in drawings:
                d.getparent().remove(d)
            r.add_picture(image_path, width=Inches(width_inches))
            print(f"[OK] Figure corrigee au paragraphe {pidx}")
            return
    raise RuntimeError(f"Aucune image trouvee au paragraphe {pidx}")

replace_image(258, SCRATCH + "/new_fig11.png")
replace_image(264, SCRATCH + "/new_fig12.png")

doc.save(DST)
print(f"=== SAUVEGARDE : {DST} ===")
