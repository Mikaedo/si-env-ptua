# -*- coding: utf-8 -*-
"""
Met a jour deux endroits du memoire pour refleter le deploiement reellement
realise :
  - §6.7 (tableau 6.4) : ajoute une ligne « environnement academique gratuit »
    detaillant la pile Render + Supabase + Cloudflare + GitHub Actions, et
    ajuste le total.
  - §6.8 : enrichit le paragraphe existant avec les URL de production reelles
    et les resultats de validation, sans allonger significativement.

Contrainte : garder le corps a 51 pages max. Compression legere ailleurs si
besoin.
"""
import copy
import docx
from docx.text.paragraph import Paragraph

CHEMIN = r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx"
d = docx.Document(CHEMIN)


def par(debut):
    for p in d.paragraphs:
        if p.text.strip().startswith(debut):
            return p
    raise SystemExit("PARAGRAPHE INTROUVABLE : " + debut[:70])


def remplacer(p, vieux, neuf):
    if vieux not in p.text:
        raise SystemExit("FRAGMENT : " + vieux[:70])
    p.runs[0].text = p.text.replace(vieux, neuf)
    for r in p.runs[1:]:
        r.text = ""


def ecrire_cellule(cellule, texte):
    p = cellule.paragraphs[0]
    if p.runs:
        p.runs[0].text = texte
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(texte)
    for extra in cellule.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)


# ═════════════════════════════════════════════════════════════════════════
# 1. Tableau 6.4 : injecter une ligne « pile gratuite academique »
# ═════════════════════════════════════════════════════════════════════════

# Trouve le tableau 6.4 (colonne 0 = "Poste")
t64 = None
for t in d.tables:
    if t.rows and t.rows[0].cells and t.rows[0].cells[0].text.strip() == "Poste":
        t64 = t
        break
if t64 is None:
    raise SystemExit("Tableau 6.4 introuvable")

# Trouve la ligne "Hebergement (VPS Systalink)" pour inserer juste apres
from docx.table import _Row
i_ref = None
for i, ligne in enumerate(t64.rows):
    if ligne.cells[0].text.strip().startswith("Hébergement (VPS Systalink)"):
        i_ref = i
        break
if i_ref is None:
    raise SystemExit("Ligne Systalink introuvable dans le tableau 6.4")

# Clone la ligne pour heriter du style, insere apres i_ref
modele = t64.rows[i_ref]
neuf_tr = copy.deepcopy(modele._tr)
modele._tr.addnext(neuf_tr)
nouvelle = _Row(neuf_tr, t64)
ecrire_cellule(nouvelle.cells[0], "Hébergement cloud gratuit (validation académique)")
ecrire_cellule(nouvelle.cells[1], (
    "Pile réellement déployée pour la validation du mémoire : Render (backend "
    "FastAPI en Docker), Supabase (PostgreSQL 16 + PostGIS + stockage des "
    "photos), Cloudflare Pages (tableau de bord Angular), GitHub Actions "
    "(intégration et livraison continue, watchdog 24/7). Aucune carte bancaire "
    "requise, aucune limite pratique de bande passante ; endormissement des "
    "services au bout de 15 min à 48 h contré par un dispositif de surveillance "
    "actif (paragraphe 6.8)."
))
ecrire_cellule(nouvelle.cells[2], "0 FCFA / mois")

# Ajuste le libelle et le montant du total de fonctionnement
for ligne in t64.rows[1:]:
    if ligne.cells[0].text.strip().startswith("Coût de fonctionnement annuel"):
        ecrire_cellule(ligne.cells[1],
            "Environnement de validation (Render + Supabase + Cloudflare) : "
            "gratuit. Cible de production (VPS Systalink) : hébergement 12 mois "
            "+ domaine."
        )
        ecrire_cellule(ligne.cells[2],
            "0 FCFA / an (validation) ; 51 500 à 969 500 FCFA / an (production)")
        break


# ═════════════════════════════════════════════════════════════════════════
# 2. §6.7 : ajouter une phrase apres le paragraphe introductif
# ═════════════════════════════════════════════════════════════════════════
p67_intro = par("Cette section chiffre le coût de développement")
remplacer(
    p67_intro,
    "Cette section chiffre le coût de développement et de fonctionnement du SI-ENV à partir de références publiques datées d'août 2026 (tableau 6.4), puis situe cet ordre de grandeur face à des plateformes commerciales de gestion HSE (tableau 6.5). Le taux de conversion retenu, 1 USD ≈ 600 FCFA, est une approximation d'août 2026 et non un taux officiel.",
    "Cette section chiffre le coût de développement et de fonctionnement du "
    "SI-ENV à partir de références publiques datées d'août 2026 (tableau 6.4). "
    "Deux scénarios y coexistent : la cible de production sur serveur privé "
    "virtuel Systalink, retenue pour un déploiement pilote AGEROUTE, et "
    "l'environnement de validation académique effectivement déployé (Render, "
    "Supabase, Cloudflare Pages, GitHub Actions), gratuit et sans engagement. "
    "Le taux de conversion retenu, 1 USD ≈ 600 FCFA, est une approximation "
    "d'août 2026 et non un taux officiel."
)


# ═════════════════════════════════════════════════════════════════════════
# 3. §6.8 : enrichir avec les URL de production reelles
# ═════════════════════════════════════════════════════════════════════════
# Ajouter une phrase a la fin du paragraphe "Ce dispositif ne se substitue pas"
p68_fin = par("Ce dispositif ne se substitue pas")
remplacer(
    p68_fin,
    "Ce dispositif ne se substitue pas à un hébergement de production. Il "
    "atteint en revanche une disponibilité effective de l'ordre de 99,7 pour "
    "cent, avec un temps de réparation moyen inférieur à cinq minutes en cas "
    "d'incident détecté. Le passage à la cible de production, le serveur privé "
    "virtuel du tableau 6.4, se fait par simple changement de la variable "
    "DATABASE_URL et de la commande de lancement, sans modification du code.",
    "Ce dispositif ne se substitue pas à un hébergement de production. Il "
    "atteint en revanche une disponibilité effective de l'ordre de 99,7 pour "
    "cent, avec un temps de réparation moyen inférieur à cinq minutes en cas "
    "d'incident détecté. Le tableau de bord est publié à l'adresse "
    "si-env-ptua.pages.dev et le backend à si-env-ptua.onrender.com ; le "
    "health check profond, joué toutes les heures, valide la chaîne complète "
    "de bout en bout (authentification, requête PostGIS, accès au bucket, "
    "réponse Google Earth Engine). Le passage à la cible de production, le "
    "serveur privé virtuel du tableau 6.4, se fait par simple changement de "
    "la variable DATABASE_URL et de la commande de lancement, sans "
    "modification du code."
)


# ═════════════════════════════════════════════════════════════════════════
# Sauvegarde
# ═════════════════════════════════════════════════════════════════════════
d.save(CHEMIN)
print("Tableau 6.4 enrichi (+1 ligne), §6.7 et §6.8 mis a jour.")
