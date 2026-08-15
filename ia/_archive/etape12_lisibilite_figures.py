# -*- coding: utf-8 -*-
"""
Etape 12 : rendre aux figures du corps une taille confortable.

Le corps est descendu a 47 pages, soit trois pages sous le plafond de 50. Cette
marge est rendue aux figures, ramenees de 45 % a 60 % de leur taille d'origine :
les diagrammes UML et les captures d'ecran redeviennent lisibles a l'impression,
ce qui compte davantage qu'un gain de pages devenu inutile.

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os

from docx import Document

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')

A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

TAILLE_VOULUE = 0.60
TAILLE_ACTUELLE = 0.45
FACTEUR = TAILLE_VOULUE / TAILLE_ACTUELLE

#: En dessous de cette largeur, l'image est un logo de page de garde ou de
#: tableau des technologies : on n'y touche pas.
SEUIL_LOGO_CM = 3.0
#: Largeur utile d'une page A4 avec les marges du reglement.
LARGEUR_UTILE_CM = 16.0


def bornes_corps(doc):
    debut = fin = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if debut is None and t.startswith('Introduction g') and i > 250:
            debut = i
        if debut is not None and t.upper().startswith('BIBLIOGRAPHIE'):
            fin = i
            break
    return debut, fin or len(doc.paragraphs)


def main():
    doc = Document(CIBLE)
    debut, fin = bornes_corps(doc)
    print("  corps : paragraphes %d a %d" % (debut, fin))

    agrandies = plafonnees = 0
    for i in range(debut, fin):
        for ext in doc.paragraphs[i]._element.iter('{%s}ext' % A_NS):
            cx, cy = ext.get('cx'), ext.get('cy')
            if not (cx and cy):
                continue
            l_cm, h_cm = int(cx) / 360000.0, int(cy) / 360000.0
            if l_cm < SEUIL_LOGO_CM:
                continue
            f = FACTEUR
            # Ne jamais depasser la largeur utile de la page.
            if l_cm * f > LARGEUR_UTILE_CM:
                f = LARGEUR_UTILE_CM / l_cm
                plafonnees += 1
            ext.set('cx', str(int(round(l_cm * f * 360000))))
            ext.set('cy', str(int(round(h_cm * f * 360000))))
            agrandies += 1

    print("  figures portees a %d %% de l'original : %d (dont %d plafonnees a "
          "la largeur de page)" % (TAILLE_VOULUE * 100, agrandies, plafonnees))

    doc.save(CIBLE)
    print("\nEnregistre : %s" % CIBLE)


if __name__ == "__main__":
    main()
