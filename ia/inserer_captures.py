# -*- coding: utf-8 -*-
"""
Insere dans le corps du memoire les captures qui demontrent.

Le critere retenu est simple : une capture merite le corps lorsqu'elle etablit
ce que le texte affirme sans pouvoir le montrer. Celle de la vue de
consultation en est l'exemple le plus net. Le memoire soutient que l'agence de
tutelle et le bailleur voient les memes ecrans depourvus de toute commande
d'ecriture ; trois paragraphes peinent a en convaincre, une image y suffit.

Les captures qui disent seulement « cela existe » restent en annexe. Un memoire
de licence n'est pas juge sur le nombre d'ecrans montres mais sur la justesse
des choix exposes, et une accumulation de copies d'ecran se lit comme du
remplissage.

Chaque figure est posee juste apres le passage qu'elle appuie, non regroupee en
fin de section : une image separee de son propos oblige le lecteur a un
va-et-vient qui lui fait perdre le fil.
"""
import copy
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_captures.docx")
DOSSIER = Path(r"D:\etude_soutenance\SI-ENV\ia\captures_memoire")

# Chaque entree associe le debut du paragraphe apres lequel inserer, le
# fichier, la legende et la hauteur souhaitee.
INSERTIONS = [
    (
        "Le rattachement au chantier est déduit de la position",
        ["citoyen_accueil.png", "citoyen_depot.png"],
        "Figure 5.3 : Application citoyenne. À gauche, l'écran d'accueil qui "
        "expose la raison pour laquelle la position est demandée. À droite, le "
        "dépôt d'une doléance, réduit à une catégorie et à une description.",
        7.6,
    ),
    (
        "Afin de matérialiser le cloisonnement des responsabilités",
        ["web_consultation_ande.png"],
        "Figure 4.10 : Tableau de bord vu par l'Agence Nationale de "
        "l'Environnement. La barre latérale ne propose ni les plaintes ni "
        "l'administration, et la liste des signalements ne comporte aucune "
        "commande de modification.",
        8.2,
    ),
]


def style_legende(doc):
    """Reprend le style des legendes deja presentes dans le document."""
    for p in doc.paragraphs:
        if p.text.strip().startswith("Figure 4.2 :"):
            return p.style
    return doc.styles["Normal"]


def inserer_apres(ancre, modele):
    el = copy.deepcopy(modele._element)
    ancre._element.addnext(el)
    p = Paragraph(el, ancre._parent)
    for seg in list(p.runs):
        seg._element.getparent().remove(seg._element)
    return p


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    modele_legende = style_legende(doc)

    modele = None
    for p in doc.paragraphs:
        if p.style.name == "Normal" and len(p.text.strip()) > 60:
            modele = p
            break

    for debut, fichiers, legende, hauteur in INSERTIONS:
        ancre = None
        for p in doc.paragraphs:
            if p.text.strip().startswith(debut):
                ancre = p
                break
        if ancre is None:
            print(f"  MANQUE : {debut[:44]}")
            continue

        # La legende vient en premier dans l'ordre d'insertion, puisque chaque
        # element s'insere immediatement apres l'ancre et repousse le suivant.
        p_leg = inserer_apres(ancre, modele)
        p_leg.style = modele_legende
        p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_leg.add_run(legende)
        r.font.size = Pt(9)
        r.italic = True

        p_img = inserer_apres(ancre, modele)
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        for i, f in enumerate(fichiers):
            chemin = DOSSIER / f
            if not chemin.exists():
                print(f"  fichier absent : {f}")
                continue
            run.add_picture(str(chemin), height=Cm(hauteur))
            if i < len(fichiers) - 1:
                run.add_text("  ")

        print(f"  ok : {legende[:46]}")

    doc.save(SOURCE)
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
