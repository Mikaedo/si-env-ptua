# -*- coding: utf-8 -*-
"""
Etape 1 sur 4 de la remise en forme du memoire.

  a) Les bordures oranges de page ne subsistent que sur la couverture. Le
     document compte trois sections, la premiere correspondant exactement a la
     page de garde : il suffit de retirer la bordure des deux autres.

  b) Ajout d'une couverture en noir et blanc juste apres la couverture en
     couleur, pour la version imprimee. Elle reprend l'integralite du contenu
     de la garde - y compris les DEUX TABLEAUX qui portent les logos et les
     encadrants - avec toutes les couleurs ramenees au noir.

Le contenu est copie au niveau des elements du corps du document, et non des
seuls paragraphes : la page de garde melange paragraphes et tableaux, et une
copie limitee aux paragraphes perdrait les logos.

Sortie : MEMOIRE_ETAPE1.docx
"""
import copy
import os
import re
import shutil

from docx import Document
from docx.oxml.ns import qn

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_DIBY_logos.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_ETAPE1.docx")


def index_fin_couverture(body):
    """Index de l'element portant le sectPr de la premiere section."""
    for i, el in enumerate(body):
        if el.tag.endswith('}p'):
            pPr = el.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                return i
    raise SystemExit("Fin de la premiere section introuvable")


A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'


def griser_images(element):
    """Applique l'effet « niveaux de gris » aux images de l'element.

    On agit sur le rendu (balise a:grayscl dans a:blip) plutot que sur le
    fichier image : les deux couvertures partagent les memes ressources, et
    convertir l'image elle-meme desaturerait aussi la version en couleur.
    """
    n = 0
    for blip in element.iter('{%s}blip' % A_NS):
        if blip.find('{%s}grayscl' % A_NS) is not None:
            continue
        blip.append(blip.makeelement('{%s}grayscl' % A_NS, {}))
        n += 1
    return n


def noircir(element):
    """Ramene au noir tout ce qui porte une couleur dans l'element."""
    n = 0
    for rPr in element.iter(qn('w:rPr')):
        couleur = rPr.find(qn('w:color'))
        if couleur is not None:
            couleur.set(qn('w:val'), '000000')
            n += 1
    # Bordures de paragraphe (encadre du theme) et de tableau
    for balise in (qn('w:pBdr'), qn('w:tblBorders'), qn('w:tcBorders')):
        for bordures in element.iter(balise):
            for cote in bordures:
                if cote.get(qn('w:color')):
                    cote.set(qn('w:color'), '000000')
                    n += 1
    # Aplats de couleur : un fond colore devient un gris sale a l'impression,
    # on le blanchit plutot que de le convertir.
    for shd in element.iter(qn('w:shd')):
        fill = shd.get(qn('w:fill'))
        if fill and fill not in ('auto', 'FFFFFF'):
            shd.set(qn('w:fill'), 'FFFFFF')
            n += 1
    return n


def main():
    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)
    body = doc.element.body

    # ── a) Bordures de page : uniquement la couverture ───────────────────────
    retirees = 0
    for i, section in enumerate(doc.sections):
        if i == 0:
            continue
        bordures = section._sectPr.find(qn('w:pgBorders'))
        if bordures is not None:
            section._sectPr.remove(bordures)
            retirees += 1
    print("  bordures de page retirees (hors couverture) : %d" % retirees)

    # ── b) Couverture noir et blanc ──────────────────────────────────────────
    fin = index_fin_couverture(body)
    elements = list(body)[:fin + 1]
    nb_tableaux = sum(1 for e in elements if e.tag.endswith('}tbl'))
    print("  couverture : %d elements dont %d tableau(x)"
          % (len(elements), nb_tableaux))

    # sectPr de la couverture : sert de modele, bordure passee au noir
    pPr = elements[-1].find(qn('w:pPr'))
    sectPr_nb = copy.deepcopy(pPr.find(qn('w:sectPr')))
    for cote in sectPr_nb.find(qn('w:pgBorders')) or []:
        if cote.get(qn('w:color')):
            cote.set(qn('w:color'), '000000')

    # Insertion des copies apres la couverture couleur. On parcourt a l'envers
    # en s'ancrant toujours sur le meme element : l'ordre d'origine est ainsi
    # conserve.
    ancre = elements[-1]
    copies = []
    couleurs = grises = 0
    for el in reversed(elements):
        clone = copy.deepcopy(el)
        # La copie ne doit pas embarquer le sectPr de la couverture couleur
        if clone.tag.endswith('}p'):
            p_pr = clone.find(qn('w:pPr'))
            if p_pr is not None:
                s = p_pr.find(qn('w:sectPr'))
                if s is not None:
                    p_pr.remove(s)
        couleurs += noircir(clone)
        grises += griser_images(clone)
        ancre.addnext(clone)
        copies.append(clone)
    print("  elements colores ramenes au noir/blanc : %d" % couleurs)
    print("  logos passes en niveaux de gris        : %d" % grises)

    # Le dernier paragraphe de la copie porte le sectPr noir : cela cree la
    # section dediee et empeche la bordure noire de deborder sur la suite.
    dernier = copies[0]
    if not dernier.tag.endswith('}p'):
        dernier = dernier.makeelement(qn('w:p'), {})
        copies[0].addnext(dernier)
    p_pr = dernier.find(qn('w:pPr'))
    if p_pr is None:
        p_pr = dernier.makeelement(qn('w:pPr'), {})
        dernier.insert(0, p_pr)
    p_pr.append(sectPr_nb)

    doc.save(SORTIE)

    # ── Verification ─────────────────────────────────────────────────────────
    verif = Document(SORTIE)
    B = './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    print("\n  sections : %d" % len(verif.sections))
    for i, s in enumerate(verif.sections):
        x = s._sectPr.xml
        c = sorted(set(re.findall(r'w:color=.([0-9A-Fa-f]{6})', x))) if 'pgBorders' in x else []
        print("    section %d | bordure %-5s | %s" % (i, 'pgBorders' in x, c or '-'))
    # Images presentes dans les deux couvertures
    body2 = verif.element.body
    idx = index_fin_couverture(body2)
    img_couleur = sum(len(e.findall(B)) for e in list(body2)[:idx + 1])
    print("\n  images dans la couverture couleur : %d" % img_couleur)
    print("  paragraphes : %d | tableaux : %d"
          % (len(verif.paragraphs), len(verif.tables)))
    print("\nEnregistre : %s" % SORTIE)


if __name__ == "__main__":
    main()
