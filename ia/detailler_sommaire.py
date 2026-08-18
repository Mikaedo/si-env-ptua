# -*- coding: utf-8 -*-
"""
Fait figurer les sous-chapitres, avec leur page, dans le sommaire.

Le sommaire s'arretait aux chapitres. L'ecole demande que leurs subdivisions y
apparaissent aussi, chacune avec sa pagination.

Les numeros ne sont pas recalcules : ils sont repris de la table des matieres,
champ Word remis a jour a chaque export, donc deja exacte. Les recalculer a la
main reviendrait a entretenir deux sources de verite pour une meme information,
avec la certitude qu'elles divergent un jour.

Seules les subdivisions numerotees sont retenues, du type 4.2 ou 5.3 bis. Les
intertitres non numerotes qui structurent certains chapitres, « Analyse de
l'existant » ou « Module d'intelligence artificielle », restent en dehors : les
faire entrer transformerait le sommaire en doublon de la table des matieres,
alors que sa fonction est de donner une vue d'ensemble tenant d'un regard.

L'allongement du sommaire ne decale pas la pagination reprise ici. Les pages
liminaires portent des chiffres romains dans une section distincte, et le corps
recommence a un : ajouter des lignes en amont ne deplace donc aucune page du
corps.
"""
import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm, Pt

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_sommaire_detaille.docx")

LARGEUR_UTILE = Cm(15.5)
RETRAIT = Cm(0.8)
NUMEROTEE = re.compile(r"^(\d+)\.\d+")


def subdivisions(doc):
    """Sous-chapitres numerotes, groupes par chapitre, avec leur page."""
    depart = next(i for i, p in enumerate(doc.paragraphs)
                  if p.text.strip() == "TABLE DES MATIÈRES")
    groupes = {}
    for p in doc.paragraphs[depart + 1:]:
        if p.style.name != "toc 2":
            continue
        morceaux = p.text.strip().split("\t")
        if len(morceaux) < 2:
            continue
        libelle, page = morceaux[0].strip(), morceaux[-1].strip()
        accroche = NUMEROTEE.match(libelle)
        if not accroche or not page.isdigit():
            continue
        groupes.setdefault(int(accroche.group(1)), []).append((libelle, page))
    return groupes


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    groupes = subdivisions(doc)
    total = sum(len(v) for v in groupes.values())
    print(f"{total} subdivisions reparties sur {len(groupes)} chapitres")

    paras = doc.paragraphs
    debut = next(i for i, p in enumerate(paras) if p.text.strip() == "SOMMAIRE")
    fin = next(i for i, p in enumerate(paras)
               if i > debut and p.text.strip() == "LISTE DES FIGURES")

    # Les lignes de chapitre servent de points d'ancrage ; le modele de mise en
    # forme est pris sur une ligne existante non grasse, pour que les nouvelles
    # lignes heritent police, taille et interligne sans rien fixer en dur.
    modele = next(p for p in paras[debut + 1:fin]
                  if p.text.strip() and not (p.runs and p.runs[0].bold))

    ancres = []
    for p in paras[debut + 1:fin]:
        accroche = re.match(r"^Chapitre (\d+)", p.text.strip())
        if accroche:
            ancres.append((int(accroche.group(1)), p))

    poses = 0
    for numero, ancre in ancres:
        precedent = ancre
        for libelle, page in groupes.get(numero, []):
            element = copy.deepcopy(modele._element)
            precedent._element.addnext(element)
            from docx.text.paragraph import Paragraph
            ligne = Paragraph(element, modele._parent)

            for fragment in list(ligne.runs)[1:]:
                fragment._element.getparent().remove(fragment._element)
            texte = f"{libelle}\t{page}"
            if ligne.runs:
                ligne.runs[0].text = texte
                ligne.runs[0].bold = False
            else:
                ligne.add_run(texte)

            forme = ligne.paragraph_format
            forme.left_indent = RETRAIT
            taquets = forme.tab_stops
            for _ in range(len(taquets)):
                taquets._pPr.tabs.remove(taquets._pPr.tabs[0])
            taquets.add_tab_stop(LARGEUR_UTILE - RETRAIT,
                                 WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            for fragment in ligne.runs:
                fragment.font.size = Pt(11)

            precedent = ligne
            poses += 1
        print(f"  chapitre {numero} : {len(groupes.get(numero, []))} lignes")

    doc.save(SOURCE)
    print(f"\n{poses} lignes ajoutees au sommaire")
    print(f"Sauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
