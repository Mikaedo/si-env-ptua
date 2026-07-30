"""
inserer_captures_annexes.py
===========================
Insere les captures de code (PNG) dans les Annexes A et B du memoire.
"""
from docx import Document
from docx.shared import Inches, Pt
import os

CHEMIN_ENTREE = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v48.docx"
CHEMIN_SORTIE = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v51.docx"
DOSSIER_CAPTURES = r"D:\etude_soutenance\SI-ENV\ia\captures_annexes"

doc = Document(CHEMIN_ENTREE)

# Trouver les paragraphes des annexes
annexe_a_idx = None
annexe_b_idx = None
annexe_c_idx = None
placeholder_idx = None

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if "Annexe A" in t and "Extraits de code" in t:
        annexe_a_idx = i
    elif "Annexe B" in t:
        annexe_b_idx = i
    elif "Annexe C" in t:
        annexe_c_idx = i
    elif "contenu des annexes" in p.text:
        placeholder_idx = i

print(f"Annexe A: paragraphe {annexe_a_idx}")
print(f"Annexe B: paragraphe {annexe_b_idx}")
print(f"Annexe C: paragraphe {annexe_c_idx}")
print(f"Placeholder: paragraphe {placeholder_idx}")

# Captures a inserer
captures = [
    # (apres quel paragraphe, fichier image, legende)
    (annexe_a_idx, "annexe_a1_yolov8n.png", "Figure A.1 : Script d'entrainement YOLOv8n (1_entrainer_detection.py)"),
    (annexe_a_idx, "annexe_a2_mobilenet_p1.png", "Figure A.2 : Script d'entrainement MobileNetV2 - Partie 1 (2_entrainer_classification.py)"),
    (annexe_a_idx, "annexe_a2_mobilenet_p2.png", "Figure A.3 : Script d'entrainement MobileNetV2 - Partie 2"),
    (annexe_b_idx, "annexe_b1_metriques_yolo.png", "Figure B.1 : Calcul des metriques de detection YOLOv8n (mAP, precision, rappel, F1)"),
    (annexe_b_idx, "annexe_b2_metriques_mobilenet.png", "Figure B.2 : Calcul des metriques de classification MobileNetV2 (classification_report + matrice de confusion)"),
]

# Inserer les captures apres les paragraphes correspondants
# On insere en ordre inverse pour ne pas decaler les index

for ref_idx, img_file, legende in reversed(captures):
    ref_para = doc.paragraphs[ref_idx]
    
    # Creer un paragraphe pour l'image
    img_path = os.path.join(DOSSIER_CAPTURES, img_file)
    if not os.path.exists(img_path):
        print(f"ATTENTION: {img_path} non trouve")
        continue
    
    # Paragraphe legende
    legende_p = doc.add_paragraph()
    legende_run = legende_p.add_run(legende)
    legende_run.font.name = 'Times New Roman'
    legende_run.font.size = Pt(10)
    legende_run.font.italic = True
    ref_para._element.addnext(legende_p._element)
    
    # Paragraphe image
    img_p = doc.add_paragraph()
    img_p.alignment = 1  # Center
    run = img_p.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    ref_para._element.addnext(img_p._element)
    
    # Paragraphe vide (espacement)
    empty_p = doc.add_paragraph()
    ref_para._element.addnext(empty_p._element)
    
    print(f"  Insere: {img_file} apres paragraphe {ref_idx}")

# Supprimer le placeholder
if placeholder_idx is not None:
    placeholder = doc.paragraphs[placeholder_idx]
    for run in placeholder.runs:
        run.text = ""
    print(f"Placeholder vide au paragraphe {placeholder_idx}")

doc.save(CHEMIN_SORTIE)
print(f"\n>> Document sauvegarde : {CHEMIN_SORTIE}")
