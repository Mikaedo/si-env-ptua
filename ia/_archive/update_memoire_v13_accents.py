# -*- coding: utf-8 -*-
"""
Restitue les accents dans les textes ajoutes au memoire.

Les legendes des figures d'annexes, les introductions d'annexes, l'annexe D et
les entrees correspondantes de la liste des figures avaient ete redigees sans
accents. Le reste du memoire etant accentue, l'ecart etait visible.

Le remplacement est volontairement RESTREINT aux paragraphes ajoutes, reperes
par des marqueurs surs : sans cette restriction, un mot comme « traces »
(legitime au sens de « traces ») serait transforme a tort dans le corps du
texte.

Sortie : MEMOIRE_N'GUESSAN_v13.docx
"""
import os
import re
import shutil

from docx import Document

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v12.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v13.docx")

#: Un paragraphe est traite s'il contient l'un de ces marqueurs.
MARQUEURS = (
    'Figure A.1', 'Figure B.1', 'Figure B.2', 'Figure C.1', 'Figure C.2',
    'Figure D.1', 'Figure D.2', 'Figure D.3', 'Figure D.4', 'Figure D.5',
    'Annexe D :', 'figures B.1 et B.2', 'figures C.1 et C.2',
    'La figure A.1 reproduit', 'Les figures D.1',
)

#: Mots a re-accentuer. L'ordre importe : les formes les plus longues d'abord,
#: pour eviter qu'un remplacement partiel ne casse un mot plus long.
MOTS = [
    ("hyperparametres", "hyperparamètres"),
    ("parametrage", "paramétrage"),
    ("parametres", "paramètres"),
    ("d'entrainement", "d'entraînement"),
    ("metriques", "métriques"),
    ("d'evaluation", "d'évaluation"),
    ("evaluation", "évaluation"),
    ("detection", "détection"),
    ("precision", "précision"),
    ("criticite", "criticité"),
    ("modeles", "modèles"),
    ("utilises", "utilisés"),
    ("presentent", "présentent"),
    ("apres", "après"),
    ("erosion", "érosion"),
    ("Generation", "Génération"),
    ("generation", "génération"),
    ("generees", "générées"),
    ("creees", "créées"),
    ("d'ecran", "d'écran"),
    ("recenses", "recensés"),
    ("caractere", "caractère"),
    ("verifiable", "vérifiable"),
    ("reglementaire", "réglementaire"),
    ("deploiement", "déploiement"),
    ("donnees", "données"),
    ("affichees", "affichées"),
    ("reponse", "réponse"),
    ("retourne", "retourné"),
    ("resolu", "résolu"),
    ("geolocalises", "géolocalisés"),
    ("periode", "période"),
    ("calcules", "calculés"),
    ("interpretation", "interprétation"),
    ("precedant", "précédant"),
    ("signe ", "signé "),
    ("role ", "rôle "),
    ("les traces des axes", "les tracés des axes"),
]


def accentuer(texte):
    for sans, avec in MOTS:
        texte = texte.replace(sans, avec)
        # Variante capitalisee en debut de phrase
        if sans[0].islower():
            texte = texte.replace(sans[0].upper() + sans[1:],
                                  avec[0].upper() + avec[1:])
    return texte


def main():
    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)

    traites = mots_corriges = 0
    for par in doc.paragraphs:
        if not any(m in par.text for m in MARQUEURS):
            continue
        avant_par = par.text
        for run in par.runs:
            apres = accentuer(run.text)
            if apres != run.text:
                run.text = apres
        if par.text != avant_par:
            traites += 1
            # Compte approximatif des mots modifies
            mots_corriges += sum(
                1 for sans, _ in MOTS if sans in avant_par)

    print("  paragraphes re-accentues : %d" % traites)
    print("  occurrences corrigees    : ~%d" % mots_corriges)

    doc.save(SORTIE)

    # Verification : plus aucun mot cible sans accent dans les zones traitees
    verif = Document(SORTIE)
    restants = []
    for par in verif.paragraphs:
        if not any(m in par.text for m in MARQUEURS):
            continue
        for sans, _ in MOTS:
            if sans in par.text:
                restants.append((sans, par.text.strip()[:52]))
    print("\n  mots sans accent restants : %d" % len(restants))
    for sans, extrait in restants[:8]:
        print("    ! %s -> %s" % (sans, extrait))

    print("\nEnregistre : %s" % SORTIE)
    print("Paragraphes : %d | Tableaux : %d"
          % (len(verif.paragraphs), len(verif.tables)))


if __name__ == "__main__":
    main()
