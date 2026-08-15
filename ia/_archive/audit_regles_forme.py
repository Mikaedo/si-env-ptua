# -*- coding: utf-8 -*-
"""
Audite le memoire au regard du « Guide complet des regles de forme
MEMOIRE MIAGE / UPB ». Ne modifie rien : produit un rapport de conformite.
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

CHEMIN = sys.argv[1] if len(sys.argv) > 1 else \
    r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_N'GUESSAN_v9.docx"

BLIP = './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
doc = Document(CHEMIN)
ok, ko = [], []


def verdict(condition, libelle, detail=""):
    (ok if condition else ko).append((libelle, detail))


# ── 1. Mise en forme generale ────────────────────────────────────────────────
s = doc.sections[0]
verdict(abs(s.page_width.cm - 21) < 0.2 and abs(s.page_height.cm - 29.7) < 0.2,
        "Format A4", "%.1f x %.1f cm" % (s.page_width.cm, s.page_height.cm))
verdict(abs(s.left_margin.cm - 3) < 0.15, "Marge gauche 3 cm",
        "%.2f cm" % s.left_margin.cm)
verdict(abs(s.right_margin.cm - 2.5) < 0.15, "Marge droite 2,5 cm",
        "%.2f cm" % s.right_margin.cm)
verdict(abs(s.top_margin.cm - 2.5) < 0.15, "Marge haut 2,5 cm",
        "%.2f cm" % s.top_margin.cm)
verdict(abs(s.bottom_margin.cm - 2.5) < 0.15, "Marge bas 2,5 cm",
        "%.2f cm" % s.bottom_margin.cm)

n = doc.styles['Normal']
verdict(n.font.name == 'Times New Roman', "Police Times New Roman",
        str(n.font.name))
verdict(n.font.size and abs(n.font.size.pt - 12) < 0.1, "Taille 12 pt",
        "%s pt" % (n.font.size.pt if n.font.size else '?'))
verdict(n.paragraph_format.line_spacing and
        abs(n.paragraph_format.line_spacing - 1.5) < 0.01,
        "Interligne 1,5", str(n.paragraph_format.line_spacing))

# ── 2. Justification du corps de texte ──────────────────────────────────────
corps = [p for p in doc.paragraphs
         if len(p.text.strip()) > 120 and not p.style.name.startswith('Heading')]
justifies = [p for p in corps if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY]
taux = 100.0 * len(justifies) / len(corps) if corps else 0
verdict(taux >= 95, "Paragraphes longs justifies",
        "%.1f %% (%d / %d)" % (taux, len(justifies), len(corps)))

# ── 3. Figures : legende numerotee sous l'image, et renvoi dans le texte ────
texte_global = "\n".join(p.text for p in doc.paragraphs)
figures, sans_legende = [], []
for i, p in enumerate(doc.paragraphs):
    if not p._element.findall(BLIP):
        continue
    suivant = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
    m = re.match(r'^Figure\s+([0-9A-D]+\.[0-9]+)\s*[:.]', suivant, re.I)
    if m:
        figures.append(m.group(1))
    else:
        sans_legende.append((i, suivant[:45]))

verdict(not sans_legende,
        "Chaque figure porte une legende numerotee en dessous",
        "%d image(s) sans legende conforme" % len(sans_legende))
if sans_legende:
    for i, t in sans_legende[:12]:
        ko.append(("  -> paragraphe %d" % i, t or "(legende absente)"))

# Renvoi obligatoire depuis le texte
non_citees = []
for num in figures:
    motif = r'figure\s+' + re.escape(num)
    # On exclut la legende elle-meme du comptage
    occurrences = len(re.findall(motif, texte_global, re.I))
    if occurrences < 2:
        non_citees.append(num)
verdict(not non_citees, "Chaque figure est citee dans le texte",
        "non citees : %s" % (", ".join(non_citees) if non_citees else "aucune"))

# ── 4. Tableaux : titre AU-DESSUS et numerote ───────────────────────────────
titres_tab = re.findall(r'Tableau\s+([0-9]+\.[0-9]+)\s*[:.]', texte_global)
verdict(len(titres_tab) > 0, "Tableaux numerotes (Tableau x.y)",
        "%d titre(s) de tableau detecte(s) pour %d tableau(x)"
        % (len(titres_tab), len(doc.tables)))
non_citees_tab = [t for t in set(titres_tab)
                  if len(re.findall(r'tableau\s+' + re.escape(t), texte_global, re.I)) < 2]
verdict(not non_citees_tab, "Chaque tableau est cite dans le texte",
        "non cites : %s" % (", ".join(sorted(non_citees_tab)) if non_citees_tab else "aucun"))

# ── 5. References bibliographiques style IEEE [n] ───────────────────────────
refs = re.findall(r'\[(\d+)\]', texte_global)
verdict(len(refs) > 0, "References entre crochets presentes",
        "%d occurrence(s), %d source(s) distincte(s)"
        % (len(refs), len(set(refs))))

# ── 6. Style de redaction ───────────────────────────────────────────────────
# « on » familier, hors expressions figees
on_familier = re.findall(r"\bon\s+(?:a|va|peut|doit|utilise|voit|constate|note)\b",
                         texte_global, re.I)
verdict(not on_familier, "Absence de « on » familier",
        "%d occurrence(s) : %s" % (len(on_familier), ", ".join(on_familier[:5])))

# Signes non conformes au corps de texte
verdict('\u2192' not in texte_global, "Absence de fleche Unicode",
        "presente" if '\u2192' in texte_global else "aucune")
verdict('\u2014' not in texte_global, "Absence de tiret cadratin",
        "%d occurrence(s)" % texte_global.count('\u2014'))

# ── 7. Ordre des remerciements ──────────────────────────────────────────────
bloc = ""
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().upper().startswith('REMERCIEMENT'):
        bloc = "\n".join(x.text for x in doc.paragraphs[i:i + 25])
        break
if bloc:
    attendu = ['UPB', 'encadr', 'jury', 'AGEROUTE', 'famille']
    positions = []
    for mot in attendu:
        m = re.search(mot, bloc, re.I)
        positions.append(m.start() if m else -1)
    presentes = [(a, p) for a, p in zip(attendu, positions) if p >= 0]
    ordonne = all(presentes[i][1] < presentes[i + 1][1]
                  for i in range(len(presentes) - 1))
    verdict(ordonne, "Ordre des remerciements",
            " < ".join(a for a, _ in presentes) if presentes else "non detecte")
    manquants = [a for a, p in zip(attendu, positions) if p < 0]
    if manquants:
        ko.append(("  -> destinataires non detectes", ", ".join(manquants)))
else:
    ko.append(("Section Remerciements", "non trouvee"))

# ── Rapport ─────────────────────────────────────────────────────────────────
print("=" * 74)
print("AUDIT DES REGLES DE FORME - %s" % CHEMIN.split("\\")[-1])
print("=" * 74)
print("\nCONFORME (%d)" % len(ok))
for libelle, detail in ok:
    print("  [ok] %-48s %s" % (libelle, detail))
print("\nA CORRIGER (%d)" % len(ko))
for libelle, detail in ko:
    print("  [KO] %-48s %s" % (libelle, detail))
print("\n%d/%d regles verifiees conformes" % (len(ok), len(ok) + len(ko)))
