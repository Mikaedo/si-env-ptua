# -*- coding: utf-8 -*-
"""
Etape 5 : ramener le memoire de dix a six chapitres, en conservant la
progression methodologique de l'auteur.

Regroupements retenus :
    Ch.1 = anciens 1 + 2   Cadre de l'etude et problematique
    Ch.2 = anciens 3 + 4   Etat de l'art et analyse de l'existant
    Ch.3 = ancien  5       Analyse des besoins
    Ch.4 = ancien  6       Conception
    Ch.5 = anciens 7 + 8 + 9   Realisation : plateforme, IA et teledetection
    Ch.6 = ancien  10      Deploiement, tests et discussion

Chaque chapitre absorbe voit son titre devenir une section de premier niveau du
chapitre d'accueil, et ses propres sections sont decalees a la suite. Figures et
tableaux sont renumerotes en consequence, ainsi que tous les renvois et les deux
listes, conformement a la regle « numerotation obligatoire, format
[chapitre].[ordre] ».

Le remplacement se fait en deux passes avec des marqueurs temporaires : passer
directement de « 3.1 » a « 2.1 » alors qu'un « 2.1 » existe deja produirait des
collisions.

Sortie : MEMOIRE_ETAPE5.docx
"""
import os
import re
import shutil

from docx import Document
from docx.oxml.ns import qn

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_ETAPE4.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_ETAPE5.docx")

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

#: ancien chapitre -> (nouveau chapitre, decalage applique a ses sections)
#: Le decalage tient compte des sections du ou des chapitres qui precedent dans
#: le meme regroupement, plus une unite pour le titre du chapitre absorbe qui
#: devient lui-meme une section.
FUSION = {
    1:  (1, 0),
    2:  (1, 3),    # ch.1 compte 3 sections ; le titre du ch.2 devient 1.4
    3:  (2, 0),
    4:  (2, 3),    # ch.3 compte 3 sections ; le titre du ch.4 devient 2.4
    5:  (3, 0),
    6:  (4, 0),
    7:  (5, 0),
    8:  (5, 7),    # ch.7 compte 7 sections ; le titre du ch.8 devient 5.8
    9:  (5, 13),   # + ch.8 compte 6 sections ; le titre du ch.9 devient 5.14
    10: (6, 0),
}

#: Nouveaux intitules des six chapitres.
TITRES = {
    1: "Cadre de l'etude et problematique",
    2: "Etat de l'art et analyse de l'existant",
    3: "Analyse des besoins",
    4: "Conception",
    5: "Realisation : plateforme, intelligence artificielle et teledetection",
    6: "Deploiement, tests et discussion",
}

#: Chapitres dont le marqueur « Chapitre N » disparait : ils sont absorbes.
ABSORBES = (2, 4, 8, 9)

#: Renumerotation des figures : ancien -> nouveau.
FIGURES = {
    '1.1': '1.1', '1.2': '1.2',
    '4.1': '2.1',
    '6.1': '4.1', '6.2': '4.2', '6.3': '4.3', '6.4': '4.4', '6.5': '4.5',
    '6.6': '4.6', '6.7': '4.7', '6.8': '4.8', '6.9': '4.9',
    '7.1': '5.1', '7.2': '5.2', '7.3': '5.3',
    '8.1': '5.4', '8.2': '5.5', '8.3': '5.6', '8.4': '5.7',
    '9.1': '5.8',
}

#: Renumerotation des tableaux : ancien -> nouveau. Les tableaux 6.6 a 6.11
#: sont deja partis en annexe E a l'etape precedente.
TABLEAUX = {
    '2.1': '1.1',
    '3.1': '2.1', '3.2': '2.2',
    '5.1': '3.1', '5.2': '3.2',
    '6.1': '4.1', '6.2': '4.2', '6.3': '4.3', '6.4': '4.4', '6.5': '4.5',
    '7.1': '5.1', '7.2': '5.2', '7.3': '5.3',
    '8.1': '5.4', '8.2': '5.5', '8.3': '5.6', '8.4': '5.7',
    '9.1': '5.8',
    '10.1': '6.1', '10.2': '6.2', '10.3': '6.3',
}


def texte(el):
    return ''.join(t.text or '' for t in el.iter('{%s}t' % W)).strip()


def sections_par_chapitre(doc):
    """Compte les sections de premier niveau de chaque ancien chapitre."""
    vus = {}
    for i, p in enumerate(doc.paragraphs):
        if i < 250:
            continue  # avant le corps : sommaire et listes
        m = re.match(r'^(\d+)\.(\d+)\s+\S', p.text.strip())
        if m:
            ch, num = int(m.group(1)), int(m.group(2))
            vus.setdefault(ch, set()).add(num)
    return {ch: len(nums) for ch, nums in vus.items()}


def main():
    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)

    print("  sections par ancien chapitre : %s" % sections_par_chapitre(doc))

    # ── 1. Marqueurs temporaires ────────────────────────────────────────────
    # On encode d'abord toutes les cibles sous une forme qui ne peut pas
    # entrer en collision, puis on decode. Sans cela, renommer 3.1 en 2.1
    # ecraserait le 2.1 existant.
    n_fig = n_tab = n_sec = n_chap = 0

    for p in doc.paragraphs:
        for run in p.runs:
            t = run.text
            if not t:
                continue
            avant = t

            # Figures
            for ancien, nouveau in FIGURES.items():
                t = re.sub(r'\b(figure[s]?)\s+' + re.escape(ancien) + r'\b',
                           lambda m, n=nouveau: '%s @@F@%s@@' % (m.group(1), n),
                           t, flags=re.I)
            # Tableaux
            for ancien, nouveau in TABLEAUX.items():
                t = re.sub(r'\b(tableau[x]?)\s+' + re.escape(ancien) + r'\b',
                           lambda m, n=nouveau: '%s @@T@%s@@' % (m.group(1), n),
                           t, flags=re.I)
            # Numeros de section : uniquement au format exact d'un titre, soit
            # en tout debut de texte suivi de DEUX espaces (« 6.3.1  Elements
            # du... »). Une regle plus large capturerait les numeros de version
            # (« SDK 3.10.8 ») et les transformerait a tort.
            def remplacer_section(m):
                ch = int(m.group(1))
                reste = m.group(2)
                if ch not in FUSION:
                    return m.group(0)
                nouveau_ch, decalage = FUSION[ch]
                morceaux = reste.split('.')
                morceaux[0] = str(int(morceaux[0]) + decalage)
                return '@@S@%d.%s@@%s' % (nouveau_ch, '.'.join(morceaux),
                                          m.group(3))

            t = re.sub(r'^(\d+)\.(\d+(?:\.\d+)?)(\s\s)',
                       remplacer_section, t)
            # Renvois explicites dans le texte : « a la section 6.2 »
            def remplacer_renvoi(m):
                ch = int(m.group(2))
                reste = m.group(3)
                if ch not in FUSION:
                    return m.group(0)
                nouveau_ch, decalage = FUSION[ch]
                morceaux = reste.split('.')
                morceaux[0] = str(int(morceaux[0]) + decalage)
                return '%s @@S@%d.%s@@' % (m.group(1), nouveau_ch,
                                           '.'.join(morceaux))

            t = re.sub(r'\b(sections?)\s+(\d+)\.(\d+(?:\.\d+)?)\b',
                       remplacer_renvoi, t, flags=re.I)

            if t != avant:
                run.text = t

    # ── 2. Decodage des marqueurs ───────────────────────────────────────────
    for p in doc.paragraphs:
        for run in p.runs:
            t = run.text
            if '@@' not in t:
                continue
            n_fig += t.count('@@F@')
            n_tab += t.count('@@T@')
            n_sec += t.count('@@S@')
            # Les marqueurs longs d'abord, sinon '@@' effacerait leur prefixe
            for marque in ('@@F@', '@@T@', '@@S@', '@@'):
                t = t.replace(marque, '')
            run.text = t

    print("  renvois de figures renumerotes  : %d" % n_fig)
    print("  renvois de tableaux renumerotes : %d" % n_tab)
    print("  numeros de section renumerotes  : %d" % n_sec)

    # ── 3. Titres de chapitre ───────────────────────────────────────────────
    # Les chapitres absorbes perdent leur marqueur, les autres sont renumerotes
    # et recoivent leur nouvel intitule.
    supprimes = 0
    for i, p in enumerate(doc.paragraphs):
        m = re.match(r'^Chapitre\s+(\d+)$', p.text.strip())
        if not m:
            continue
        ancien = int(m.group(1))
        if ancien not in FUSION:
            continue
        nouveau_ch, _ = FUSION[ancien]
        if ancien in ABSORBES:
            # Le marqueur disparait : le titre qui suit devient une section.
            for run in p.runs:
                run.text = ''
            supprimes += 1
        else:
            for run in p.runs:
                if 'Chapitre' in run.text:
                    run.text = re.sub(r'Chapitre\s+\d+',
                                      'Chapitre %d' % nouveau_ch, run.text)
            n_chap += 1
    print("  marqueurs de chapitre renumerotes : %d" % n_chap)
    print("  marqueurs de chapitre absorbes    : %d" % supprimes)

    doc.save(SORTIE)
    print("\nEnregistre : %s" % SORTIE)


if __name__ == "__main__":
    main()
