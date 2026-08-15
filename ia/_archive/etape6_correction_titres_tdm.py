# -*- coding: utf-8 -*-
"""
Etape 6 : corriger ce que l'etape precedente a manque.

Deux defauts constates :

  1. Certains titres de chapitre n'ont pas ete renumerotes parce que leur texte
     est reparti sur plusieurs runs (« Chapitre » puis « 6 ») : un remplacement
     applique run par run ne peut pas reconnaitre le motif complet. On travaille
     donc au niveau du paragraphe, en reecrivant le texte dans le premier run.

  2. La table des matieres, saisie en texte statique et non en champ Word,
     conserve les anciens numeros. Elle est renumerotee avec le meme mapping que
     le corps, et les entrees des chapitres absorbes deviennent de simples
     intitules de section.

Les nouveaux intitules de chapitre sont egalement poses, dans le corps comme
dans la table des matieres.

Sortie : MEMOIRE_ETAPE6.docx
"""
import os
import re
import shutil

from docx import Document

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_ETAPE5.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_ETAPE6.docx")

#: ancien numero de chapitre -> nouveau numero (None = chapitre absorbe)
CHAPITRES = {1: 1, 2: None, 3: 2, 4: None, 5: 3, 6: 4, 7: 5, 8: None,
             9: None, 10: 6}

#: Nouveaux intitules, poses sur le paragraphe qui suit « Chapitre N ».
TITRES = {
    1: "Cadre de l'etude et problematique",
    2: "Etat de l'art et analyse de l'existant",
    3: "Analyse des besoins",
    4: "Conception",
    5: "Realisation : plateforme, intelligence artificielle et teledetection",
    6: "Deploiement, tests et discussion",
}

#: Anciens intitules, pour reconnaitre le paragraphe de titre a remplacer.
ANCIENS_TITRES = {
    1: "Presentation de la structure d'accueil",
    3: "Etat de l'art",
    5: "Analyse des besoins",
    6: "Conception",
    7: "Implementation et environnement technique",
    10: "Deploiement, tests et discussion",
}

#: Decalage des sections par ancien chapitre, identique a l'etape 5.
DECALAGE = {1: (1, 0), 2: (1, 3), 3: (2, 0), 4: (2, 3), 5: (3, 0),
            6: (4, 0), 7: (5, 0), 8: (5, 7), 9: (5, 13), 10: (6, 0)}


W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def ecrire(paragraphe, texte):
    """Remplace le texte d'un paragraphe, hyperliens compris.

    Les entrees d'une table des matieres sont enveloppees dans des elements
    w:hyperlink : leurs runs n'apparaissent pas dans `paragraphe.runs`. Ecrire
    via cette propriete ajouterait du texte sans effacer l'existant, d'ou des
    libelles concatenes. On agit donc sur tous les noeuds de texte du
    paragraphe.
    """
    noeuds = list(paragraphe._element.iter('{%s}t' % W))
    if not noeuds:
        paragraphe.add_run(texte)
        return
    noeuds[0].text = texte
    for n in noeuds[1:]:
        n.text = ''


def main():
    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)
    paras = doc.paragraphs

    # ── 1. Titres de chapitre dans le corps ─────────────────────────────────
    corriges = 0
    for i, p in enumerate(paras):
        m = re.match(r'^Chapitre\s+(\d+)$', p.text.strip())
        if not m:
            continue
        ancien = int(m.group(1))
        # A ce stade le corps porte deja des numeros partiellement corriges :
        # on se fie a l'intitule qui suit pour identifier le chapitre d'origine.
        suivant = ''
        pos_titre = None
        for k in range(1, 4):
            if i + k < len(paras) and paras[i + k].text.strip():
                suivant = paras[i + k].text.strip()
                pos_titre = i + k
                break
        origine = None
        for num, titre in ANCIENS_TITRES.items():
            if _normaliser(titre) == _normaliser(suivant):
                origine = num
                break
        if origine is None:
            continue
        nouveau = CHAPITRES[origine]
        if nouveau is None:
            continue
        if ancien != nouveau:
            ecrire(p, "Chapitre %d" % nouveau)
            corriges += 1
        # Nouvel intitule
        if pos_titre is not None:
            ecrire(paras[pos_titre], TITRES[nouveau])
    print("  numeros de chapitre corriges dans le corps : %d" % corriges)

    # ── 2. Table des matieres ───────────────────────────────────────────────
    # La mention « Table des matieres » apparait deux fois : une premiere comme
    # simple entree du sommaire, en tete de document, et une seconde comme
    # veritable table detaillee. On retient la derniere, et on verifie qu'elle
    # est bien suivie d'entrees paginees (reperables a leur tabulation).
    debut_tdm = None
    for i, p in enumerate(paras):
        if not p.text.strip().upper().startswith('TABLE DES MATI'):
            continue
        suite = ' '.join(x.text for x in paras[i + 1:i + 8])
        if '\t' in suite and debut_tdm is None:
            debut_tdm = i  # premiere table reellement paginee
    if debut_tdm is None:
        print("  ! table des matieres introuvable")
        doc.save(SORTIE)
        return

    n_chap = n_sec = supprimees = 0
    i = debut_tdm + 1
    while i < len(paras):
        t = paras[i].text
        if not t.strip():
            i += 1
            continue
        if t.strip().upper().startswith(('RESUME', 'R\u00c9SUM\u00c9', 'ABSTRACT',
                                         'ANNEXES', 'BIBLIOGRAPHIE')):
            break

        # Entree de chapitre : « Chapitre N<tab>page »
        m = re.match(r'^Chapitre\s+(\d+)(\t.*)?$', t.strip())
        if m:
            ancien = int(m.group(1))
            suffixe = m.group(2) or ''
            nouveau = CHAPITRES.get(ancien)
            if nouveau is None:
                # Chapitre absorbe : l'entree disparait, son intitule reste
                # comme simple ligne de section.
                ecrire(paras[i], '')
                supprimees += 1
            else:
                ecrire(paras[i], "Chapitre %d%s" % (nouveau, suffixe))
                n_chap += 1
                # L'intitule suit : on pose le nouveau
                if i + 1 < len(paras):
                    suite = paras[i + 1].text
                    tab = suite[suite.find('\t'):] if '\t' in suite else ''
                    ecrire(paras[i + 1], TITRES[nouveau] + tab)
            i += 1
            continue

        # Entree de section : « 6.3.1  Intitule<tab>page »
        m = re.match(r'^(\d+)\.(\d+(?:\.\d+)?)(\s\s.*)$', t)
        if m:
            ancien_ch = int(m.group(1))
            if ancien_ch in DECALAGE:
                nouveau_ch, decalage = DECALAGE[ancien_ch]
                morceaux = m.group(2).split('.')
                morceaux[0] = str(int(morceaux[0]) + decalage)
                ecrire(paras[i], "%d.%s%s"
                       % (nouveau_ch, '.'.join(morceaux), m.group(3)))
                n_sec += 1
        i += 1

    print("  entrees de chapitre renumerotees : %d" % n_chap)
    print("  entrees de chapitre supprimees   : %d" % supprimees)
    print("  entrees de section renumerotees  : %d" % n_sec)

    doc.save(SORTIE)
    print("\nEnregistre : %s" % SORTIE)


def _normaliser(s):
    """Comparaison insensible aux accents et a la casse."""
    import unicodedata
    s = unicodedata.normalize('NFD', s)
    s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
    return re.sub(r'[^a-z0-9]', '', s.lower())


if __name__ == "__main__":
    main()
