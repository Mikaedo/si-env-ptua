# -*- coding: utf-8 -*-
"""
Ramene le corps du memoire sous le plafond de pages.

Les ajouts consacres au volet citoyen ont porte le corps de cinquante et une a
cinquante-trois pages. Plutot que d'entamer le texte existant, la compression
porte sur les paragraphes que je viens moi-meme d'ecrire : c'est la ou se
trouve le plus de redondance, un texte redige d'un trait comportant toujours
des reprises que la relecture resserre.

Le principe est de ne rien retirer de ce qui est demontre, mais de dire la
meme chose plus court. Les deux points que le jury peut interroger, la
justification du rayon d'influence et la portee reelle du controle de position,
sont conserves intacts : ce sont eux qui protegent le candidat.
"""
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_resserrement.docx")

# Chaque entree associe un debut de paragraphe a sa version resserree.
REECRITURES = [
    (
        "Le mécanisme de gestion des plaintes du PTUA reposait",
        "Le mécanisme de gestion des plaintes du PTUA reposait jusqu'ici sur un "
        "recueil au guichet ou lors des réunions de quartier, ce qui suppose qu'un "
        "habitant se déplace et tombe sur une permanence ouverte : beaucoup de "
        "nuisances ne remontaient donc jamais. L'application citoyenne ouvre un "
        "second canal depuis le téléphone de la personne concernée, sans rien "
        "changer au traitement en aval. Les doléances rejoignent la file déjà "
        "instruite par le spécialiste du suivi du Plan d'Action de Réinstallation, "
        "et le canal de saisie est conservé afin que l'apport du téléphone soit "
        "mesurable dans les rapports remis au bailleur.",
    ),
    (
        "L'accès est conditionné à la proximité géographique",
        "L'accès est conditionné à la proximité géographique : une plainte "
        "environnementale n'a de sens que si elle émane de quelqu'un qui subit la "
        "nuisance, et le dispositif serait vite saturé s'il acceptait des dépôts "
        "émis de n'importe où. La vérification s'appuie sur un rayon d'influence "
        "propre à chaque chantier, notion que tout PGES manipule sous le nom de "
        "zone d'influence directe : un terrassement lourd dérange plus loin qu'une "
        "reprise de chaussée, et le spécialiste du suivi environnemental fixe cette "
        "étendue ouvrage par ouvrage. Le chantier retenu est le plus proche parmi "
        "ceux dont la zone englobe effectivement la position, et non le plus proche "
        "dans l'absolu, un ouvrage voisin au périmètre resserré pouvant sinon "
        "écarter un site plus lointain mais réellement couvrant.",
    ),
    (
        "Deux choix méritent d'être justifiés",
        "Le rattachement au chantier est déduit de la position et non choisi dans "
        "une liste, un habitant n'ayant aucune raison de connaître les "
        "dénominations administratives des ouvrages. Aucune reconnaissance "
        "automatique n'est par ailleurs sollicitée dans cette application, "
        "contrairement à celle des agents : un riverain décrit une gêne, il ne pose "
        "pas de diagnostic, et la qualification revient au spécialiste qui dispose "
        "du contexte du chantier.",
    ),
    (
        "La portée de ce contrôle doit être énoncée avec exactitude",
        "La portée de ce contrôle doit enfin être énoncée avec exactitude. La "
        "position transmise atteste d'une localisation au moment de l'inscription, "
        "elle ne prouve pas la qualité de riverain : une personne de passage "
        "satisfait la condition, et un dispositif de production devrait lui "
        "adjoindre une vérification d'identité ou une validation par le comité de "
        "quartier. Ce filtre écarte les dépôts manifestement extérieurs à la zone "
        "du projet, ce qui répond à l'objectif poursuivi, sans constituer une "
        "authentification de résidence.",
    ),
    (
        "L'application fonctionne hors connexion d'abord",
        None,  # traite a part : on ne resserre que le fragment ajoute
    ),
]

# Fragment ajoute en 5.3, remplace par une version plus courte.
FRAGMENT_53_LONG = (
    " Le projet produit en réalité deux applications distinctes, bâties sur "
    "un socle de code commun et différenciées au moyen des variantes de "
    "production Android. La première s'adresse aux agents de l'AGEROUTE et "
    "correspond à la description qui précède. La seconde vise les riverains "
    "des chantiers ; elle partage le service d'accès à l'API, les modèles de "
    "données, la géolocalisation et la charte graphique, mais ne conserve "
    "de la première ni la détection d'objets ni la cartographie. Ce découpage "
    "répond à une réalité de terrain autant qu'à une contrainte technique : "
    "les deux publics ne relèvent pas du même canal de distribution, "
    "l'application des agents se déployant en interne quand celle des "
    "riverains a vocation à être téléchargée librement, si bien qu'elles ne "
    "peuvent pas partager un identifiant applicatif. Les deux paquets "
    "s'installent côte à côte sur un même terminal."
)
FRAGMENT_53_COURT = (
    " Le projet produit en réalité deux applications distinctes, bâties sur un "
    "socle commun et différenciées par les variantes de production Android. "
    "Celle des agents correspond à la description qui précède ; celle des "
    "riverains partage le service d'API, les modèles, la géolocalisation et la "
    "charte, sans la détection d'objets ni la cartographie. Les deux publics ne "
    "relevant pas du même canal de distribution, elles ne peuvent pas partager "
    "un identifiant applicatif et s'installent côte à côte sur un même terminal."
)

FRAGMENT_62_LONG = (
    " La couverture s'est étoffée au fil des évolutions : elle porte "
    "aujourd'hui sur le cloisonnement des profils de consultation, sur la "
    "condition de proximité géographique du volet citoyen, sur la traçabilité "
    "des rapports transmis et sur la complétude du jeu de données initial. "
    "Cette dernière vérification parcourt l'énumération des rôles définie dans "
    "le modèle plutôt qu'une liste tenue à la main, de sorte qu'un profil "
    "ajouté sans compte de démonstration fasse échouer la suite."
)
FRAGMENT_62_COURT = (
    " La couverture porte notamment sur le cloisonnement des profils de "
    "consultation, la condition de proximité du volet citoyen, la traçabilité "
    "des rapports transmis et la complétude du jeu de données initial, cette "
    "dernière vérification parcourant l'énumération des rôles du modèle plutôt "
    "qu'une liste tenue à la main."
)


def reecrire(paragraphe, texte):
    for segment in list(paragraphe.runs)[1:]:
        segment._element.getparent().remove(segment._element)
    if paragraphe.runs:
        paragraphe.runs[0].text = texte
    else:
        paragraphe.add_run(texte)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    avant = sum(len(p.text) for p in doc.paragraphs)
    faits = []

    for p in doc.paragraphs:
        t = p.text.strip()

        for debut, remplacement in REECRITURES:
            if remplacement and t.startswith(debut):
                reecrire(p, remplacement)
                faits.append(debut[:44])
                break

        if FRAGMENT_53_LONG.strip() in p.text:
            reecrire(p, p.text.replace(FRAGMENT_53_LONG, FRAGMENT_53_COURT))
            faits.append("fragment 5.3")

        if FRAGMENT_62_LONG.strip() in p.text:
            reecrire(p, p.text.replace(FRAGMENT_62_LONG, FRAGMENT_62_COURT))
            faits.append("fragment 6.2")

    apres = sum(len(p.text) for p in doc.paragraphs)
    doc.save(SOURCE)

    for f in faits:
        print(f"  resserre : {f}")
    print(f"\n{avant - apres} caracteres retires ({len(faits)} passages)")
    print(f"Sauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
