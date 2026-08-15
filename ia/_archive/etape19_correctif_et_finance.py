# -*- coding: utf-8 -*-
"""
Etape 19 : quatre corrections regroupees en un seul passage.

  a) Collision de numerotation. Le tableau « Composition du jeu de donnees »
     partage le numero 5.2 avec le tableau « Packages Flutter », alors que le
     texte qui l'introduit dit deja « le tableau 5.3 detaille la composition du
     corpus ». Les renvois en prose anticipaient donc deja la bonne
     numerotation ; seules les legendes des tableaux 5.2 (le second), 5.3, 5.4,
     5.5 et 5.6 sont decalees d'un cran (5.3 a 5.7), du plus grand numero vers
     le plus petit pour ne jamais creer de doublon transitoire.

  b) Deux paragraphes issus de la condensation de l'etape 10 avaient echappe
     au correcteur d'accents general (mots absents du dictionnaire) : ils sont
     reecrits integralement, accents compris.

  c) Ajout d'une section 6.7 « Etude financiere et analyse cout-benefice »,
     avec un tableau 6.4 donnant une estimation indicative des couts de
     developpement et de fonctionnement du SI-ENV. Les montants sont explicitement
     qualifies d'indicatifs : ce sont des hypotheses de calcul, pas des factures
     reelles, a ajuster par l'auteur selon la gratification de stage et l'offre
     d'hebergement effectivement retenues.

  d) Reconstruction de la « LISTE DES TABLEAUX » : le bloc actuel contient une
     sequence dupliquee (tableaux 6.1 a 6.3 lister deux fois) et garde les
     anciens numeros non corriges. Il est regenere integralement a partir de
     l'etat reel du corps, une fois les corrections a), c) appliquees.

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os
import re
import sys
import copy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def ecrire(paragraphe, contenu):
    noeuds = list(paragraphe._element.iter(W + 't'))
    if not noeuds:
        paragraphe.add_run(contenu)
        return
    noeuds[0].text = contenu
    for n in noeuds[1:]:
        n.text = ''


def renommer_caption_unique(doc, titre_distinctif, nouveau_numero):
    """Renomme la legende dont le titre (apres les deux-points) est unique."""
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith('Tableau') and titre_distinctif in t:
            m = re.match(r'^Tableau\s+([\w.]+)\s*:\s*(.+)$', t)
            if not m:
                continue
            ecrire(p, 'Tableau %s : %s' % (nouveau_numero, m.group(2)))
            return True
    return False


def main():
    doc = Document(CIBLE)

    # ── a) Collision 5.2 : decalage en cascade, du plus grand au plus petit ──
    cascade = [
        ('Indices environnementaux calculés via GEE', '5.7'),
        ('Hyperparamètres optimisés pour YOLOv8 et MobileNetV2', '5.6'),
        ('Benchmark des modèles de classification (phase 2)', '5.5'),
        ('Benchmark des modèles de détection (phase 1)', '5.4'),
        ("Composition du jeu de données pour l'entraînement (Dataset)", '5.3'),
    ]
    n_renum = 0
    for titre, num in cascade:
        if renommer_caption_unique(doc, titre, num):
            n_renum += 1
    print("  legendes renumerotees (correction de la collision 5.2) : %d / %d"
          % (n_renum, len(cascade)))

    # ── b) Paragraphes non accentues restants ────────────────────────────────
    CORRECTIONS_PARAGRAPHES = [
        (
            "Le SI-ENV repond aux six lacunes",
            "Le SI-ENV répond aux six lacunes du chapitre 2 : signalement "
            "instantané avec diagnostic IA, subjectivité encadrée, "
            "géolocalisation automatique, données centralisées, rapports en "
            "secondes, alertes quasi temps réel (tableau de bord rafraîchi "
            "toutes les 10 à 15 secondes ; la synchronisation mobile reste "
            "déclenchée manuellement). Les limites : tests exécutés en "
            "environnement de développement local, dataset Recycle Trash non "
            "encore validé sur des photographies réelles des chantiers du "
            "PTUA, propagation d'erreur inhérente au pipeline en cascade (une "
            "détection manquée par YOLOv8 dégrade la classification par "
            "MobileNetV2), dépendance Internet pour le satellite, résolution "
            "Sentinel-5P trop grossière au niveau chantier."
        ),
        (
            "Perspectives : au-del",
            "Perspectives : au-delà du déploiement pilote et de "
            "l'enrichissement du dataset (5 000 images), trois axes "
            "prolongeraient le SI-ENV. D'abord l'extension des capteurs IoT "
            "aux nuisances mal couvertes par la vision par ordinateur : "
            "sondes de turbidité et de niveau d'eau pour les eaux "
            "stagnantes, sonomètres pour le bruit, capteurs PM2.5/PM10 pour "
            "les poussières, couplés à des drones pour les zones "
            "inaccessibles au sol, en complément des indices Sentinel dont "
            "la résolution reste trop grossière à l'échelle d'un chantier. "
            "Ensuite le remplacement des seuils empiriques actuels par un "
            "modèle prédictif entraîné sur l'historique des signalements, "
            "des indices satellitaires et des données météorologiques, afin "
            "d'anticiper un risque plutôt que de le constater. Enfin "
            "l'extension à d'autres projets AGEROUTE. Le SI-ENV reste par "
            "ailleurs open source et fonctionne hors ligne, contrairement "
            "aux solutions commerciales (Enablon, Cority) qui exigent "
            "licences et connectivité permanente."
        ),
    ]
    n_acc = 0
    for debut, nouveau in CORRECTIONS_PARAGRAPHES:
        for p in doc.paragraphs:
            if p.text.strip().startswith(debut):
                ecrire(p, nouveau)
                n_acc += 1
                break
    print("  paragraphes ré-accentués : %d / %d"
          % (n_acc, len(CORRECTIONS_PARAGRAPHES)))

    # ── c) Section 6.7 : étude financière ────────────────────────────────────
    ancre = None
    for p in doc.paragraphs:
        if p.text.strip().startswith('Cette troisième partie a présenté'):
            ancre = p
            break
    if ancre is None:
        raise SystemExit("Point d'insertion (conclusion partielle) introuvable")

    # Style de titre 6.6 a reproduire pour le titre 6.7
    ref_h2 = None
    for p in doc.paragraphs:
        if p.text.strip().startswith('6.6') and 'Perspectives' in p.text:
            ref_h2 = p
            break

    ancre_el = ancre._element

    def nouveau_paragraphe(style_source, texte, gras_partiel=None):
        p_el = OxmlElement('w:p')
        ancre_el.addprevious(p_el)
        from docx.text.paragraph import Paragraph
        par = Paragraph(p_el, ancre._parent)
        if style_source is not None:
            par.style = style_source.style
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if style_source is None else par.alignment
        par.add_run(texte)
        return par

    titre67 = nouveau_paragraphe(ref_h2, "6.7  Étude financière et analyse "
                                 "coût-bénéfice")

    intro = nouveau_paragraphe(None,
        "Cette section propose une estimation indicative des coûts de "
        "développement et de fonctionnement du SI-ENV. Les montants "
        "reposent sur des hypothèses de calcul explicites (tableau 6.4) et "
        "sont à ajuster selon le montant réel de la gratification de stage "
        "et l'offre d'hébergement effectivement retenue.")
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    caption = nouveau_paragraphe(None,
        "Tableau 6.4 : Estimation indicative des coûts de développement et "
        "de fonctionnement.")
    for r in caption.runs:
        r.font.size = Pt(10)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    # ── Tableau 6.4 ──
    donnees = [
        ("Poste", "Hypothèse retenue", "Montant estimé"),
        ("Ressources humaines (stage)",
         "3 mois de stage (03 mai – 03 août 2026), valorisés au tarif "
         "usuel de gratification de stage en Côte d'Ivoire",
         "≈ 450 000 FCFA"),
        ("Encadrement académique et technique",
         "Suivi assuré dans le cadre de la mission institutionnelle de "
         "l'encadrant UPB et du maître de stage AGEROUTE",
         "Non facturé"),
        ("Licences logicielles",
         "Stack entièrement open source : FastAPI, Angular, Flutter, "
         "PostgreSQL/PostGIS, Docker, ONNX Runtime",
         "0 FCFA"),
        ("Hébergement (VPS)",
         "Offre VPS avec au moins 2 Go de RAM, accès Docker et PostGIS",
         "60 000 à 120 000 FCFA / an"),
        ("Nom de domaine",
         "Enregistrement et renouvellement annuel",
         "≈ 10 000 FCFA / an"),
        ("Total investissement initial (hors valorisation du stage)",
         "Licences + hébergement + nom de domaine",
         "70 000 à 130 000 FCFA"),
        ("Coût de fonctionnement annuel récurrent",
         "Hébergement + nom de domaine",
         "70 000 à 130 000 FCFA / an"),
    ]

    # Style de reference : tableau 6.3
    style_ref = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith('Tableau 6.3'):
            for tbl in doc.tables:
                pass
            break
    ref_table = None
    els = list(doc.element.body)
    for i, el in enumerate(els):
        if el.tag.endswith('}p'):
            txt = ''.join(t.text or '' for t in el.iter(W + 't')).strip()
            if txt.startswith('Tableau 6.3'):
                for j in range(i + 1, min(i + 3, len(els))):
                    if els[j].tag.endswith('}tbl'):
                        from docx.table import Table
                        ref_table = Table(els[j], doc)
                        break
                break

    table = doc.add_table(rows=len(donnees), cols=3)
    if ref_table is not None and ref_table.style is not None:
        table.style = ref_table.style
    table.autofit = True
    for i, (a, b, c) in enumerate(donnees):
        cells = table.rows[i].cells
        for cell, valeur in zip(cells, (a, b, c)):
            cell.text = ''
            par = cell.paragraphs[0]
            run = par.add_run(valeur)
            run.font.size = Pt(9)
            if i == 0 or a.startswith('Total') or a.startswith('Coût de fonctionnement'):
                run.bold = True
            par.paragraph_format.space_before = Pt(0)
            par.paragraph_format.space_after = Pt(1)

    caption._element.addnext(table._tbl)

    discussion = nouveau_paragraphe(None,
        "Contrairement aux plateformes commerciales de gestion HSE "
        "(Enablon, Cority), qui facturent un abonnement annuel par "
        "utilisateur, le SI-ENV ne supporte aucun coût de licence : "
        "l'ensemble de la chaîne technique repose sur des briques open "
        "source. Le coût récurrent se limite à l'hébergement et au nom de "
        "domaine, ce qui rend la solution reproductible pour d'autres "
        "projets AGEROUTE sans surcoût de licence.")
    discussion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    print("  section 6.7 « Étude financière » insérée avec le tableau 6.4")

    doc.save(CIBLE)
    print("\nPremière sauvegarde effectuée (avant reconstruction de la liste "
          "des tableaux)")

    # ── d) Reconstruction de la liste des tableaux ──────────────────────────
    doc = Document(CIBLE)
    paras = doc.paragraphs

    debut = fin = None
    for i, p in enumerate(paras):
        if p.text.strip().upper() == 'LISTE DES TABLEAUX':
            debut = i
        elif debut is not None and p.text.strip().upper().startswith(
                'LISTE DES SIGLES'):
            fin = i
            break
    if debut is None or fin is None:
        raise SystemExit("Bloc « Liste des tableaux » introuvable")
    print("  ancien bloc « Liste des tableaux » : paragraphes %d a %d"
          % (debut, fin))

    # Recomposition de la verite a partir du corps (legende suivie d'un vrai
    # tableau), dans l'ordre d'apparition.
    els = list(doc.element.body)
    vrais = []
    for i, el in enumerate(els):
        if not el.tag.endswith('}p'):
            continue
        txt = ''.join(t.text or '' for t in el.iter(W + 't')).strip()
        m = re.match(r'^Tableau\s+([\w.]+)\s*:\s*(.+)$', txt)
        if not m:
            continue
        suivi = False
        for j in range(i + 1, min(i + 3, len(els))):
            if els[j].tag.endswith('}tbl'):
                suivi = True
                break
            if els[j].tag.endswith('}p'):
                t2 = ''.join(t.text or '' for t in els[j].iter(W + 't')).strip()
                if t2:
                    break
        if suivi:
            vrais.append((m.group(1), m.group(2)))
    print("  tableaux reels recenses dans le corps : %d" % len(vrais))

    # Modele de mise en forme : le premier paragraphe existant du bloc.
    modele = paras[debut + 2] if debut + 2 < fin else None

    # Suppression de l'ancien contenu (hors titre lui-meme)
    for i in range(fin - 1, debut, -1):
        el = paras[i]._element
        el.getparent().remove(el)

    ancre2 = paras[debut]._element
    for num, titre in vrais:
        p_el = OxmlElement('w:p')
        ancre2.addnext(p_el)
        ancre2 = p_el
        from docx.text.paragraph import Paragraph
        par = Paragraph(p_el, paras[debut]._parent)
        if modele is not None:
            par.style = modele.style
        run = par.add_run("Tableau %s : %s" % (num, titre))
        if modele is not None and modele.runs:
            run.font.size = modele.runs[0].font.size

    print("  liste des tableaux reconstruite : %d entrées" % len(vrais))

    doc.save(CIBLE)
    print("\nEnregistré : %s" % CIBLE)


if __name__ == "__main__":
    main()
