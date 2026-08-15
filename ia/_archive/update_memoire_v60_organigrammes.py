# -*- coding: utf-8 -*-
"""
v60 -> v61 : remplace les Figures 1.1 (organigramme AGEROUTE) et 1.2
(organigramme PTUA) par des versions vectorielles propres (generees en
HTML/CSS puis capturees), fidèles a la structure des documents officiels
fournis mais sans les artefacts de photo (main, reflets, signatures,
tampons) ni le trait parasite au-dessus de "Coordonnateur".
"""
from docx import Document
from docx.oxml.ns import qn

SRC = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v60.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v61.docx"
SCRATCH = r"C:\Users\DELL\AppData\Local\Temp\claude\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25\scratchpad"

doc = Document(SRC)

from docx.shared import Inches

def replace_image_correct(pidx, image_path, width_inches=6.1):
    p = doc.paragraphs[pidx]
    for r in p.runs:
        drawings = r.element.findall(qn('w:drawing'))
        if drawings:
            for d in drawings:
                d.getparent().remove(d)
            r.add_picture(image_path, width=Inches(width_inches))
            print(f"[OK] Figure remplacee au paragraphe {pidx} (largeur {width_inches}in)")
            return
    raise RuntimeError(f"Aucune image trouvee au paragraphe {pidx}")

replace_image_correct(258, f"{SCRATCH}\\new_fig11.png")
replace_image_correct(264, f"{SCRATCH}\\new_fig12.png")

doc.save(DST)
print(f"\n=== MEMOIRE SAUVEGARDE : {DST} ===")
