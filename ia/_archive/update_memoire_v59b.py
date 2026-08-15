# -*- coding: utf-8 -*-
"""
v59 -> v60 : complete le tableau 6.1 (decoupage des sprints) avec deux
sprints supplementaires couvrant la phase de validation/tests reels et la
finalisation, jusqu'au 05/08/2026 (au lieu de s'arreter le 09/07/2026 sans
couvrir le reste du stage).
"""
import copy
from docx import Document

SRC = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v59.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v60.docx"

doc = Document(SRC)
t = doc.tables[9]

# Verifier qu'on a bien la bonne table avant de la modifier
assert t.rows[0].cells[0].text.strip() == "Sprint", f"Table inattendue : {t.rows[0].cells[0].text!r}"

new_rows = [
    ["Sprint 6", "10/07/2026\n23/07/2026",
     "Tests d'integration et fiabilisation.",
     "Deploiement Docker complet (3 conteneurs), cablage du declenchement automatique des alertes (seuil -> alerte -> email), correction du pipeline de detection mobile (NMS), fiabilisation de la synchronisation satellite GEE."],
    ["Sprint 7", "24/07/2026\n05/08/2026",
     "Validation finale et finalisation.",
     "Execution de la suite de tests fonctionnels (32 scenarios), mesure des performances reelles (API, synchronisation, inference), correction des ecarts entre resultats attendus et resultats mesures, redaction finale du memoire."],
]

# On clone la derniere ligne existante (memes styles de police/bordures) puis
# on remplace son texte, plutot que d'utiliser add_row() qui perd la mise en forme.
template_row = t.rows[-1]._tr
for row_data in new_rows:
    new_tr = copy.deepcopy(template_row)
    template_row.addnext(new_tr)
    template_row = new_tr
    new_row = t.rows[-1]
    for j, val in enumerate(row_data):
        cell = new_row.cells[j]
        for p in cell.paragraphs:
            for run in list(p.runs)[1:]:
                run.text = ""
            if p.runs:
                p.runs[0].text = val.split("\n")[0]
                for extra_line in val.split("\n")[1:]:
                    p.runs[0].text += "\n" + extra_line
            else:
                p.add_run(val)

print(f"[OK] {len(new_rows)} sprints ajoutes au tableau 6.1, fin au 05/08/2026")

doc.save(DST)
print(f"\n=== MEMOIRE SAUVEGARDE : {DST} ===")
