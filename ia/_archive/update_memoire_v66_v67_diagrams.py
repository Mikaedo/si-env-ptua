# -*- coding: utf-8 -*-
"""Remplace 3 figures du chapitre 6 par des versions corrigees conformement
au cours UML (INTRODUCTION_A_UML) :
  - Figure 6.2 (cas d'utilisation) : routage orthogonal des associations
    acteur->cas d'utilisation pour qu'aucune ligne ne traverse plus une
    ellipse tierce (defaut du diagramme original).
  - Figure 6.5 (sequence authentification) : le message "rediriger vers
    l'espace applicatif" etait place APRES la fermeture du fragment alt
    (donc execute meme en cas d'echec) ; deplace dans la seule branche
    de succes.
  - Figure 6.6 (sequence generation rapport PGES) : meme bug corrige pour
    le message "telecharger le rapport".
Le diagramme de classe (Figure 6.3) a ete verifie conforme (aucune
redondance d'attribut/methode) et n'est pas touche."""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

SRC = r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_N'GUESSAN_v6.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_N'GUESSAN_v7.docx"
SCRATCH = r"C:\Users\DELL\AppData\Local\Temp\claude\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25\scratchpad"

doc = Document(SRC)
paras = doc.paragraphs

def replace_image(pidx, image_path, width_inches=6.1):
    p = paras[pidx]
    for r in p.runs:
        drawings = r.element.findall(qn('w:drawing'))
        if drawings:
            for d in drawings:
                d.getparent().remove(d)
            r.add_picture(image_path, width=Inches(width_inches))
            print(f"[OK] Figure remplacee au paragraphe {pidx}")
            return
    raise RuntimeError(f"Aucune image trouvee au paragraphe {pidx}")

# Indices retrouves par detection de dessin dans v6 (voir recherche prealable)
replace_image(372, SCRATCH + "/usecase_fixed.png", width_inches=6.3)
replace_image(400, SCRATCH + "/seq65_fixed.png", width_inches=6.1)
replace_image(404, SCRATCH + "/seq66_fixed.png", width_inches=6.3)

doc.save(DST)
print(f"\n=== MEMOIRE SAUVEGARDE : {DST} ===")
