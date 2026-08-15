# -*- coding: utf-8 -*-
"""
Etape 7 : deplacer en annexe F les tableaux de reference, et renumeroter en
cascade ceux qui restent dans le corps.

Criteres de tri retenus :
  - partent en annexe les tableaux qui servent de REFERENCE, consultes
    ponctuellement et non necessaires au fil du raisonnement : decoupage des
    sprints, dictionnaire de donnees, identification des acteurs, liste des
    points d'entree de l'API ;
  - restent dans le corps ceux qui portent la DEMONSTRATION : besoins
    fonctionnels (coeur de l'analyse), benchmarks des modeles, matrice des
    habilitations, synthese des tests.

Les tableaux restants sont renumerotes pour ne laisser aucun trou dans les
sequences, conformement a la regle de numerotation continue.

Sortie : MEMOIRE_ETAPE7.docx
"""
import os
import re
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_ETAPE6.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_ETAPE7.docx")

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

#: Tableaux deplaces : ancien numero -> numero d'annexe.
DEPLACES = {
    '4.1': 'F.1',   # decoupage des sprints
    '4.2': 'F.2',   # dictionnaire de donnees
    '4.4': 'F.3',   # acteurs principaux
    '4.5': 'F.4',   # acteurs secondaires
    '5.2': 'F.5',   # points d'entree de l'API
}

#: Renumerotation des tableaux restants, pour combler les trous laisses.
RESTANTS = {
    '4.3': '4.1',           # seul tableau conserve au chapitre 4
    '5.1': '5.1',
    '5.3': '5.2', '5.4': '5.3', '5.5': '5.4',
    '5.6': '5.5', '5.7': '5.6', '5.8': '5.7',
}


def texte(el):
    return ''.join(t.text or '' for t in el.iter('{%s}t' % W)).strip()


def ecrire(paragraphe, contenu):
    """Ecrit dans un paragraphe, hyperliens compris."""
    noeuds = list(paragraphe._element.iter('{%s}t' % W))
    if not noeuds:
        paragraphe.add_run(contenu)
        return
    noeuds[0].text = contenu
    for n in noeuds[1:]:
        n.text = ''


def main():
    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)
    body = doc.element.body

    # ── Reperage des couples titre + tableau a deplacer ─────────────────────
    elements = list(body)
    couples = []
    for i, el in enumerate(elements):
        if not el.tag.endswith('}p'):
            continue
        m = re.match(r'^Tableau\s+(\d+\.\d+)\s*[:.]', texte(el))
        if not m or m.group(1) not in DEPLACES:
            continue
        tbl = None
        for j in range(i + 1, min(i + 4, len(elements))):
            if elements[j].tag.endswith('}tbl'):
                tbl = elements[j]
                break
        if tbl is not None:
            couples.append((m.group(1), el, tbl))

    print("  couples reperes dans le corps : %d / %d"
          % (len(couples), len(DEPLACES)))
    if len(couples) != len(DEPLACES):
        raise SystemExit("Reperage incomplet, aucune modification appliquee")

    # ── Renvoi laisse dans le corps ─────────────────────────────────────────
    premier = couples[0][1]
    renvoi = premier.makeelement('{%s}p' % W, {})
    premier.addprevious(renvoi)
    par = Paragraph(renvoi, doc.paragraphs[0]._parent)
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.add_run(
        "Le decoupage des sprints, le dictionnaire de donnees, "
        "l'identification detaillee des acteurs et la liste des points "
        "d'entree de l'API sont presentes en annexe F, aux tableaux F.1 a F.5.")

    # ── Creation de l'annexe F ──────────────────────────────────────────────
    doc.add_page_break()
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    titre.add_run("Annexe F : Tableaux de reference de la conception et de "
                  "l'implementation").bold = True

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "Les tableaux F.1 a F.5 rassemblent les elements de reference "
        "consultes ponctuellement : organisation du developpement en sprints "
        "(tableau F.1), dictionnaire de donnees (tableau F.2), identification "
        "des acteurs principaux et secondaires (tableaux F.3 et F.4) et "
        "principaux points d'entree de l'interface de programmation "
        "(tableau F.5).")

    ancre = doc.paragraphs[-1]._element
    for ancien, titre_el, tbl_el in couples:
        nouveau = DEPLACES[ancien]
        for t in titre_el.iter('{%s}t' % W):
            if t.text and 'Tableau' in t.text:
                t.text = t.text.replace('Tableau %s' % ancien,
                                        'Tableau %s' % nouveau)
        ancre.addnext(tbl_el)
        ancre.addnext(titre_el)
        ancre = tbl_el
        print("    Tableau %-5s -> Tableau %s" % (ancien, nouveau))

    # ── Mise a jour des renvois : deplaces puis restants ───────────────────
    # Deux passes avec marqueurs : renommer 5.3 en 5.2 alors qu'un 5.2 existe
    # encore produirait une collision.
    n_dep = n_rest = 0
    for p in doc.paragraphs:
        for run in p.runs:
            t = run.text
            if not t:
                continue
            avant = t
            for ancien, nouveau in DEPLACES.items():
                t = re.sub(r'\b(tableau[x]?)\s+' + re.escape(ancien) + r'\b',
                           lambda m, n=nouveau: '%s @@%s@@' % (m.group(1), n),
                           t, flags=re.I)
            for ancien, nouveau in RESTANTS.items():
                if ancien == nouveau:
                    continue
                t = re.sub(r'\b(tableau[x]?)\s+' + re.escape(ancien) + r'\b',
                           lambda m, n=nouveau: '%s @@%s@@' % (m.group(1), n),
                           t, flags=re.I)
            if t != avant:
                run.text = t

    for p in doc.paragraphs:
        for run in p.runs:
            if '@@' in run.text:
                n_dep += 1
                run.text = run.text.replace('@@', '')

    # Titres des tableaux restants
    for p in doc.paragraphs:
        m = re.match(r'^Tableau\s+(\d+\.\d+)\s*[:.]', p.text.strip())
        if not m:
            continue
        ancien = m.group(1)
        if ancien in RESTANTS and RESTANTS[ancien] != ancien:
            contenu = p.text.replace('Tableau %s' % ancien,
                                     'Tableau %s' % RESTANTS[ancien], 1)
            ecrire(p, contenu)
            n_rest += 1

    print("  renvois mis a jour               : %d" % n_dep)
    print("  titres de tableaux renumerotes   : %d" % n_rest)

    doc.save(SORTIE)
    print("\nEnregistre : %s" % SORTIE)


if __name__ == "__main__":
    main()
