# -*- coding: utf-8 -*-
"""
Etape 10 : condenser les paragraphes les plus longs et corriger les renvois de
chapitre devenus faux apres la fusion.

Deux renvois pointaient encore vers l'ancienne numerotation : « comme detaille
au chapitre 9 » (devenu le chapitre 5) et « les six lacunes du chapitre 4 »
(absorbe par le chapitre 2). Ils sont corriges ici, en meme temps que la
reecriture plus dense de ces paragraphes.

La condensation preserve l'integralite de l'information : seules les formules
redondantes et les redites sont retirees.

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os

from docx import Document

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

#: Reecritures : fragment reconnaissable de l'original -> version condensee.
REECRITURES = [
    (
        "Perspectives : au-del",
        "Perspectives : au-dela du deploiement pilote et de l'enrichissement du "
        "dataset (5 000 images), trois axes prolongeraient le SI-ENV. D'abord "
        "l'extension des capteurs IoT aux nuisances mal couvertes par la vision "
        "par ordinateur : sondes de turbidite et de niveau d'eau pour les eaux "
        "stagnantes, sonometres pour le bruit, capteurs PM2.5/PM10 pour les "
        "poussieres, couples a des drones pour les zones inaccessibles au sol, "
        "en complement des indices Sentinel dont la resolution reste trop "
        "grossiere a l'echelle d'un chantier. Ensuite le remplacement des "
        "seuils empiriques actuels par un modele predictif entraine sur "
        "l'historique des signalements, des indices satellitaires et des "
        "donnees meteorologiques, afin d'anticiper un risque plutot que de le "
        "constater. Enfin l'extension a d'autres projets AGEROUTE. Le SI-ENV "
        "reste par ailleurs open source et fonctionne hors ligne, "
        "contrairement aux solutions commerciales (Enablon, Cority) qui "
        "exigent licences et connectivite permanente."
    ),
    (
        "Pour l'eau stagnante, une observation satellite directe",
        "Pour l'eau stagnante, l'observation satellite directe a ete ecartee : "
        "Sentinel-2 depend trop de la couverture nuageuse et de sa revisite de "
        "5 jours pour un phenomene qui apparait et disparait en quelques heures "
        "apres une pluie. Un indice de risque predictif a ete retenu, croisant "
        "deux couches disponibles dans GEE : le cumul de pluie sur 48 heures "
        "(CHIRPS, UCSB-CHG/CHIRPS/DAILY) et la pente du terrain autour du "
        "chantier (SRTM, USGS/SRTMGL1_003, via ee.Terrain.slope()). La regle "
        "retenue est la suivante : une pluie forte sur pente faible signale un "
        "risque eleve de flaques persistantes ; sur pente marquee, l'ecoulement "
        "naturel limite ce risque ; une pluie faible maintient un risque faible. "
        "Le relief a ete prefere a un indice fonde sur l'historique des "
        "signalements, indisponible avant deploiement (demarrage a froid). Il "
        "s'agit deliberement d'un indice a seuils calibres empiriquement et non "
        "d'un modele entraine, choix realiste au regard du temps de stage."
    ),
    (
        "La teledetection satellitaire permet d'observer",
        "La teledetection satellitaire permet d'observer la surface terrestre "
        "depuis l'espace. Les satellites Sentinel, operes par l'Agence Spatiale "
        "Europeenne dans le cadre du programme Copernicus, fournissent des "
        "donnees gratuites et libres d'acces. Sentinel-2 [15] capture des "
        "images multispectrales a 10 metres de resolution avec une revisite de "
        "5 jours, permettant le calcul du NDWI [16] pour les eaux de surface et "
        "du NDVI pour la vegetation. Sentinel-5P [17] embarque le spectrometre "
        "TROPOMI et mesure la qualite de l'air, notamment le NO2. Google Earth "
        "Engine [18] traite ces donnees sans telechargement local. Le tableau "
        "2.2 compare les deux satellites etudies. Comme detaille au chapitre 5, "
        "le SI-ENV retient Sentinel-5P pour le NO2 mais ecarte Sentinel-2/NDWI "
        "au profit d'un indice de risque pluie/relief (CHIRPS + SRTM), plus "
        "robuste a la couverture nuageuse pour un phenomene aussi ponctuel."
    ),
    (
        "Le SI-ENV r",  # « Le SI-ENV repond aux six lacunes... »
        "Le SI-ENV repond aux six lacunes du chapitre 2 : signalement instantane "
        "avec diagnostic IA, subjectivite encadree, geolocalisation automatique, "
        "donnees centralisees, rapports en secondes, alertes quasi temps reel "
        "(tableau de bord rafraichi toutes les 10 a 15 secondes ; la "
        "synchronisation mobile reste declenchee manuellement). Les limites : "
        "tests executes en environnement de developpement local, dataset Recycle "
        "Trash non encore valide sur des photographies reelles des chantiers du "
        "PTUA, propagation d'erreur inherente au pipeline en cascade (une "
        "detection manquee par YOLOv8 degrade la classification par "
        "MobileNetV2), dependance Internet pour le satellite, resolution "
        "Sentinel-5P trop grossiere au niveau chantier."
    ),
]


def ecrire(paragraphe, contenu):
    noeuds = list(paragraphe._element.iter('{%s}t' % W))
    if not noeuds:
        paragraphe.add_run(contenu)
        return
    noeuds[0].text = contenu
    for n in noeuds[1:]:
        n.text = ''


def main():
    doc = Document(CIBLE)
    gain = 0
    appliques = 0
    for debut, nouveau in REECRITURES:
        for p in doc.paragraphs:
            t = p.text.strip()
            if len(t) < 700 or not t.startswith(debut):
                continue
            gain += len(t) - len(nouveau)
            ecrire(p, nouveau)
            appliques += 1
            break
    print("  paragraphes condenses : %d / %d" % (appliques, len(REECRITURES)))
    print("  caracteres economises : %d" % gain)

    # Renvois de chapitre encore obsoletes ailleurs dans le corps
    corrections = {
        'au chapitre 9': 'au chapitre 5',
        'du chapitre 4': 'du chapitre 2',
        'au chapitre 8': 'au chapitre 5',
        'au chapitre 10': 'au chapitre 6',
        'au chapitre 7': 'au chapitre 5',
    }
    n = 0
    for p in doc.paragraphs:
        for run in p.runs:
            for ancien, nouveau in corrections.items():
                if ancien in run.text:
                    run.text = run.text.replace(ancien, nouveau)
                    n += 1
    print("  renvois de chapitre corriges : %d" % n)

    doc.save(CIBLE)
    print("\nEnregistre : %s" % CIBLE)


if __name__ == "__main__":
    main()
