# -*- coding: utf-8 -*-
"""
Etape 21 : remplacer l'etude financiere par une version fondee sur des
references publiques reelles, datees d'aout 2026.

L'ancienne section 6.7 posait des hypotheses non sourcees. Celle-ci s'appuie
sur quatre recherches web menees pour cette etape :

  - Salaire d'un developpeur web junior a Abidjan : Simoon CV, guide des
    salaires Cote d'Ivoire 2026 (200 000 FCFA/mois en debut de carriere).
  - Hebergement VPS : Systalink, page « Solutions d'hebergement VPS Cote
    d'Ivoire » (a partir de 3 500 FCFA/mois pour l'entree de gamme, 40 000 a
    80 000 FCFA/mois pour les plans hautes performances ; le palier
    intermediaire adapte a Docker + PostGIS n'est pas publie en detail et
    reste a confirmer par devis).
  - Licences des plateformes HSE commerciales : ITQlick et PricingNow pour
    Enablon (a partir de 500 USD/utilisateur/mois, licence annuelle de base
    autour de 20 000 USD/utilisateur) ; Cority et comparables (40-50
    USD/utilisateur/mois plus 15 000 a 200 000 USD de mise en oeuvre).
  - Nom de domaine .ci : NIC-CI via lenomdedomaine.ci (9 000 a 9 500 FCFA/an).

Conversion utilisee : 1 USD ~= 600 FCFA (ordre de grandeur aout 2026, taux
approximatif, a ne pas presenter comme un taux officiel).

Deux tableaux remplacent l'ancien tableau 6.4 :
  - 6.4 : couts de developpement et de fonctionnement du SI-ENV, avec source
    de chaque hypothese ;
  - 6.5 : comparaison indicative avec des plateformes commerciales de gestion
    HSE, pour situer un ordre de grandeur (non un cout reellement evite par
    AGEROUTE, qui n'a jamais souscrit ce type de licence).

Les quatre sources sont ajoutees a la bibliographie, references [24] a [27],
dans le meme style que les entrees existantes.

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table

sys.stdout.reconfigure(encoding='utf-8')

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def ecrire(paragraphe, contenu):
    noeuds = list(paragraphe._element.iter(W + 't'))
    if not noeuds:
        paragraphe.add_run(contenu)
        return
    noeuds[0].text = contenu
    for n in noeuds[1:]:
        n.text = ''


def main():
    doc = Document(CIBLE)
    paras = doc.paragraphs

    # ── Reperage de l'ancienne section 6.7 (titre, intro, legende, tableau,
    #    discussion) et de son ancrage de sortie ──────────────────────────────
    i_titre = i_intro = i_legende = i_discussion = i_sortie = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith('6.7') and 'financière' in t:
            i_titre = i
        elif i_titre is not None and i_intro is None and t.startswith(
                'Cette section propose'):
            i_intro = i
        elif i_titre is not None and t.startswith('Tableau 6.4'):
            i_legende = i
        elif i_titre is not None and t.startswith(
                'Contrairement aux plateformes'):
            i_discussion = i
        elif i_titre is not None and t.startswith(
                'Cette troisième partie a présenté'):
            i_sortie = i
            break
    assert None not in (i_titre, i_intro, i_legende, i_discussion, i_sortie), \
        "Repérage incomplet : %r" % [i_titre, i_intro, i_legende,
                                     i_discussion, i_sortie]
    print("  ancienne section 6.7 : titre=%d intro=%d légende=%d "
          "discussion=%d sortie=%d" % (i_titre, i_intro, i_legende,
                                       i_discussion, i_sortie))

    ancre_sortie_el = paras[i_sortie]._element
    body = doc.element.body

    # Style de reference (titre 6.6, deja noir, deja Heading 2)
    ref_h2_style = paras[i_titre].style

    # Style de la legende de tableau (police 10, keep_with_next) : on relit
    # les proprietes du modele avant de le supprimer.
    modele_legende = paras[i_legende]
    taille_legende = (modele_legende.runs[0].font.size
                      if modele_legende.runs else Pt(10))

    # Style de la table de reference (6.3) pour reproduire l'apparence
    els = list(body)
    ref_table = None
    for i, el in enumerate(els):
        if el.tag.endswith('}p'):
            txt = ''.join(t.text or '' for t in el.iter(W + 't')).strip()
            if txt.startswith('Tableau 6.3'):
                for j in range(i + 1, min(i + 3, len(els))):
                    if els[j].tag.endswith('}tbl'):
                        ref_table = Table(els[j], doc)
                        break
                break

    # ── Suppression de l'ancien contenu (titre a discussion inclus) ─────────
    for i in range(i_discussion, i_titre - 1, -1):
        el = paras[i]._element
        el.getparent().remove(el)
    # Suppression de l'ancien tableau 6.4 (entre l'ex-legende et l'ex-discussion,
    # deja retire des paragraphes mais toujours dans le corps XML)
    for el in list(body):
        if el.tag.endswith('}tbl'):
            tbl_test = Table(el, doc)
            if tbl_test.rows and tbl_test.rows[0].cells[0].text.strip() == 'Poste':
                el.getparent().remove(el)
                break

    print("  ancienne section 6.7 supprimée")

    # ── Reconstruction ────────────────────────────────────────────────────
    ancre_el = ancre_sortie_el

    def nouveau_paragraphe(style, texte, justifie=True):
        p_el = OxmlElement('w:p')
        ancre_el.addprevious(p_el)
        par = Paragraph(p_el, paras[i_sortie]._parent)
        if style is not None:
            par.style = style
        if justifie:
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return par

    def ajouter_table(donnees, gras_lignes=()):
        table = doc.add_table(rows=len(donnees), cols=len(donnees[0]))
        if ref_table is not None and ref_table.style is not None:
            table.style = ref_table.style
        table.autofit = True
        for i, ligne in enumerate(donnees):
            cellules = table.rows[i].cells
            for cell, valeur in zip(cellules, ligne):
                cell.text = ''
                par = cell.paragraphs[0]
                run = par.add_run(valeur)
                run.font.size = Pt(9)
                if i == 0 or i in gras_lignes:
                    run.bold = True
                par.paragraph_format.space_before = Pt(0)
                par.paragraph_format.space_after = Pt(1)
        ancre_el.addprevious(table._tbl)
        return table

    def caption(texte):
        p = nouveau_paragraphe(None, None)
        run = p.add_run(texte)
        run.font.size = taille_legende
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        return p

    # Titre 6.7
    titre = nouveau_paragraphe(ref_h2_style,
                               "6.7  Étude financière et analyse coût-bénéfice")
    for r in titre.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
    if not titre.runs:
        r = titre.add_run("6.7  Étude financière et analyse coût-bénéfice")
        r.font.color.rgb = RGBColor(0, 0, 0)

    # Intro
    nouveau_paragraphe(None,
        "Cette section chiffre le coût de développement et de fonctionnement "
        "du SI-ENV à partir de références publiques datées d'août 2026 "
        "(tableau 6.4), puis situe cet ordre de grandeur face à des "
        "plateformes commerciales de gestion HSE (tableau 6.5). Le taux de "
        "conversion retenu, 1 USD ≈ 600 FCFA, est une approximation d'août "
        "2026 et non un taux officiel.")

    # Tableau 6.4
    caption("Tableau 6.4 : Coûts de développement et de fonctionnement du "
           "SI-ENV, sources publiques (août 2026).")
    donnees_64 = [
        ("Poste", "Hypothèse retenue et source", "Montant estimé"),
        ("Ressources humaines (stage)",
         "3 mois (03 mai – 03 août 2026), valorisés au tarif d'entrée d'un "
         "développeur web junior à Abidjan, 200 000 FCFA/mois [24]",
         "600 000 FCFA"),
        ("Encadrement académique et technique",
         "Mission institutionnelle de l'encadrant UPB et du maître de "
         "stage AGEROUTE",
         "Non facturé"),
        ("Licences logicielles",
         "Stack entièrement open source : FastAPI, Angular, Flutter, "
         "PostgreSQL/PostGIS, Docker, ONNX Runtime",
         "0 FCFA"),
        ("Hébergement (VPS Systalink)",
         "Entrée de gamme à partir de 3 500 FCFA/mois ; un plan avec "
         "≥ 2 Go de RAM, Docker et PostGIS se situe au-dessus, jusqu'à "
         "40 000-80 000 FCFA/mois pour les offres hautes performances [25] "
         "— palier exact à confirmer par devis",
         "3 500 à 80 000 FCFA / mois"),
        ("Nom de domaine .ci",
         "Enregistrement NIC-CI via bureau accrédité, 9 000 à 9 500 "
         "FCFA/an [26]",
         "≈ 9 500 FCFA / an"),
        ("Total investissement initial (hors valorisation du stage)",
         "Licences + premier mois d'hébergement + domaine",
         "13 000 à 89 500 FCFA"),
        ("Coût de fonctionnement annuel récurrent",
         "Hébergement (12 mois) + domaine",
         "51 500 à 969 500 FCFA / an"),
    ]
    ajouter_table(donnees_64, gras_lignes={6, 7})

    nouveau_paragraphe(None,
        "L'écart entre les deux bornes de l'hébergement tient à l'absence de "
        "grille tarifaire publique pour le palier intermédiaire chez "
        "Systalink : seuls le tarif d'entrée et le tarif haute performance "
        "sont publiés [25]. Un devis exact, tenant compte de la charge "
        "réelle des conteneurs Docker (FastAPI, PostgreSQL/PostGIS, Nginx), "
        "reste nécessaire avant le déploiement en production.")

    # Tableau 6.5
    caption("Tableau 6.5 : Comparaison indicative avec des plateformes "
           "commerciales de gestion HSE (août 2026).")
    donnees_65 = [
        ("Solution", "Modèle de licence publié", "Coût annuel indicatif"),
        ("SI-ENV",
         "Open source, aucune licence récurrente",
         "≈ 0,05 à 1 million FCFA/an (hébergement seul)"),
        ("Enablon CSR Suite",
         "À partir de 500 USD/utilisateur/mois [27]",
         "≈ 20 000 USD/utilisateur/an, soit ≈ 12 millions FCFA"),
        ("Enablon Risk Management",
         "Déploiement entreprise [27]",
         "À partir de 50 000 USD/an, soit ≈ 30 millions FCFA"),
        ("Cority et plateformes EHS comparables",
         "40 à 50 USD/utilisateur/mois + 15 000 à 200 000 USD de mise en "
         "œuvre [27]",
         "Plusieurs dizaines de millions de FCFA la première année"),
    ]
    ajouter_table(donnees_65, gras_lignes={1})

    nouveau_paragraphe(None,
        "Ces montants, cités à titre de comparaison d'ordre de grandeur et "
        "non comme un coût réellement évité par AGEROUTE — qui n'a jamais "
        "souscrit ce type de licence —, montrent que le choix d'une pile "
        "entièrement open source évite un poste de dépense récurrent "
        "généralement compris entre plusieurs millions et plusieurs dizaines "
        "de millions de FCFA par an pour une solution commerciale "
        "équivalente. Le coût réel du SI-ENV se limite à l'hébergement et au "
        "nom de domaine, ce qui le rend reproductible pour d'autres projets "
        "AGEROUTE sans surcoût de licence.")

    print("  nouvelle section 6.7 insérée (tableaux 6.4 et 6.5)")

    # ── Bibliographie : ajout des sources [24] à [27] ────────────────────────
    i_biblio_fin = None
    for i, p in enumerate(paras):
        if p.text.strip().startswith('[23]'):
            i_biblio_fin = i
            break
    assert i_biblio_fin is not None, "Référence [23] introuvable"
    ancre_biblio = paras[i_biblio_fin]._element
    modele_biblio = paras[i_biblio_fin]

    nouvelles_refs = [
        "[24]Simoon CV, Salaires moyens par métier en Côte d'Ivoire 2026, "
        "en ligne : simoon-cv.com, consulté en août 2026.",
        "[25]Systalink, Solutions d'hébergement VPS Côte d'Ivoire, en "
        "ligne : systalink.com, consulté en août 2026.",
        "[26]NIC-CI, Prix nom de domaine .ci, en ligne : "
        "lenomdedomaine.ci, consulté en août 2026.",
        "[27]ITQlick et PricingNow, Enablon CSR Suite Pricing 2026 et "
        "Cority EHS Software Pricing, en ligne : itqlick.com et "
        "pricingnow.com, consulté en août 2026.",
    ]
    for ref in nouvelles_refs:
        p_el = OxmlElement('w:p')
        ancre_biblio.addnext(p_el)
        ancre_biblio = p_el
        par = Paragraph(p_el, paras[i_biblio_fin]._parent)
        par.style = modele_biblio.style
        run = par.add_run(ref)
        if modele_biblio.runs:
            run.font.size = modele_biblio.runs[0].font.size

    print("  bibliographie : 4 références ajoutées ([24] à [27])")

    doc.save(CIBLE)
    print("\nEnregistré : %s" % CIBLE)


if __name__ == "__main__":
    main()
