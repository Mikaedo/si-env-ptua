"""
inspect_doc_format.py
---------------------
Inspecte le formatage du memoire (marges, styles, polices).
"""
from docx import Document
from docx.shared import Cm, Pt, Emu

doc_path = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v56.docx"
doc = Document(doc_path)

# Marges
for section in doc.sections:
    print(f"=== Section ===")
    print(f"  Page width: {section.page_width} EMU = {section.page_width / 914400:.2f} inches = {section.page_width / 360000:.2f} cm")
    print(f"  Page height: {section.page_height} EMU = {section.page_height / 914400:.2f} inches")
    print(f"  Left margin: {section.left_margin} EMU = {section.left_margin / 360000:.2f} cm")
    print(f"  Right margin: {section.right_margin} EMU = {section.right_margin / 360000:.2f} cm")
    print(f"  Top margin: {section.top_margin} EMU = {section.top_margin / 360000:.2f} cm")
    print(f"  Bottom margin: {section.bottom_margin} EMU = {section.bottom_margin / 360000:.2f} cm")
    print(f"  Header distance: {section.header_distance / 360000:.2f} cm")
    print(f"  Footer distance: {section.footer_distance / 360000:.2f} cm")

# Styles utilises
print(f"\n=== Styles par defaut ===")
style = doc.styles['Normal']
print(f"  Police: {style.font.name}")
print(f"  Taille: {style.font.size}")

# Quelques paragraphes pour voir les styles
print(f"\n=== Derniers paragraphes (pour voir les styles d'annexes existantes) ===")
for i, p in enumerate(doc.paragraphs[-30:]):
    idx = len(doc.paragraphs) - 30 + i
    text_preview = p.text[:80] if p.text else "(vide)"
    style_name = p.style.name if p.style else "None"
    align = p.alignment
    print(f"  [{idx}] style={style_name} align={align} | {text_preview}")

# Voir s'il y a deja des annexes
print(f"\n=== Recherche d'annexes existantes ===")
for i, p in enumerate(doc.paragraphs):
    if "annexe" in p.text.lower() or "Annexe" in p.text:
        print(f"  [{i}] {p.text[:100]}")

# Voir les images existantes dans le doc
print(f"\n=== Images dans le document ===")
from docx.opc.constants import RELATIONSHIP_TYPE as RT
img_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        img_count += 1
print(f"  Nombre d'images: {img_count}")
