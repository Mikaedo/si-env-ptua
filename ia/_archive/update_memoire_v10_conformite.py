# -*- coding: utf-8 -*-
"""
Met le memoire en conformite avec le « Guide complet des regles de forme
MEMOIRE MIAGE / UPB », et ajoute l'annexe des captures de tests.

Corrections appliquees :
  1. suppression de la fleche Unicode du chapitre 7 ;
  2. duree du jeton portee de 1 h a 12 h (chapitre 7 et tableau 10.2) ;
  3. legendes numerotees placees SOUS chaque figure des annexes A, B et C,
     et suppression des descriptifs qui figuraient AU-DESSUS des images de
     l'annexe C (le guide impose le titre en dessous pour les figures) ;
  4. phrase d'introduction dans chaque annexe renvoyant a ses figures, le
     guide imposant de toujours referencer une figure dans le texte ;
  5. justification des paragraphes de corps qui ne l'etaient pas ;
  6. ajout de l'annexe D (5 captures de tests) au meme format.

Les modifications de structure sont appliquees en ordre DECROISSANT d'index :
inserer ou supprimer un paragraphe decale tous les suivants.

Sortie : MEMOIRE_N'GUESSAN_v10.docx
"""
import copy
import os
import shutil

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v9.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v10.docx")
CAPTURES = r"D:\etude_soutenance\SI-ENV\ia\captures_annexes"

GRAY = RGBColor(0x47, 0x55, 0x69)

#: Index (dans le fichier source) du paragraphe portant l'image -> legende.
LEGENDES = {
    688: "Figure A.1 : Extrait de code - hyperparametres et parametres "
         "d'entrainement des modeles YOLOv8n et MobileNetV2.",
    693: "Figure B.1 : Extrait de code - calcul des metriques d'evaluation de "
         "la detection (precision, rappel, F1-score, mAP).",
    695: "Figure B.2 : Extrait de code - calcul des metriques d'evaluation de "
         "la classification de criticite.",
    699: "Figure C.1 : Concentrations moyennes de NO2 avant, pendant et apres "
         "travaux (Sentinel-5P/TROPOMI, chantier Rocade Y4).",
    702: "Figure C.2 : Indice de risque pluie/erosion avant, pendant et apres "
         "travaux (CHIRPS + SRTM, chantier Rocade Y4).",
}

#: Index du titre d'annexe -> phrase d'introduction citant ses figures.
INTROS = {
    686: "La figure A.1 reproduit les hyperparametres et les parametres "
         "d'entrainement effectivement utilises pour les deux modeles.",
    691: "Les figures B.1 et B.2 reproduisent le code de calcul des metriques "
         "d'evaluation, respectivement pour la detection et pour la "
         "classification de criticite.",
    696: "Les figures C.1 et C.2 presentent les cartes satellitaires produites "
         "via Google Earth Engine, en comparaison avant, pendant et apres "
         "travaux.",
}

#: Descriptifs situes AU-DESSUS d'une image : a supprimer, la legende
#: numerotee etant desormais placee en dessous.
A_SUPPRIMER = (698, 701)

#: Paragraphes de corps a justifier (identifies par l'audit).
A_JUSTIFIER = (310, 577)

PREUVES = [
    ("test_T01_auth_jwt.png",
     "Figure D.1 : Test T01 - Authentification JWT. Appel de POST /auth/login "
     "depuis la documentation interactive : reponse HTTP 200, jeton signe "
     "retourne et role SPEC_ENV correctement resolu."),
    ("test_T06_carte_filtres.png",
     "Figure D.2 : Test T06 - Carte interactive avec filtres. Tableau de bord "
     "affichant les marqueurs de signalements geolocalises, les traces des axes "
     "du PTUA et les filtres par chantier, commune, periode et statut."),
    ("test_T07_alertes_seuil.png",
     "Figure D.3 : Test T07 - Alertes generees par franchissement de seuil. "
     "Liste des alertes creees automatiquement par le service d'evaluation, "
     "avec leur niveau et leur horodatage."),
    ("test_T08_satellite_gee.png",
     "Figure D.4 : Test T08 - Analyse satellitaire Google Earth Engine. "
     "Restitution des quatre indices calcules (NO2, NDVI, NDWI, risque "
     "pluie/erosion) par chantier, avec leur interpretation."),
    ("test_T05_rapport_pges.png",
     "Figure D.5 : Test T05 - Generation du rapport PGES. Ecran de parametrage "
     "du rapport (chantiers, periode, destinataire) precedant la production du "
     "fichier PDF."),
]


def _inserer_apres(paragraphe, texte, italique=True, taille=11, centre=True):
    """Insere un nouveau paragraphe juste apres celui fourni."""
    nouveau = copy.deepcopy(paragraphe._element)
    # On vide la copie de tout contenu avant d'y ecrire la legende
    for enfant in list(nouveau):
        if not enfant.tag.endswith('}pPr'):
            nouveau.remove(enfant)
    paragraphe._element.addnext(nouveau)

    from docx.text.paragraph import Paragraph
    par = Paragraph(nouveau, paragraphe._parent)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER if centre else WD_ALIGN_PARAGRAPH.JUSTIFY
    r = par.add_run(texte)
    r.italic = italique
    r.bold = False
    r.font.size = Pt(taille)
    r.font.color.rgb = GRAY if italique else None
    return par


def main():
    for fichier, _ in PREUVES:
        if not os.path.exists(os.path.join(CAPTURES, fichier)):
            raise SystemExit("Capture manquante : %s" % fichier)

    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)

    # ── 1 et 2 : nettoyage du corps du texte (aucun impact sur les index) ───
    fleches = jetons = 0
    for par in doc.paragraphs:
        for run in par.runs:
            if "\u2192" in run.text:
                run.text = run.text.replace(
                    "routes \u2192 services \u2192 repositories",
                    "routes, services et repositories")
                run.text = run.text.replace(" \u2192 ", ", ").replace("\u2192", ",")
                fleches += 1
            if "JWT (1h" in run.text:
                run.text = run.text.replace("JWT (1h", "JWT (12h")
                jetons += 1
    cellules = 0
    for tab in doc.tables:
        for ligne in tab.rows:
            for cell in ligne.cells:
                if "Jeton valide 1h" in cell.text:
                    for par in cell.paragraphs:
                        for run in par.runs:
                            if "1h" in run.text:
                                run.text = run.text.replace("1h", "12h")
                                cellules += 1
    print("  fleches supprimees / jetons corriges : %d / %d (+%d en tableau)"
          % (fleches, jetons, cellules))

    # ── 4 : justification (aucun impact sur les index) ──────────────────────
    for i in A_JUSTIFIER:
        if i < len(doc.paragraphs):
            doc.paragraphs[i].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    print("  paragraphes justifies                : %d" % len(A_JUSTIFIER))

    # ── 3 : structure des annexes, en ordre DECROISSANT ─────────────────────
    operations = []
    for i, texte in LEGENDES.items():
        operations.append((i, 'legende', texte))
    for i, texte in INTROS.items():
        operations.append((i, 'intro', texte))
    for i in A_SUPPRIMER:
        operations.append((i, 'supprimer', None))
    operations.sort(key=lambda o: o[0], reverse=True)

    poses = intros = supprimes = 0
    for index, action, texte in operations:
        par = doc.paragraphs[index]
        if action == 'legende':
            _inserer_apres(par, texte)
            poses += 1
        elif action == 'intro':
            _inserer_apres(par, texte, italique=False, taille=12, centre=False)
            intros += 1
        else:
            par._element.getparent().remove(par._element)
            supprimes += 1
    print("  legendes d'annexes posees            : %d" % poses)
    print("  introductions d'annexes ajoutees     : %d" % intros)
    print("  descriptifs mal places supprimes     : %d" % supprimes)

    # ── 6 : annexe D ────────────────────────────────────────────────────────
    doc.add_page_break()
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    titre.add_run(
        "Annexe D : Captures d'ecran des tests fonctionnels (justification du "
        "tableau 10.2)").bold = True

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "Les figures D.1 a D.5 documentent cinq des tests fonctionnels recenses "
        "au tableau 10.2, choisis pour leur caractere directement verifiable a "
        "l'ecran : authentification (figure D.1), cartographie interactive "
        "(figure D.2), alertes automatiques (figure D.3), analyse satellitaire "
        "(figure D.4) et generation du rapport reglementaire (figure D.5). Ces "
        "captures ont ete prises sur le deploiement Docker local du SI-ENV ; "
        "les donnees affichees sont des donnees de test.")

    for fichier, legende in PREUVES:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(os.path.join(CAPTURES, fichier),
                                    width=Cm(15.5))
        p_lg = doc.add_paragraph()
        p_lg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_lg.add_run(legende)
        r.italic = True
        r.font.size = Pt(11)
        r.font.color.rgb = GRAY

    doc.save(SORTIE)
    print("\nEnregistre : %s" % SORTIE)
    print("Paragraphes : %d | Tableaux : %d"
          % (len(doc.paragraphs), len(doc.tables)))


if __name__ == "__main__":
    main()
