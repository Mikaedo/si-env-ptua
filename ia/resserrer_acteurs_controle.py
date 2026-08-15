# -*- coding: utf-8 -*-
"""
Limite les acteurs de controle a ceux que le systeme sert reellement.

Le paragraphe 1.5 citait quatre institutions : l'ANDE, le CIAPOL et l'OIPR, la
BAD n'apparaissant qu'ailleurs. Or le SI-ENV n'ouvre un acces qu'a deux
d'entre elles, l'ANDE et la BAD, qui sont les deux profils en consultation
seule du dispositif. Annoncer en introduction des interlocuteurs auxquels rien
n'est ensuite propose expose a une question previsible et sans bonne reponse :
pourquoi le CIAPOL n'a-t-il pas de compte ?

Le paragraphe est donc aligne sur ce qui existe. Les sigles devenus inutiles
quittent la liste, pour la meme raison : un sigle repertorie mais jamais
employe attire l'oeil sans rien apporter.

L'arrete cite en reference conserve son intitule complet, numero d'acte
officiel qui ne se modifie pas.
"""
import re
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_acteurs.docx")
TEXTE = Path(r"C:\Users\DELL\AppData\Local\Temp\claude"
             r"\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25"
             r"\scratchpad\acteurs.txt")

# Sigles a retirer de la liste : ils ne servent plus le propos.
A_RETIRER = ("CIAPOL", "OIPR", "MINEDD")


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    paras = doc.paragraphs

    # 1. Le paragraphe des acteurs de controle.
    cible = next(p for p in paras
                 if p.text.strip().startswith("Au-delà de la CC-PTUA"))
    for seg in list(cible.runs)[1:]:
        seg._element.getparent().remove(seg._element)
    cible.runs[0].text = TEXTE.read_text(encoding="utf-8").strip()
    print("paragraphe 1.5 reecrit")

    # 2. Les entrees de la liste des sigles.
    debut = next(i for i, p in enumerate(paras)
                 if p.text.strip() == "LISTE DES SIGLES ET ABRÉVIATIONS")
    fin = next(i for i, p in enumerate(paras)
               if i > debut and p.style.name.startswith("Heading 1"))
    for p in list(paras[debut + 1:fin]):
        libelle = p.text.strip()
        if any(re.match(rf"^{s}\s*:", libelle) for s in A_RETIRER):
            p._element.getparent().remove(p._element)
            print(f"  sigle retire : {libelle[:52]}")

    doc.save(SOURCE)

    # 3. Controle : plus aucune mention hors references bibliographiques.
    reste = [p.text.strip()[:70] for p in Document(SOURCE).paragraphs
             if any(s in p.text for s in A_RETIRER)]
    print("\nmentions restantes :")
    for r in reste:
        print(f"  {r}")
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
