# -*- coding: utf-8 -*-
"""
Ferme les deux dernieres portes ouvertes du memoire.

La premiere est un defaut de sourcage. Le tableau 5.7 donne quatre seuils sans
dire d'ou ils viennent, ce qui expose a une question sans reponse ecrite. La
note ajoutee distingue trois niveaux : ce qui est reglementaire, ce qui releve
de la convention en teledetection, et ce qui est calibre empiriquement. Assumer
la troisieme categorie vaut mieux que de laisser croire a une norme.

La seconde est une apparente contradiction. Le temps d'inference figure a 4,3
millisecondes au tableau 5.4 et a 23,8 au tableau 6.3, sans qu'aucune phrase ne
rapproche les deux. Un lecteur attentif y voit une incoherence ; il s'agit en
realite de deux materiels differents, le banc d'essai sur processeur graphique
et l'execution ONNX sur processeur d'ordinateur.

Les notes se placent apres la legende du tableau concerne et non avant, pour
que le lecteur ait vu les valeurs avant d'en lire la portee.

Le texte est lu depuis un fichier en UTF-8 : compose en ligne de commande, il
perdrait ses accents.
"""
import copy
import shutil
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_deux_notes.docx")
TEXTE = Path(r"C:\Users\DELL\AppData\Local\Temp\claude"
             r"\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25"
             r"\scratchpad\deux_notes.txt")

ANCRES = {
    "SEUILS": "Tableau 5.7 : Indices environnementaux calculés",
    "INFERENCE": "Tableau 6.3 : Performances mesurées du SI-ENV",
}


def decouper(brut):
    """Separe le fichier en sections reperees par leur mot-cle en majuscules."""
    sections, cle, lignes = {}, None, []
    for ligne in brut.splitlines():
        if ligne.strip() in ANCRES:
            if cle:
                sections[cle] = " ".join(lignes).strip()
            cle, lignes = ligne.strip(), []
        elif ligne.strip():
            lignes.append(ligne.strip())
    if cle:
        sections[cle] = " ".join(lignes).strip()
    return sections


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    notes = decouper(TEXTE.read_text(encoding="utf-8"))

    # Modele de mise en forme : un paragraphe de corps existant, pour heriter
    # justification, interligne et espacement sans rien fixer en dur.
    modele = next(p for p in doc.paragraphs
                  if p.style.name == "Normal" and len(p.text.strip()) > 250
                  and not p.text.strip().startswith(("Figure", "Tableau")))

    for cle, debut in ANCRES.items():
        # La legende du corps, non celle de la liste des tableaux en tete de
        # document : on ne retient donc que les occurrences tardives.
        paras = doc.paragraphs
        candidats = [p for i, p in enumerate(paras)
                     if p.text.strip().startswith(debut) and i > 300]
        if not candidats:
            print(f"  legende introuvable : {debut[:44]}")
            continue
        legende = candidats[-1]

        element = copy.deepcopy(modele._element)
        legende._element.addnext(element)
        note = Paragraph(element, modele._parent)
        for fragment in list(note.runs)[1:]:
            fragment._element.getparent().remove(fragment._element)
        note.runs[0].text = notes[cle]
        print(f"  note posee apres « {debut[:40]} » : "
              f"{len(notes[cle])} caracteres")

    doc.save(SOURCE)

    controle = Document(SOURCE)
    for repere in ("seuils de vigilance", "ne mesurent pas la même chose"):
        present = any(repere in p.text for p in controle.paragraphs)
        print(f"  presence de « {repere} » : {present}")
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
