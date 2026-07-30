"""
inserer_figures_ia.py
=====================
Insere les figures d'entrainement IA dans le memoire Word.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

MEMOIRE = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v42.docx"
IA_DIR = r"D:\etude_soutenance\SI-ENV\ia"

# Figures a inserer : (chemin_image, numero_figure, legende, mots_cles_recherche)
FIGURES = [
    # Detection YOLOv8
    (
        os.path.join(IA_DIR, "runs_detection", "dechets_yolov8n", "results.png"),
        "Figure 8.1",
        "Figure 8.1 - Courbes d'apprentissage du modele YOLOv8n (detection de dechets)",
        ["Figure 8.1", "courbes d'apprentissage", "YOLOv8"]
    ),
    (
        os.path.join(IA_DIR, "runs_detection", "dechets_yolov8n", "confusion_matrix.png"),
        "Figure 8.2",
        "Figure 8.2 - Matrice de confusion du modele YOLOv8n (detection)",
        ["Figure 8.2", "matrice de confusion", "YOLOv8", "detection"]
    ),
    (
        os.path.join(IA_DIR, "runs_detection", "dechets_yolov8n", "PR_curve.png"),
        "Figure 8.3",
        "Figure 8.3 - Courbe Precision-Recall du modele YOLOv8n",
        ["Figure 8.3", "Precision-Recall", "PR curve"]
    ),
    # Classification MobileNetV2
    (
        os.path.join(IA_DIR, "matrice_confusion_classification.png"),
        "Figure 8.4",
        "Figure 8.4 - Matrice de confusion du modele MobileNetV2 (classification de criticite)",
        ["Figure 8.4", "matrice de confusion", "MobileNetV2", "classification"]
    ),
]

def trouver_paragraphe_avec_texte(doc, mots_cles):
    """Cherche un paragraphe contenant un des mots-cles."""
    for i, para in enumerate(doc.paragraphs):
        texte = para.text.lower()
        for mot in mots_cles:
            if mot.lower() in texte:
                return i
    return None

def inserer_figure(doc, index_para, chemin_image, legende):
    """Insere une image + legende apres le paragraphe donne."""
    if not os.path.exists(chemin_image):
        print(f"  [!] Image introuvable : {chemin_image}")
        return False

    para = doc.paragraphs[index_para]
    # Inserer l'image dans un nouveau paragraphe apres
    new_para = para.insert_paragraph_before("") if para.text == "" else None

    # Creer un paragraphe pour l'image
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(chemin_image, width=Inches(5.5))

    # Deplacer le paragraphe image apres le paragraphe cible
    para._element.addnext(img_para._element)

    # Creer un paragraphe pour la legende
    legende_para = doc.add_paragraph()
    legende_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_leg = legende_para.add_run(legende)
    run_leg.font.size = Pt(10)
    run_leg.font.italic = True

    # Deplacer la legende apres l'image
    img_para._element.addnext(legende_para._element)

    print(f"  [OK] Insere : {legende}")
    return True

def main():
    print(f">> Ouverture du memoire : {MEMOIRE}")
    doc = Document(MEMOIRE)
    print(f">> {len(doc.paragraphs)} paragraphes trouves")

    figures_inserees = 0

    for chemin, num_fig, legende, mots_cles in FIGURES:
        print(f"\n>> Recherche pour {num_fig}...")
        index = trouver_paragraphe_avec_texte(doc, mots_cles)

        if index is not None:
            print(f"  Paragraphe trouve a l'index {index}: '{doc.paragraphs[index].text[:80]}...'")
            if inserer_figure(doc, index, chemin, legende):
                figures_inserees += 1
        else:
            print(f"  [!] Aucun paragraphe trouve pour {num_fig}. Insertion a la fin.")
            # Inserer a la fin du document
            doc.add_picture(chemin, width=Inches(5.5))
            leg_para = doc.add_paragraph(legende)
            leg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in leg_para.runs:
                run.font.size = Pt(10)
                run.font.italic = True
            figures_inserees += 1
            print(f"  [OK] Insere a la fin : {legende}")

    # Sauvegarder
    output = MEMOIRE.replace("v42", "v43")
    doc.save(output)
    print(f"\n>> Memoire sauvegarde : {output}")
    print(f">> {figures_inserees}/{len(FIGURES)} figures inserees")

if __name__ == "__main__":
    main()
