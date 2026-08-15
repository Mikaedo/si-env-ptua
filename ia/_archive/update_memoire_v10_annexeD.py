# -*- coding: utf-8 -*-
"""
Met a jour le memoire :
  1. remplace la fleche Unicode "->" du chapitre 7 par une formulation textuelle
     (demande de l'auteur : aucun signe de ce type dans le corps du texte) ;
  2. corrige la duree du jeton annoncee au chapitre 7 (1 h -> 12 h), restee
     incoherente avec la configuration reelle du backend ;
  3. ajoute une Annexe D contenant 5 captures d'ecran qui justifient le
     tableau 10.2 (synthese des tests fonctionnels).

Sortie : MEMOIRE_N'GUESSAN_v10.docx (l'original n'est pas modifie).
"""
import os
import shutil

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v9.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v10.docx")
CAPTURES = r"D:\etude_soutenance\SI-ENV\ia\captures_annexes"

GRAY = RGBColor(0x47, 0x55, 0x69)

#: Captures a joindre, dans l'ordre. Les legendes respectent les regles de
#: forme UPB : numerotation obligatoire au format [chapitre].[ordre] et titre
#: place SOUS la figure.
PREUVES = [
    ("test_T01_auth_jwt.png",
     "Figure D.1 : Test T01 - Authentification JWT. Appel de POST /auth/login "
     "depuis la documentation interactive : reponse HTTP 200, jeton signe "
     "retourne et role SPEC_ENV correctement resolu."),
    ("test_T06_carte_filtres.png",
     "Figure D.2 : Test T06 - Carte interactive avec filtres. Tableau de bord "
     "affichant les marqueurs de signalements geolocalises, les traces des axes "
     "du PTUA et les filtres par chantier, commune, periode et statut."),
    ("test_T07_alertes_seuil.png",
     "Figure D.3 : Test T07 - Alertes generees par franchissement de seuil. "
     "Liste des alertes creees automatiquement par le service d'evaluation, "
     "avec leur niveau et leur horodatage."),
    ("test_T08_satellite_gee.png",
     "Figure D.4 : Test T08 - Analyse satellitaire Google Earth Engine. "
     "Restitution des quatre indices calcules (NO2, NDVI, NDWI, risque "
     "pluie/erosion) par chantier, avec leur interpretation."),
    ("test_T05_rapport_pges.png",
     "Figure D.5 : Test T05 - Generation du rapport PGES. Ecran de parametrage "
     "du rapport (chantiers, periode, destinataire) precedant la production du "
     "fichier PDF."),
]


def main():
    for fichier, _ in PREUVES:
        chemin = os.path.join(CAPTURES, fichier)
        if not os.path.exists(chemin):
            raise SystemExit("Capture manquante : %s" % chemin)

    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)

    # ── 1 et 2 : nettoyage du corps du texte ────────────────────────────────
    fleches = jetons = 0
    for par in doc.paragraphs:
        for run in par.runs:
            if "\u2192" in run.text:
                # "(routes -> services -> repositories)" devient une enumeration
                run.text = run.text.replace(
                    "routes \u2192 services \u2192 repositories",
                    "routes, services et repositories")
                # Toute fleche residuelle est remplacee par une virgule
                run.text = run.text.replace(" \u2192 ", ", ")
                run.text = run.text.replace("\u2192", ",")
                fleches += 1
            if "JWT (1h" in run.text:
                run.text = run.text.replace("JWT (1h", "JWT (12h")
                jetons += 1
    print("  runs contenant une fleche traites : %d" % fleches)
    print("  mentions de duree du jeton corrigees : %d" % jetons)

    # Le tableau des tests annonce aussi la duree du jeton
    cellules = 0
    for tab in doc.tables:
        for ligne in tab.rows:
            for cell in ligne.cells:
                if "Jeton valide 1h" in cell.text:
                    for par in cell.paragraphs:
                        for run in par.runs:
                            if "1h" in run.text:
                                run.text = run.text.replace("1h", "12h")
                                cellules += 1
    print("  cellules de tableau corrigees : %d" % cellules)

    # ── 3 : Annexe D ────────────────────────────────────────────────────────
    doc.add_page_break()

    # Titre d'annexe : meme forme que les annexes A, B et C deja presentes
    # (style Normal, gras, justifie, taille heritee de 12 pt).
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    titre.add_run(
        "Annexe D : Captures d'ecran des tests fonctionnels (justification du "
        "tableau 10.2)").bold = True

    # Le guide de forme impose de renvoyer aux figures depuis le texte :
    # « Toujours referencer dans le texte ». L'introduction cite donc
    # explicitement les figures D.1 a D.5.
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "Les figures D.1 a D.5 documentent cinq des tests fonctionnels recenses "
        "au tableau 10.2, choisis pour leur caractere directement verifiable a "
        "l'ecran : authentification (figure D.1), cartographie interactive "
        "(figure D.2), alertes automatiques (figure D.3), analyse satellitaire "
        "(figure D.4) et generation du rapport reglementaire (figure D.5). Ces "
        "captures ont ete prises sur le deploiement Docker local du SI-ENV ; "
        "les donnees affichees sont des donnees de test.")

    for fichier, legende in PREUVES:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(os.path.join(CAPTURES, fichier),
                                    width=Cm(15.5))
        # Titre SOUS la figure, conformement aux regles de forme.
        p_lg = doc.add_paragraph()
        p_lg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_lg.add_run(legende)
        r.italic = True
        r.font.size = Pt(11)
        r.font.color.rgb = GRAY

    doc.save(SORTIE)
    print("\nEnregistre : %s" % SORTIE)
    print("Paragraphes : %d | Tableaux : %d"
          % (len(doc.paragraphs), len(doc.tables)))


if __name__ == "__main__":
    main()
