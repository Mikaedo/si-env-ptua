# -*- coding: utf-8 -*-
"""
Repartit les colonnes du tableau des huit profils selon leur contenu.

Les trois colonnes faisaient 5,17 cm chacune, partage arithmetique et non
typographique : la colonne Interface ne recoit que « Mobile » ou « Web » et
gaspillait sa largeur, pendant que le perimetre, seule colonne redigee,
retombait a la ligne sur presque chaque rang. Le tableau paraissait donc
desequilibre alors que ses colonnes etaient rigoureusement egales.

Les largeurs sont reprises a proportion du texte le plus long de chaque
colonne, la somme restant fixee a 15,5 cm, largeur utile entre les marges.

python-docx ne connait pas de largeur de colonne : la valeur doit etre posee
sur chaque cellule, faute de quoi Word retombe sur une repartition automatique.
"""
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Cm

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_colonnes.docx")

ENTETES = ["Rôle", "Interface", "Périmètre principal"]
LARGEURS = [4.6, 2.6, 8.3]          # 15,5 cm au total


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    # La derniere mention hors bibliographie d'un organisme desormais ecarte.
    for p in doc.paragraphs:
        if "et le CIAPOL" in p.text:
            for seg in p.runs:
                if "CIAPOL" in seg.text:
                    seg.text = seg.text.replace("et le CIAPOL", "et l'ANDE")
            print("mention corrigee :", p.text.strip()[-72:])

    cible = None
    for t in doc.tables:
        if [c.text.strip() for c in t.rows[0].cells] == ENTETES:
            cible = t
            break
    if cible is None:
        raise SystemExit("tableau des profils introuvable")

    print(f"tableau trouve : {len(cible.rows)} lignes")
    cible.autofit = False
    for ligne in cible.rows:
        for cellule, largeur in zip(ligne.cells, LARGEURS):
            cellule.width = Cm(largeur)

    doc.save(SOURCE)
    controle = next(t for t in Document(SOURCE).tables
                    if [c.text.strip() for c in t.rows[0].cells] == ENTETES)
    print("largeurs :", [round(c.width.cm, 2) for c in controle.rows[0].cells])
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
