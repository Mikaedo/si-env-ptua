# -*- coding: utf-8 -*-
"""
Redige les conclusions partielles manquantes et reprend la conclusion generale.

La conclusion de la troisieme partie etait vide : le titre figurait, le texte
non. Celle de la deuxieme partie decrivait une architecture qui a evolue
depuis. La conclusion generale, elle, restait allusive la ou elle devrait etre
la page la plus nette du document, puisque c'est souvent la seule qu'un membre
du jury relit juste avant de poser sa premiere question.

Le ton recherche est celui d'un etudiant qui rend compte de son travail :
phrases de longueur inegale, jugements assumes, limites reconnues sans
contrition. On evite la tournure de plaquette commerciale, qui sonne faux dans
un memoire, et l'accumulation d'adjectifs, qui trahit un texte ecrit pour
remplir plutot que pour dire.
"""
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_conclusions.docx")


CONCLUSION_2 = (
    "Cette deuxième partie a traduit les besoins recueillis auprès de la CC-PTUA "
    "en une architecture modulaire, modélisée en UML et déclinée en un schéma de "
    "base de données spatiale. Le travail de conception a surtout consisté à "
    "trancher : décider qui voit quoi, où s'arrête la responsabilité de chacun, "
    "et ce que le système refuse de faire. La matrice des habilitations est le "
    "résultat le plus directement contestable de cette partie, et donc le plus "
    "utile à discuter, car elle engage une lecture de la gouvernance du "
    "programme. Y faire figurer l'agence de tutelle et le bailleur en "
    "consultation seule n'est pas un détail technique : c'est reconnaître que "
    "celui qui contrôle ne doit pas pouvoir modifier ce qu'il examine. La "
    "troisième partie présente l'implémentation de ce système et les résultats "
    "obtenus."
)

CONCLUSION_3 = (
    "Cette troisième partie a montré le système en fonctionnement plutôt qu'en "
    "intention. Les modèles de détection et de classification ont été entraînés "
    "puis mesurés sur des données réelles, l'analyse satellitaire s'appuie sur "
    "des images effectivement téléchargées, et le déploiement a été éprouvé dans "
    "deux environnements distincts. Les chiffres présentés sont ceux qui ont été "
    "relevés, y compris lorsqu'ils s'écartent de ce qui était espéré."
    "  "
    "Trois enseignements se dégagent de cette phase. Le premier tient à l'écart "
    "entre un modèle qui fonctionne sur un jeu de test et un modèle qui "
    "fonctionne sur un chantier : la lumière, la poussière et l'angle de prise "
    "de vue dégradent des performances qui semblaient acquises, ce qui a conduit "
    "à traiter le diagnostic automatique comme une aide au classement et non "
    "comme une décision. Le deuxième porte sur la donnée satellitaire, dont la "
    "richesse ne dispense pas d'une interprétation prudente : un indice de "
    "végétation dégradé peut signaler un chantier mal maîtrisé comme une simple "
    "saison sèche. Le troisième concerne le déploiement, où la contrainte "
    "budgétaire a imposé des choix qui se sont révélés instructifs, la migration "
    "d'un hébergeur à un autre n'ayant demandé aucune modification du code parce "
    "que l'application avait été conteneurisée dès le départ."
    "  "
    "Reste ce que cette partie ne démontre pas. Le système a été éprouvé "
    "techniquement, il n'a pas été confronté à un usage quotidien par des agents "
    "qui ne l'ont pas conçu. C'est la limite principale de ce travail, et la "
    "conclusion générale y revient."
)

CONCLUSION_GENERALE = (
    "Ce mémoire est parti d'un constat simple : le suivi environnemental des "
    "chantiers du Projet de Transport Urbain d'Abidjan reposait sur des fiches "
    "papier et des tableurs, alors que le Plan de Gestion Environnementale et "
    "Sociale impose un rendu régulier à l'agence de tutelle et au bailleur. "
    "Entre ce qui était exigé et ce qui était outillé, l'écart se comblait à la "
    "main, au prix d'un travail de ressaisie que personne ne pouvait tenir "
    "durablement."
    "  "
    "Le système construit répond à cet écart sur trois plans. Il permet de "
    "saisir un signalement sur le terrain, y compris sans réseau, avec une "
    "position et une photographie horodatées. Il consolide ces remontées et les "
    "croise avec des indices satellitaires, ce qui donne au spécialiste du suivi "
    "environnemental une vue qu'aucun classeur ne pouvait offrir. Il produit "
    "enfin le rapport réglementaire et garde la trace de sa remise, ce qui "
    "transforme une obligation administrative en une opération vérifiable."
    "  "
    "Deux décisions ont davantage compté que les autres. La première a été de "
    "faire tourner la reconnaissance d'images sur le téléphone plutôt que sur un "
    "serveur : elle coûte en précision, mais elle permet de travailler là où le "
    "réseau manque, c'est-à-dire précisément sur les chantiers. La seconde a été "
    "d'ouvrir le dispositif aux riverains par une application distincte, en "
    "conditionnant l'accès à la proximité d'un chantier. Le mécanisme de gestion "
    "des plaintes existait déjà sur le papier ; il supposait qu'un habitant se "
    "déplace jusqu'à une permanence, ce qui suffisait à décourager la plupart "
    "des signalements."
    "  "
    "Ce travail comporte des limites qu'il serait malhonnête de passer sous "
    "silence. Le système n'a pas été utilisé en conditions réelles par des "
    "agents qui ne l'ont pas conçu, et c'est le seul essai qui permettrait de "
    "juger de son ergonomie. Le corpus d'entraînement, constitué à partir de "
    "sources ouvertes, ne reflète pas parfaitement les chantiers d'Abidjan. La "
    "vérification de position atteste d'une localisation, non d'une résidence. "
    "Et l'hébergement retenu, gratuit, convient à une validation académique mais "
    "pas à une mise en service."
    "  "
    "Ces limites dessinent la suite. Une phase pilote sur un chantier, avec les "
    "équipes en place, permettrait de mesurer ce qu'aucun test automatisé ne "
    "mesure : si l'outil est réellement utilisé quand personne ne regarde. Un "
    "corpus enrichi de photographies prises sur les sites du PTUA améliorerait "
    "la reconnaissance là où elle est faible. Le passage à un hébergement de "
    "production, enfin, ne demanderait qu'un changement de paramètres, ce que le "
    "chapitre six établit."
    "  "
    "Sur le plan personnel, ce projet m'a moins appris à écrire du code qu'à "
    "décider. Choisir entre deux modèles dont l'un est plus précis et l'autre "
    "plus rapide, accepter qu'une fonctionnalité utile ne soit pas prioritaire, "
    "reconnaître qu'une mesure flatteuse ne prouve rien : ces arbitrages ne "
    "s'apprennent pas dans un cours. Ils m'ont aussi appris qu'un système "
    "d'information ne vaut que par l'usage qu'on en fait, et qu'un outil "
    "techniquement irréprochable dont personne ne se sert n'a rien résolu."
)


def reecrire(p, texte):
    for seg in list(p.runs)[1:]:
        seg._element.getparent().remove(seg._element)
    if p.runs:
        p.runs[0].text = texte
    else:
        p.add_run(texte)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    paras = doc.paragraphs

    # ── Conclusion de la deuxieme partie ──────────────────────────────────
    for p in paras:
        if p.text.strip().startswith("Cette deuxième partie a traduit"):
            reecrire(p, CONCLUSION_2)
            print("  ok  conclusion de la deuxieme partie")
            break

    # ── Conclusion de la troisieme partie, restee vide ────────────────────
    # Le titre existe sans texte : on ecrit dans le paragraphe qui le suit,
    # en le creant si necessaire.
    import copy
    from docx.text.paragraph import Paragraph

    modele = None
    for p in paras:
        if p.text.strip().startswith("Cette première partie a permis"):
            modele = p
            break

    titres_concl = [i for i, p in enumerate(paras)
                    if "onclusion partielle" in p.text.strip().lower()
                    and p.style.name.startswith("Heading")]
    if len(titres_concl) >= 3 and modele is not None:
        idx = titres_concl[-1]
        suivant = paras[idx + 1] if idx + 1 < len(paras) else None
        if suivant is not None and not suivant.text.strip():
            reecrire(suivant, CONCLUSION_3)
            suivant.style = modele.style
            print("  ok  conclusion de la troisieme partie (paragraphe existant)")
        else:
            nouveau = copy.deepcopy(modele._element)
            paras[idx]._element.addnext(nouveau)
            p = Paragraph(nouveau, paras[idx]._parent)
            reecrire(p, CONCLUSION_3)
            print("  ok  conclusion de la troisieme partie (paragraphe insere)")

    # ── Conclusion generale ───────────────────────────────────────────────
    doc2 = Document(SOURCE) if False else doc
    dans_conclusion = False
    remplacee = False
    for p in doc2.paragraphs:
        t = p.text.strip()
        if p.style.name.startswith("Heading 1"):
            dans_conclusion = (t == "Conclusion générale")
            continue
        if dans_conclusion and t and not remplacee:
            reecrire(p, CONCLUSION_GENERALE)
            remplacee = True
        elif dans_conclusion and t and remplacee:
            # Les paragraphes suivants de l'ancienne conclusion sont vides de
            # leur contenu : tout est repris dans le texte unique ci-dessus.
            reecrire(p, "")
    print(f"  {'ok' if remplacee else 'MANQUE'}  conclusion generale")

    doc.save(SOURCE)
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
