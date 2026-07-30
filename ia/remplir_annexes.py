"""
remplir_annexes.py
=================
Remplace le placeholder des annexes dans MEMOIRE_SI-ENV_v48.docx
par le contenu reel des annexes A, B et C.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

CHEMIN_ENTREE = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v48.docx"
CHEMIN_SORTIE = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v49.docx"

doc = Document(CHEMIN_ENTREE)

# Trouver le paragraph placeholder
placeholder_idx = None
for i, p in enumerate(doc.paragraphs):
    if "contenu des annexes" in p.text or "à insérer manuellement" in p.text:
        placeholder_idx = i
        print(f"Placeholder trouve au paragraphe {i}: {p.text[:80]}")
        break

if placeholder_idx is None:
    print("ERREUR: Placeholder non trouve")
    exit(1)

# Fonction pour inserer un paragraphe apres un index donne
def inserer_apres(doc, ref_para, texte, style_name=None, bold=False, italic=False, size=11, color=None, align=None):
    new_p = doc.add_paragraph()
    if style_name:
        new_p.style = doc.styles[style_name]
    if align:
        new_p.alignment = align
    run = new_p.add_run(texte)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Deplacer le nouveau paragraphe apres la reference
    ref_para._element.addnext(new_p._element)
    return new_p

# Recuperer le paragraphe placeholder
placeholder = doc.paragraphs[placeholder_idx]

# On va inserer les annexes AVANT le placeholder, en remplacant son contenu
# Strategy: vider le placeholder et inserer le contenu apres

# Vider le placeholder
for run in placeholder.runs:
    run.text = ""

# Inserer le contenu dans l'ordre (on insere apres le placeholder, donc en ordre inverse)

# --- ANNEXE C : Cartes satellite ---
contenu_c = [
    ("Annexe C : Cartes satellite : risque pluie/relief et NO2 avant/après travaux", True, 13),
    ("", False, 11),
    ("Les cartes ci-dessous sont produites avec Google Earth Engine (GEE) en utilisant les données CHIRPS (précipitations 48h), SRTM (relief) et Sentinel-5P TROPOMI (NO2).", False, 11),
    ("", False, 11),
    ("Figure C.1 : Carte de risque pluie/relief – Zone du PTUA (CHIRPS 48h + SRTM). Les zones en rouge indiquent un fort cumul pluviométrique sur terrain plat (pente < 3°), propice à la formation d'eaux stagnantes.", False, 11),
    ("", False, 11),
    ("Figure C.2 : Carte de concentration NO2 – Avant travaux (Sentinel-5P TROPOMI, moyenne mensuelle). Les valeurs sont exprimées en µmol/m². Les zones de chantier apparaissent en jaune-orange.", False, 11),
    ("", False, 11),
    ("Figure C.3 : Carte de concentration NO2 – Pendant travaux (Sentinel-5P TROPOMI, moyenne mensuelle). Une augmentation du NO2 est visible sur les zones de construction (engins, circulation).", False, 11),
    ("", False, 11),
    ("Figure C.4 : Carte NDWI – Détection des surfaces d'eau (Sentinel-2, composite mensuel). Les pixels bleus indiquent les plans d'eau persistants, susceptibles de constituer des gîtes larvaires.", False, 11),
    ("", False, 11),
    ("Méthodologie :", True, 11),
    ("• CHIRPS : cumul des précipitations sur 48 heures, seuil > 50 mm", False, 11),
    ("• SRTM : pente calculée par dérivée du MNT, seuil < 3° (terrain plat)", False, 11),
    ("• Sentinel-5P : NO2 troposphérique, filtrage qualité > 0.5, agrégation mensuelle", False, 11),
    ("• Sentinel-2 : NDWI = (Green - NIR) / (Green + NIR), seuil > 0.3", False, 11),
    ("• Superposition : intersection des zones à forte pluie + faible pente = risque élevé", False, 11),
]

# --- ANNEXE B : Extraits de code métriques ---
contenu_b = [
    ("Annexe B : Extraits de code : calcul des métriques d'évaluation", True, 13),
    ("", False, 11),
    ("Les métriques d'évaluation (précision, rappel, F1-score, mAP) sont calculées à l'aide des bibliothèques Ultralytics (détection) et scikit-learn (classification).", False, 11),
    ("", False, 11),
    ("B.1 – Évaluation YOLOv8n (détection) :", True, 11),
    ("", False, 11),
    ("from ultralytics import YOLO", False, 10),
    ("model = YOLO('runs_detection/dechets_yolov8n/weights/best.pt')", False, 10),
    ("metriques = model.val(data='dataset/data.yaml')", False, 10),
    ("", False, 10),
    ("# Métriques calculées automatiquement :", False, 10),
    ("# mAP@0.5      = metriques.box.map50   # 0.365", False, 10),
    ("# mAP@0.5:0.95 = metriques.box.map      # 0.217", False, 10),
    ("# Précision    = metriques.box.mp       # 0.522", False, 10),
    ("# Rappel       = metriques.box.mr       # 0.370", False, 10),
    ("# F1 = 2 * (P * R) / (P + R)            # 0.434", False, 10),
    ("", False, 11),
    ("B.2 – Évaluation MobileNetV2 (classification) :", True, 11),
    ("", False, 11),
    ("from sklearn.metrics import classification_report, confusion_matrix", False, 10),
    ("import seaborn as sns, matplotlib.pyplot as plt", False, 10),
    ("", False, 10),
    ("report = classification_report(vrais, predits, target_names=classes)", False, 10),
    ("print(report)", False, 10),
    ("# Output :", False, 10),
    ("#               precision  recall  f1-score  support", False, 10),
    ("#       faible       0.77     0.85      0.81       115", False, 10),
    ("#      modere       0.42     0.35      0.38        45", False, 10),
    ("#    important       0.12     0.08      0.10        14", False, 10),
    ("#    accuracy                           0.61       174", False, 10),
    ("#   macro avg       0.44     0.43      0.43       174", False, 10),
    ("# weighted avg       0.56     0.61      0.57       174", False, 10),
    ("", False, 11),
    ("B.3 – Matrice de confusion :", True, 11),
    ("", False, 11),
    ("cm = confusion_matrix(vrais, predits)", False, 10),
    ("sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',", False, 10),
    ("            xticklabels=classes, yticklabels=classes)", False, 10),
    ("plt.xlabel('Prédit'); plt.ylabel('Réel')", False, 10),
    ("plt.title('Matrice de confusion - MobileNetV2')", False, 10),
    ("plt.savefig('matrice_confusion.png', dpi=200)", False, 10),
]

# --- ANNEXE A : Extraits de code hyperparamètres ---
contenu_a = [
    ("Annexe A : Extraits de code : hyperparamètres et paramètres d'entraînement", True, 13),
    ("", False, 11),
    ("Les extraits ci-dessous présentent la configuration des hyperparamètres pour les deux modèles d'intelligence artificielle.", False, 11),
    ("", False, 11),
    ("A.1 – Entraînement YOLOv8n (détection) :", True, 11),
    ("", False, 11),
    ("from ultralytics import YOLO", False, 10),
    ("", False, 10),
    ("CHEMIN_DATA_YAML = 'dataset/data.yaml'", False, 10),
    ("NB_EPOCHS = 10", False, 10),
    ("TAILLE_IMAGE = 320", False, 10),
    ("BATCH = 8", False, 10),
    ("", False, 10),
    ("model = YOLO('yolov8n.pt')", False, 10),
    ("resultats = model.train(", False, 10),
    ("    data=CHEMIN_DATA_YAML,", False, 10),
    ("    epochs=NB_EPOCHS,", False, 10),
    ("    imgsz=TAILLE_IMAGE,", False, 10),
    ("    batch=BATCH,", False, 10),
    ("    lr0=0.0001,", False, 10),
    ("    optimizer='AdamW',", False, 10),
    ("    patience=20,", False, 10),
    ("    augment=True,", False, 10),
    ("    project='runs_detection',", False, 10),
    ("    name='dechets_yolov8n',", False, 10),
    (")", False, 10),
    ("", False, 10),
    ("# Export ONNX pour déploiement mobile", False, 10),
    ("model.export(format='onnx')", False, 10),
    ("", False, 11),
    ("A.2 – Entraînement MobileNetV2 (classification) :", True, 11),
    ("", False, 11),
    ("import torch, torch.nn as nn", False, 10),
    ("from torchvision import datasets, transforms, models", False, 10),
    ("", False, 10),
    ("DOSSIER = 'dataset_criticite'", False, 10),
    ("NB_EPOCHS = 10", False, 10),
    ("BATCH = 8", False, 10),
    ("LR = 0.0001", False, 10),
    ("", False, 10),
    ("# Augmentation de données", False, 10),
    ("transform = transforms.Compose([", False, 10),
    ("    transforms.Resize((224, 224)),", False, 10),
    ("    transforms.RandomHorizontalFlip(),", False, 10),
    ("    transforms.RandomRotation(15),", False, 10),
    ("    transforms.ColorJitter(brightness=0.2),", False, 10),
    ("    transforms.ToTensor(),", False, 10),
    ("    transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225]),", False, 10),
    ("])", False, 10),
    ("", False, 10),
    ("# Transfer learning : MobileNetV2 pré-entraîné sur ImageNet", False, 10),
    ("model = models.mobilenet_v2(weights=models.MobileNet_V2_Weights.DEFAULT)", False, 10),
    ("model.classifier[1] = nn.Linear(model.last_channel, 3)  # 3 classes", False, 10),
    ("", False, 10),
    ("critere = nn.CrossEntropyLoss()", False, 10),
    ("optim = torch.optim.AdamW(model.parameters(), lr=LR)", False, 10),
    ("", False, 10),
    ("# Boucle d'entraînement", False, 10),
    ("for epoch in range(NB_EPOCHS):", False, 10),
    ("    model.train()", False, 10),
    ("    for images, labels in train_dl:", False, 10),
    ("        optim.zero_grad()", False, 10),
    ("        sorties = model(images)", False, 10),
    ("        perte = critere(sorties, labels)", False, 10),
    ("        perte.backward()", False, 10),
    ("        optim.step()", False, 10),
    ("", False, 10),
    ("# Export ONNX", False, 10),
    ("exemple = torch.randn(1, 3, 224, 224)", False, 10),
    ("torch.onnx.export(model, exemple, 'mobilenetv2_criticite.onnx')", False, 10),
]

# Inserer dans l'ordre : A, puis B, puis C (apres le placeholder vide)
# On insere en ordre inverse car addnext met juste apres

ref = placeholder

# D'abord inserer Annexe C (en dernier, donc on l'insere en premier car addnext)
for texte, bold, size in reversed(contenu_c):
    new_p = doc.add_paragraph()
    run = new_p.add_run(texte)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if size == 10:
        run.font.name = 'Consolas'
    ref._element.addnext(new_p._element)
    ref = new_p

# Inserer un separateur
sep = doc.add_paragraph()
sep_run = sep.add_run("─" * 60)
sep_run.font.name = 'Times New Roman'
sep_run.font.size = Pt(11)
sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
ref._element.addnext(sep._element)
ref = sep

# Inserer Annexe B
for texte, bold, size in reversed(contenu_b):
    new_p = doc.add_paragraph()
    run = new_p.add_run(texte)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if size == 10:
        run.font.name = 'Consolas'
    ref._element.addnext(new_p._element)
    ref = new_p

# Separateure
sep2 = doc.add_paragraph()
sep2_run = sep2.add_run("─" * 60)
sep2_run.font.name = 'Times New Roman'
sep2_run.font.size = Pt(11)
sep2_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
ref._element.addnext(sep2._element)
ref = sep2

# Inserer Annexe A
for texte, bold, size in reversed(contenu_a):
    new_p = doc.add_paragraph()
    run = new_p.add_run(texte)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if size == 10:
        run.font.name = 'Consolas'
    ref._element.addnext(new_p._element)
    ref = new_p

doc.save(CHEMIN_SORTIE)
print(f">> Document sauvegarde : {CHEMIN_SORTIE}")

# Verifier
doc2 = Document(CHEMIN_SORTIE)
for i in range(695, min(760, len(doc2.paragraphs))):
    p = doc2.paragraphs[i]
    t = p.text.strip()
    label = t[:120] if t else '(vide)'
    print(f'{i}: {label}')
