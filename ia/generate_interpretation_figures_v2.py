# -*- coding: utf-8 -*-
"""
Genere la version actualisee du document d'interpretation des figures
statistiques du SI-ENV.

Par rapport a la version precedente, ce document couvre :
  - les figures du chapitre 5 avec les chiffres du modele REELLEMENT
    reentraine (6 classes, dataset Recycle Trash) ;
  - les optimisations d'inference ajoutees ensuite (letterbox, TTA, NMS par
    classe) et le gain mesure sur echantillon ;
  - les statistiques du tableau de bord et leur logique de calcul ;
  - les indices satellitaires et leurs seuils d'interpretation.

Ecrase le fichier existant dans le dossier MEMOIRE.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os

SORTIE = r"C:\Users\DELL\Downloads\MEMOIRE\INTERPRETATION_FIGURES_STATISTIQUES_SI-ENV.docx"
FIGURES = r"D:\etude_soutenance\SI-ENV\ia\captures_annexes"

ORANGE = RGBColor(0xE8, 0x6C, 0x00)
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GRAY = RGBColor(0x47, 0x55, 0x69)
GREEN = RGBColor(0x16, 0x7C, 0x3C)
RED = RGBColor(0xB9, 0x1C, 0x1C)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin = sec.bottom_margin = Cm(2.5)


def _police(style, nom='Times New Roman'):
    rpr = style.element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(a), nom)
    rpr.append(rf)


normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
_police(normal)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(8)

for nom, taille, couleur in [('Heading 1', 20, ORANGE),
                             ('Heading 2', 15, NAVY),
                             ('Heading 3', 12.5, GRAY)]:
    st = doc.styles[nom]
    st.font.name = 'Times New Roman'
    st.font.size = Pt(taille)
    st.font.bold = True
    st.font.color.rgb = couleur
    _police(st)

encadre = doc.styles.add_style('PointCle', 1)
encadre.base_style = doc.styles['Normal']
encadre.font.italic = True
encadre.font.color.rgb = NAVY
encadre.font.size = Pt(11.5)
encadre.paragraph_format.left_indent = Cm(0.5)
encadre.paragraph_format.space_before = Pt(4)
encadre.paragraph_format.space_after = Pt(10)


def h1(t):
    doc.add_heading(level=1).add_run(t)


def h2(t):
    doc.add_heading(level=2).add_run(t)


def h3(t):
    doc.add_heading(level=3).add_run(t)


def para(t, gras=False, italique=False, taille=None, couleur=None, centre=False):
    p = doc.add_paragraph()
    if centre:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t)
    r.bold, r.italic = gras, italique
    if taille:
        r.font.size = Pt(taille)
    if couleur:
        r.font.color.rgb = couleur
    return p


def puce(t):
    doc.add_paragraph(t, style='List Bullet')


def cle(t):
    doc.add_paragraph(style='PointCle').add_run(t)


def figure(fichier, legende, largeur_cm=15.5):
    """Insere une figure centree avec sa legende.

    Le document explique comment lire chaque figure : sans l'image en regard,
    l'explication reste abstraite. Une figure absente est signalee plutot
    qu'ignoree silencieusement, pour que le manque soit visible.
    """
    chemin = os.path.join(FIGURES, fichier)
    if not os.path.exists(chemin):
        para("[Figure manquante : %s]" % fichier, italique=True,
             taille=10, couleur=RED, centre=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(chemin, width=Cm(largeur_cm))
    lg = doc.add_paragraph()
    lg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = lg.add_run(legende)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY


def tableau(entetes, lignes, largeurs=None):
    t = doc.add_table(rows=1, cols=len(entetes))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, e in enumerate(entetes):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(e)
        r.bold = True
        r.font.size = Pt(10.5)
    for ligne in lignes:
        cells = t.add_row().cells
        for i, v in enumerate(ligne):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(10.5)
    if largeurs:
        for ligne in t.rows:
            for i, l in enumerate(largeurs):
                ligne.cells[i].width = Cm(l)
    doc.add_paragraph()
    return t


# ═══════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ═══════════════════════════════════════════════════════════════════════════
para("INTERPRÉTATION DES FIGURES ET DES STATISTIQUES", gras=True, taille=19,
     couleur=ORANGE, centre=True)
para("SI-ENV — Système d'Information Environnemental du PTUA", gras=True,
     taille=13, couleur=NAVY, centre=True)
para("Lecture guidée des figures du chapitre 5 (intelligence artificielle), "
     "des indicateurs du tableau de bord et des indices satellitaires",
     italique=True, taille=11, couleur=GRAY, centre=True)
para("Document de préparation à la soutenance — version actualisée",
     italique=True, taille=10.5, couleur=GRAY, centre=True)

doc.add_paragraph()
para("Ce document répond à une question simple : si un membre du jury pointe "
     "une figure et demande « qu'est-ce que cela signifie, et comment l'avez-vous "
     "obtenu ? », que faut-il répondre ? Chaque section donne la lecture de la "
     "figure, le chiffre à retenir, et la limite à assumer honnêtement.")

cle("Règle d'or : ne jamais présenter un chiffre sans dire d'où il vient. "
    "Toutes les valeurs de ce document proviennent soit d'un entraînement "
    "reproductible sur Google Colab, soit d'une mesure effectuée sur "
    "l'application elle-même.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("1. L'architecture en cascade : deux modèles, deux rôles")
# ═══════════════════════════════════════════════════════════════════════════

para("Le module d'intelligence artificielle du SI-ENV ne fait pas une seule "
     "chose mais deux, dans cet ordre. C'est le point le plus souvent mal "
     "formulé en soutenance, il faut donc être très clair.")

puce("YOLOv8n (détection) répond à la question « où sont les déchets sur cette "
     "photo, et de quelle matière sont-ils ? ». Il renvoie des boîtes "
     "englobantes, chacune associée à une classe et à un score de confiance.")
puce("MobileNetV2 (classification) reçoit ensuite chaque zone recadrée par "
     "YOLOv8n et répond à une autre question : « quelle criticité "
     "environnementale attribuer à ce déchet ? » — faible, modéré ou élevé.")

para("Cette séparation est un choix d'architecture assumé : détecter et "
     "apprécier la gravité sont deux tâches de nature différente. Un seul "
     "modèle qui ferait les deux serait plus difficile à entraîner et surtout "
     "impossible à corriger indépendamment.")

h2("Où tourne chaque modèle ?")
para("Les deux modèles sont exportés au format ONNX et exécutés directement sur "
     "le téléphone, sans appel réseau. C'est une contrainte de terrain : les "
     "chantiers du PTUA n'ont pas de couverture réseau garantie, et l'agent "
     "doit pouvoir travailler hors connexion.")

cle("Conséquence technique importante : l'inférence tourne sur un isolate "
    "séparé de l'interface. La bibliothèque ONNX Runtime expose une fonction "
    "d'exécution synchrone ; l'appeler sur le fil principal gèlerait "
    "complètement l'application pendant l'analyse.")

# ═══════════════════════════════════════════════════════════════════════════
h1("2. Le jeu de données et le protocole d'entraînement")
# ═══════════════════════════════════════════════════════════════════════════

para("Le mémoire détaille la composition du corpus dans son tableau 5.3. Le "
     "total est de 2 462 images, réparties en trois sous-ensembles :")

tableau(
    ["Sous-ensemble", "Effectif", "Rôle"],
    [
        ["Recycle Trash — entraînement", "1 970", "Apprentissage des poids"],
        ["Recycle Trash — validation", "247", "Suivi pendant l'entraînement, "
                                              "arrêt anticipé"],
        ["Recycle Trash — test", "245", "Évaluation finale, jamais vue à "
                                        "l'entraînement"],
        ["Augmentation de données", "intégrée", "Mosaic, MixUp et transformations "
                                                "appliquées à la volée"],
    ],
    largeurs=[6.0, 2.6, 7.4],
)

cle("Distinction à maîtriser absolument : la validation sert à décider quand "
    "arrêter l'entraînement, le test sert à mesurer la performance finale. "
    "Confondre les deux ferait douter de la rigueur du protocole. Les métriques "
    "annoncées (mAP@50 = 0,807) portent sur les 247 images de validation, "
    "représentant 568 objets annotés.")

para("Les hyperparamètres retenus figurent dans le tableau 5.6 du mémoire :")

tableau(
    ["Hyperparamètre", "Valeur retenue", "Justification"],
    [
        ["Taux d'apprentissage", "0,001 (AdamW)", "Convergence rapide sur GPU"],
        ["Taille de lot", "32", "Compromis mémoire GPU / stabilité"],
        ["Époques", "100 avec arrêt anticipé", "Convergence atteinte avant la fin"],
        ["Patience", "30", "Arrêt si aucune amélioration sur 30 époques"],
        ["Augmentation", "Flip + rotation 15° + Mosaic",
         "Simule les variations de prise de vue sur le terrain"],
        ["Résolution", "320 × 320 pixels", "Contrainte d'inférence sur téléphone"],
        ["Matériel", "GPU NVIDIA T4 (Google Colab)", "—"],
        ["Export", "ONNX, opset 12", "Exécution embarquée sur mobile"],
    ],
    largeurs=[4.4, 5.4, 6.2],
)

para("Le choix de YOLOv8n et d'une résolution de 320 pixels n'est pas un "
     "hasard : c'est le compromis qui permet une inférence en quelques "
     "centaines de millisecondes sur un téléphone de milieu de gamme. Un "
     "modèle plus gros donnerait de meilleurs scores mais serait inutilisable "
     "sur le terrain.")

cle("Question probable du jury : « pourquoi ne pas avoir pris YOLOv8s ou "
    "YOLOv8m ? » Réponse : parce que la contrainte de conception est le "
    "téléphone de l'agent, pas le score maximal. C'est un arbitrage "
    "performance / faisabilité, à assumer explicitement.")

h2("2.1 Les benchmarks : pourquoi ces deux modèles et pas d'autres")

para("C'est un point central du chapitre 5, et probablement la première "
     "question du jury sur cette partie : le choix des modèles n'est pas "
     "arbitraire, il résulte d'une comparaison chiffrée.")

h3("Tableau 8.2 — Comparaison des modèles de détection")

tableau(
    ["Modèle", "mAP@0,5", "Précision", "Rappel", "F1", "Inférence"],
    [
        ["YOLOv8n (retenu)", "0,807", "0,797", "0,717", "0,755", "4,3 ms"],
        ["Faster R-CNN", "0,685", "0,710", "0,640", "0,673", "312,5 ms"],
        ["SSD300", "0,612", "0,640", "0,580", "0,608", "185,3 ms"],
    ],
    largeurs=[4.0, 2.2, 2.4, 2.2, 1.8, 3.4],
)

para("La lecture est sans ambiguïté : YOLOv8n est à la fois le plus précis et "
     "le plus rapide, avec un temps d'inférence 70 fois inférieur à celui de "
     "Faster R-CNN.")

cle("Argument à retenir : « le modèle retenu n'est pas un compromis entre "
    "qualité et vitesse — il domine les deux autres sur les deux critères. "
    "C'est le résultat d'un benchmark, pas d'une préférence. »")

h3("Tableau 8.3 — Comparaison des modèles de classification")

tableau(
    ["Modèle", "Précision", "Rappel", "F1", "Taille", "Inférence"],
    [
        ["MobileNetV2 (retenu)", "0,88", "0,85", "0,86", "8,9 Mo", "15,2 ms"],
        ["ResNet50", "0,82", "0,80", "0,81", "98 Mo", "112,4 ms"],
        ["VGG16", "0,79", "0,78", "0,78", "138 Mo", "245,8 ms"],
    ],
    largeurs=[4.4, 2.4, 2.0, 1.8, 2.4, 3.0],
)

para("Ici encore, le modèle retenu est le meilleur sur tous les critères. La "
     "colonne « taille » est déterminante : MobileNetV2 pèse 8,9 Mo contre "
     "138 Mo pour VGG16. Embarquer VGG16 dans une application mobile serait "
     "difficilement acceptable, pour un résultat inférieur.")

cle("Chiffre à connaître : MobileNetV2 atteint 0,88 de précision pour la "
    "classification de criticité, soit un score sensiblement meilleur que "
    "celui de la détection. C'est logique : classer une zone déjà recadrée en "
    "trois niveaux est une tâche plus simple que localiser des objets dans une "
    "scène complète.")

# ═══════════════════════════════════════════════════════════════════════════
h1("3. Figure 5.5 — Métriques par classe du modèle de détection")
# ═══════════════════════════════════════════════════════════════════════════

figure("yolo_results_v2.png",
       "Figure 5.5 — Métriques par classe du modèle de détection YOLOv8n "
       "(précision, rappel, mAP@50 et mAP@50-95 pour les 6 matières).")

h2("Comment lire la figure")
para("Le graphique compare, pour chacune des 6 matières, quatre grandeurs : la "
     "précision, le rappel, la mAP@50 et la mAP@50-95. Il faut savoir "
     "expliquer ces quatre termes sans hésiter.")

tableau(
    ["Métrique", "Ce qu'elle mesure", "Comment l'expliquer simplement"],
    [
        ["Précision",
         "Part des détections annoncées qui sont correctes",
         "« Quand le modèle dit "'"'"c'est du métal"'"'", a-t-il raison ? »"],
        ["Rappel",
         "Part des déchets réellement présents qui ont été trouvés",
         "« Combien de déchets le modèle a-t-il laissé passer ? »"],
        ["mAP@50",
         "Qualité globale, boîte considérée juste si elle recouvre 50 % de la vérité",
         "La métrique de référence, celle que l'on cite en premier"],
        ["mAP@50-95",
         "Même chose, moyennée sur des exigences de recouvrement croissantes",
         "Beaucoup plus sévère : mesure la précision du cadrage"],
    ],
    largeurs=[2.8, 6.2, 7.0],
)

h2("Les chiffres réels obtenus")

tableau(
    ["Classe", "Précision", "Rappel", "mAP@50", "mAP@50-95", "Instances"],
    [
        ["metal", "0,896", "0,952", "0,970", "0,868", "125"],
        ["cardboard", "0,762", "0,800", "0,869", "0,654", "95"],
        ["paper", "0,847", "0,646", "0,795", "0,572", "77"],
        ["glass", "0,811", "0,694", "0,778", "0,625", "62"],
        ["organic", "0,706", "0,653", "0,741", "0,552", "72"],
        ["plastic", "0,761", "0,559", "0,690", "0,495", "137"],
        ["GLOBAL", "0,797", "0,717", "0,807", "0,628", "568"],
    ],
    largeurs=[2.8, 2.4, 2.2, 2.2, 2.6, 2.2],
)

h2("Ce qu'il faut dire sur cette figure")

para("Le chiffre à retenir est la mAP@50 globale de 0,807, soit 80,7 %. C'est "
     "un résultat honnête pour un modèle nano entraîné en 100 époques sur un "
     "jeu de données de taille modeste.")

para("Mais la lecture intéressante est celle des écarts entre classes :")

puce("Le métal est de loin le mieux détecté (mAP@50 = 0,970). Explication "
     "physique : les canettes et tôles ont des reflets spéculaires et des "
     "formes régulières, très caractéristiques visuellement.")
puce("Le plastique est le moins bien détecté (mAP@50 = 0,690) alors qu'il est "
     "la classe la plus représentée (137 instances). C'est contre-intuitif et "
     "il faut savoir l'expliquer : les sachets plastique sont transparents, "
     "déformables, sans forme stable — ils épousent le support et se fondent "
     "dans l'arrière-plan.")
puce("Le rappel du plastique (0,559) est nettement plus faible que sa précision "
     "(0,761). Autrement dit : quand le modèle annonce du plastique il se "
     "trompe rarement, mais il en laisse passer beaucoup.")

cle("Piège classique du jury : « votre classe la plus fréquente est la moins "
    "bien reconnue, n'est-ce pas la preuve que votre modèle est mauvais ? » "
    "Non : cela montre que la difficulté ne vient pas du volume de données "
    "mais de la nature visuelle de l'objet. Le plastique souple est un cas "
    "difficile reconnu dans la littérature.")

# ═══════════════════════════════════════════════════════════════════════════
h1("4. Figure 5.6 — Matrice de confusion")
# ═══════════════════════════════════════════════════════════════════════════

figure("yolo_confusion_matrix_v2.png",
       "Figure 5.6 — Matrice de confusion du modèle de détection. La diagonale "
       "correspond aux bonnes réponses ; la colonne « background » aux déchets "
       "non détectés.")

h2("Comment lire la figure")
para("La matrice croise ce que le modèle a prédit (colonnes) avec ce qui était "
     "réellement présent (lignes). La diagonale correspond aux bonnes réponses ; "
     "tout ce qui est hors diagonale est une confusion.")

para("Deux zones méritent une attention particulière en soutenance :")
puce("La colonne « arrière-plan » : ce sont les déchets réellement présents que "
     "le modèle n'a pas vus du tout. Elle explique directement les valeurs de "
     "rappel faibles.")
puce("Les confusions entre matières proches : papier et carton se confondent "
     "assez naturellement, ce qui est cohérent puisqu'il s'agit du même "
     "matériau de base sous deux formes.")

cle("Formulation à retenir : « les confusions du modèle ne sont pas "
    "aléatoires, elles suivent une logique matérielle. Papier et carton se "
    "confondent parce qu'ils sont physiquement proches. C'est rassurant : cela "
    "montre que le modèle a appris des caractéristiques visuelles pertinentes "
    "et non du bruit. »")

# ═══════════════════════════════════════════════════════════════════════════
h1("5. Figure 5.7 — Courbe précision-rappel")
# ═══════════════════════════════════════════════════════════════════════════

figure("yolo_PR_curve_v2.png",
       "Figure 5.7 — Courbe précision-rappel par classe. L'aire sous la courbe "
       "correspond à la mAP@50 : 0,807 pour l'ensemble des classes.")

para("Cette courbe montre l'arbitrage fondamental de toute détection : plus on "
     "abaisse le seuil de confiance, plus on trouve de déchets (le rappel "
     "monte) mais plus on produit de fausses alertes (la précision baisse).")

para("L'aire sous la courbe est précisément la mAP@50 : 0,807. Une courbe qui "
     "reste haute et plate longtemps avant de s'effondrer est le signe d'un bon "
     "modèle.")

h2("Le choix du seuil dans l'application : une décision documentée")

para("Le seuil de confiance retenu dans l'application est de 0,25, et non 0,5 "
     "comme on le lit souvent. Ce choix découle directement de la figure 5.5.")

tableau(
    ["Seuil", "Conséquence", "Adapté à quel usage ?"],
    [
        ["0,50",
         "Peu de fausses alertes, mais beaucoup de déchets manqués "
         "(le rappel du plastique est de 0,559)",
         "Un usage où une fausse alerte coûte cher"],
        ["0,25",
         "Davantage de détections, quelques faux positifs à faible score",
         "Le suivi environnemental : mieux vaut signaler un déchet douteux "
         "que passer à côté d'une pollution réelle"],
    ],
    largeurs=[2.0, 7.5, 6.5],
)

cle("Argument à préparer : le seuil n'est pas un réglage arbitraire, c'est une "
    "décision métier. En suivi environnemental, l'oubli d'un déchet réel est "
    "plus grave qu'une vérification inutile de l'agent. Le seuil de 0,25 "
    "traduit ce choix, et il est cohérent avec le rappel mesuré au chapitre 5.")

# ═══════════════════════════════════════════════════════════════════════════
h1("6. Figure 5.8 — Matrice de confusion de la classification de criticité")
# ═══════════════════════════════════════════════════════════════════════════

figure("matrice_confusion_classification.png",
       "Figure 5.8 — Matrice de confusion du modèle MobileNetV2 "
       "(classification de criticité : faible, modéré, important).",
       largeur_cm=12.5)

h2("D'où viennent les trois niveaux de criticité")

para("Point essentiel à maîtriser, car c'est la première question qu'un "
     "examinateur attentif posera. Le corpus d'origine ne contient aucune "
     "étiquette de criticité : elle a été construite par une règle "
     "déterministe appliquée aux annotations, à savoir le nombre d'objets "
     "présents sur l'image.")

tableau(
    ["Nombre d'objets annotés", "Criticité attribuée"],
    [["1 à 2 objets", "Faible"],
     ["3 à 5 objets", "Modérée"],
     ["6 objets ou plus", "Importante"]],
    largeurs=[8.0, 8.0],
)

para("Cette règle a deux mérites : elle est reproductible, et n'importe qui "
     "peut la vérifier. Mais elle a une conséquence qu'il faut énoncer "
     "soi-même : la frontière entre cinq et six objets n'a aucune signature "
     "visuelle globale. La tâche demandée au réseau relève du dénombrement "
     "plus que de la reconnaissance de forme, ce qui borne par construction "
     "la performance atteignable sur les niveaux intermédiaires. Les F1 de "
     "0,59 sur « modérée » et 0,67 sur « importante » s'expliquent donc par "
     "la définition de l'étiquette, et pas seulement par le déséquilibre du "
     "corpus.")

cle("Formulation à employer : « la criticité a été définie par comptage "
    "d'objets ; séparer cinq objets de six revient à dénombrer, ce pour quoi "
    "un classifieur d'image entière n'est pas conçu. C'est une limite de "
    "conception que j'assume et dont je tire trois pistes concrètes. »")

para("Trois pistes découlent de cette analyse, et elles figurent en "
     "perspectives du mémoire :")

puce("Dériver la criticité du comptage des objets retournés par le détecteur "
     "plutôt que d'entraîner un second réseau : plus économe, et exactement "
     "aussi fidèle à la définition retenue. C'est le comparateur naturel du "
     "classifieur.")
puce("Ramener la décision à sa forme utile pour l'agent, qui est binaire : "
     "faut-il intervenir ou non. Regrouper « modérée » et « importante » "
     "écarte la frontière la plus instable.")
puce("Introduire une option de rejet : sous un seuil de confiance, le modèle "
     "s'abstient et laisse l'agent trancher. On annonce alors un taux de "
     "couverture et une justesse sur les cas couverts, ce qui décrit un outil "
     "utilisable plutôt qu'un score moyen.")

para("Attention à ne pas confondre cette figure avec la figure 5.6. Toutes "
     "deux sont des matrices de confusion, mais elles portent sur deux modèles "
     "et deux tâches différentes :")

puce("La figure 5.6 concerne la détection : elle croise les 6 matières et "
     "comporte une colonne « arrière-plan » pour les déchets non détectés.")
puce("La figure 5.8 concerne la criticité : elle croise 3 niveaux ordonnés, et "
     "tous les objets y sont déjà détectés puisqu'ils sortent de la phase 1.")

h2("Comment lire cette matrice")

para("Sa lecture est plus indulgente que celle de la détection, parce que les "
     "trois niveaux sont ordonnés. Confondre « modéré » et « important » — deux "
     "niveaux adjacents — est une erreur d'un cran. Confondre « faible » et "
     "« important » serait un saut de deux crans, bien plus problématique.")

para("Il faut donc regarder si les erreurs restent concentrées autour de la "
     "diagonale, et non dispersées dans les coins opposés.")

cle("Formulation à retenir : « les confusions se concentrent entre niveaux "
    "adjacents, ce qui est le comportement attendu sur une échelle ordonnée. "
    "Une erreur d'un cran reste exploitable pour le suivi, d'autant que l'agent "
    "peut corriger la proposition du modèle. »")

h2("Le lien avec le tableau 5.5")

para("Cette matrice illustre visuellement les chiffres du benchmark : "
     "précision de 0,88 et rappel de 0,85 pour MobileNetV2. Savoir relier la "
     "figure au tableau montre que vous avez compris les deux, et pas seulement "
     "recopié l'un ou l'autre.")

cle("Question possible : « votre classification est-elle plus fiable que "
    "votre détection ? » Réponse honnête : globalement elle paraît meilleure, "
    "mais uniquement grâce à la classe « faible », largement majoritaire. Sur "
    "les deux niveaux qui déclenchent une action, elle est nettement moins "
    "sûre, pour la raison de conception exposée plus haut. Ne présentez donc "
    "pas ce module comme un instrument de mesure de la gravité : la criticité "
    "enregistrée reste celle que l'agent déclare, celle de l'IA étant "
    "conservée à côté pour mesurer plus tard le taux d'accord.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("7. Les optimisations d'inférence et leur gain mesuré")
# ═══════════════════════════════════════════════════════════════════════════

para("Cette section est nouvelle et constitue un excellent sujet à mettre en "
     "avant : elle montre qu'un bon score d'entraînement ne suffit pas, et que "
     "la façon de préparer l'image avant de la donner au modèle a un effet "
     "considérable.")

h2("7.1 Le letterbox : préserver les proportions")

para("Un défaut a été identifié dans le prétraitement initial. L'image était "
     "redimensionnée directement en 320 × 320 pixels. Or une photo de "
     "téléphone est au format 3:4 : la redimensionner en carré l'écrase et "
     "déforme tous les objets qu'elle contient.")

para("Le modèle, lui, a été entraîné avec un letterbox — un redimensionnement "
     "qui conserve les proportions et complète les bords par une teinte neutre. "
     "L'application ne reproduisait donc pas les conditions d'entraînement.")

para("La correction consiste à appliquer le même letterbox dans l'application, "
     "puis à annuler la transformation pour replacer les boîtes détectées aux "
     "bonnes coordonnées.")

h2("7.2 L'augmentation au moment du test (TTA)")

para("Le modèle analyse l'image, puis son image miroir, et les détections des "
     "deux passes sont fusionnées. Le miroir horizontal fait partie des "
     "augmentations vues pendant l'entraînement : les vraies détections se "
     "confirment sur les deux passes, les détections instables s'écartent.")

h2("7.3 La suppression des doublons classe par classe")

para("La suppression des détections redondantes était appliquée globalement. "
     "Conséquence : une bouteille en plastique posée sur un carton, deux objets "
     "qui se recouvrent largement, faisait disparaître l'un des deux. La "
     "suppression se fait désormais matière par matière, comme dans "
     "l'implémentation de référence d'Ultralytics.")

h2("7.4 Gain mesuré")

para("Les trois corrections ont été évaluées sur un échantillon de 9 images, en "
     "comparant les trois modes de prétraitement sur le même modèle.")

tableau(
    ["Image / objet", "Redim. brut", "Letterbox", "Letterbox + TTA"],
    [
        ["Photo terrain (3:4) — papier", "32 %", "45 %", "77 %"],
        ["Lot de validation — plastique", "42 %", "70 %", "70 %"],
        ["Lot de validation — papier", "37 %", "57 %", "59 %"],
        ["Lot de validation (autre) — objets trouvés", "5", "7", "7"],
        ["Lot d'entraînement — métal", "non détecté", "non détecté", "58 %"],
        ["TOTAL objets détectés (9 images)", "22", "20", "25"],
    ],
    largeurs=[6.8, 3.0, 3.0, 3.2],
)

para("Deux observations à savoir défendre :")
puce("Sur une image déjà carrée, les trois modes donnent un résultat "
     "rigoureusement identique. C'est la preuve que le gain vient bien de la "
     "correction de la déformation, et non d'un effet de hasard.")
puce("Le nombre total de détections baisse parfois tout en faisant monter la "
     "confiance. Cela signifie que le redimensionnement brut produisait des "
     "faux positifs à faible score : moins de détections mais plus fiables est "
     "une amélioration, pas une régression.")

cle("C'est le point le plus valorisant du chapitre 5 : il montre une démarche "
    "d'ingénieur. Le modèle n'a pas changé, seule la façon de lui présenter "
    "l'image a été corrigée, et la confiance sur les vraies détections a "
    "progressé de moitié à plus du double selon les cas.")

h2("7.5 Ce que ces optimisations ne font pas")

para("Il est indispensable d'être honnête sur ce point, car le jury peut le "
     "soulever. Ces corrections améliorent la qualité pratique des détections "
     "mais ne modifient pas la mAP intrinsèque du modèle, qui reste à 0,807. "
     "Pour dépasser 90 %, il faudrait réentraîner avec une résolution de 640 "
     "pixels, une architecture plus grande et davantage d'époques — au prix de "
     "la faisabilité sur téléphone.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("8. Les statistiques du tableau de bord")
# ═══════════════════════════════════════════════════════════════════════════

figure("web_dashboard_v2.png",
       "Illustration A — Tableau de bord du Spécialiste Suivi Environnemental : "
       "indicateurs cliquables, cartographie des tracés PTUA, répartition par "
       "statut et par type de nuisance.")

para("Le tableau de bord destiné au Spécialiste Suivi Environnemental présente "
     "des indicateurs agrégés. Leur logique de calcul doit pouvoir être "
     "expliquée, car une incohérence apparente serait immédiatement relevée.")

h2("8.1 Les compteurs de statut")

tableau(
    ["Indicateur", "Définition exacte"],
    [
        ["Total", "Nombre de signalements enregistrés sur le périmètre"],
        ["En attente", "Statut NOUVEAU : pas encore pris en charge"],
        ["En cours", "Statut EN_TRAITEMENT : instruction engagée"],
        ["Clôturés", "Statut CLOTURE : traitement mené à terme"],
        ["Rejetés", "Statut REJETE : signalement écarté avec motif"],
    ],
    largeurs=[4.0, 12.0],
)

para("Ces cinq compteurs vérifient une propriété importante : la somme des "
     "quatre statuts est égale au total. Cette cohérence est vérifiable en "
     "direct devant le jury.")

cle("Point d'honnêteté méthodologique : dans une version antérieure, les "
    "statuts « en traitement » et « rejeté » n'étaient comptés dans aucun "
    "indicateur. Des signalements existaient donc sans apparaître nulle part, "
    "et la somme ne retombait pas sur le total. Savoir raconter la détection "
    "et la correction de ce défaut est plus convaincant que de prétendre que "
    "tout a toujours été juste.")

h2("8.2 Les deux taux, et pourquoi il en faut deux")

puce("Taux de traitement = clôturés / total. Il mesure les dossiers menés à "
     "terme.")
puce("Taux de prise en charge = (clôturés + rejetés) / total. Un rejet motivé "
     "est un dossier instruit : l'ignorer sous-estimerait le travail réalisé.")

h2("8.3 Les urgences")

para("L'indicateur d'urgences ne compte que les signalements de criticité "
     "élevée encore ouverts. Un signalement grave mais déjà clôturé n'est plus "
     "une action à mener ; le compter gonflerait artificiellement l'alerte.")

h2("8.4 La répartition par type de nuisance")

para("Un signalement peut cumuler plusieurs types de nuisance. Les types sont "
     "donc éclatés avant comptage, faute de quoi un regroupement produirait des "
     "catégories artificielles du genre « Déchets de chantier, Bruit ».")

cle("Conséquence assumée à annoncer avant qu'on ne la relève : la somme des "
    "parts de la répartition peut dépasser le nombre de signalements, "
    "puisqu'un même signalement alimente plusieurs types. Ce n'est pas une "
    "erreur de calcul mais la conséquence directe du multi-typage.")

h2("8.5 L'évolution mensuelle")

para("La courbe couvre les six derniers mois. Le calcul recule de mois "
     "calendaire en mois calendaire.")

cle("Défaut corrigé, intéressant à citer si l'on demande comment la qualité du "
    "code a été contrôlée : la version antérieure reculait par tranches de "
    "30 jours. Comme les mois n'ont pas tous 30 jours, un 31 mars produisait la "
    "séquence mars, mars, janvier — le mois de février disparaissait purement "
    "et simplement de la courbe.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("9. Les indices satellitaires et leurs seuils")
# ═══════════════════════════════════════════════════════════════════════════

figure("web_satellite_v2.png",
       "Illustration B — Page d'analyse satellitaire (cf. figure 9.1 du mémoire) : "
       "quatre indices avec leur interprétation, puis détail par chantier.")

para("Le module de télédétection s'appuie sur Google Earth Engine. Quatre "
     "indices sont calculés sur l'emprise des chantiers. Il faut savoir "
     "expliquer ce que chacun mesure et justifier les seuils d'alerte.")

tableau(
    ["Indice", "Source", "Ce qu'il mesure", "Seuils retenus"],
    [
        ["NO₂", "Sentinel-5P / TROPOMI",
         "Dioxyde d'azote : pollution liée au trafic et aux engins",
         "> 50 dégradée ; > 30 modérée"],
        ["NDVI", "Sentinel-2 (rouge / proche infrarouge)",
         "Vigueur du couvert végétal",
         "< 0,30 dégradé ; < 0,40 moyen"],
        ["NDWI", "Sentinel-2 (proche infrarouge / moyen infrarouge)",
         "Teneur en eau de surface",
         "< 0,20 sécheresse ; < 0,30 modérée"],
        ["Risque pluie", "CHIRPS + relief SRTM",
         "Risque d'érosion combinant pluviométrie et pente",
         "> 7/10 élevé ; > 5/10 modéré"],
    ],
    largeurs=[2.0, 4.0, 5.5, 4.5],
)

h2("Le piège du NDWI, à connaître absolument")

para("Le NDWI est le seul indice dont un niveau BAS constitue l'alerte : une "
     "valeur faible signale une sécheresse et un stress hydrique, pas un excès "
     "d'eau. Les trois autres indices fonctionnent en sens inverse.")

cle("Erreur à ne pas commettre devant le jury : interpréter un NDWI faible "
    "comme un risque d'eaux stagnantes. C'est exactement l'inverse. Cette "
    "inversion de sens a d'ailleurs provoqué une incohérence détectée puis "
    "corrigée dans l'interface, où une même valeur était qualifiée "
    "différemment à deux endroits de la page.")

h2("La synthèse en tête de page")

para("Une bannière de synthèse retient le point le plus préoccupant parmi les "
     "quatre indices et l'accompagne de l'action attendue. L'objectif est "
     "qu'un spécialiste puisse décider sans convertir mentalement des "
     "micromoles par mètre carré en niveau de risque.")

h2("La question de la fraîcheur des données")

para("Les appels à Earth Engine prennent une trentaine de secondes pour "
     "l'ensemble des chantiers. Un cache persistant de sept jours a donc été "
     "mis en place. La durée est justifiée par la nature des données : les "
     "indices reposent sur des composites de plusieurs jours et ne varient pas "
     "d'une heure à l'autre.")

cle("Si l'on demande « vos données satellitaires sont-elles en temps réel ? », "
    "la réponse honnête est non, et c'est volontaire : Sentinel-2 repasse tous "
    "les cinq jours au même endroit, et la couverture nuageuse tropicale impose "
    "de travailler sur des composites. Prétendre au temps réel serait une "
    "faute scientifique.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
h1("10. Récapitulatif des chiffres à connaître par cœur")
# ═══════════════════════════════════════════════════════════════════════════

tableau(
    ["Grandeur", "Valeur", "Où elle apparaît"],
    [
        ["mAP@50 globale (détection)", "0,807 (80,7 %)", "Figure 5.5, tableau 8.x"],
        ["mAP@50-95 globale", "0,628", "Figure 5.5"],
        ["Précision globale", "0,797", "Figure 5.5"],
        ["Rappel global", "0,717", "Figure 5.5"],
        ["Meilleure classe", "metal — 0,970", "Figure 5.5"],
        ["Classe la plus difficile", "plastic — 0,690", "Figure 5.5"],
        ["Images de validation", "247 (568 instances)", "Protocole, chapitre 5"],
        ["Nombre de classes", "6 matières", "Jeu de données Recycle Trash"],
        ["Résolution d'inférence", "320 × 320", "Contrainte téléphone"],
        ["Seuil de confiance", "0,25", "Choix métier documenté"],
        ["Seuil de recouvrement (NMS)", "0,45", "Valeur usuelle Ultralytics"],
        ["Durée de validité du jeton", "12 heures", "Configuration du backend"],
        ["Cache satellitaire", "7 jours", "Justifié par les composites"],
    ],
    largeurs=[5.5, 4.5, 6.0],
)

para("Si vous ne devez retenir qu'une seule phrase :", gras=True)
para("« Le modèle de détection atteint 80,7 % de mAP@50 sur six classes de "
     "matières, avec un écart marqué entre le métal, très bien reconnu, et le "
     "plastique souple, qui reste le cas le plus difficile. »",
     italique=True, couleur=NAVY)

doc.add_paragraph()
para("Fin du document — SI-ENV / PTUA", italique=True, taille=10,
     couleur=GRAY, centre=True)

doc.save(SORTIE)
print("Enregistre :", SORTIE)
print("Paragraphes :", len(doc.paragraphs), "| Tableaux :", len(doc.tables))
