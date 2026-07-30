# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import os

doc = Document(r'C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v56.docx')
FIG = r'D:\etude_soutenance\SI-ENV\ia\captures_annexes'

def replace_image_in_para(para, image_path, width_inches=5.5):
    for run in para.runs:
        drawings = run.element.findall(qn('w:drawing'))
        if drawings:
            for d in drawings:
                d.getparent().remove(d)
            run.add_picture(image_path, width=Inches(width_inches))
            return True
    return False

# Trouver toutes les images dans les annexes (para >= 690)
print("=== Images trouvees dans les annexes ===")
for i, p in enumerate(doc.paragraphs):
    if i >= 690:
        for run in p.runs:
            if run.element.findall(qn('w:drawing')):
                print(f'  Para {i}: "{p.text[:80]}"')

# Remplacer les images
replacements = {
    699: os.path.join(FIG, 'annexe_a1_output_yolo.png'),
    704: os.path.join(FIG, 'annexe_a2_output_mobilenet.png'),
    706: os.path.join(FIG, 'annexe_b1_output_metriques_yolo.png'),
}

for idx, img_path in replacements.items():
    if os.path.exists(img_path):
        para = doc.paragraphs[idx]
        if replace_image_in_para(para, img_path):
            print(f'[OK] Image annexe remplacee (para {idx})')
        else:
            print(f'[WARN] Pas d image dans para {idx}')

doc.save(r'C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v56.docx')
print('[OK] Memoire v56 sauvegarde avec nouvelles annexes')
