# -*- coding: utf-8 -*-
"""
Etape 16 : porter les figures du corps a leur taille de confort.

Le corps est stabilise a 46 pages, quatre sous le plafond de 50. Les figures,
descendues a 60 % pendant la reduction, sont portees a la valeur passee en
argument. Le passage de 45 a 60 % n'avait coute aucune page : la marge est donc
reelle et vaut mieux d'etre investie dans la lisibilite des diagrammes UML et
des captures d'ecran, que le jury examine.

Usage : python etape16_taille_figures.py <taille_actuelle> <taille_voulue>
        python etape16_taille_figures.py 0.60 0.80

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os
import re
import sys

from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')

A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

SEUIL_LOGO_CM = 3.0
LARGEUR_UTILE_CM = 16.0
#: Une figure plus haute que cela occupe la page entiere et en gaspille le bas.
HAUTEUR_MAX_CM = 20.0


def bornes_corps(doc):
    """Bornes du corps, table des matieres exclue.

    Les entrees de la table des matieres contiennent une tabulation avant leur
    numero de page : c'est ce qui les distingue des vrais titres.
    """
    debut = fin = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if '\t' in t:
            continue
        if debut is None and t == 'Introduction générale' and i > 200:
            debut = i
        if debut is not None and t == 'BIBLIOGRAPHIE':
            fin = i
            break
    return debut, fin


def main():
    actuelle = float(sys.argv[1]) if len(sys.argv) > 1 else 0.60
    voulue = float(sys.argv[2]) if len(sys.argv) > 2 else 0.80
    facteur = voulue / actuelle

    doc = Document(CIBLE)
    debut, fin = bornes_corps(doc)
    if debut is None or fin is None:
        raise SystemExit("Bornes du corps introuvables, rien modifie")
    print("  corps : paragraphes %d a %d" % (debut, fin))

    n = plafonnees = 0
    for i in range(debut, fin):
        for ext in doc.paragraphs[i]._element.iter('{%s}ext' % A_NS):
            cx, cy = ext.get('cx'), ext.get('cy')
            if not (cx and cy):
                continue
            l_cm, h_cm = int(cx) / 360000.0, int(cy) / 360000.0
            if l_cm < SEUIL_LOGO_CM:
                continue
            f = facteur
            if l_cm * f > LARGEUR_UTILE_CM:
                f = min(f, LARGEUR_UTILE_CM / l_cm)
                plafonnees += 1
            if h_cm * f > HAUTEUR_MAX_CM:
                f = min(f, HAUTEUR_MAX_CM / h_cm)
            ext.set('cx', str(int(round(l_cm * f * 360000))))
            ext.set('cy', str(int(round(h_cm * f * 360000))))
            n += 1

    print("  figures portees de %d %% a %d %% : %d (dont %d plafonnees)"
          % (actuelle * 100, voulue * 100, n, plafonnees))

    doc.save(CIBLE)
    print("\nEnregistre : %s" % CIBLE)


if __name__ == "__main__":
    main()
