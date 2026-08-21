# -*- coding: utf-8 -*-
"""
Ajoute au dictionnaire de donnees l'entite qui lui manquait.

Le dictionnaire documentait onze entites, dont Rapport, et passait sous silence
TransmissionRapport. Or c'est l'inverse de la realite du schema : le rapport
produit est un fichier, sans table propre, tandis que sa remise est bel et bien
enregistree. Le modele conceptuel presente les deux entites ; le dictionnaire
n'en decrivait qu'une.

Les douze attributs ajoutes sont ceux de la table, releves dans le modele et
non deduits : c'est ce qui distingue un dictionnaire de donnees d'une
paraphrase du diagramme.

Une precision merite d'y figurer et s'y trouve : l'emetteur est conserve sous
forme d'adresse et non de cle etrangere, afin que la trace survive a la
suppression du compte qui l'a produite. Un lecteur qui verrait une colonne
emetteur_email sans explication y soupconnerait une denormalisation par
negligence.
"""
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_dictionnaire.docx")

ENTITE = "TransmissionRapport"
LIGNES = [
    ("id", "Entier (PK)", "Identifiant de la transmission"),
    ("transmis_le", "Horodatage", "Date et heure de la remise"),
    ("emetteur_email", "Texte",
     "Adresse de l'auteur, conservée en clair pour que la trace survive à "
     "la suppression du compte"),
    ("destinataire_email", "Texte", "Adresse de l'organisme destinataire"),
    ("organisme", "Texte", "Organisme visé : ANDE, BAD"),
    ("periode_debut", "Texte", "Début de la période couverte (AAAA-MM-JJ)"),
    ("periode_fin", "Texte", "Fin de la période couverte (AAAA-MM-JJ)"),
    ("chantiers", "Texte", "Chantiers couverts par le rapport transmis"),
    ("nom_fichier", "Texte", "Nom du fichier PDF remis"),
    ("taille_octets", "Entier", "Taille du fichier transmis"),
    ("succes", "Booléen", "Acheminement réussi ou non"),
    ("detail_erreur", "Texte", "Motif de l'échec, le cas échéant"),
]


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    cible = None
    for t in doc.tables:
        if [c.text.strip() for c in t.rows[0].cells][:2] == ["Entité", "Attribut"]:
            cible = t
            break
    if cible is None:
        raise SystemExit("dictionnaire de donnees introuvable")

    avant = len({l.cells[0].text.strip() for l in cible.rows[1:]
                 if l.cells[0].text.strip()})
    print(f"{avant} entites avant ajout, {len(cible.rows) - 1} lignes")

    # Les lignes sont ajoutees en fin de tableau, a la suite de la derniere
    # entite : le dictionnaire suit l'ordre d'apparition, non l'alphabet.
    modele = cible.rows[-1]
    for attribut, type_, description in LIGNES:
        ligne = cible.add_row()
        for cellule, texte in zip(ligne.cells,
                                  (ENTITE, attribut, type_, description)):
            p = cellule.paragraphs[0]
            if p.runs:
                p.runs[0].text = texte
            else:
                p.add_run(texte)
        # Reprend la mise en forme de la derniere ligne existante.
        for source, destination in zip(modele.cells, ligne.cells):
            destination.width = source.width

    doc.save(SOURCE)

    controle = Document(SOURCE)
    for t in controle.tables:
        if [c.text.strip() for c in t.rows[0].cells][:2] == ["Entité", "Attribut"]:
            entites = sorted({l.cells[0].text.strip() for l in t.rows[1:]
                              if l.cells[0].text.strip()})
            print(f"{len(entites)} entites apres ajout, "
                  f"{len(t.rows) - 1} lignes")
            print(f"  {', '.join(entites)}")
            break

    # La legende annonce un nombre : il doit suivre.
    for p in controle.paragraphs:
        if "Dictionnaire de données" in p.text:
            print(f"\nlegende actuelle : {p.text.strip()}")
    print(f"Sauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
