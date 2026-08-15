# -*- coding: utf-8 -*-
"""
Etape 14 : accentuer tout le texte ajoute par les scripts de restructuration.

Les etapes 4 a 11 ont insere du texte sans accents : titres des six chapitres
fusionnes, sommaire, annexes E et F, paragraphes condenses. Cette etape corrige
l'ensemble.

Le dictionnaire ne retient que des formes dont la version sans accent n'est
jamais un mot francais valide : « probleme », « donnees », « systeme »... Le
remplacement peut donc s'appliquer a tout le document sans risque d'abimer le
texte d'origine. Les formes ambigues sont volontairement ecartees : « eleve »
(eleve / elevee), « procede », « repere », « trace », « cote », dont le sens
depend du contexte et qu'un remplacement automatique fausserait.

La casse initiale est preservee, de sorte qu'un mot en debut de phrase ou dans
un titre reste capitalise.

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os
import re

from docx import Document

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

#: Formes sans accent -> forme correcte. Uniquement des cas non ambigus.
MOTS = {
    # structure du document
    'etude': 'étude', 'etudes': 'études', 'etudie': 'étudie',
    'problematique': 'problématique', 'problematiques': 'problématiques',
    'realisation': 'réalisation', 'realisations': 'réalisations',
    'deploiement': 'déploiement', 'teledetection': 'télédétection',
    'generale': 'générale', 'general': 'général', 'generales': 'générales',
    'generaux': 'généraux', 'generation': 'génération',
    'premiere': 'première', 'premieres': 'premières',
    'deuxieme': 'deuxième', 'troisieme': 'troisième',
    'quatrieme': 'quatrième', 'cinquieme': 'cinquième', 'sixieme': 'sixième',
    'etat': 'état', 'etats': 'états',
    'etape': 'étape', 'etapes': 'étapes',
    # informatique et methode
    'donnees': 'données', 'donnee': 'donnée',
    'modele': 'modèle', 'modeles': 'modèles',
    'systeme': 'système', 'systemes': 'systèmes',
    'methodologie': 'méthodologie', 'methode': 'méthode',
    'methodes': 'méthodes', 'methodologique': 'méthodologique',
    'resultat': 'résultat', 'resultats': 'résultats',
    'detection': 'détection', 'detections': 'détections',
    'critere': 'critère', 'criteres': 'critères',
    'parametre': 'paramètre', 'parametres': 'paramètres',
    'probleme': 'problème', 'problemes': 'problèmes',
    'periode': 'période', 'periodes': 'périodes',
    'reference': 'référence', 'references': 'références',
    'numero': 'numéro', 'numeros': 'numéros', 'numerique': 'numérique',
    'numeriques': 'numériques',
    'scenario': 'scénario', 'scenarios': 'scénarios',
    'categorie': 'catégorie', 'categories': 'catégories',
    'strategie': 'stratégie', 'strategique': 'stratégique',
    'implementation': 'implémentation',
    'developpement': 'développement', 'developpements': 'développements',
    'interpretation': 'interprétation',
    'requete': 'requête', 'requetes': 'requêtes',
    'reseau': 'réseau', 'reseaux': 'réseaux',
    'schema': 'schéma', 'schemas': 'schémas',
    'repartition': 'répartition', 'reponse': 'réponse', 'reponses': 'réponses',
    'representation': 'représentation', 'presentation': 'présentation',
    'pretraitement': 'prétraitement', 'procedure': 'procédure',
    'procedures': 'procédures', 'recuperation': 'récupération',
    'redaction': 'rédaction', 'reduction': 'réduction', 'reduire': 'réduire',
    'resolution': 'résolution', 'selection': 'sélection',
    'specification': 'spécification', 'specifications': 'spécifications',
    'specifique': 'spécifique', 'specifiques': 'spécifiques',
    'synthese': 'synthèse', 'verification': 'vérification',
    'evaluation': 'évaluation', 'evolution': 'évolution',
    'element': 'élément', 'elements': 'éléments',
    'metrique': 'métrique', 'metriques': 'métriques',
    'mecanisme': 'mécanisme', 'mecanismes': 'mécanismes',
    'materiel': 'matériel', 'theorique': 'théorique', 'theorie': 'théorie',
    'predictif': 'prédictif', 'prediction': 'prédiction',
    'precision': 'précision', 'precis': 'précis',
    'prevu': 'prévu', 'prevoir': 'prévoir', 'prevision': 'prévision',
    'regle': 'règle', 'regles': 'règles', 'reglement': 'règlement',
    'reglementaire': 'réglementaire', 'reglementation': 'réglementation',
    'serie': 'série', 'series': 'séries',
    'operationnel': 'opérationnel', 'operation': 'opération',
    'operations': 'opérations', 'operateur': 'opérateur',
    # qualites et mesures
    'securite': 'sécurité', 'qualite': 'qualité', 'fiabilite': 'fiabilité',
    'efficacite': 'efficacité', 'capacite': 'capacité',
    'activite': 'activité', 'activites': 'activités',
    'priorite': 'priorité', 'priorites': 'priorités',
    'autorite': 'autorité', 'communaute': 'communauté',
    'difficulte': 'difficulté', 'difficultes': 'difficultés',
    'realite': 'réalité', 'unite': 'unité', 'unites': 'unités',
    'entite': 'entité', 'entites': 'entités', 'identite': 'identité',
    'propriete': 'propriété', 'proprietaire': 'propriétaire',
    'severite': 'sévérité', 'stabilite': 'stabilité', 'validite': 'validité',
    'visibilite': 'visibilité', 'lisibilite': 'lisibilité',
    'sensibilite': 'sensibilité', 'majorite': 'majorité',
    'quantite': 'quantité', 'quantites': 'quantités',
    'mobilite': 'mobilité', 'modalite': 'modalité', 'modalites': 'modalités',
    'responsabilite': 'responsabilité', 'possibilite': 'possibilité',
    'possibilites': 'possibilités', 'necessite': 'nécessité',
    'necessaire': 'nécessaire', 'necessaires': 'nécessaires',
    'fonctionnalite': 'fonctionnalité', 'fonctionnalites': 'fonctionnalités',
    'integrite': 'intégrité', 'integration': 'intégration',
    'tolerance': 'tolérance', 'verite': 'vérité', 'variete': 'variété',
    # metier et environnement
    'specialiste': 'spécialiste', 'specialistes': 'spécialistes',
    'specialite': 'spécialité', 'geolocalisation': 'géolocalisation',
    'geographique': 'géographique', 'geographiques': 'géographiques',
    'meteorologique': 'météorologique', 'meteorologiques': 'météorologiques',
    'vegetation': 'végétation', 'temperature': 'température',
    'phenomene': 'phénomène', 'phenomenes': 'phénomènes',
    'poussiere': 'poussière', 'poussieres': 'poussières',
    'prelevement': 'prélèvement', 'prelevements': 'prélèvements',
    'emission': 'émission', 'emissions': 'émissions',
    'energie': 'énergie', 'energetique': 'énergétique',
    'ecosysteme': 'écosystème', 'ecosystemes': 'écosystèmes',
    'economique': 'économique', 'economie': 'économie',
    'vehicule': 'véhicule', 'vehicules': 'véhicules',
    'telephone': 'téléphone', 'ecran': 'écran', 'ecrans': 'écrans',
    'echelle': 'échelle', 'echange': 'échange', 'echanges': 'échanges',
    'echantillon': 'échantillon', 'echantillons': 'échantillons',
    'ecart': 'écart', 'ecarts': 'écarts', 'duree': 'durée', 'durees': 'durées',
    'enquete': 'enquête', 'enquetes': 'enquêtes',
    'fenetre': 'fenêtre', 'fenetres': 'fenêtres',
    'tache': 'tâche', 'taches': 'tâches', 'role': 'rôle', 'roles': 'rôles',
    'legende': 'légende', 'legendes': 'légendes',
    'repertoire': 'répertoire', 'reunion': 'réunion',
    'societe': 'société', 'region': 'région', 'regions': 'régions',
    'routiere': 'routière', 'routieres': 'routières',
    'sphere': 'sphère', 'espece': 'espèce', 'especes': 'espèces',
    'evenement': 'événement', 'evenements': 'événements',
    'experience': 'expérience', 'frequence': 'fréquence',
    'sequence': 'séquence', 'consequence': 'conséquence',
    'consequences': 'conséquences', 'presence': 'présence',
    'difference': 'différence', 'differences': 'différences',
    'different': 'différent', 'differents': 'différents',
    'differente': 'différente', 'differentes': 'différentes',
    'independant': 'indépendant', 'independance': 'indépendance',
    'dependance': 'dépendance', 'dependances': 'dépendances',
    'inferieur': 'inférieur', 'inferieure': 'inférieure',
    'superieur': 'supérieur', 'superieure': 'supérieure',
    'precedent': 'précédent', 'precedente': 'précédente',
    'particuliere': 'particulière', 'particulieres': 'particulières',
    'maniere': 'manière', 'manieres': 'manières',
    'matiere': 'matière', 'matieres': 'matières',
    'hierarchie': 'hiérarchie', 'interet': 'intérêt', 'interets': 'intérêts',
    'liberte': 'liberté', 'degre': 'degré', 'degres': 'degrés',
    'cle': 'clé', 'cles': 'clés', 'acces': 'accès', 'apres': 'après',
    'succes': 'succès', 'tres': 'très', 'egalement': 'également',
    'immediat': 'immédiat', 'immediatement': 'immédiatement',
    'regulier': 'régulier', 'reguliere': 'régulière',
    'regulierement': 'régulièrement',
    'reel': 'réel', 'reelle': 'réelle', 'reels': 'réels', 'reelles': 'réelles',
    'reellement': 'réellement', 'detail': 'détail', 'details': 'détails',
    'etroite': 'étroite', 'etendue': 'étendue', 'ecole': 'école',
    'edition': 'édition', 'education': 'éducation',
    'electrique': 'électrique', 'electronique': 'électronique',
    'extreme': 'extrême', 'extremite': 'extrémité',
    'entrainement': 'entraînement', 'entrainements': 'entraînements',
    'portee': 'portée', 'pieces': 'pièces', 'piece': 'pièce',
    'reparation': 'réparation', 'reussite': 'réussite',
    'revision': 'révision', 'ulterieur': 'ultérieur',
}

#: Expressions a corriger en bloc, la substitution mot a mot ne suffisant pas.
EXPRESSIONS = {
    "l'etude": "l'étude",
    "d'etude": "d'étude",
    "l'existant": "l'existant",       # sans accent, present pour memoire
    "l'etat": "l'état",
    "d'etat": "d'état",
    "l'element": "l'élément",
    "l'echelle": "l'échelle",
    "l'ecran": "l'écran",
    "l'etape": "l'étape",
    "d'etape": "d'étape",
    "l'evaluation": "l'évaluation",
    "l'evolution": "l'évolution",
    "l'energie": "l'énergie",
    "l'echange": "l'échange",
    "l'ecart": "l'écart",
    "l'enquete": "l'enquête",
    "l'entite": "l'entité",
    "l'identite": "l'identité",
    "l'integration": "l'intégration",
    "l'integrite": "l'intégrité",
    "l'interet": "l'intérêt",
    "l'operation": "l'opération",
    "l'implementation": "l'implémentation",
    "l'interpretation": "l'interprétation",
    "l'entrainement": "l'entraînement",
    "l'evenement": "l'événement",
    "l'experience": "l'expérience",
    "l'ecosysteme": "l'écosystème",
    "l'economie": "l'économie",
}


def casse_de(source, cible):
    """Reporte la capitalisation de `source` sur `cible`."""
    if source.isupper():
        return cible.upper()
    if source[:1].isupper():
        return cible[:1].upper() + cible[1:]
    return cible


def corriger(texte):
    def remplace(m):
        mot = m.group(0)
        bon = MOTS.get(mot.lower())
        return casse_de(mot, bon) if bon else mot

    # Le mot doit etre entier ; \b seul laisserait passer les mots deja
    # accentues dont la racine figure dans le dictionnaire.
    nouveau = re.sub(r"\b[A-Za-zÀ-ÿ']+\b", remplace, texte)
    for avant, apres in EXPRESSIONS.items():
        if avant == apres:
            continue
        nouveau = re.sub(re.escape(avant), apres, nouveau, flags=re.I)
    return nouveau


def main():
    doc = Document(CIBLE)

    n_runs = n_par = 0
    vus = set()

    def traiter(paragraphe):
        nonlocal n_runs, n_par
        change = False
        for run in paragraphe.runs:
            if not run.text:
                continue
            nouveau = corriger(run.text)
            if nouveau != run.text:
                run.text = nouveau
                n_runs += 1
                change = True
        # Les runs de python-docx ignorent le contenu des hyperliens : on
        # complete au niveau des noeuds de texte pour le sommaire et la table
        # des matieres.
        for t in paragraphe._element.iter('{%s}t' % W):
            if not t.text or id(t) in vus:
                continue
            vus.add(id(t))
            nouveau = corriger(t.text)
            if nouveau != t.text:
                t.text = nouveau
                change = True
        if change:
            n_par += 1

    for p in doc.paragraphs:
        traiter(p)
    for table in doc.tables:
        for ligne in table.rows:
            for cellule in ligne.cells:
                for p in cellule.paragraphs:
                    traiter(p)

    print("  paragraphes corriges : %d" % n_par)
    print("  fragments de texte modifies : %d" % n_runs)

    doc.save(CIBLE)
    print("\nEnregistre : %s" % CIBLE)


if __name__ == "__main__":
    main()
