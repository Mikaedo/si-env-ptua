# -*- coding: utf-8 -*-
"""
Ramene les legendes de figure a une longueur lisible.

Une legende nomme ce que la figure montre ; elle ne la commente pas. Six
d'entre elles avaient derive vers le paragraphe explicatif, jusqu'a deux cent
vingt caracteres, ce qui alourdit la page et deborde sur trois lignes sous
l'image.

Le contenu n'est pas perdu pour autant. Ce qui disparait d'une legende est soit
deja dit dans le texte qui l'entoure, soit accessoire : la liste exhaustive des
sources satellitaires, l'enumeration des ecrans d'une capture. Ce qui porte le
sens, en revanche, est conserve : pour la figure 4.10, le fait que l'ecran soit
depourvu de commande de modification est precisement ce que la figure
demontre, et reste donc ecrit.

Chaque legende apparait deux fois, dans la liste des figures et sous l'image.
La substitution porte sur les deux, faute de quoi la liste annoncerait un
libelle que le corps ne porte plus.
"""
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_legendes.docx")

REECRITURES = [
    # (ancienne legende, nouvelle)
    ("Figure 1.1 : Organigramme de la Direction des Affaires Juridiques, "
     "Moyens Généraux et Système d'Information (AGEROUTE).",
     "Figure 1.1 : Organigramme de la direction d'accueil à l'AGEROUTE."),

    ("Figure 4.10 : Tableau de bord vu par l'Agence Nationale de "
     "l'Environnement. La barre latérale ne propose ni les plaintes ni "
     "l'administration, et la liste des signalements ne comporte aucune "
     "commande de modification.",
     "Figure 4.10 : Tableau de bord vu par l'ANDE, dépourvu de toute commande "
     "de modification."),

    ("Figure 5.2 : Captures d'écran de l'application mobile (liste des "
     "signalements, carte des chantiers PTUA, détail d'un signalement).",
     "Figure 5.2 : Application mobile agent : liste, carte et détail d'un "
     "signalement."),

    ("Figure 5.2 bis : Application citoyenne. À gauche, l'écran d'accueil qui "
     "expose la raison pour laquelle la position est demandée. À droite, le "
     "dépôt d'une doléance, réduit à une catégorie et à une description.",
     "Figure 5.2 bis : Application citoyenne : écran d'accueil et dépôt d'une "
     "doléance."),

    ("Figure 5.9 : Indices environnementaux calculés via Google Earth Engine "
     "(NO2, NDVI, NDWI, risque pluie/relief) pour les six chantiers PTUA "
     "(données réelles Sentinel-5P, Sentinel-2, CHIRPS et SRTM, capture du "
     "10 août 2026).",
     "Figure 5.9 : Indices satellitaires des six chantiers PTUA, calculés via "
     "Google Earth Engine (août 2026)."),

    ("Figure 5.3 : Captures d'écran du tableau de bord (vue d'ensemble, "
     "statistiques et carte des traces PTUA).",
     "Figure 5.3 : Tableau de bord : vue d'ensemble, statistiques et carte."),
]


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    for avant, apres in REECRITURES:
        touches = 0
        for p in doc.paragraphs:
            if avant not in p.text:
                continue
            entier = p.text.replace(avant, apres)
            for fragment in list(p.runs)[1:]:
                fragment._element.getparent().remove(fragment._element)
            if p.runs:
                p.runs[0].text = entier
            else:
                p.add_run(entier)
            touches += 1
        etat = f"{touches} emplacement(s)" if touches else "INTROUVABLE"
        print(f"  [{len(apres):>3}] {apres[:58]}  ({etat})")

    doc.save(SOURCE)

    # Controle : plus aucune legende au-dela de la limite retenue.
    import re
    controle = Document(SOURCE)
    longues = []
    vues = set()
    for p in controle.paragraphs:
        t = p.text.strip()
        m = re.match(r"^(Figure [\d.]+(?: bis)?)\s*:", t)
        if m and len(t) > 110 and m.group(1) not in vues:
            vues.add(m.group(1))
            longues.append(f"{len(t)} : {t[:70]}")
    print(f"\nlegendes encore au-dela de 110 caracteres : {len(longues)}")
    for l in longues:
        print(f"  {l}")
    print(f"Sauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
