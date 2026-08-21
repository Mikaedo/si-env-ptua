# -*- coding: utf-8 -*-
"""
Rattache les trois captures orphelines a une annexe et les renumerote.

Les figures 5.1, 5.3 et 5.4 avaient ete deplacees du corps vers les annexes
sans que leur numerotation suive. Elles se retrouvaient apres l'annexe F, sans
titre d'annexe, en portant une numerotation de chapitre au milieu de figures
numerotees par lettre. Une figure ainsi posee n'appartient a rien.

Elles forment desormais l'annexe G, et deviennent G.1, G.2 et G.3. La
correction porte sur quatre endroits, car un numero de figure vit a quatre
places : la legende sous l'image, l'entree dans la liste des figures, et les
renvois du corps, ici au nombre de trois. En oublier un laisserait le texte
appeler une figure qui n'existe plus.

L'ordre est conserve : G.1 pour la documentation des services, G.2 pour le
tableau de bord, G.3 pour le journal d'audit.
"""
import copy
import shutil
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_annexe_g.docx")

TITRE_G = ("Annexe G : Captures d'écran des interfaces livrées "
           "(services web, tableau de bord, journal d'audit)")

# Ancien numero -> nouveau. L'ordre importe : traiter 5.1 avant 5.10 eviterait
# une collision, mais aucune figure 5.10 n'existe ici.
RENUMEROTATION = [("Figure 5.1", "Figure G.1"),
                  ("Figure 5.3", "Figure G.2"),
                  ("Figure 5.4", "Figure G.3")]
RENVOIS = [("figure 5.1", "figure G.1"),
           ("figure 5.3", "figure G.2"),
           ("figure 5.4", "figure G.3"),
           ("La figure 5.3", "La figure G.2")]


def reecrire(paragraphe, texte):
    for fragment in list(paragraphe.runs)[1:]:
        fragment._element.getparent().remove(fragment._element)
    if paragraphe.runs:
        paragraphe.runs[0].text = texte
    else:
        paragraphe.add_run(texte)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    # 1. Le titre de l'annexe, calque sur celui de l'annexe F.
    modele = next(p for p in doc.paragraphs
                  if p.text.strip().startswith("Annexe F :"))
    premiere = next(p for p in doc.paragraphs
                    if p.text.strip().startswith("Figure 5.1 :"))
    element = copy.deepcopy(modele._element)
    premiere._element.addprevious(element)
    titre = Paragraph(element, modele._parent)
    reecrire(titre, TITRE_G)
    print(f"  annexe creee : {TITRE_G[:56]}")

    # 2. Les legendes et l'entree de la liste des figures.
    for avant, apres in RENUMEROTATION:
        touches = 0
        for p in doc.paragraphs:
            if p.text.strip().startswith(avant + " "):
                reecrire(p, p.text.replace(avant, apres, 1))
                touches += 1
        print(f"  {avant} devient {apres} : {touches} emplacement(s)")

    # 3. Les renvois du corps.
    for avant, apres in RENVOIS:
        for p in doc.paragraphs:
            if avant in p.text and not p.text.strip().startswith("Figure"):
                reecrire(p, p.text.replace(avant, apres))
                print(f"  renvoi corrige : {apres}")

    doc.save(SOURCE)

    # Controle : plus aucune trace des anciens numeros.
    import re
    controle = Document(SOURCE)
    restes = []
    for p in controle.paragraphs:
        for m in re.finditer(r"[Ff]igure (5\.1|5\.3|5\.4)(?![\d])", p.text):
            restes.append(p.text[max(0, m.start() - 40):m.end() + 40])
    print(f"\nmentions residuelles de 5.1, 5.3 ou 5.4 : {len(restes)}")
    for r in restes:
        print(f"  {r}")
    print(f"Sauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
