# -*- coding: utf-8 -*-
"""
Ajoute une colonne « Logo » au tableau 7.1 (Technologies et versions retenues).

python-docx ne sait pas inserer une colonne : il faut agir sur le XML, en
ajoutant une definition de colonne dans la grille du tableau puis une cellule
dans chaque ligne, a la position voulue. La colonne est placee juste apres
« Technologie » : le logo est ainsi accole au nom qu'il illustre.

Les largeurs sont redistribuees pour que le tableau tienne dans la largeur
utile de la page (16 cm avec les marges du memoire).

Sortie : MEMOIRE_N'GUESSAN_DIBY_logos.docx
"""
import copy
import os
import shutil

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_DIBY.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_DIBY_logos.docx")
LOGOS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logos_tech")

#: Index du tableau des technologies et position d'insertion de la colonne.
TABLEAU = 21
POSITION = 2  # apres « Composant » (0) et « Technologie » (1)

#: Ligne du tableau -> fichier de logo. La ligne 0 est l'en-tete.
CORRESPONDANCE = {
    1: "flutter.png",      # Application mobile   : Flutter (Dart)
    2: "fastapi.png",      # Serveur d'application: FastAPI (Python)
    3: "postgresql.png",   # Base de donnees      : PostgreSQL + PostGIS
    4: "swagger.png",      # Services web         : REST + Swagger/OpenAPI
    5: "jwt.png",          # Securite             : python-jose + JWT
    6: "angular.png",      # Tableau de bord      : Angular (TypeScript)
    7: "leaflet.png",      # Cartographie         : Flutter Map / Leaflet
}

#: Largeurs finales des 4 colonnes, en centimetres (total 16 cm).
LARGEURS = (4.2, 5.6, 2.0, 4.2)


def inserer_colonne(table, position):
    """Ajoute une colonne vide a la position donnee, grille comprise."""
    grille = table._tbl.find(qn('w:tblGrid'))
    colonnes = grille.findall(qn('w:gridCol'))
    nouvelle = copy.deepcopy(colonnes[-1])
    colonnes[position - 1].addnext(nouvelle)

    for ligne in table.rows:
        cellules = ligne._tr.findall(qn('w:tc'))
        # On duplique une cellule existante pour heriter de ses bordures et de
        # son ombrage, puis on la vide de son contenu.
        modele = copy.deepcopy(cellules[position - 1])
        for enfant in list(modele):
            if not enfant.tag.endswith('}tcPr'):
                modele.remove(enfant)
        # Une cellule doit contenir au moins un paragraphe pour rester valide
        paragraphe = modele.makeelement(qn('w:p'), {})
        modele.append(paragraphe)
        cellules[position - 1].addnext(modele)


def main():
    manquants = [f for f in CORRESPONDANCE.values()
                 if not os.path.exists(os.path.join(LOGOS, f))]
    if manquants:
        raise SystemExit("Logos manquants : %s" % manquants)

    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)
    table = doc.tables[TABLEAU]
    print("  avant : %d lignes x %d colonnes"
          % (len(table.rows), len(table.columns)))

    inserer_colonne(table, POSITION)
    print("  apres : %d lignes x %d colonnes"
          % (len(table.rows), len(table.columns)))

    # En-tete de la nouvelle colonne, dans le style des autres en-tetes
    cellule_entete = table.rows[0].cells[POSITION]
    modele_run = None
    for run in table.rows[0].cells[1].paragraphs[0].runs:
        modele_run = run
        break
    par = cellule_entete.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run("Logo")
    if modele_run is not None:
        run.bold = modele_run.bold
        run.font.size = modele_run.font.size
        run.font.name = modele_run.font.name
        if modele_run.font.color and modele_run.font.color.rgb:
            run.font.color.rgb = modele_run.font.color.rgb
    else:
        run.bold = True

    # Insertion des logos
    poses = 0
    for ligne, fichier in CORRESPONDANCE.items():
        cellule = table.rows[ligne].cells[POSITION]
        par = cellule.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(os.path.join(LOGOS, fichier), width=Cm(1.1))
        poses += 1
        print("    ligne %d <- %s" % (ligne, fichier))
    print("  logos inseres : %d" % poses)

    # Redistribution des largeurs
    table.autofit = False
    for ligne in table.rows:
        for i, largeur in enumerate(LARGEURS):
            if i < len(ligne.cells):
                ligne.cells[i].width = Cm(largeur)

    doc.save(SORTIE)

    # Verification
    verif = Document(SORTIE)
    t = verif.tables[TABLEAU]
    B = './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    images = sum(len(c._tc.findall(B)) for r in t.rows for c in r.cells)
    print("\n  colonnes finales : %d | images dans le tableau : %d"
          % (len(t.columns), images))
    print("  en-tetes : %s"
          % " | ".join(c.text.strip() for c in t.rows[0].cells))
    print("\nEnregistre : %s" % SORTIE)


if __name__ == "__main__":
    main()
