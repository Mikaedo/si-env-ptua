# -*- coding: utf-8 -*-
"""
Genere le guide de preparation a la soutenance : cours + questions/reponses
(du plus simple au plus difficile), organise par modules, Times New Roman.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRATCH = r"C:\Users\DELL\AppData\Local\Temp\claude\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25\scratchpad"
OUT = r"C:\Users\DELL\Downloads\GUIDE_PREPARATION_SOUTENANCE_SI-ENV.docx"

ORANGE = RGBColor(0xE8, 0x6C, 0x00)
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GRAY = RGBColor(0x47, 0x55, 0x69)

doc = Document()

# ------------------------------------------------------------------
# Styles de base : tout en Times New Roman
# ------------------------------------------------------------------
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
rpr = normal.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts'); rpr.append(rFonts)
for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'):
    rFonts.set(qn(a), 'Times New Roman')
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(8)

for hname, size, color in [('Heading 1', 20, ORANGE), ('Heading 2', 15, NAVY), ('Heading 3', 12.5, GRAY)]:
    st = doc.styles[hname]
    st.font.name = 'Times New Roman'
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    rpr = st.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rpr.append(rFonts)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'):
        rFonts.set(qn(a), 'Times New Roman')

# Style pour les questions
q_style = doc.styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH)
q_style.base_style = doc.styles['Normal']
q_style.font.bold = True
q_style.font.color.rgb = NAVY
q_style.font.size = Pt(12)
q_style.paragraph_format.space_before = Pt(10)
q_style.paragraph_format.space_after = Pt(3)

diff_colors = {'Basique': RGBColor(0x16,0x7C,0x3C), 'Intermédiaire': RGBColor(0xB4,0x5B,0x00), 'Difficile': RGBColor(0xB9,0x1C,0x1C), 'Piège': RGBColor(0x7C,0x2D,0x92)}

# ------------------------------------------------------------------
# Fonctions d'aide
# ------------------------------------------------------------------

def add_title_page():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(200)
    r = p.add_run("GUIDE DE PRÉPARATION À LA SOUTENANCE")
    r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = ORANGE
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("SI-ENV — Système d'Information Environnemental du PTUA")
    r2.font.size = Pt(16); r2.font.color.rgb = NAVY; r2.font.bold = True
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(20)
    r3 = p3.add_run("Cours + Questions/Réponses, du plus basique au plus difficile\nOrganisé par modules — pour maîtriser aussi bien le fond technique que la défense du mémoire")
    r3.font.size = Pt(12); r3.italic = True; r3.font.color.rgb = GRAY
    doc.add_page_break()

def add_sommaire():
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("SOMMAIRE"); r.font.size=Pt(20); r.font.bold=True; r.font.color.rgb=ORANGE
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin); run._r.append(instr)
    r2 = p.add_run(); r2._r.append(fld_sep)
    txt = p.add_run("Clic droit → Mettre à jour les champs pour générer le sommaire.")
    txt.italic = True
    r3 = p.add_run(); r3._r.append(fld_end)
    doc.add_page_break()

def add_module(num, title, intro):
    doc.add_heading(f"Module {num} — {title}", level=1)
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(10)

def add_sub(title):
    doc.add_heading(title, level=2)

def add_qa(question, answer, level='Basique'):
    p = doc.add_paragraph(style='Question')
    tag = p.add_run(f"[{level}]  ")
    tag.font.color.rgb = diff_colors.get(level, NAVY)
    tag.font.size = Pt(10); tag.bold = True
    p.add_run("Q : " + question)
    a = doc.add_paragraph(answer)
    a.paragraph_format.left_indent = Pt(14)

def add_image(path, width_in=6.0, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = c.add_run(caption)
        cr.italic = True; cr.font.size = Pt(10.5); cr.font.color.rgb = GRAY

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(14)
    r = p.add_run("💡 " + text)
    r.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x66, 0x33)

# ==================================================================
# GENERATION DU DOCUMENT
# ==================================================================
add_title_page()
add_sommaire()

# ==================================================================
# MODULE 1 — CONTEXTE, PROBLEMATIQUE, METHODOLOGIE DE PROJET
# ==================================================================
add_module(1, "Contexte, problématique et méthodologie de projet",
    "Ce module couvre le cadre institutionnel (AGEROUTE, PTUA), la justification du sujet, "
    "et la méthodologie de gestion de projet (Scrum). Ce sont typiquement les toutes premières "
    "questions du jury — elles servent à vérifier que vous maîtrisez le contexte avant d'entrer "
    "dans le technique.")

add_sub("1.1 Le contexte institutionnel")
doc.add_paragraph(
    "AGEROUTE (Agence de Gestion des Routes) est un établissement public ivoirien créé en 2001, "
    "chargé de l'entretien et du développement du réseau routier national. Le PTUA (Projet de "
    "Transport Urbain d'Abidjan) est un grand projet d'infrastructure (657,8 milliards de FCFA) "
    "financé notamment par la Banque Africaine de Développement (BAD), classé Catégorie 1 (impact "
    "environnemental et social élevé) — ce qui impose un suivi environnemental rigoureux et "
    "documenté. La CC-PTUA (Cellule de Coordination) pilote ce projet ; vous avez été affecté à la "
    "Cellule Informatique et Logistique, rattachée au SEDA (Service Études et Développement "
    "Applicatif) de la DSI d'AGEROUTE."
)
add_qa("Qu'est-ce qu'AGEROUTE et quel est son rôle ?",
    "AGEROUTE est l'agence publique ivoirienne responsable de la gestion, l'entretien et le "
    "développement du réseau routier national. Elle intervient sur tout le cycle de vie des "
    "infrastructures routières : planification, construction, entretien.", "Basique")
add_qa("Pourquoi le PTUA nécessite-t-il un suivi environnemental particulier ?",
    "Parce qu'il est classé Catégorie 1 par la BAD (impact environnemental et social potentiellement "
    "élevé, du fait de l'ampleur des travaux en milieu urbain dense). Les bailleurs de fonds exigent "
    "un Plan de Gestion Environnementale et Sociale (PGES) et un reporting régulier et traçable — "
    "c'est précisément ce que SI-ENV vient outiller.", "Basique")
add_qa("Quelle est la problématique exacte de votre mémoire ?",
    "Le suivi environnemental du PTUA reposait sur des outils bureautiques classiques (carnets, "
    "e-mails), ce qui posait six problèmes : détection tardive des nuisances, subjectivité de "
    "l'évaluation de criticité, absence de géolocalisation précise, données dispersées, rapports "
    "manuels chronophages, absence d'alertes proactives. La question de recherche est : peut-on, "
    "avec les technologies disponibles, transformer ce suivi artisanal en un dispositif fiable, "
    "géolocalisé et réactif ?", "Basique")
add_qa("Pourquoi ne pas avoir simplement amélioré les outils Excel/Word existants au lieu de développer un système complet ?",
    "Les outils bureautiques ne résolvent pas les six limites structurelles identifiées : ils ne "
    "peuvent pas géolocaliser automatiquement, croiser des données satellite, fonctionner hors ligne "
    "sur le terrain, ni déclencher des alertes automatiques. Un tableur reste un stockage passif ; "
    "SI-ENV est un système actif qui capture, analyse et alerte.", "Intermédiaire")

add_sub("1.2 Méthodologie Scrum")
doc.add_paragraph(
    "Scrum est une méthode agile de gestion de projet. Le travail est découpé en itérations courtes "
    "et régulières appelées sprints (ici, environ une semaine chacun), à l'issue desquelles un "
    "incrément fonctionnel du produit est livré et évalué. Cela s'oppose au cycle en V (ou cascade), "
    "où toutes les phases — analyse, conception, réalisation, tests — se déroulent une seule fois, "
    "séquentiellement, sans retour en arrière possible avant la fin."
)
add_qa("Qu'est-ce qu'un sprint ?",
    "Une période de temps fixe et courte (ici environ une semaine) pendant laquelle une liste "
    "d'objectifs précis est réalisée, aboutissant à un incrément livrable et testable du système.",
    "Basique")
add_qa("Pourquoi Scrum plutôt qu'un cycle en V, plus classique pour un projet institutionnel ?",
    "Le cycle en V exige de figer les besoins dès le départ, ce qui est risqué sur un projet où les "
    "besoins métier (critères de criticité, workflow de validation) ont émergé progressivement au "
    "contact des utilisateurs (Spécialiste Environnement, Expert HSE) pendant le stage. Scrum permet "
    "de valider chaque brique avec les utilisateurs avant de passer à la suivante, réduisant le "
    "risque de développer une fonctionnalité inadaptée.", "Intermédiaire")
add_qa("Un stage de 3 mois avec des sprints d'une semaine, n'est-ce pas artificiel pour une seule personne ?",
    "C'est une adaptation assumée : les cérémonies Scrum (mêlées quotidiennes, revue de sprint) sont "
    "allégées, mais le principe structurant — livrer un incrément testable chaque semaine plutôt "
    "qu'attendre la fin du stage pour tout tester — a été conservé car il a permis de détecter tôt "
    "des incompréhensions de besoin (ex. les critères de criticité) plutôt qu'à la fin du stage.",
    "Piège")

# ==================================================================
# MODULE 2 — CONCEPTION : UML ET MERISE
# ==================================================================
add_module(2, "Conception — UML et MERISE",
    "UML et MERISE sont deux méthodes de modélisation, utilisées ici de façon complémentaire : "
    "UML pour décrire le comportement et l'architecture du système, MERISE pour modéliser la base "
    "de données relationnelle. Le jury adore vérifier que vous savez vraiment lire vos propres "
    "diagrammes — révisez-les avant l'oral.")

add_sub("2.1 UML — les diagrammes utilisés")
doc.add_paragraph(
    "UML (Unified Modeling Language) est un langage de modélisation orienté objet, standardisé, "
    "utilisé pour représenter différentes vues d'un système. Quatre diagrammes ont été utilisés :"
)
doc.add_paragraph("• Diagramme de cas d'utilisation : représente les interactions entre les acteurs (utilisateurs, systèmes externes) et les fonctionnalités du système, sans détailler le « comment ».", style='List Bullet')
doc.add_paragraph("• Diagramme de classes : représente la structure statique — les entités du système, leurs attributs, et les relations entre elles (associations, héritage).", style='List Bullet')
doc.add_paragraph("• Diagramme de séquence : représente le déroulement chronologique des échanges de messages entre acteurs et composants pour un scénario précis.", style='List Bullet')
doc.add_paragraph("• Diagramme de déploiement : représente la répartition physique des composants logiciels sur les machines/serveurs (ici : téléphone, serveur AGEROUTE avec Docker, cloud GEE).", style='List Bullet')

add_qa("Quelle est la différence entre un diagramme de cas d'utilisation et un diagramme de séquence ?",
    "Le cas d'utilisation dit QUOI le système fait pour chaque acteur (une vue fonctionnelle, statique, "
    "en haut niveau). Le diagramme de séquence dit COMMENT cela se déroule dans le temps, message par "
    "message, entre les objets/composants concrets (une vue dynamique, détaillée).", "Basique")
add_qa("Qu'est-ce qu'un acteur en UML ? Un acteur est-il forcément une personne ?",
    "Un acteur est une entité externe qui interagit avec le système sans en faire partie. Ce n'est "
    "pas forcément une personne : dans SI-ENV, Google Earth Engine est un acteur (système externe) "
    "au même titre que le Spécialiste Environnement.", "Basique")
add_qa("Dans votre diagramme de classes, comment est modélisée la relation entre Signalement et Utilisateur ?",
    "Un Utilisateur SAISIT un Signalement : association 1,1 côté Signalement (un signalement est "
    "saisi par exactement un utilisateur) et 0,n côté Utilisateur (un utilisateur peut saisir zéro "
    "à plusieurs signalements). C'est une association simple, pas un héritage.", "Intermédiaire")
add_qa("Pourquoi ne pas avoir modélisé les rôles (RESP_ENV, EXPERT_HSE, etc.) comme des sous-classes d'Utilisateur en UML (héritage) plutôt que comme un simple attribut roleSysteme ?",
    "Un héritage aurait été pertinent si les rôles avaient des attributs ou comportements "
    "structurellement différents (par exemple, des tables séparées). Ici, les cinq rôles partagent "
    "exactement les mêmes attributs et ne se distinguent que par leurs permissions (gérées par le "
    "RBAC applicatif, pas par la structure de données) — un attribut énuméré est donc plus simple et "
    "suffisant, conformément au principe de ne pas complexifier un modèle sans bénéfice réel.",
    "Difficile")

add_sub("2.2 MERISE — MCD et MLD")
doc.add_paragraph(
    "MERISE est une méthode d'analyse et de conception de systèmes d'information, très utilisée en "
    "France et en Afrique francophone pour la modélisation des bases de données. Elle distingue "
    "plusieurs niveaux d'abstraction, dont deux ont été produits dans ce mémoire :"
)
doc.add_paragraph("• Le MCD (Modèle Conceptuel de Données) décrit le QUOI : les entités du monde réel (Signalement, Chantier, Alerte...), leurs propriétés, et les relations (associations) entre elles, avec leurs cardinalités — indépendamment de toute technologie.", style='List Bullet')
doc.add_paragraph("• Le MLD (Modèle Logique de Données) traduit le MCD en un schéma relationnel concret : les entités deviennent des tables, les associations deviennent des clés étrangères (ou des tables de jointure pour les associations n,n).", style='List Bullet')

add_qa("C'est quoi une cardinalité, par exemple 0,n ou 1,1 ?",
    "Une cardinalité indique combien de fois, au minimum et au maximum, une occurrence d'une entité "
    "peut participer à une association. Exemple : Utilisateur (0,n) —SAISIT— Signalement (1,1) "
    "signifie qu'un utilisateur peut saisir de 0 à N signalements (minimum 0 : un utilisateur peut "
    "n'avoir jamais rien saisi), mais qu'un signalement est toujours saisi par exactement 1 "
    "utilisateur (ni 0, ni plusieurs).", "Basique")
add_qa("Comment une association many-to-many (n,n) du MCD se traduit-elle dans le MLD ?",
    "Elle devient une table de jointure (table d'association) à part entière, contenant en clé "
    "primaire composite les clés étrangères des deux entités liées, plus éventuellement des attributs "
    "propres à la relation elle-même.", "Intermédiaire")
add_qa("Pourquoi utiliser à la fois UML (objet) et MERISE (entité-association), alors que ce sont deux écoles de pensée différentes et parfois jugées redondantes ?",
    "Elles répondent à des besoins différents et complémentaires ici : UML documente le comportement "
    "et l'architecture logicielle (utile pour le code, orienté objet, adapté à FastAPI/Flutter qui "
    "sont eux-mêmes orientés objet) ; MERISE, en revanche, reste la méthode la plus lisible et la "
    "plus enseignée pour concevoir un schéma relationnel PostgreSQL avec des non-informaticiens du "
    "métier (Spécialiste Environnement) — le MCD peut être relu et validé par eux sans connaître "
    "UML. Ce n'est pas une redondance mais un choix pragmatique : le bon outil pour chaque livrable.",
    "Piège")
add_qa("Le champ geom (type Geometry) du Signalement, comment est-il représenté au niveau conceptuel puis logique ?",
    "Au niveau MCD, c'est simplement une propriété de l'entité Signalement (sa position). Au niveau "
    "MLD/physique, PostGIS l'implémente comme une colonne de type geometry(Point, 4326) — 4326 étant "
    "le code EPSG du système de coordonnées géographiques standard (WGS 84, utilisé par le GPS).",
    "Difficile")

add_image(f"{SCRATCH}\\illus_archi.png", 6.0, "Illustration — architecture 3-tiers de SI-ENV (à connaître par cœur)")

# ==================================================================
# MODULE 3 — BACKEND, API REST, SECURITE, BASE DE DONNEES
# ==================================================================
add_module(3, "Backend, API REST, sécurité et base de données",
    "Le backend est le cœur technique du système. C'est le module le plus susceptible de questions "
    "techniques pointues, car c'est souvent la partie la mieux maîtrisée par les jurys MIAGE.")

add_sub("3.1 API REST — les bases")
doc.add_paragraph(
    "Une API REST (Representational State Transfer) est un style d'architecture pour exposer des "
    "services web via HTTP. Chaque ressource (un signalement, un chantier...) est identifiée par une "
    "URL, et on agit dessus avec les verbes HTTP standards : GET (lire), POST (créer), PATCH/PUT "
    "(modifier), DELETE (supprimer). Les réponses utilisent des codes de statut normalisés : 200 "
    "(succès), 201 (créé), 400 (requête invalide), 401 (non authentifié), 403 (authentifié mais non "
    "autorisé), 404 (introuvable), 500 (erreur serveur)."
)
add_qa("Qu'est-ce que le principe \"stateless\" (sans état) en REST, et pourquoi est-il important ici ?",
    "Le serveur ne conserve aucune information de session entre deux requêtes : chaque requête doit "
    "contenir toutes les informations nécessaires à son traitement (ici, le jeton JWT). C'est "
    "essentiel pour un système mobile hors-ligne : l'application peut envoyer ses requêtes de "
    "synchronisation à n'importe quel moment, sans dépendre d'une session serveur qui aurait pu "
    "expirer entre-temps.", "Basique")
add_qa("Quelle est la différence entre les codes 401 et 403 ?",
    "401 (Unauthorized) signifie que la requête n'est pas authentifiée du tout (jeton absent ou "
    "invalide). 403 (Forbidden) signifie que l'utilisateur EST authentifié, mais que son rôle ne lui "
    "donne pas le droit d'effectuer cette action précise (ex. un Responsable Environnement qui tente "
    "de créer un compte utilisateur, réservé à l'Administrateur).", "Intermédiaire")
add_qa("Pourquoi FastAPI plutôt que Django ou Flask ?",
    "FastAPI a été choisi pour trois raisons concrètes : (1) validation automatique des données "
    "entrantes/sortantes via Pydantic, ce qui réduit les erreurs de type sur un projet manipulant "
    "beaucoup de données géospatiales et de dates ; (2) documentation interactive générée "
    "automatiquement (Swagger UI, figure 7.1), utile pour un stage où le temps de rédaction manuelle "
    "de la documentation était limité ; (3) performances asynchrones natives (basé sur Starlette/"
    "ASGI), pertinentes pour des appels externes lents comme Google Earth Engine.", "Intermédiaire")

add_sub("3.2 Authentification JWT")
doc.add_paragraph(
    "JWT (JSON Web Token) est un standard de jeton d'authentification. Après connexion (email + mot "
    "de passe), le serveur génère un jeton signé cryptographiquement (ici HS256, valable 1 heure) "
    "contenant l'identité et le rôle de l'utilisateur. Ce jeton est renvoyé par le client à chaque "
    "requête (en-tête Authorization: Bearer <jeton>) ; le serveur vérifie sa signature sans avoir "
    "besoin d'interroger la base de données ni de conserver une session."
)
add_qa("En quoi un jeton JWT diffère-t-il d'un système de session classique (cookie de session) ?",
    "Une session classique stocke l'état de connexion côté serveur (dans une mémoire ou une base), "
    "et le client ne détient qu'un identifiant de session. Un JWT, lui, contient directement les "
    "informations (identité, rôle, expiration) et est auto-suffisant, signé, vérifiable sans état "
    "serveur — ce qui simplifie le passage à l'échelle et convient à un client mobile qui peut être "
    "hors-ligne.", "Basique")
add_qa("Que se passe-t-il si un jeton JWT est volé (interception réseau) ?",
    "L'attaquant peut l'utiliser jusqu'à son expiration (ici 1 heure maximum), en usurpant "
    "l'identité de l'utilisateur avec ses droits exacts (pas plus). C'est pourquoi (1) toutes les "
    "communications transitent en HTTPS (chiffrement empêchant l'interception en clair), et (2) la "
    "durée de validité est volontairement courte plutôt qu'illimitée.", "Intermédiaire")
add_qa("Le mot de passe est-il stocké en clair dans la base ?",
    "Non, jamais. Il est haché avec bcrypt, un algorithme de hachage à sens unique et volontairement "
    "lent (protection contre les attaques par force brute), avec un sel intégré automatiquement. "
    "Seul le hash est stocké ; il est impossible de retrouver le mot de passe original à partir du "
    "hash.", "Basique")
add_qa("Pourquoi HS256 (clé symétrique partagée) plutôt que RS256 (clé publique/privée asymétrique) pour signer les JWT ?",
    "HS256 utilise une seule clé secrète partagée pour signer ET vérifier — plus simple à gérer pour "
    "une architecture où c'est le même serveur FastAPI qui émet et vérifie les jetons. RS256 "
    "(asymétrique) se justifie surtout quand plusieurs services indépendants doivent vérifier des "
    "jetons sans pouvoir en émettre eux-mêmes (ex. microservices multiples) — ce qui n'est pas le "
    "cas ici avec un backend monolithique unique.", "Difficile")

add_sub("3.3 RBAC — contrôle d'accès par rôle")
doc.add_paragraph(
    "RBAC (Role-Based Access Control) restreint les actions possibles selon le rôle de "
    "l'utilisateur. SI-ENV définit cinq rôles : Responsable Environnement, Expert HSE, Spécialiste "
    "Suivi Environnemental, Spécialiste Suivi du P.A.R, Administrateur — chacun avec des permissions "
    "précises, vérifiées côté serveur à chaque requête sensible (jamais seulement côté client)."
)
add_qa("Pourquoi la vérification des droits doit-elle se faire côté serveur, et pas seulement côté application mobile/web ?",
    "Parce que le client (application mobile, navigateur) est entièrement sous le contrôle de "
    "l'utilisateur : il peut être modifié, contourné, ou ses requêtes rejouées directement via un "
    "outil comme curl. Si le contrôle n'existait que côté client, n'importe qui pourrait appeler "
    "l'API directement sans passer par l'interface et contourner les restrictions. Le serveur est la "
    "seule autorité de confiance.", "Intermédiaire")

add_sub("3.4 PostgreSQL et PostGIS")
doc.add_paragraph(
    "PostgreSQL est un système de gestion de base de données relationnelle. PostGIS est une "
    "extension qui ajoute des types de données géométriques (points, polygones...) et des fonctions "
    "spatiales (distance, intersection, buffer...) directement utilisables en SQL."
)
add_qa("Qu'est-ce que PostGIS apporte concrètement que PostgreSQL seul n'a pas ?",
    "PostGIS permet de stocker nativement des coordonnées géographiques (au lieu de deux colonnes "
    "numériques latitude/longitude séparées) et surtout d'exécuter des requêtes spatiales directement "
    "en SQL : trouver tous les signalements dans un rayon de 2,5 km d'un point (ST_DWithin), calculer "
    "une zone tampon (ST_Buffer), etc. Sans PostGIS, ces calculs devraient être faits manuellement "
    "côté application, plus lentement et avec plus de risques d'erreur.", "Basique")
add_qa("Pourquoi une base relationnelle (SQL) plutôt qu'une base NoSQL type MongoDB, alors que les données géospatiales sont parfois citées comme un bon cas d'usage NoSQL ?",
    "Les données de SI-ENV sont fortement structurées et relationnelles par nature (un signalement "
    "appartient à un chantier, est saisi par un utilisateur, peut déclencher une alerte...) — "
    "exactement le cas d'usage pour lequel le modèle relationnel excelle, avec des contraintes "
    "d'intégrité référentielle (clés étrangères) garanties nativement. PostGIS offre en plus un "
    "support géospatial mature et standardisé (OGC), tandis qu'un NoSQL aurait nécessité de "
    "réimplémenter manuellement ces contraintes et relations, augmentant le risque d'incohérence.",
    "Difficile")

# ==================================================================
# MODULE 4 — MOBILE, ARCHITECTURE HORS-LIGNE
# ==================================================================
add_module(4, "Application mobile et architecture hors-ligne",
    "L'application terrain doit fonctionner sans réseau (chantiers isolés). C'est un point fort du "
    "projet — le jury appréciera que vous en maîtrisiez les subtilités techniques.")

add_sub("4.1 Flutter et le pattern BLoC")
doc.add_paragraph(
    "Flutter est un framework de développement mobile multiplateforme (Google), utilisant le "
    "langage Dart, qui permet de produire une seule base de code pour Android et iOS. BLoC "
    "(Business Logic Component) est un patron de conception qui sépare strictement la logique "
    "métier (états, événements) de l'affichage (widgets) : l'interface envoie des événements "
    "(ex. StartSync), le BLoC les traite et émet des états (Loading, Success, Error) que "
    "l'interface se contente d'afficher."
)
add_qa("Pourquoi Flutter plutôt qu'une application native (Kotlin/Swift) ou React Native ?",
    "Flutter permet un seul code source pour Android et iOS avec des performances proches du natif "
    "(compilation en code machine, pas d'interprétation JavaScript comme React Native), et surtout "
    "un écosystème mature pour l'usage hors-ligne (SQLite, gestion d'état). Le choix natif aurait "
    "doublé l'effort de développement pour un stage de 3 mois.", "Basique")
add_qa("Quel est l'intérêt du pattern BLoC plutôt que de mettre la logique directement dans les widgets ?",
    "Cela rend le code testable indépendamment de l'interface (on peut tester la logique de "
    "synchronisation sans lancer l'affichage), et réutilisable : plusieurs écrans peuvent réagir au "
    "même BLoC. Cela évite aussi le mélange de responsabilités qui rend un code difficile à maintenir "
    "à mesure que l'application grossit.", "Intermédiaire")

add_sub("4.2 Fonctionnement hors-ligne et synchronisation")
doc.add_paragraph(
    "Un signalement créé sans réseau est d'abord stocké localement dans une base SQLite embarquée "
    "sur le téléphone, avec un statut PENDING_SYNC et un identifiant unique généré côté client "
    "(uuid_mobile). Au retour du réseau, l'utilisateur déclenche la synchronisation : chaque "
    "signalement en attente est envoyé au serveur, qui vérifie l'uuid_mobile pour éviter les doublons "
    "avant de l'insérer dans PostgreSQL."
)
add_qa("Pourquoi générer l'identifiant unique (UUID) côté mobile plutôt que de laisser le serveur l'attribuer, comme c'est l'usage classique avec un ID auto-incrémenté ?",
    "Parce que le téléphone doit pouvoir créer un signalement complet et cohérent SANS jamais avoir "
    "contacté le serveur (il peut rester des jours hors ligne). Si l'identifiant venait du serveur, "
    "l'enregistrement local n'aurait pas d'identité stable tant que la synchronisation n'a pas eu "
    "lieu. Le UUID garantit une unicité quasi certaine générée localement, sans coordination réseau.",
    "Intermédiaire")
add_qa("Deux agents créent chacun un signalement hors ligne, sur des chantiers différents, en même temps. Y a-t-il un risque de conflit lors de la synchronisation ?",
    "Non : chaque signalement a son propre UUID généré indépendamment (collision quasiment "
    "impossible statistiquement), et un signalement n'est jamais modifié par deux utilisateurs à la "
    "fois dans ce système (il n'y a pas d'édition collaborative simultanée d'un même signalement). Le "
    "risque de conflit réel est donc structurellement très faible — c'est explicitement discuté dans "
    "le mémoire (section 6.6).", "Difficile")
add_qa("Pourquoi la synchronisation est-elle déclenchée manuellement (bouton) plutôt qu'automatiquement dès que le réseau revient ?",
    "C'est une limite assumée et documentée, pas un choix délibérément optimal : une synchronisation "
    "automatique en arrière-plan (via connectivity_plus et un Timer.periodic, par exemple) serait "
    "une amélioration naturelle listée en perspective. Le déclenchement manuel actuel donne "
    "cependant à l'agent terrain un contrôle explicite sur le moment de l'envoi, ce qui évite une "
    "consommation de données mobiles non maîtrisée sur un chantier isolé où la connexion peut être "
    "facturée à l'usage.", "Piège")

# ==================================================================
# MODULE 5 — INTELLIGENCE ARTIFICIELLE
# ==================================================================
add_module(5, "Intelligence artificielle : détection et classification",
    "Module le plus technique et le plus scruté. Le jury va vérifier que vous comprenez vos "
    "modèles au-delà des noms (YOLOv8, MobileNetV2) — les bases du deep learning et vos métriques.")

add_sub("5.1 Réseaux de neurones convolutifs (CNN) — les bases")
doc.add_paragraph(
    "Un CNN (Convolutional Neural Network) est un type de réseau de neurones spécialisé dans le "
    "traitement d'images. Il applique successivement des filtres (noyaux de convolution) qui "
    "détectent des motifs de plus en plus complexes : les premières couches détectent des contours "
    "et des textures simples, les couches profondes détectent des formes et objets complets. C'est "
    "l'entraînement (rétropropagation du gradient) qui ajuste automatiquement les valeurs de ces "
    "filtres à partir d'exemples annotés."
)
add_qa("Pourquoi un CNN plutôt qu'un réseau de neurones classique (fully connected) pour traiter des images ?",
    "Un réseau classique traiterait chaque pixel indépendamment, avec un nombre de paramètres "
    "explosif pour une image de taille réaliste, et sans tenir compte du fait qu'un motif visuel "
    "(un bord, une texture) a la même signification où qu'il apparaisse dans l'image. Le CNN "
    "partage les mêmes filtres sur toute l'image (invariance par translation), ce qui réduit "
    "drastiquement le nombre de paramètres et exploite la structure spatiale de l'image.", "Basique")
add_qa("Qu'est-ce que le transfer learning (apprentissage par transfert), et pourquoi l'avez-vous utilisé ?",
    "Plutôt que d'entraîner un réseau depuis zéro (ce qui nécessite des millions d'images et des "
    "jours de calcul), on part d'un modèle déjà entraîné sur un grand jeu de données généraliste "
    "(ImageNet pour MobileNetV2, COCO pour YOLOv8), et on ré-entraîne seulement les dernières "
    "couches sur nos propres données (déchets de chantier). Le modèle réutilise les motifs visuels "
    "de bas niveau déjà appris (bords, textures, formes) et se spécialise rapidement avec beaucoup "
    "moins de données et de temps de calcul.", "Intermédiaire")
add_qa("Qu'est-ce que le surapprentissage (overfitting), et comment le détecte-t-on ?",
    "Le surapprentissage survient quand le modèle mémorise les exemples d'entraînement au lieu "
    "d'apprendre des motifs généralisables — il devient excellent sur les données déjà vues mais "
    "mauvais sur des données nouvelles. On le détecte en comparant la perte (loss) sur "
    "l'entraînement et sur la validation : si la perte d'entraînement continue de baisser alors que "
    "la perte de validation stagne ou remonte, c'est un signe de surapprentissage. C'est un risque "
    "réel avec un petit jeu de données comme le vôtre.", "Difficile")

add_sub("5.2 Le pipeline en cascade : YOLOv8 puis MobileNetV2")
add_image(f"{SCRATCH}\\illus_ia.png", 6.0, "Illustration — pipeline IA en cascade utilisé pour le diagnostic terrain")
doc.add_paragraph(
    "Le diagnostic se fait en deux étapes successives. YOLOv8n (nano, la plus légère variante de "
    "YOLOv8) détecte d'abord la présence et la position des déchets dans la photo (une ou plusieurs "
    "boîtes englobantes). Chaque zone détectée est ensuite recadrée et transmise à MobileNetV2, qui "
    "classe sa criticité (faible / modérée / importante). C'est une architecture en cascade : deux "
    "modèles spécialisés plutôt qu'un seul modèle généraliste."
)
add_qa("Pourquoi deux modèles séparés plutôt qu'un seul modèle qui détecterait ET classerait la criticité directement ?",
    "Ce sont deux tâches de nature différente : la détection (localiser un objet) et la "
    "classification fine de gravité (évaluer une accumulation) demandent des représentations "
    "différentes. Séparer permet aussi d'entraîner et d'améliorer chaque modèle indépendamment, et "
    "de réutiliser MobileNetV2 même si la méthode de détection change plus tard.", "Intermédiaire")
add_qa("Quel est l'inconvénient principal d'un pipeline en cascade comme celui-ci ?",
    "La propagation d'erreur : si YOLOv8 rate une détection (faux négatif) ou détecte imprécisément "
    "la zone, la classification de criticité par MobileNetV2 en aval est nécessairement dégradée, "
    "puisqu'elle travaille sur une image recadrée potentiellement erronée. C'est explicitement "
    "identifié comme une limite dans la discussion du mémoire (chapitre 10).", "Difficile")
add_qa("Qu'est-ce que le NMS (Non-Maximum Suppression), et pourquoi est-il nécessaire après YOLO ?",
    "YOLO propose souvent plusieurs boîtes qui se chevauchent pour un même objet réel (redondance "
    "du quadrillage interne du modèle). Le NMS élimine les boîtes redondantes en ne conservant, "
    "parmi les boîtes qui se chevauchent fortement (IoU élevé), que celle de plus forte confiance. "
    "Sans NMS, un seul déchet pourrait apparaître détecté plusieurs fois.", "Intermédiaire")

add_sub("5.3 Les métriques d'évaluation")
doc.add_paragraph(
    "Quatre métriques reviennent constamment en vision par ordinateur : la précision, le rappel, "
    "le F1-score et le mAP."
)
doc.add_paragraph("• Précision = Vrais Positifs / (Vrais Positifs + Faux Positifs) — parmi ce que le modèle a détecté/prédit, quelle proportion est correcte.", style='List Bullet')
doc.add_paragraph("• Rappel (Recall) = Vrais Positifs / (Vrais Positifs + Faux Négatifs) — parmi ce qui existait réellement, quelle proportion a été détectée/trouvée.", style='List Bullet')
doc.add_paragraph("• F1-score = moyenne harmonique de la précision et du rappel — un compromis unique qui pénalise fortement un déséquilibre entre les deux.", style='List Bullet')
doc.add_paragraph("• mAP (mean Average Precision) = moyenne de la précision calculée à différents seuils de rappel, pour toutes les classes — la métrique standard en détection d'objets. mAP@0.5 signifie qu'une détection est considérée correcte si le chevauchement (IoU) avec la vraie boîte dépasse 50%.", style='List Bullet')

add_qa("Dans le contexte de SI-ENV, pourquoi le rappel est-il plus critique que la précision pour la détection de déchets ?",
    "Un faux négatif (déchet réel non détecté) signifie qu'une pollution passe inaperçue sur le "
    "terrain — conséquence directe sur la fiabilité du suivi environnemental. Un faux positif "
    "(fausse alerte) est gênant mais sans conséquence grave : l'agent peut simplement l'ignorer "
    "après vérification visuelle. On préfère donc, à qualité égale, un modèle qui rate moins de vrais "
    "cas (meilleur rappel), quitte à avoir quelques fausses alertes.", "Difficile")
add_qa("Qu'est-ce que l'IoU (Intersection over Union) ?",
    "C'est le rapport entre la surface d'intersection et la surface d'union de deux boîtes "
    "englobantes (la boîte prédite et la boîte réelle annotée). Un IoU de 1 signifie une "
    "correspondance parfaite, un IoU de 0 signifie aucun chevauchement. Le seuil de 0,5 dans mAP@0.5 "
    "signifie qu'on tolère une imprécision de position tant que le chevauchement reste supérieur à "
    "50%.", "Intermédiaire")

add_sub("5.4 ONNX et le déploiement mobile")
doc.add_paragraph(
    "ONNX (Open Neural Network Exchange) est un format standard et interopérable pour représenter "
    "des modèles de deep learning, indépendamment du framework d'entraînement (ici PyTorch/"
    "Ultralytics). Une fois exporté au format .onnx, le modèle est exécuté sur le téléphone via "
    "ONNX Runtime, sans dépendre d'un serveur — le diagnostic fonctionne donc hors ligne."
)
add_qa("Pourquoi exporter en ONNX plutôt que d'exécuter le modèle PyTorch original directement sur le téléphone ?",
    "PyTorch est une bibliothèque volumineuse pensée pour l'entraînement, pas optimisée pour "
    "l'inférence embarquée sur mobile. ONNX Runtime est une bibliothèque légère et optimisée "
    "spécifiquement pour l'inférence (exécution) sur des appareils aux ressources limitées, "
    "indépendante du framework d'origine.", "Intermédiaire")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("⚠ Question quasi certaine du jury — voir Module 8 pour la réponse complète et honnête : ")
r.bold = True; r.font.color.rgb = RGBColor(0xB9,0x1C,0x1C)
r2 = p.add_run("« Le modèle réellement embarqué dans votre application correspond-il à celui dont vous présentez les résultats au chapitre 8 ? »")
r2.italic = True

# ==================================================================
# MODULE 6 — ANALYSE SATELLITAIRE / TELEDETECTION
# ==================================================================
add_module(6, "Analyse satellitaire et télédétection",
    "Le module le moins « informatique classique » — le jury peut tester votre culture générale en "
    "sciences de l'environnement autant que la partie technique (API GEE).")

add_sub("6.1 Notions de base de télédétection")
doc.add_paragraph(
    "La télédétection consiste à observer la surface terrestre à distance, depuis un satellite, en "
    "mesurant le rayonnement électromagnétique réfléchi ou émis par les surfaces, dans différentes "
    "bandes (longueurs d'onde) du spectre : visible, infrarouge, etc. Chaque type de surface "
    "(végétation, eau, sol nu, zone urbaine) a une signature spectrale caractéristique, ce qui "
    "permet de la distinguer automatiquement par le calcul."
)
add_qa("Qu'est-ce qu'un indice spectral, par exemple le NDVI ?",
    "Un indice spectral est une combinaison mathématique simple de plusieurs bandes spectrales, "
    "conçue pour faire ressortir une caractéristique précise. Le NDVI (Normalized Difference "
    "Vegetation Index) se calcule ainsi : NDVI = (PIR − Rouge) / (PIR + Rouge), où PIR est la "
    "réflectance dans le proche infrarouge et Rouge la réflectance dans le rouge visible. La "
    "végétation en bonne santé réfléchit fortement le proche infrarouge et absorbe le rouge "
    "(photosynthèse), donnant un NDVI élevé (proche de 1) ; l'eau ou le sol nu donnent un NDVI proche "
    "de 0 ou négatif.", "Intermédiaire")
add_qa("Que mesure le NDWI, et pourquoi l'avoir utilisé pour les eaux stagnantes ?",
    "Le NDWI (Normalized Difference Water Index) met en évidence la présence d'eau en exploitant la "
    "forte absorption de l'eau dans le proche infrarouge par rapport au vert. Dans SI-ENV, il a "
    "finalement été écarté au profit de l'observation terrain directe pour les eaux stagnantes, car "
    "Sentinel-2 dépend trop de la couverture nuageuse et de sa fréquence de revisite (5 jours) pour "
    "un phénomène qui peut apparaître et disparaître en quelques heures après une pluie.",
    "Difficile")
add_qa("Pourquoi utiliser le NO2 comme indicateur de qualité de l'air plutôt que les particules fines (PM2.5) directement ?",
    "Le NO2 troposphérique est mesuré directement et de façon fiable par Sentinel-5P/TROPOMI, un "
    "instrument dédié à la chimie atmosphérique avec une bonne résolution temporelle (revisite "
    "quotidienne). Les PM2.5 ne sont pas mesurées directement par satellite avec la même fiabilité ; "
    "elles sont généralement estimées indirectement (AOD — épaisseur optique des aérosols), ce qui "
    "introduit plus d'incertitude. Le NO2 constitue par ailleurs un traceur pertinent des activités "
    "de chantier (engins, groupes électrogènes).", "Difficile")

add_sub("6.2 Google Earth Engine (GEE)")
doc.add_paragraph(
    "Google Earth Engine est une plateforme cloud de traitement de données satellitaires à grande "
    "échelle. Plutôt que de télécharger et traiter localement des images satellite volumineuses, on "
    "envoie le calcul à GEE via son API Python (earthengine-api), qui exécute le traitement "
    "côté serveur et ne renvoie que le résultat (une valeur, une image réduite)."
)
add_qa("Pourquoi Google Earth Engine plutôt que de télécharger les images Sentinel directement (via Copernicus par exemple) ?",
    "GEE évite de gérer une infrastructure de stockage et de calcul lourde : les images satellite "
    "brutes représentent des téraoctets de données. GEE héberge déjà l'intégralité des archives "
    "Sentinel/Landsat et exécute les calculs côté cloud, ne renvoyant que le résultat final. Cela "
    "s'intègre nativement dans un backend Python (FastAPI) via son SDK.", "Basique")

# ==================================================================
# MODULE 7 — DEPLOIEMENT, DOCKER, TESTS
# ==================================================================
add_module(7, "Déploiement, conteneurisation et tests",
    "Module généralement plus court à l'oral, mais des questions de culture DevOps de base "
    "tombent souvent pour vérifier une compréhension au-delà du copier-coller de commandes.")

add_sub("7.1 Docker et conteneurisation")
doc.add_paragraph(
    "Un conteneur est un environnement d'exécution isolé et léger, empaquetant une application "
    "avec toutes ses dépendances (bibliothèques, configuration), garantissant qu'elle s'exécute "
    "identiquement sur n'importe quelle machine. SI-ENV utilise trois conteneurs orchestrés par "
    "Docker Compose : la base PostGIS, le backend FastAPI, et nginx (reverse proxy)."
)
add_qa("Quelle est la différence entre un conteneur et une machine virtuelle (VM) ?",
    "Une VM virtualise un système d'exploitation complet (avec son propre noyau), ce qui est lourd "
    "en ressources. Un conteneur partage le noyau du système hôte et n'isole que l'espace "
    "applicatif (processus, fichiers, réseau) — beaucoup plus léger et rapide à démarrer, mais avec "
    "une isolation un peu moins forte qu'une VM.", "Basique")
add_qa("Pourquoi trois conteneurs séparés plutôt qu'une seule machine avec tout installé dessus ?",
    "Séparer la base de données, le backend et le reverse proxy permet de les faire évoluer, "
    "redémarrer ou mettre à l'échelle indépendamment, et rapproche l'environnement de développement "
    "de production (chacun a ses propres dépendances isolées, sans conflit de versions). C'est "
    "aussi la pratique standard pour une architecture destinée à évoluer vers plusieurs instances "
    "du backend derrière le même reverse proxy si la charge augmente.", "Intermédiaire")
add_qa("Quel est le rôle exact de nginx dans cette architecture ?",
    "nginx agit comme reverse proxy : il reçoit toutes les requêtes HTTPS entrantes, gère le "
    "certificat TLS (chiffrement), et redirige chaque requête vers le bon service interne (API "
    "backend ou fichiers statiques du dashboard). Cela évite d'exposer directement le serveur "
    "applicatif FastAPI sur Internet.", "Intermédiaire")

add_sub("7.2 Les types de tests")
doc.add_paragraph(
    "Trois niveaux de tests sont distingués classiquement : les tests unitaires vérifient une "
    "fonction ou un composant isolé (ex. le calcul de l'indice de risque pluie) ; les tests "
    "d'intégration vérifient que plusieurs composants fonctionnent correctement ensemble (ex. la "
    "chaîne mobile → backend → base de données) ; les tests fonctionnels vérifient qu'un scénario "
    "utilisateur complet aboutit au résultat attendu (ex. « un agent crée un signalement hors ligne, "
    "le synchronise, il apparaît sur le dashboard »)."
)
add_qa("Comment vos tests ont-ils été exécutés concrètement ?",
    "Via pytest et TestClient de FastAPI, qui simule des requêtes HTTP sans lancer un vrai serveur "
    "réseau, avec une base de données de test isolée. 32 tests automatisés couvrent 12 scénarios "
    "fonctionnels (authentification, création de signalement, synchronisation, RBAC, etc.), "
    "exécutés à la fois en environnement de développement et sur le déploiement Docker complet.",
    "Basique")

# ==================================================================
# MODULE 8 — QUESTIONS PIEGES SUR LE MEMOIRE
# ==================================================================
add_module(8, "Questions pièges sur votre mémoire — soyez prêt",
    "Ce module rassemble les points que le jury est le plus susceptible de creuser, parce que ce "
    "sont des points de vigilance réels, identifiés lors d'une relecture rigoureuse du mémoire et "
    "du code. La bonne stratégie n'est pas de les cacher, mais d'y répondre avec assurance et "
    "honnêteté — un jury respecte davantage un candidat qui assume une limite qu'un candidat pris "
    "en défaut à essayer de la masquer.")

add_qa("Le modèle de détection réellement embarqué dans votre application mobile est-il celui dont vous présentez les résultats (mAP 0,798) au chapitre 8 ?",
    "Non, et c'est assumé explicitement dans la discussion du chapitre 10 : le modèle validé au "
    "chapitre 8 a été entraîné sur GPU (Google Colab) sur le jeu de données Recycle Trash (6 "
    "classes). Le modèle actuellement embarqué dans l'application est une version d'entraînement "
    "rapide, mono-classe, utilisée pour valider le pipeline technique (préprocessing, inférence "
    "ONNX, NMS) pendant le développement. Le ré-export du modèle 6 classes vers l'application "
    "mobile — qui nécessite d'adapter le décodage de sortie du modèle (6 scores de classe au lieu "
    "d'une simple confiance) — reste une étape d'intégration finale identifiée comme perspective "
    "immédiate.", "Piège")
add_qa("Vos résultats de classification (F1 = 0,86) sont-ils fiables ? Le jury peut-il vous demander de les reproduire en direct ?",
    "Les chiffres du chapitre 8 proviennent d'un entraînement réel sur GPU, cohérent avec les "
    "courbes et matrices de confusion présentées en annexe. Si on vous demande de le reproduire en "
    "direct sur votre machine locale (sans GPU), attendez-vous à des chiffres plus modestes sur un "
    "entraînement rapide (c'est un effet attendu de la taille du dataset et du nombre d'époques, pas "
    "une contradiction) — expliquez calmement la différence d'environnement d'entraînement plutôt "
    "que de vous laisser déstabiliser par l'écart.", "Piège")
add_qa("Votre système est-il vraiment « temps réel » comme vous l'écrivez à certains endroits ?",
    "Il faut être précis sur ce terme : le tableau de bord web se rafraîchit automatiquement toutes "
    "les 10 à 15 secondes (polling HTTP), ce qui est un quasi temps réel perceptible par "
    "l'utilisateur, mais ce n'est pas un push serveur→client instantané (comme un WebSocket). Côté "
    "mobile, la synchronisation est déclenchée manuellement par l'agent, pas automatique. Le mémoire "
    "a été corrigé pour employer « quasi temps réel » plutôt que « temps réel » sans nuance, "
    "précisément pour rester rigoureux sur ce point.", "Piège")
add_qa("Pourquoi les performances (API, synchronisation) ont-elles été mesurées en local et non sur le VPS de production visé ?",
    "Parce que le VPS de production (2 vCPU / 4 Go RAM) n'a pas encore été provisionné à ce stade du "
    "stage — le déploiement réel est une étape post-soutenance. Les mesures ont donc été prises sur "
    "un déploiement Docker local strictement identique en configuration (mêmes 3 conteneurs, mêmes "
    "images), ce qui valide l'architecture et le comportement applicatif, mais les temps réseau/API "
    "pourront différer légèrement une fois sur l'infrastructure cible — ce qui est indiqué "
    "explicitement dans le mémoire plutôt que dissimulé.", "Piège")
add_qa("Le déclenchement automatique d'alerte (seuil dépassé → alerte → notification) fonctionne-t-il réellement, ou juste l'affichage des alertes existantes ?",
    "Il fonctionne réellement de bout en bout : lorsqu'un indice satellite dépasse un seuil "
    "configuré (ou, à défaut de seuil personnalisé, un statut « mauvais » calculé par les règles "
    "métier), une alerte est automatiquement créée en base et un e-mail est envoyé aux responsables "
    "concernés (Spécialiste Suivi Environnemental, Expert HSE), avec une déduplication sur 24h pour "
    "éviter le spam. La notification « push » mobile, en revanche, n'existe pas — seul l'e-mail est "
    "implémenté ; c'est une nuance à connaître si la question porte précisément sur le canal de "
    "notification.", "Difficile")
add_qa("Pourquoi le dataset de détection (Recycle Trash) et non des photos réelles des chantiers du PTUA ?",
    "Faute d'autorisation de photographier extensivement les chantiers durant le stage (contrainte "
    "d'accès et de temps), un jeu de données public de déchets recyclables (Recycle Trash, "
    "Roboflow/GitHub) a été utilisé, dont les contextes visuels (extérieur, surfaces mixtes) sont "
    "proches des conditions de chantier. C'est une limite assumée : le modèle n'a pas encore été "
    "validé sur des photographies réelles du PTUA, ce qui est explicitement indiqué dans la "
    "discussion comme perspective de validation terrain.", "Intermédiaire")
add_qa("Si le jury vous met devant une incohérence ou une imprécision que vous n'aviez pas anticipée, que faire ?",
    "Ne jamais improviser une justification qui sonne faux. Reconnaître calmement le point "
    "(« vous avez raison, c'est une limite que je peux préciser »), replacer la réponse dans le "
    "contexte du travail réalisé (stage de 3 mois, contraintes de temps et d'accès terrain), et si "
    "possible, indiquer comment vous le corrigeriez avec plus de temps. Un jury valorise la "
    "capacité d'auto-critique bien plus qu'une défense rigide d'un point intenable.", "Piège")

# ==================================================================
# MODULE 9 — CULTURE GENERALE INFORMATIQUE
# ==================================================================
add_module(9, "Culture générale en informatique",
    "Questions transversales que le jury pose parfois pour vérifier des fondamentaux, "
    "indépendamment du projet — à réviser rapidement si le temps le permet.")

add_qa("Qu'est-ce qu'une architecture 3-tiers (ou client-serveur à 3 niveaux) ?",
    "Une organisation en trois couches distinctes : la couche présentation (interfaces mobile et "
    "web), la couche métier/applicative (le backend, qui contient la logique) et la couche données "
    "(la base de données). Chaque couche ne communique qu'avec la couche adjacente, ce qui facilite "
    "la maintenance et permet de faire évoluer une couche sans impacter les autres.", "Basique")
add_qa("Quelle est la différence entre HTTP et HTTPS ?",
    "HTTPS est HTTP encapsulé dans une couche de chiffrement (TLS/SSL). Sans HTTPS, toutes les "
    "données (y compris les identifiants et jetons JWT) circulent en clair sur le réseau et "
    "peuvent être interceptées. HTTPS garantit la confidentialité (chiffrement), l'intégrité "
    "(détection de toute modification en transit) et l'authenticité du serveur (certificat).",
    "Basique")
add_qa("Qu'est-ce qu'une clé étrangère, et à quoi sert-elle ?",
    "Une clé étrangère est une colonne (ou un ensemble de colonnes) d'une table qui référence la "
    "clé primaire d'une autre table, matérialisant une relation entre deux entités et garantissant "
    "l'intégrité référentielle : impossible d'insérer un signalement pointant vers un chantier "
    "inexistant, par exemple.", "Basique")
add_qa("Quelle est la différence entre un test unitaire et un test d'intégration ?",
    "Un test unitaire isole une seule fonction ou un seul composant (avec des dépendances "
    "simulées/mockées). Un test d'intégration vérifie que plusieurs composants réels fonctionnent "
    "correctement ensemble (ex. le backend et une vraie base de données).", "Intermédiaire")
add_qa("Qu'est-ce que l'idempotence, et où intervient-elle dans votre système ?",
    "Une opération est idempotente si l'exécuter plusieurs fois produit le même résultat "
    "qu'une seule exécution. C'est le cas de la synchronisation d'un signalement : si le réseau "
    "coupe après l'envoi mais avant la réception de la confirmation, l'application peut renvoyer le "
    "même signalement (même UUID) sans risque de le dupliquer en base, car le serveur détecte "
    "l'UUID déjà existant et renvoie l'enregistrement existant au lieu d'en créer un nouveau.",
    "Difficile")
add_qa("Qu'est-ce qu'une injection SQL, et comment votre système s'en protège-t-il ?",
    "Une injection SQL consiste à insérer du code SQL malveillant dans une donnée utilisateur non "
    "filtrée, pour manipuler la requête exécutée par le serveur. SQLAlchemy (l'ORM utilisé par "
    "FastAPI) protège nativement contre ce risque en utilisant systématiquement des requêtes "
    "paramétrées, où les valeurs utilisateur ne sont jamais concaténées directement dans le texte "
    "SQL.", "Intermédiaire")

# ------------------------------------------------------------------
# Derniere page : conseils de posture
# ------------------------------------------------------------------
doc.add_page_break()
doc.add_heading("Derniers conseils avant l'oral", level=1)
doc.add_paragraph(
    "• Reformulez toujours la question avant de répondre si elle est ambiguë — cela vous laisse le "
    "temps de structurer votre réponse et montre que vous écoutez vraiment.\n"
    "• Une réponse courte et exacte vaut mieux qu'une réponse longue et floue. Si vous ne savez pas, "
    "dites-le simplement et proposez une piste de raisonnement plutôt que d'inventer.\n"
    "• Reliez toujours vos réponses techniques à un choix concret de votre projet (« dans SI-ENV, "
    "cela se traduit par... ») — le jury évalue votre projet, pas un cours théorique récité.\n"
    "• Sur les points faibles (module 8), n'attendez pas d'être acculé : les mentionner "
    "spontanément dans votre présentation orale, avec les perspectives associées, désamorce la "
    "question avant qu'elle ne devienne un piège.\n"
    "• Respirez, parlez lentement. Un jury \"bouche bée\" se construit par la clarté et la maîtrise "
    "calme, pas par la vitesse de débit."
)

doc.save(OUT)
print(f"\n=== GUIDE GENERE : {OUT} ===")

