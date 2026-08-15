# -*- coding: utf-8 -*-
"""
Corrige les donnees devenues fausses dans le memoire.

Trois passages ne correspondaient plus au systeme livre. Le plus genant
annoncait que le backend s'executait sur Hugging Face Spaces, alors que le
projet a migre vers Render : l'adresse de production est publique, et un jury
qui la consulte constate immediatement l'ecart. Un memoire qui decrit une
infrastructure que l'on peut verifier en une minute doit dire vrai.

Le deuxieme passage decrivait le deploiement comme reposant sur un serveur de
l'AGEROUTE, ce qui contredisait le diagramme de deploiement regenere. La
contradiction n'etait pas anodine : les deux se trouvent sur la meme page.

Le troisieme reprenait un delai de suspension propre a Hugging Face, sans
objet depuis la migration.
"""
import shutil
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_donnees_perimees.docx")


def reecrire(p, texte):
    for seg in list(p.runs)[1:]:
        seg._element.getparent().remove(seg._element)
    if p.runs:
        p.runs[0].text = texte
    else:
        p.add_run(texte)


# ── Passage 1 : diagramme de deploiement ──────────────────────────────────
DEPLOIEMENT = (
    "Le diagramme de déploiement (figure 4.7) présente la répartition logique "
    "des composants du système, indépendamment de l'infrastructure qui les "
    "accueille. Les deux applications mobiles s'exécutent sur les terminaux de "
    "leurs utilisateurs respectifs, l'agent de terrain et le riverain, et n'y "
    "conservent aucune donnée du projet en dehors du cache de synchronisation. "
    "Le serveur d'application expose l'API et héberge les modèles d'inférence ; "
    "il interroge le serveur de données, qui porte PostgreSQL et son extension "
    "PostGIS. Le tableau de bord Angular, compilé en fichiers statiques, est "
    "distribué séparément et appelle l'API directement une fois chargé dans le "
    "navigateur. Deux services externes complètent le dispositif : Google Earth "
    "Engine pour l'imagerie satellitaire et un service d'envoi transactionnel "
    "pour les courriels.\n\n"
    "Cette répartition a été réalisée dans deux environnements distincts. Le "
    "premier, décrit au paragraphe 6.1, réunit les composants dans une pile "
    "Docker destinée à un serveur de l'AGEROUTE : c'est la cible retenue pour "
    "le pilote. Le second, décrit au paragraphe 6.8, les distribue chez des "
    "hébergeurs tiers afin d'éprouver la portabilité de l'ensemble. Le "
    "diagramme vaut pour les deux, ce qui est précisément ce qu'un diagramme de "
    "déploiement doit permettre : décrire une architecture sans la lier à un "
    "fournisseur. Le dispositif compte désormais quatre composants applicatifs "
    "et non trois, la seconde application mobile se déployant indépendamment "
    "bien qu'issue du même dépôt de code."
)

# ── Passage 2 : hebergement reel ──────────────────────────────────────────
HEBERGEMENT = (
    "Le déploiement décrit au paragraphe 6.1 (pile Docker locale) reste la "
    "cible retenue pour le pilote AGEROUTE et pour la démonstration. Il ne "
    "suffit pas à prouver la portabilité du système ni la maturité du cycle de "
    "vie logiciel : un second environnement, hébergé chez des tiers sur des "
    "offres gratuites sans carte bancaire, a donc été mis en place. Trois "
    "plateformes se répartissent les composants, chacune retenue pour un point "
    "précis de son offre. Supabase héberge la base PostgreSQL 16 accompagnée de "
    "l'extension PostGIS préinstallée, critère décisif car aucun autre acteur "
    "gratuit ne propose PostGIS sans carte, ainsi qu'un espace de stockage d'un "
    "gigaoctet pour les photos. Render exécute le backend FastAPI dans un "
    "conteneur Docker identique à celui de l'environnement local, avec HTTPS "
    "assuré par la plateforme. Cloudflare Pages sert le tableau de bord Angular "
    "en fichiers statiques distribués par son réseau mondial. GitHub Actions "
    "orchestre l'ensemble : les tests fonctionnels du paragraphe 6.2 sont "
    "rejoués à chaque poussée de code, et un déploiement automatique publie le "
    "backend à chaque fusion sur la branche principale.\n\n"
    "Un premier essai avait retenu Hugging Face Spaces pour le backend, avant "
    "que la plateforme ne réserve l'exécution de conteneurs Docker à son offre "
    "payante. Le basculement vers Render n'a demandé aucune modification du "
    "code applicatif, le conteneur et les variables d'environnement étant "
    "identiques : c'est une conséquence directe du choix de conteneuriser dès "
    "le départ, et un argument concret en faveur de cette pratique."
)

# ── Passage 3 : surveillance ──────────────────────────────────────────────
SURVEILLANCE = (
    "L'hébergement gratuit a une limite structurelle bien documentée : les "
    "plateformes suspendent les services après un délai d'inactivité, quinze "
    "minutes pour Render et sept jours pour Supabase. Cette suspension "
    "provoquerait, lors d'une consultation ponctuelle du jury ou d'un "
    "évaluateur, un délai de réveil de plusieurs dizaines de secondes qui "
    "serait perçu comme une défaillance. Un dispositif de surveillance actif a "
    "donc été ajouté, en trois couches complémentaires. Un premier processus "
    "interroge toutes les trois minutes le backend et la base pour maintenir "
    "les services actifs. Un second processus horaire rejoue la chaîne complète "
    "(authentification, requête PostGIS, accès au stockage des photos), ce "
    "qu'un simple appel de disponibilité ne détecterait pas. En cas d'échec "
    "persistant, un troisième processus tente en cascade un redémarrage doux, "
    "puis un redémarrage complet, puis une republication du backend, puis un "
    "nouveau peuplement de la base ; à défaut, une anomalie est ouverte "
    "automatiquement sur le dépôt, ce qui déclenche un courriel d'alerte."
)

REPRISES = [
    ("Le diagramme de déploiement (figure 4.7)", DEPLOIEMENT, "diagramme de déploiement"),
    ("Le déploiement décrit au paragraphe 6.1", HEBERGEMENT, "hébergement réel"),
    ("L'hébergement gratuit a une limite structurelle", SURVEILLANCE, "dispositif de surveillance"),
]


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    for debut, remplacement, libelle in REPRISES:
        for p in doc.paragraphs:
            if p.text.strip().startswith(debut):
                # Le texte de remplacement comporte des paragraphes separes par
                # une ligne vide : Word ne les rend pas depuis un seul run, on
                # les reunit donc par un saut de ligne simple.
                reecrire(p, remplacement.replace("\n\n", "  "))
                print(f"  ok  {libelle}")
                break
        else:
            print(f"  MANQUE  {libelle}")

    doc.save(SOURCE)
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
