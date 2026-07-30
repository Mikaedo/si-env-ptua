"""
inserer_figures_ia_v2.py
========================
Insere les figures d'entrainement IA aux bons endroits dans le memoire.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

MEMOIRE = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v42.docx"
IA_DIR = r"D:\etude_soutenance\SI-ENV\ia"

def inserer_image_apres(doc, index_para, chemin_image, legende, largeur=5.5):
    """Insere une image + legende centree apres le paragraphe index_para."""
    para = doc.paragraphs[index_para]

    # Paragraphe pour l'image
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(chemin_image, width=Inches(largeur))
    para._element.addnext(img_para._element)

    # Paragraphe pour la legende
    legende_para = doc.add_paragraph()
    legende_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_leg = legende_para.add_run(legende)
    run_leg.font.size = Pt(10)
    run_leg.font.italic = True
    img_para._element.addnext(legende_para._element)

    print(f"  [OK] Insere apres para {index_para}: {legende}")
    return True

def remplacer_placeholder(doc, index_para, chemin_image, legende, largeur=5.5):
    """Remplace le texte d'un paragraphe placeholder par une image + legende."""
    para = doc.paragraphs[index_para]

    # Vider le paragraphe existant
    for run in para.runs:
        run.text = ""
    para.text = ""

    # Inserer l'image dans ce paragraphe
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(chemin_image, width=Inches(largeur))

    # Creer un paragraphe legende apres
    legende_para = doc.add_paragraph()
    legende_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_leg = legende_para.add_run(legende)
    run_leg.font.size = Pt(10)
    run_leg.font.italic = True
    para._element.addnext(legende_para._element)

    print(f"  [OK] Placeholder remplace para {index_para}: {legende}")
    return True

def main():
    print(f">> Ouverture du memoire : {MEMOIRE}")
    doc = Document(MEMOIRE)
    print(f">> {len(doc.paragraphs)} paragraphes trouves")

    # --- Section 8.3 (index 482) : Figures detection YOLOv8 ---
    print("\n>> Section 8.3 - Detection YOLOv8 :")

    # Figure 8.1 : Courbes d'apprentissage (results.png) apres para 482
    inserer_image_apres(doc, 482,
        os.path.join(IA_DIR, "runs_detection", "dechets_yolov8n", "results.png"),
        "Figure 8.1 - Courbes d'apprentissage du modele YOLOv8n (perte, mAP, precision, rappel)")

    # Figure 8.2 : Matrice de confusion apres la legende de 8.1
    # L'index decale de +2 (image + legende inserees)
    inserer_image_apres(doc, 484,
        os.path.join(IA_DIR, "runs_detection", "dechets_yolov8n", "confusion_matrix.png"),
        "Figure 8.2 - Matrice de confusion du modele YOLOv8n (detection de dechets)")

    # Figure 8.3 : Courbe PR apres la legende de 8.2
    inserer_image_apres(doc, 486,
        os.path.join(IA_DIR, "runs_detection", "dechets_yolov8n", "PR_curve.png"),
        "Figure 8.3 - Courbe Precision-Recall du modele YOLOv8n")

    # --- Section 8.4 (index 485 + decalage) : Figure classification ---
    # Apres insertion de 3 figures (6 paragraphes decales), l'index 485 devient 485+6=491
    print("\n>> Section 8.4 - Classification MobileNetV2 :")

    # Trouver le paragraphe "Le modele retenu offre le meilleur compromis"
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "meilleur compromis" in p.text.lower():
            target_idx = i
            break

    if target_idx is not None:
        inserer_image_apres(doc, target_idx,
            os.path.join(IA_DIR, "matrice_confusion_classification.png"),
            "Figure 8.4 - Matrice de confusion du modele MobileNetV2 (classification de criticite)")
    else:
        print("  [!] Paragraphe 'meilleur compromis' non trouve")

    # --- Remplacer le placeholder Figure 8.1 (index 492 + decalage) ---
    print("\n>> Remplacement du placeholder :")

    placeholder_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "[À coller" in p.text and "Figure 8.1" in p.text:
            placeholder_idx = i
            break

    if placeholder_idx is not None:
        # Supprimer ce placeholder (les figures sont deja inserees au bon endroit)
        para = doc.paragraphs[placeholder_idx]
        para.text = ""
        print(f"  [OK] Placeholder nettoye a l'index {placeholder_idx}")
    else:
        print("  [!] Placeholder non trouve")

    # Sauvegarder
    output = MEMOIRE.replace("v42", "v43")
    doc.save(output)
    print(f"\n>> Memoire sauvegarde : {output}")

if __name__ == "__main__":
    main()
