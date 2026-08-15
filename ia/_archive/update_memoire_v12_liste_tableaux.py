# -*- coding: utf-8 -*-
"""
Restaure les 4 entrees disparues de la LISTE DES TABLEAUX.

Les tableaux 9.1, 10.1, 10.2 et 10.3 existent bien dans le corps du memoire,
correctement titres et cites, mais leurs entrees ont ete perdues de la liste
entre les versions v5 et v6. La liste s'arretait donc au chapitre 8 alors que
le document va jusqu'au chapitre 10.

Les libelles sont reprises telles quelles depuis la version v5, qui les
contenait encore, afin de ne rien reformuler.

Sortie : MEMOIRE_N'GUESSAN_v12.docx
"""
import copy
import os
import re
import shutil

from docx import Document
from docx.text.paragraph import Paragraph

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v11.docx")
REFERENCE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v5.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v12.docx")

A_RESTAURER = ('9.1', '10.1', '10.2', '10.3')


def entrees_liste(chemin):
    """Renvoie (index du titre, [(numero, libelle, index)]) de la liste."""
    doc = Document(chemin)
    debut = None
    for i, p in enumerate(doc.paragraphs[:250]):
        if p.text.strip().upper().startswith('LISTE DES TABLEAUX'):
            debut = i
            break
    if debut is None:
        raise SystemExit("Section LISTE DES TABLEAUX introuvable dans %s" % chemin)
    trouvees = []
    for i in range(debut + 1, min(debut + 80, len(doc.paragraphs))):
        t = doc.paragraphs[i].text.strip()
        m = re.match(r'^Tableau\s+(\d+\.\d+)', t)
        if m:
            trouvees.append((m.group(1), t, i))
        elif t and trouvees:
            break
    return doc, debut, trouvees


def main():
    # Libelles d'origine, repris de la version qui les contenait encore
    _, _, ref = entrees_liste(REFERENCE)
    libelles = {num: texte for num, texte, _ in ref}
    manquants = [n for n in A_RESTAURER if n not in libelles]
    if manquants:
        raise SystemExit("Libelles introuvables dans v5 : %s" % manquants)

    shutil.copy2(SOURCE, SORTIE)
    doc, debut, presentes = entrees_liste(SORTIE)
    nums_presents = [n for n, _, _ in presentes]
    print("  liste des tableaux : p%d | %d entree(s) presente(s)"
          % (debut, len(presentes)))

    a_ajouter = [n for n in A_RESTAURER if n not in nums_presents]
    if not a_ajouter:
        print("  rien a restaurer, la liste est deja complete")
        return

    # On insere apres la derniere entree existante, en ordre inverse pour que
    # l'ordre final suive la numerotation des chapitres.
    dernier_index = presentes[-1][2]
    modele = doc.paragraphs[dernier_index]
    print("  derniere entree : %s (p%d)"
          % (modele.text.strip()[:44], dernier_index))

    for num in reversed(a_ajouter):
        element = copy.deepcopy(modele._element)
        for enfant in list(element):
            if not enfant.tag.endswith('}pPr'):
                element.remove(enfant)
        modele._element.addnext(element)
        par = Paragraph(element, modele._parent)
        # Le formatage est herite du modele : meme style, meme taille.
        par.add_run(libelles[num])
        print("    + %s" % libelles[num][:62])

    doc.save(SORTIE)

    # Verification
    _, _, apres = entrees_liste(SORTIE)
    verif = Document(SORTIE)
    txt = "\n".join(p.text for p in verif.paragraphs)
    reels = sorted(set(re.findall(r'^Tableau\s+(\d+\.\d+)\s*[:.]', txt, re.M)),
                   key=lambda x: (int(x.split('.')[0]), int(x.split('.')[1])))
    listes = [n for n, _, _ in apres]
    absents = [n for n in reels if n not in listes]
    print("\n  tableaux dans le document : %d" % len(reels))
    print("  tableaux listes           : %d" % len(listes))
    print("  absents de la liste       : %s" % (absents or "aucun"))
    print("\nEnregistre : %s" % SORTIE)
    print("Paragraphes : %d | Tableaux : %d"
          % (len(verif.paragraphs), len(verif.tables)))


if __name__ == "__main__":
    main()
