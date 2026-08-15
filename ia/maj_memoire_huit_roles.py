# -*- coding: utf-8 -*-
"""
Met le memoire en accord avec le systeme livre.

Le document decrivait cinq profils utilisateurs, une seule application mobile
et trente-deux tests. Le dispositif en compte desormais huit, deux et cent
dix-neuf. Un memoire dont les figures et les tableaux ne correspondent plus au
code expose son auteur a la question la plus embarrassante qui soit en
soutenance : celle qui porte sur un ecart que le jury constate avant lui.

Le script reprend donc les passages concernes plutot que de reecrire le
document. Chaque remplacement vise un texte precis et signale son echec, de
sorte qu'une formulation qui aurait change entre deux executions ne passe pas
inapercue.
"""
import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_huit_roles.docx")

journal = []


def remplacer(paragraphe, ancien, nouveau, etiquette):
    """Remplace un fragment en preservant la mise en forme du paragraphe.

    Word decoupe un paragraphe en segments de mise en forme, et un texte
    recherche chevauche souvent plusieurs d'entre eux. On reconstruit donc le
    paragraphe a partir de son premier segment, dont on herite le style.
    """
    if ancien not in paragraphe.text:
        journal.append(f"  MANQUE  {etiquette}")
        return False
    complet = paragraphe.text.replace(ancien, nouveau)
    for segment in list(paragraphe.runs)[1:]:
        segment._element.getparent().remove(segment._element)
    if paragraphe.runs:
        paragraphe.runs[0].text = complet
    else:
        paragraphe.add_run(complet)
    journal.append(f"  ok      {etiquette}")
    return True


def cellule(tableau, ligne, colonne, texte, gras=False, taille=9):
    """Ecrit une cellule en reprenant la typographie du tableau."""
    c = tableau.cell(ligne, colonne)
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run(texte)
    r.bold = gras
    r.font.size = Pt(taille)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    # ── 3.4 Profils d'utilisateurs ────────────────────────────────────────
    journal.append("3.4 Profils d'utilisateurs")
    for p in doc.paragraphs:
        if p.text.strip().startswith("Cinq r"):
            remplacer(
                p,
                "Cinq rôles distincts ont été retenus dans le SI-ENV, alignés sur la chaîne "
                "de responsabilité institutionnelle du PTUA.",
                "Huit rôles distincts ont été retenus dans le SI-ENV, alignés sur la chaîne "
                "de responsabilité institutionnelle du PTUA. Les cinq premiers correspondent "
                "aux intervenants opérationnels, ceux qui produisent ou traitent "
                "l'information. Les trois derniers traduisent la gouvernance réelle du "
                "programme : l'agence de tutelle et le bailleur consultent le dispositif "
                "sans jamais y écrire, et les riverains l'alimentent au titre du mécanisme "
                "de gestion des plaintes.",
                "phrase d'introduction des profils",
            )
            break
    for p in doc.paragraphs:
        if p.text.strip().startswith("Tableau 3.3"):
            remplacer(p, "Périmètre des cinq rôles utilisateurs.",
                      "Périmètre des huit rôles utilisateurs.", "légende du tableau 3.3")
            break

    # Tableau 3.3 : trois lignes a ajouter.
    t33 = doc.tables[9]
    nouveaux_profils = [
        ("Agence Nationale de l'Environnement",
         "Web (consultation)",
         "Autorité de tutelle environnementale. Vérifie la conformité "
         "réglementaire des chantiers et reçoit les rapports périodiques. "
         "Aucun droit d'écriture."),
        ("Banque Africaine de Développement",
         "Web (consultation)",
         "Bailleur du programme. Contrôle le respect des sauvegardes "
         "opérationnelles, volet social compris. Aucun droit d'écriture."),
        ("Riverain",
         "Mobile citoyen",
         "Habitant de la zone d'influence d'un chantier. Dépose des doléances "
         "qui alimentent le mécanisme de gestion des plaintes."),
    ]
    for nom, interface, perimetre in nouveaux_profils:
        ligne = t33.add_row()
        for i, valeur in enumerate((nom, interface, perimetre)):
            c = ligne.cells[i]
            c.text = ""
            r = c.paragraphs[0].add_run(valeur)
            r.font.size = Pt(9)
    journal.append(f"  ok      3 profils ajoutes au tableau 3.3")

    # ── 4.3.2 Identification des acteurs ──────────────────────────────────
    journal.append("4.3.2 Identification des acteurs")
    for p in doc.paragraphs:
        if "sept (7) acteurs ont" in p.text:
            remplacer(
                p,
                "sept (7) acteurs ont été identifiés : cinq acteurs principaux "
                "interagissant directement avec le système, et deux acteurs secondaires",
                "dix (10) acteurs ont été identifiés : huit acteurs principaux "
                "interagissant directement avec le système, et deux acteurs secondaires",
                "nombre d'acteurs",
            )
            break

    # ── 4.5 Matrice des habilitations ─────────────────────────────────────
    journal.append("4.5 Matrice des habilitations (RBAC)")
    rbac = doc.tables[11]

    # Trois colonnes s'ajoutent a la matrice. python-docx ne sait pas inserer
    # de colonne : on reconstruit le tableau ligne par ligne en repartant de
    # son contenu, ce qui preserve le style applique au tableau lui-meme.
    ancien = [[c.text.strip() for c in r.cells] for r in rbac.rows]
    colonnes_ajoutees = ["ANDE", "BAD", "Riverain"]
    #             ANDE      BAD       Riverain
    valeurs = {
        "Saisie sur terrain (Mobile)":          ("Non", "Non", "Doléances"),
        "Traitement d'un signalement":          ("Non", "Non", "Non"),
        "Réception et revue des alertes":       ("Lecture", "Lecture", "Non"),
        "Gestion Plaintes/P.A.R":               ("Non", "Lecture", "Dépôt"),
        "Lancement Satellitaire":               ("Lecture", "Lecture", "Non"),
        "Génération rapport PGES":              ("Lecture", "Lecture", "Non"),
        "Statistiques globales":                ("Lecture", "Lecture", "Non"),
    }

    for _ in colonnes_ajoutees:
        rbac.add_column(rbac.columns[-1].width)

    depart = len(ancien[0])
    for j, titre in enumerate(colonnes_ajoutees):
        cellule(rbac, 0, depart + j, titre, gras=True, taille=8)

    for i, ligne in enumerate(ancien[1:], start=1):
        intitule = ligne[0]
        trio = None
        for cle, v in valeurs.items():
            if cle.lower().startswith(intitule.lower()[:18]):
                trio = v
                break
        trio = trio or ("Non", "Non", "Non")
        for j, valeur in enumerate(trio):
            cellule(rbac, i, depart + j, valeur, taille=8)
    journal.append("  ok      3 colonnes ajoutees a la matrice RBAC")

    # Une precision s'impose sous la matrice : la lecture seule n'est pas une
    # convention d'affichage, elle est refusee par le serveur.
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Afin de matérialiser le cloisonnement"):
            remplacer(
                p,
                "le tableau 4.1 synthétise les habilitations accordées à chaque profil "
                "au travers du système de contrôle d'accès basé sur les rôles (RBAC).",
                "le tableau 4.1 synthétise les habilitations accordées à chaque profil "
                "au travers du système de contrôle d'accès basé sur les rôles (RBAC). "
                "La mention « Lecture » y désigne une consultation sans aucun droit "
                "d'écriture : toute requête de modification émise avec un jeton de "
                "l'agence de tutelle ou du bailleur est rejetée par le serveur avant "
                "même d'atteindre le traitement métier, et non simplement masquée dans "
                "l'interface. Cette distinction est essentielle, un rapport de conformité "
                "perdant toute valeur si celui qui le contrôle pouvait retoucher les "
                "données qu'il examine.",
                "precision sur la lecture seule",
            )
            break

    doc.save(SOURCE)
    print("\n".join(journal))
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
