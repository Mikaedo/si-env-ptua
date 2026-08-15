# -*- coding: utf-8 -*-
"""
Ajoute au GUIDE_PREPARATION_SOUTENANCE une section 11 sur le deploiement
cloud, la demarche cycle de vie logiciel et la strategie de repli.

Le guide couvrait bien le fond scientifique (chapitres 1 a 6, IA, securite
conceptuelle) mais rien sur ce qui a ete construit ensuite : Render,
Supabase, Cloudflare Pages, GitHub Actions, watchdog, self-heal, URL
vivantes, demo web live, demo mobile en cascade.
"""
import copy
import docx
from docx.text.paragraph import Paragraph

CHEMIN = r"C:\Users\DELL\Downloads\MEMOIRE\GUIDE_PREPARATION_SOUTENANCE_SI-ENV.docx"
d = docx.Document(CHEMIN)


def _dernier_h1():
    """Retrouve le dernier titre H1 du document pour cloner son style."""
    ref = None
    for p in d.paragraphs:
        if p.style and p.style.name == "Heading 1":
            ref = p
    return ref


def _dernier_h2():
    ref = None
    for p in d.paragraphs:
        if p.style and p.style.name == "Heading 2":
            ref = p
    return ref


def _dernier_p():
    """Un paragraphe de style Normal, pour cloner le style corps."""
    ref = None
    for p in d.paragraphs:
        style = p.style.name if p.style else ""
        if style in ("Normal", "Body Text", "Paragraphe standard"):
            ref = p
    return ref


modele_h1 = _dernier_h1()
modele_h2 = _dernier_h2()
modele_p = _dernier_p()

if not (modele_h1 and modele_p):
    raise SystemExit("Styles introuvables dans le guide")


def _ajouter(texte: str, modele: Paragraph):
    """Ajoute un nouveau paragraphe a la fin du document, style clone."""
    corps = d.element.body
    neuf = copy.deepcopy(modele._p)
    # Nettoie tous les runs sauf le premier, puis y met le texte.
    p = Paragraph(neuf, modele._parent)
    if p.runs:
        p.runs[0].text = texte
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(texte)
    corps.append(neuf)
    return p


def h1(t): return _ajouter(t, modele_h1)
def h2(t): return _ajouter(t, modele_h2 or modele_h1)
def para(t): return _ajouter(t, modele_p)


# ══════════════════════════════════════════════════════════════════════
# SECTION 11 : demonstration cloud et cycle de vie logiciel
# ══════════════════════════════════════════════════════════════════════

h1("11. Démonstration cloud et cycle de vie logiciel")

para(
    "Cette section complète la partie « démonstration » : outre le mode local "
    "(Docker + adb reverse sur câble USB), le système est en réalité déployé "
    "en production sur trois plateformes gratuites, accessibles depuis "
    "n'importe quel navigateur, dans le monde entier, 24 h sur 24. C'est le "
    "point qui distingue le mieux ce mémoire d'un projet étudiant classique "
    "et il faut savoir le mettre en avant."
)


h2("11.1 Les URL à retenir par cœur")

para(
    "Trois adresses définissent la version publique du SI-ENV. À chaque push "
    "sur la branche principale du dépôt GitHub, le backend et le tableau de "
    "bord sont automatiquement reconstruits et redéployés."
)

# Petit tableau texte, sans creer un vrai tableau Word (evite les soucis de style)
para("Tableau de bord (Angular) : https://si-env-ptua.pages.dev")
para("Backend et documentation Swagger (FastAPI) : https://si-env-ptua.onrender.com/docs")
para("Dépôt et intégration continue (GitHub Actions) : https://github.com/Mikaedo/si-env-ptua")

para(
    "Ces trois adresses sont vivantes en permanence : un service de "
    "surveillance appelé par GitHub Actions les interroge toutes les trois "
    "minutes pour empêcher la mise en veille automatique du plan gratuit."
)


h2("11.2 Déroulé de la démonstration en ligne (5 minutes)")

para(
    "La démonstration se fait de préférence sur le déploiement cloud, plus "
    "impressionnant qu'un simple Docker local et qui prouve la portabilité "
    "réelle du système. Ordre recommandé :"
)

para(
    "Un : ouvrez https://si-env-ptua.pages.dev dans un onglet préparé à "
    "l'avance. Connectez-vous avec le compte Spécialiste Environnement "
    "(spec.env@ageroute.ci / spec123). Le tableau de bord affiche les "
    "indicateurs consolidés en direct depuis la base."
)

para(
    "Deux : montrez la carte des six chantiers PTUA géolocalisés, les "
    "compteurs par statut, la répartition par type de nuisance. Insistez sur "
    "le fait que toutes ces valeurs viennent d'une base PostgreSQL avec "
    "extension spatiale PostGIS, hébergée chez Supabase."
)

para(
    "Trois : cliquez sur « Analyse satellitaire ». Les quatre indicateurs "
    "(NO2, NDVI, NDWI, risque pluie/relief) sont calculés en direct via "
    "Google Earth Engine sur les six chantiers du PTUA. Comptez vingt "
    "secondes pour le premier appel, puis instantané ensuite."
)

para(
    "Quatre : ouvrez la page « Rapports PGES », choisissez une période, "
    "générez un PDF. Le rapport agrège en quelques secondes ce qui prenait "
    "auparavant plusieurs jours au format papier."
)

para(
    "Cinq : ouvrez la documentation Swagger sur "
    "https://si-env-ptua.onrender.com/docs. Elle est générée automatiquement "
    "à partir du code Python, ce qui prouve la cohérence entre le code livré "
    "et la documentation."
)


h2("11.3 Démonstration mobile : le circuit déchets 100 % automatique")

para(
    "Sur le téléphone, ouvrez l'application SI-ENV (icône AGEROUTE bleue avec "
    "repère cartographique et feuille). Connectez-vous avec le même compte, "
    "puis appuyez sur « Nouveau signalement » puis « Déchets de chantier »."
)

para(
    "Un viseur vidéo s'ouvre : YOLOv8n analyse chaque image en temps réel et "
    "dessine un cadre vert autour des déchets détectés, avec l'étiquette de "
    "la classe (métal, plastique, papier, carton, verre, organique) et le "
    "pourcentage de confiance. Appuyez sur le bouton de capture."
)

para(
    "De retour à l'écran de saisie, les champs « type de nuisance » et "
    "« criticité » ont disparu : le type est fixé à « déchets de chantier », "
    "la criticité est déduite du nombre d'objets détectés selon la règle "
    "documentée au paragraphe 5.8 du mémoire. L'agent n'a plus qu'à valider "
    "et envoyer. C'est la traduction concrète de la subjectivité encadrée "
    "que le chapitre 2 identifiait comme deuxième lacune du système existant."
)


h2("11.4 Défendre la démarche « cycle de vie logiciel »")

para(
    "Le paragraphe 6.8 du mémoire décrit le dispositif de déploiement continu "
    "et de haute disponibilité. Si le jury vous interroge, la réponse tient "
    "en quatre points :"
)

para(
    "Un : intégration continue. À chaque proposition de modification sur "
    "GitHub, les trente-deux tests fonctionnels sont rejoués par GitHub "
    "Actions. Aucun code cassé ne peut atteindre la branche principale."
)

para(
    "Deux : livraison continue. Une fois la fusion sur la branche principale, "
    "Render reconstruit et redéploie le backend automatiquement, sans "
    "intervention humaine. Cloudflare fait de même pour le tableau de bord."
)

para(
    "Trois : surveillance continue. Un processus Watchdog interroge le "
    "backend et la base toutes les trois minutes. Un contrôle profond "
    "horaire rejoue la chaîne complète d'authentification et d'accès aux "
    "données pour détecter les défaillances qu'un simple ping laisserait "
    "passer."
)

para(
    "Quatre : réparation automatique. Si trois échecs consécutifs sont "
    "détectés, un troisième processus tente en cascade un redémarrage doux, "
    "puis un redémarrage complet, puis une republication du backend. À "
    "défaut, une issue GitHub est ouverte automatiquement, ce qui déclenche "
    "un courriel d'alerte."
)

para(
    "Cette combinaison intégration, livraison, surveillance et réparation "
    "couvre le cycle de vie logiciel complet tel que l'ISO 12207 le décrit, "
    "ce qui est rarement mis en œuvre dans un projet étudiant."
)


h2("11.5 Questions du jury sur le déploiement et leur réponse")

para(
    "« Combien coûte votre déploiement actuel ? » — Zéro franc CFA par mois. "
    "Trois plateformes retenues sur leur offre gratuite sans carte bancaire, "
    "avec les limites documentées au tableau 6.4 du mémoire. Un passage à un "
    "serveur privé virtuel Systalink pour une exploitation AGEROUTE réelle "
    "coûterait de l'ordre de dix mille francs CFA par mois, sans "
    "modification du code."
)

para(
    "« Que se passe-t-il si Render tombe pendant la démonstration ? » — Le "
    "Watchdog détecte l'incident en moins de trois minutes, le self-healer "
    "tente une réparation automatique en cascade. Si le service reste "
    "indisponible malgré tout, un courriel d'alerte est envoyé et la "
    "démonstration bascule sur le mode Docker local préparé sur mon poste. "
    "Aucun scénario ne me laisse sans démonstration."
)

para(
    "« Avez-vous un plan de reprise d'activité ? » — Les données sont "
    "sauvegardées quotidiennement par Supabase (politique par défaut du plan "
    "gratuit). Le code source complet est versionné sur GitHub avec "
    "historique. Un déploiement complet à partir de zéro sur une autre "
    "plateforme prend une trentaine de minutes."
)

para(
    "« Comment gérez-vous les mots de passe des utilisateurs ? » — Hachage "
    "bcrypt avec facteur de coût par défaut, jamais de stockage en clair. "
    "L'administrateur crée les comptes sans mot de passe : l'utilisateur "
    "définit le sien à la première connexion. La procédure de mot de passe "
    "oublié envoie un code à six chiffres par courriel, valable dix minutes."
)

para(
    "« Comment sont envoyés les courriels ? » — Trois types : mot de passe "
    "oublié, bienvenue à la création d'un compte, alerte automatique par "
    "franchissement de seuil satellite. L'envoi se fait via un service "
    "d'envoi transactionnel accessible par API HTTPS, ce qui contourne le "
    "blocage SMTP de l'hébergeur gratuit et améliore la délivrabilité."
)


h2("11.6 Stratégie de repli en trois niveaux")

para(
    "Prévoir plusieurs plans est la marque d'un ingénieur qui prend son "
    "sujet au sérieux. Voici l'ordre des replis à annoncer au jury si un "
    "incident survient."
)

para(
    "Plan A, par défaut : démonstration sur le déploiement cloud "
    "(pages.dev + onrender.com). Vous ouvrez l'URL, vous naviguez. "
    "Rien à installer, rien à démarrer. Le plus impressionnant."
)

para(
    "Plan B, si Cloudflare ou Render est temporairement en panne : "
    "démonstration sur la pile Docker locale de votre ordinateur. Vous avez "
    "lancé docker compose up une heure avant votre passage. Les trois "
    "conteneurs (backend FastAPI, base PostgreSQL/PostGIS, reverse proxy "
    "nginx) tournent en local et le tableau de bord répond sur "
    "http://localhost. Le mobile se connecte via un câble USB avec adb "
    "reverse tcp:8000 tcp:8000."
)

para(
    "Plan C, si votre ordinateur crashe complètement : démonstration à "
    "partir des captures d'écran de la partie interprétation des figures et "
    "des annexes du mémoire, puis renvoi vers l'URL cloud sur le "
    "vidéoprojecteur. Aucun scénario ne vous laisse muet."
)

para(
    "Faites au moins un test complet de chacun de ces plans la veille de la "
    "soutenance, avec le hotspot 4G de votre téléphone pour être indépendant "
    "du Wi-Fi de la salle."
)


h2("11.7 Résumé en une phrase pour le jury")

para(
    "« Le SI-ENV n'est pas qu'un prototype fonctionnel : il est déployé, "
    "surveillé et réparable automatiquement sur trois plateformes, avec "
    "trente-deux tests joués à chaque modification et une documentation "
    "technique générée en direct par le code. »"
)


d.save(CHEMIN)
print("Section 11 ajoutee au guide.")
