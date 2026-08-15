# -*- coding: utf-8 -*-
"""
Etape 20 : corriger le renommage precedent, qui avait vise la mauvaise
occurrence.

L'etape 19 renommait la premiere legende trouvee dans le document portant un
titre donne. Or le meme intitule existe deux fois : une fois dans la liste des
tableaux (avant le corps, donc rencontree en premier par une recherche
sequentielle), une fois comme vraie legende dans le corps, suivie d'un
veritable tableau. Le renommage a donc modifie la copie du front-matter, pas
la vraie legende - ce que confirme la liste reconstruite, toujours en collision
sur 5.2.

Cette etape cible exclusivement les paragraphes reellement suivis d'un tableau
(w:tbl), applique le decalage en cascade, puis reconstruit une derniere fois la
« Liste des tableaux ».

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os
import re
import sys

from docx import Document
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def ecrire_element(p_el, contenu):
    noeuds = list(p_el.iter(W + 't'))
    if not noeuds:
        return False
    noeuds[0].text = contenu
    for n in noeuds[1:]:
        n.text = ''
    return True


def renommer_vraie_legende(body, titre_distinctif, nouveau_numero):
    els = list(body)
    for i, el in enumerate(els):
        if not el.tag.endswith('}p'):
            continue
        txt = ''.join(t.text or '' for t in el.iter(W + 't')).strip()
        if titre_distinctif not in txt or not txt.startswith('Tableau'):
            continue
        suivi = False
        for j in range(i + 1, min(i + 3, len(els))):
            if els[j].tag.endswith('}tbl'):
                suivi = True
                break
            if els[j].tag.endswith('}p'):
                t2 = ''.join(t.text or '' for t in els[j].iter(W + 't')).strip()
                if t2:
                    break
        if not suivi:
            continue
        m = re.match(r'^Tableau\s+([\w.]+)\s*:\s*(.+)$', txt)
        if not m:
            continue
        ecrire_element(el, 'Tableau %s : %s' % (nouveau_numero, m.group(2)))
        return True
    return False


def main():
    doc = Document(CIBLE)
    body = doc.element.body

    cascade = [
        ('Indices environnementaux calculés via GEE', '5.7'),
        ('Hyperparamètres optimisés pour YOLOv8 et MobileNetV2', '5.6'),
        ('Benchmark des modèles de classification (phase 2)', '5.5'),
        ('Benchmark des modèles de détection (phase 1)', '5.4'),
        ("Composition du jeu de données pour l'entraînement (Dataset)", '5.3'),
    ]
    n = 0
    for titre, num in cascade:
        if renommer_vraie_legende(body, titre, num):
            n += 1
            print("  légende du corps renommée -> %s (%s)" % (num, titre[:40]))
        else:
            print("  ! non trouvée : %s" % titre[:50])
    print("  total : %d / %d" % (n, len(cascade)))

    doc.save(CIBLE)

    # ── Reconstruction finale de la liste des tableaux ──────────────────────
    doc = Document(CIBLE)
    paras = doc.paragraphs
    debut = fin = None
    for i, p in enumerate(paras):
        if p.text.strip().upper() == 'LISTE DES TABLEAUX':
            debut = i
        elif debut is not None and p.text.strip().upper().startswith(
                'LISTE DES SIGLES'):
            fin = i
            break
    if debut is None or fin is None:
        raise SystemExit("Bloc introuvable")

    els = list(doc.element.body)
    vrais = []
    for i, el in enumerate(els):
        if not el.tag.endswith('}p'):
            continue
        txt = ''.join(t.text or '' for t in el.iter(W + 't')).strip()
        m = re.match(r'^Tableau\s+([\w.]+)\s*:\s*(.+)$', txt)
        if not m:
            continue
        suivi = False
        for j in range(i + 1, min(i + 3, len(els))):
            if els[j].tag.endswith('}tbl'):
                suivi = True
                break
            if els[j].tag.endswith('}p'):
                t2 = ''.join(t.text or '' for t in els[j].iter(W + 't')).strip()
                if t2:
                    break
        if suivi:
            vrais.append((m.group(1), m.group(2)))

    modele = paras[debut + 1] if debut + 1 < fin else None
    for i in range(fin - 1, debut, -1):
        el = paras[i]._element
        el.getparent().remove(el)

    ancre = paras[debut]._element
    from docx.text.paragraph import Paragraph
    for num, titre in vrais:
        p_el = OxmlElement('w:p')
        ancre.addnext(p_el)
        ancre = p_el
        par = Paragraph(p_el, paras[debut]._parent)
        if modele is not None:
            par.style = modele.style
        run = par.add_run("Tableau %s : %s" % (num, titre))
        if modele is not None and modele.runs:
            run.font.size = modele.runs[0].font.size

    print("  liste des tableaux reconstruite : %d entrées" % len(vrais))
    for num, titre in vrais:
        print("    %-6s %s" % (num, titre[:60]))

    doc.save(CIBLE)
    print("\nEnregistré : %s" % CIBLE)


if __name__ == "__main__":
    main()
