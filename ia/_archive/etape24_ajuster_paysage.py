# -*- coding: utf-8 -*-
"""
Etape 24 : ajuster les deux figures deja passees en paysage (etape 23).

Le premier passage sous-exploitait la page paysage : tailles calculees sur une
hauteur utile trop prudente (16 cm) et images restees alignees a gauche,
laissant une bande blanche a droite et en bas. Cette etape recalcule sur la
hauteur utile reelle (18 cm de marge a marge, moins l'espace de la legende) et
centre l'image et sa legende sur la page.

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8')

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

# largeur, hauteur en cm, recalculees sur marges 1,5 cm (usable 26,7 x 18 cm)
TAILLES = {
    "Figure 4.2 : Diagramme de cas d'utilisation.": (18.15, 16.5),
    "Figure 4.3 : Diagramme de classes.": (26.7, 15.26),
}


def main():
    doc = Document(CIBLE)
    body = doc.element.body
    els = list(body)

    for i, el in enumerate(els):
        if not el.tag.endswith('}p'):
            continue
        txt = ''.join(t.text or '' for t in el.iter(W + 't')).strip()
        if txt not in TAILLES:
            continue
        larg, haut = TAILLES[txt]

        # legende : centree
        from docx.text.paragraph import Paragraph
        cap = Paragraph(el, doc.paragraphs[0]._parent)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # image : paragraphe precedent contenant le blip
        img_el = None
        for j in range(i - 1, max(0, i - 3), -1):
            if els[j].findall('.//{%s}blip' % A_NS):
                img_el = els[j]
                break
        assert img_el is not None, "image introuvable pour %r" % txt

        img_par = Paragraph(img_el, doc.paragraphs[0]._parent)
        img_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for ext in img_el.iter('{%s}ext' % A_NS):
            ext.set('cx', str(int(round(larg * 360000))))
            ext.set('cy', str(int(round(haut * 360000))))

        print("  %s -> %.2f x %.2f cm, centre" % (txt[:45], larg, haut))

    doc.save(CIBLE)
    print("\nEnregistré : %s" % CIBLE)


if __name__ == "__main__":
    main()
