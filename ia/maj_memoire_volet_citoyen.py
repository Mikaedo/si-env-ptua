# -*- coding: utf-8 -*-
"""
Complete le memoire sur le volet citoyen, le deploiement et les tests.

Trois passages restaient en decalage avec le systeme livre : la section
consacree a l'application mobile n'en decrivait qu'une, le diagramme de
deploiement ignorait la seconde, et le nombre de tests datait de fin juillet.

Le texte ajoute reste sobre. Il rend compte de ce qui a ete construit et des
raisons qui ont guide les choix, sans presenter comme acquis ce que le
dispositif ne fait pas : la verification de proximite atteste d'une position au
moment du depot, elle ne prouve pas la qualite de riverain, et le memoire doit
le dire plutot que de laisser le jury le decouvrir.
"""
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_volet_citoyen.docx")

journal = []


def reecrire(paragraphe, texte):
    """Remplace tout le contenu d'un paragraphe en gardant son style."""
    for segment in list(paragraphe.runs)[1:]:
        segment._element.getparent().remove(segment._element)
    if paragraphe.runs:
        paragraphe.runs[0].text = texte
    else:
        paragraphe.add_run(texte)


def inserer_apres(paragraphe, texte, style_source=None):
    """Insere un paragraphe juste apres celui qui est designe."""
    import copy
    nouveau = copy.deepcopy((style_source or paragraphe)._element)
    paragraphe._element.addnext(nouveau)
    from docx.text.paragraph import Paragraph
    p = Paragraph(nouveau, paragraphe._parent)
    for segment in list(p.runs)[1:]:
        segment._element.getparent().remove(segment._element)
    if p.runs:
        p.runs[0].text = texte
        p.runs[0].bold = False
        p.runs[0].italic = False
    else:
        p.add_run(texte)
    return p


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    paras = doc.paragraphs

    # ── 5.3 Application mobile : deux variantes d'un meme socle ───────────
    for p in paras:
        if p.text.strip().startswith("L'application fonctionne hors connexion d'abord"):
            ancien = p.text.strip()
            reecrire(p, ancien + (
                " Le projet produit en réalité deux applications distinctes, bâties sur "
                "un socle de code commun et différenciées au moyen des variantes de "
                "production Android. La première s'adresse aux agents de l'AGEROUTE et "
                "correspond à la description qui précède. La seconde vise les riverains "
                "des chantiers ; elle partage le service d'accès à l'API, les modèles de "
                "données, la géolocalisation et la charte graphique, mais ne conserve "
                "de la première ni la détection d'objets ni la cartographie. Ce découpage "
                "répond à une réalité de terrain autant qu'à une contrainte technique : "
                "les deux publics ne relèvent pas du même canal de distribution, "
                "l'application des agents se déployant en interne quand celle des "
                "riverains a vocation à être téléchargée librement, si bien qu'elles ne "
                "peuvent pas partager un identifiant applicatif. Les deux paquets "
                "s'installent côte à côte sur un même terminal."
            ))
            journal.append("ok  5.3 deux applications decrites")
            break
    else:
        journal.append("MANQUE  5.3")

    # ── 4.3.6 Deploiement : quatre composants ─────────────────────────────
    for p in paras:
        if p.text.strip().startswith("Le diagramme de déploiement (figure 4.7)"):
            ancien = p.text.strip()
            reecrire(p, ancien + (
                " Le dispositif compte désormais quatre composants applicatifs et non "
                "trois : au serveur d'application, à la base de données et au tableau de "
                "bord s'ajoute une seconde application mobile destinée aux riverains, "
                "déployée séparément bien qu'issue du même dépôt de code."
            ))
            journal.append("ok  4.3.6 quatrieme composant")
            break
    else:
        journal.append("MANQUE  4.3.6")

    # ── 6.2 Nombre de tests ───────────────────────────────────────────────
    for p in paras:
        if "32 tests automatisés" in p.text:
            texte = p.text
            texte = texte.replace(
                "La suite pytest (32 tests automatisés, fichier tests/test_functional.py)",
                "La suite pytest (119 tests automatisés répartis en sept fichiers)",
            )
            texte = texte.replace("32/32 test", "119/119 test")
            reecrire(p, texte + (
                " La couverture s'est étoffée au fil des évolutions : elle porte "
                "aujourd'hui sur le cloisonnement des profils de consultation, sur la "
                "condition de proximité géographique du volet citoyen, sur la traçabilité "
                "des rapports transmis et sur la complétude du jeu de données initial. "
                "Cette dernière vérification parcourt l'énumération des rôles définie dans "
                "le modèle plutôt qu'une liste tenue à la main, de sorte qu'un profil "
                "ajouté sans compte de démonstration fasse échouer la suite."
            ))
            journal.append("ok  6.2 nombre de tests actualise")
            break
    else:
        journal.append("MANQUE  6.2")

    # ── Nouvelle sous-section : le volet citoyen ──────────────────────────
    # Placee a la suite de 5.3, la ou le lecteur vient de decouvrir qu'il
    # existe deux applications.
    ancre = None
    for i, p in enumerate(paras):
        if p.text.strip().startswith("Tableau 5.2 : Packages Flutter"):
            ancre = p
            break

    if ancre is not None:
        modele = None
        for p in paras:
            if p.text.strip().startswith("L'application fonctionne hors connexion"):
                modele = p
                break

        blocs = [
            "Le mécanisme de gestion des plaintes du PTUA reposait jusqu'ici sur un "
            "recueil au guichet ou lors des réunions de quartier, ce qui suppose qu'un "
            "habitant se déplace et tombe sur une permanence ouverte. Beaucoup de "
            "nuisances ne remontaient donc jamais. L'application citoyenne ouvre un "
            "second canal, depuis le téléphone de la personne concernée, sans rien "
            "changer au traitement en aval : les doléances rejoignent la file déjà "
            "instruite par le spécialiste du suivi du Plan d'Action de Réinstallation, "
            "et le canal de saisie est conservé afin que l'apport du téléphone puisse "
            "être mesuré dans les rapports remis au bailleur.",

            "L'accès est conditionné à la proximité géographique. Une plainte "
            "environnementale n'a de sens que si elle émane de quelqu'un qui subit "
            "réellement la nuisance, et le dispositif serait rapidement saturé s'il "
            "acceptait des dépôts émis depuis n'importe où. La vérification s'appuie sur "
            "un rayon d'influence propre à chaque chantier, notion que tout PGES manipule "
            "déjà sous le nom de zone d'influence directe : un terrassement lourd dérange "
            "plus loin qu'une reprise de chaussée, et le spécialiste du suivi "
            "environnemental fixe cette étendue ouvrage par ouvrage. Le chantier retenu "
            "est le plus proche parmi ceux dont la zone englobe effectivement la "
            "position, et non le plus proche dans l'absolu, un ouvrage voisin au "
            "périmètre resserré pouvant sinon écarter un site plus lointain mais "
            "réellement couvrant.",

            "Deux choix méritent d'être justifiés. Le rattachement au chantier est déduit "
            "de la position et non choisi dans une liste, un habitant n'ayant aucune "
            "raison de connaître les dénominations administratives des ouvrages. Et "
            "aucune reconnaissance automatique n'est sollicitée dans cette application, "
            "contrairement à celle des agents : un riverain décrit une gêne, il ne pose "
            "pas de diagnostic environnemental, et la qualification relève du spécialiste "
            "qui dispose du contexte du chantier.",

            "La portée de ce contrôle doit être énoncée avec exactitude. La position "
            "transmise atteste d'une localisation au moment de l'inscription, elle ne "
            "prouve pas la qualité de riverain : une personne de passage satisfait la "
            "condition, et un dispositif de production devrait lui adjoindre une "
            "vérification d'identité ou une validation par le comité de quartier. Ce "
            "filtre écarte les dépôts manifestement extérieurs à la zone du projet, ce "
            "qui répond à l'objectif poursuivi ici, sans constituer une authentification "
            "de résidence.",
        ]

        courant = ancre
        for texte in reversed(blocs):
            courant = inserer_apres(ancre, texte, style_source=modele)

        titre = inserer_apres(ancre, "5.3 bis  Application citoyenne et canal riverain",
                              style_source=modele)
        for r in titre.runs:
            r.bold = True
        journal.append(f"ok  section 5.3 bis inseree ({len(blocs)} paragraphes)")
    else:
        journal.append("MANQUE  ancre 5.3 bis")

    doc.save(SOURCE)
    print("\n".join(journal))
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
