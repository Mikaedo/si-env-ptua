# -*- coding: utf-8 -*-
"""
Met le guide de soutenance en accord avec le systeme livre.

Le guide decrit le deroule que le candidat suivra devant le jury. Un guide qui
annonce cinq roles et une seule application mobile alors que la demonstration
en montre huit et deux met son auteur en difficulte au pire moment : il lui
faudrait improviser sur des points qu'il croyait preparés.

La section ajoutee couvre le volet citoyen et les organismes de controle, avec
les questions que ces ajouts appellent naturellement. Deux d'entre elles sont
delicates et meritent une reponse preparee : la portee reelle du controle de
position, et la raison pour laquelle les courriels ne partent que vers une
seule adresse. Mieux vaut les avoir en main que les decouvrir debout.
"""
import copy
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\GUIDE_PREPARATION_SOUTENANCE_SI-ENV.docx")
SAUVEGARDE = SOURCE.with_name("GUIDE_PREPARATION_SOUTENANCE_SI-ENV_avant_maj.docx")


def ajouter(doc, texte, style=None, gras=False, taille=None):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = style
        except KeyError:
            pass
    r = p.add_run(texte)
    r.bold = gras
    if taille:
        r.font.size = Pt(taille)
    return p


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    # ── Corrections dans le corps existant ────────────────────────────────
    corrections = [
        ("cinq rôles", "huit rôles"),
        ("Cinq rôles", "Huit rôles"),
        ("les 5 rôles", "les 8 rôles"),
        ("32 tests", "119 tests"),
        ("32/32", "119/119"),
    ]
    faits = 0
    for p in doc.paragraphs:
        for ancien, nouveau in corrections:
            if ancien in p.text:
                complet = p.text.replace(ancien, nouveau)
                for seg in list(p.runs)[1:]:
                    seg._element.getparent().remove(seg._element)
                if p.runs:
                    p.runs[0].text = complet
                faits += 1
                break
    print(f"  {faits} correction(s) dans le corps")

    # ── Nouvelle section 12 ───────────────────────────────────────────────
    ajouter(doc, "12. Les huit profils, l'application citoyenne et les organismes de contrôle",
            style="Heading 1")

    ajouter(doc,
            "Cette section couvre ce qui a été ajouté après la première version du "
            "guide. Elle remplace, sur ces points, ce que disent les sections "
            "précédentes.")

    ajouter(doc, "12.1 Ce qui a changé, en trois phrases", style="Heading 2")
    for t in [
        "Le système comptait cinq profils, il en compte huit. Les trois nouveaux "
        "traduisent la gouvernance réelle du programme : l'ANDE et la BAD "
        "consultent sans pouvoir écrire, et les riverains alimentent le mécanisme "
        "de gestion des plaintes depuis leur téléphone.",
        "Il existe désormais deux applications mobiles et non une : celle des "
        "agents AGEROUTE, et une application citoyenne destinée aux riverains. "
        "Elles partagent le même socle de code mais s'installent séparément.",
        "Le spécialiste du suivi environnemental peut transmettre officiellement "
        "un rapport à l'ANDE ou à la BAD. Chaque remise est enregistrée avec sa "
        "date et son émetteur.",
    ]:
        ajouter(doc, t)

    ajouter(doc, "12.2 Comptes de démonstration à jour", style="Heading 2")
    tableau = doc.add_table(rows=1, cols=3)
    tableau.style = "Table Grid"
    for i, titre in enumerate(("Profil", "Adresse", "Mot de passe")):
        c = tableau.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(titre)
        r.bold = True
        r.font.size = Pt(9)
    comptes = [
        ("Administrateur", "admin@sienv.ci", "admin123"),
        ("Responsable Environnement (mobile agent)", "resp.env@ageroute.ci", "env123"),
        ("Expert HSE (mobile agent)", "expert.hse@ageroute.ci", "expert123"),
        ("Spécialiste Suivi Environnemental", "spec.env@ageroute.ci", "spec123"),
        ("Spécialiste Suivi du P.A.R", "spec.par@ageroute.ci", "spec123"),
        ("ANDE (consultation)", "controle@ande.ci", "ande123"),
        ("BAD (consultation)", "mission@afdb.org", "bad123"),
        ("Riverain (mobile citoyen)", "riverain@yopougon.ci", "riverain123"),
    ]
    for ligne in comptes:
        cells = tableau.add_row().cells
        for i, v in enumerate(ligne):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(v)
            r.font.size = Pt(9)

    ajouter(doc, "12.3 Déroulé de la démonstration citoyenne (3 minutes)", style="Heading 2")
    for i, t in enumerate([
        "Ouvrez l'application SI-ENV Citoyen. L'écran d'accueil demande la "
        "position et explique pourquoi : vérifier que la personne réside près "
        "d'un chantier.",
        "Autorisez la localisation. L'application interroge le serveur, qui "
        "détermine le chantier couvrant votre position et vous y rattache.",
        "Créez un compte avec une adresse. Le rattachement est déjà décidé, le "
        "riverain n'a aucun chantier à choisir : dites-le au jury, c'est un "
        "choix de conception.",
        "Déposez une doléance : catégorie, description, envoi. La confirmation "
        "indique que le spécialiste du suivi social l'a reçue.",
        "Basculez sur le tableau de bord web, connectez-vous en spec.par et "
        "montrez la doléance qui vient d'arriver, avec sa mention « Dépôt "
        "riverain » qui la distingue d'un recueil au guichet.",
    ], start=1):
        ajouter(doc, f"{i}. {t}")

    ajouter(doc, "12.4 Démonstration des organismes de contrôle (2 minutes)", style="Heading 2")
    for t in [
        "Connectez-vous en controle@ande.ci. Faites remarquer que les écrans "
        "sont les mêmes que ceux du spécialiste, mais qu'aucun bouton d'action "
        "n'apparaît.",
        "Si le jury demande ce qui empêche vraiment l'écriture, répondez que le "
        "refus est posé côté serveur, avant même le traitement métier, et non "
        "par un masquage de l'interface. Une requête construite en dehors du "
        "tableau de bord est rejetée de la même façon. C'est ce qui donne sa "
        "valeur au contrôle : un rapport de conformité ne vaudrait rien si "
        "celui qui l'examine pouvait retoucher les données.",
        "Revenez en spec.env, ouvrez les rapports PGES, transmettez-en un. "
        "Montrez ensuite l'historique des transmissions : qui a transmis quoi, "
        "à qui, à quelle date.",
    ]:
        ajouter(doc, t)

    ajouter(doc, "12.5 Les deux questions délicates, et quoi répondre", style="Heading 2")

    ajouter(doc, "« Comment savez-vous que la personne est vraiment riveraine ? »",
            gras=True)
    ajouter(doc,
            "Répondez sans détour : je ne le sais pas, et le mémoire le dit. La "
            "position transmise atteste d'une localisation au moment de "
            "l'inscription, elle ne prouve pas la résidence. Une personne de "
            "passage satisfait la condition. Ce filtre écarte les dépôts "
            "manifestement extérieurs à la zone du projet, ce qui était "
            "l'objectif ; un dispositif en exploitation réelle devrait y "
            "adjoindre une vérification d'identité ou une validation par le "
            "comité de quartier. Assumer cette limite vaut mieux que la "
            "défendre : le jury la verra de toute façon.")

    ajouter(doc, "« Pourquoi les courriels ne partent-ils que vers votre adresse ? »",
            gras=True)
    ajouter(doc,
            "Le service d'envoi fonctionne en mode bac à sable, qui n'autorise "
            "l'expédition que vers l'adresse propriétaire du compte tant qu'aucun "
            "nom de domaine n'a été vérifié. C'est une protection anti-abus du "
            "fournisseur, pas un défaut du système : le code envoie de la même "
            "façon quel que soit le destinataire, et la trace de transmission "
            "enregistre d'ailleurs les échecs. Lever cette limite suppose "
            "d'acquérir un domaine, ce qui sortait du cadre gratuit retenu pour "
            "la validation académique.")

    ajouter(doc, "12.6 Le rayon d'influence, si on vous interroge dessus", style="Heading 2")
    ajouter(doc,
            "Chaque chantier porte une zone d'influence exprimée en mètres, que le "
            "spécialiste du suivi environnemental fixe ouvrage par ouvrage. Ce "
            "n'est pas un paramètre arbitraire : tout PGES définit une aire "
            "d'étude autour d'un ouvrage, et un terrassement lourd dérange plus "
            "loin qu'une reprise de chaussée. Le système retient le chantier le "
            "plus proche parmi ceux dont la zone couvre effectivement la "
            "position, et non le plus proche dans l'absolu, faute de quoi un "
            "ouvrage voisin au périmètre resserré écarterait un site plus "
            "lointain mais réellement couvrant.")

    ajouter(doc, "12.7 Chiffres à jour", style="Heading 2")
    for t in [
        "119 tests automatisés, contre 32 dans la première version du guide.",
        "8 profils utilisateurs, 4 composants applicatifs déployés.",
        "2 applications mobiles issues d'un socle de code commun.",
        "2 canaux de saisie pour le mécanisme de gestion des plaintes : le "
        "guichet et le téléphone.",
    ]:
        ajouter(doc, t)

    doc.save(SOURCE)
    print(f"  section 12 ajoutee")
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
