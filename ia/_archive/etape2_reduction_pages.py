# -*- coding: utf-8 -*-
"""
Etape 2 sur 3 : ramener le corps du memoire de 65 a 50 pages, sans supprimer
une ligne de texte redige.

Trois leviers, du moins au plus visible :

  a) Figures : elles occupent environ 9 pages du corps. Reduire leur hauteur de
     30 % ne nuit pas a la lisibilite (les diagrammes UML restent nets) et
     libere plusieurs pages.

  b) Pages faiblement remplies : deux pages du corps ne contiennent qu'un
     en-tete, consequence de sauts de page et de paragraphes vides accumules.

  c) Pages separatrices de parties : chaque « PREMIERE / DEUXIEME / TROISIEME
     PARTIE » occupe une page entiere pour trois lignes. Le titre est conserve,
     mais en tete de la page suivante plutot que seul sur une page.

Sortie : MEMOIRE_ETAPE2.docx
"""
import os
import re
import shutil

from docx import Document
from docx.oxml.ns import qn

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_ETAPE1.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_ETAPE2.docx")

A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

#: Facteur de reduction applique aux figures du corps.
REDUCTION = 0.70
#: Les images plus petites que ce seuil (cm) sont des logos : on n'y touche pas.
SEUIL_LARGEUR_CM = 5.0
#: Debut du corps : au-dela, on ne touche plus aux images des annexes.
DEBUT_CORPS = 250


def cm_vers_emu(v):
    return int(round(v * 360000))


def reduire_figures(doc):
    """Reduit la taille des figures du corps, hors annexes et hors logos."""
    debut_annexes = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == 'ANNEXES' and i > 400:
            debut_annexes = i
            break
    limite = debut_annexes or len(doc.paragraphs)

    reduites = gain_cm = 0
    for i, p in enumerate(doc.paragraphs):
        if not (DEBUT_CORPS <= i < limite):
            continue
        for ext in p._element.iter('{%s}ext' % A_NS):
            cx, cy = ext.get('cx'), ext.get('cy')
            if not (cx and cy):
                continue
            largeur_cm = int(cx) / 360000.0
            hauteur_cm = int(cy) / 360000.0
            if largeur_cm < SEUIL_LARGEUR_CM:
                continue  # logo ou pictogramme
            ext.set('cx', str(cm_vers_emu(largeur_cm * REDUCTION)))
            ext.set('cy', str(cm_vers_emu(hauteur_cm * REDUCTION)))
            reduites += 1
            gain_cm += hauteur_cm * (1 - REDUCTION)
    return reduites, gain_cm, limite


def compacter_paragraphes_vides(doc, limite):
    """Supprime les suites de plus de deux paragraphes vides consecutifs.

    Ces accumulations proviennent de mises en page manuelles et poussent du
    contenu sur la page suivante sans rien apporter.
    """
    body = doc.element.body
    supprimes = 0
    vides_consecutifs = []
    for el in list(body):
        if not el.tag.endswith('}p'):
            vides_consecutifs = []
            continue
        texte = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip()
        a_image = bool(el.findall('.//{%s}blip' % A_NS))
        a_saut = 'type="page"' in el.xml
        a_sectPr = (el.find(qn('w:pPr')) is not None
                    and el.find(qn('w:pPr')).find(qn('w:sectPr')) is not None)
        if texte or a_image or a_saut or a_sectPr:
            vides_consecutifs = []
            continue
        vides_consecutifs.append(el)
        # On conserve deux paragraphes vides (respiration), on retire le reste
        if len(vides_consecutifs) > 2:
            body.remove(el)
            supprimes += 1
    return supprimes


def compacter_separateurs_parties(doc):
    """Retire le saut de page qui isole chaque titre de partie sur sa page."""
    traites = 0
    for p in doc.paragraphs:
        if not re.search(r'(PREMI|DEUXI|TROISI)\wRE PARTIE', p.text.upper()):
            continue
        # Le titre reste, mais il n'occupe plus une page a lui seul : on
        # supprime le saut de page qui le precede.
        precedent = p._element.getprevious()
        while precedent is not None and precedent.tag.endswith('}p'):
            texte = ''.join(t.text or '' for t in precedent.iter(qn('w:t'))).strip()
            if texte:
                break
            if 'type="page"' in precedent.xml:
                for br in precedent.findall('.//' + qn('w:br')):
                    if br.get(qn('w:type')) == 'page':
                        br.getparent().remove(br)
                        traites += 1
                break
            precedent = precedent.getprevious()
    return traites


def main():
    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)

    reduites, gain, limite = reduire_figures(doc)
    print("  figures reduites a %d %% : %d (hauteur liberee : %.0f cm, ~%.1f pages)"
          % (REDUCTION * 100, reduites, gain, gain / 22))

    supprimes = compacter_paragraphes_vides(doc, limite)
    print("  paragraphes vides superflus supprimes : %d" % supprimes)

    sauts = compacter_separateurs_parties(doc)
    print("  sauts de page devant les titres de partie retires : %d" % sauts)

    doc.save(SORTIE)
    print("\nEnregistre : %s" % SORTIE)


if __name__ == "__main__":
    main()
