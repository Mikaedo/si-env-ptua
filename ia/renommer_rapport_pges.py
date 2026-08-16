# -*- coding: utf-8 -*-
"""
Distingue le PGES, document de planification, du rapport qui en rend compte.

Le Plan de Gestion Environnementale et Sociale s'etablit en amont du chantier,
avec l'etude d'impact : il fixe les mesures d'attenuation, les responsabilites,
le programme de surveillance et son budget. Ce que le SI-ENV produit est autre
chose : un etat, sur une periode donnee, de ce qui a ete constate et traite.
Nommer ce dernier « rapport PGES » revenait a confondre le plan et son suivi,
et exposait a une question sans reponse defendable : ou sont, dans ce rapport,
les mesures d'attenuation et leur budget ?

Le remplacement n'est donc pas mecanique. Les mentions du PGES comme
referentiel sont exactes et restent : le sigle dans sa liste, le cadre
reglementaire, les nuisances qu'il recense, la zone d'influence qu'il definit.
Seules changent les tournures qui font du document produit un PGES.

Les remplacements sont ordonnes du plus long au plus court, faute de quoi une
regle generale consommerait les cas particuliers avant qu'ils soient vus.

Une precaution technique : Word fractionne un paragraphe en fragments, et une
expression peut se trouver a cheval sur deux d'entre eux. La substitution est
donc tentee fragment par fragment, puis, si l'expression subsiste, sur le
paragraphe entier apres regroupement. Ce qui echapperait aux deux est signale
plutot que passe sous silence.
"""
import shutil
import sys
from pathlib import Path

from docx import Document

# Le guide de soutenance doit suivre le meme renommage que le memoire. Dire a
# l'oral un terme que le document ecrit ne porte plus serait la contradiction
# la plus facile a relever par un jury, et la plus evitable.
SOURCE = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(
    r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name(SOURCE.stem + "_avant_renommage.docx")

REMPLACEMENTS = [
    ("Génération rapport PGES PDF", "Génération du rapport de suivi (PDF)"),
    ("Génération rapport PGES", "Génération du rapport de suivi"),
    ("un PGES PDF en un clic", "un rapport de suivi au format PDF en un clic"),
    ("rapports PGES", "rapports de suivi environnemental"),
    ("rapport PGES", "rapport de suivi environnemental"),
]


def substituer(texte):
    for vieux, neuf in REMPLACEMENTS:
        texte = texte.replace(vieux, neuf)
    return texte


def contient_cible(texte):
    return any(vieux in texte for vieux in (v for v, _ in REMPLACEMENTS))


def traiter(paragraphe, journal):
    if not contient_cible(paragraphe.text):
        return 0

    avant = paragraphe.text
    for fragment in paragraphe.runs:
        if contient_cible(fragment.text):
            fragment.text = substituer(fragment.text)

    if contient_cible(paragraphe.text):
        # L'expression chevauche plusieurs fragments : on les regroupe.
        entier = substituer(paragraphe.text)
        for fragment in list(paragraphe.runs)[1:]:
            fragment._element.getparent().remove(fragment._element)
        if paragraphe.runs:
            paragraphe.runs[0].text = entier
        else:
            paragraphe.add_run(entier)
        journal.append(f"    (regroupe) {entier[:66]}")
        return 1

    journal.append(f"    {paragraphe.text[:72]}")
    return 1 if avant != paragraphe.text else 0


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    journal, total = [], 0

    for p in doc.paragraphs:
        total += traiter(p, journal)
    for table in doc.tables:
        for ligne in table.rows:
            for cellule in ligne.cells:
                for p in cellule.paragraphs:
                    total += traiter(p, journal)

    print(f"{total} passage(s) corrige(s) :")
    for ligne in journal:
        print(ligne)

    doc.save(SOURCE)

    restant = []
    controle = Document(SOURCE)
    for p in controle.paragraphs:
        if contient_cible(p.text):
            restant.append(p.text[:70])
    for table in controle.tables:
        for ligne in table.rows:
            for cellule in ligne.cells:
                if contient_cible(cellule.text):
                    restant.append(cellule.text[:70])

    print(f"\nmentions fautives restantes : {len(restant)}")
    for r in restant:
        print(f"  {r}")

    conserves = sum(1 for p in controle.paragraphs if "PGES" in p.text)
    print(f"mentions du PGES conservees comme referentiel : {conserves}")
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
