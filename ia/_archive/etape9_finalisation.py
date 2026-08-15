# -*- coding: utf-8 -*-
"""
Etape 9 : derniers gains pour atteindre 50 pages, puis table des matieres
paginee.

  a) Figures ramenees a 45 % de leur taille d'origine. Les diagrammes UML et
     les captures restent lisibles a cette echelle sur une page A4.
  b) Contenu des tableaux en 9 points avec marges internes reduites : plusieurs
     tableaux du corps comptent de dix a seize lignes et debordent sur la page
     suivante pour deux ou trois lignes.
  c) Table des matieres reconstruite en champ Word, ce qui produit les points
     de conduite et les numeros de page automatiquement, et les maintient a jour
     si la pagination change.

Police, interligne 1,5 et marges du reglement restent inchanges pour le texte
courant ; seuls les contenus de tableaux, plus denses par nature, sont reduits.

Sortie : MEMOIRE_ETAPE9.docx
"""
import os
import re
import shutil

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_ETAPE8.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_ETAPE9.docx")

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

CIBLE = 0.45
DEJA = 0.55
FACTEUR = CIBLE / DEJA
SEUIL_LOGO_CM = 5.0


def bornes_corps(doc):
    debut = fin = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if debut is None and t.startswith('Introduction g') and i > 250:
            debut = i
        if debut is not None and t.upper().startswith('BIBLIOGRAPHIE'):
            fin = i
            break
    return debut, fin or len(doc.paragraphs)


def reduire_figures(doc, debut, fin):
    n = 0
    for i in range(debut, fin):
        for ext in doc.paragraphs[i]._element.iter('{%s}ext' % A_NS):
            cx, cy = ext.get('cx'), ext.get('cy')
            if not (cx and cy):
                continue
            l_cm, h_cm = int(cx) / 360000.0, int(cy) / 360000.0
            if l_cm < SEUIL_LOGO_CM:
                continue
            ext.set('cx', str(int(round(l_cm * FACTEUR * 360000))))
            ext.set('cy', str(int(round(h_cm * FACTEUR * 360000))))
            n += 1
    return n


def compacter_tableaux(doc):
    """Reduit la police et les marges internes des tableaux du corps."""
    n_tab = 0
    for table in doc.tables:
        # On epargne les tableaux de la page de garde (logos, encadrants)
        if len(table.rows) <= 2 and len(table.columns) <= 2:
            continue
        for ligne in table.rows:
            for cellule in ligne.cells:
                for par in cellule.paragraphs:
                    pf = par.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(1)
                    pf.line_spacing = 1.0
                    for run in par.runs:
                        if run.font.size is None or run.font.size.pt > 9:
                            run.font.size = Pt(9)
        n_tab += 1
    return n_tab


def remplacer_tdm_par_champ(doc):
    """Remplace la table des matieres statique par un champ Word.

    Un champ produit les points de conduite et les numeros de page tout seul,
    et reste juste apres une repagination. La table saisie a la main, elle,
    devient fausse des qu'une page se decale.
    """
    paras = doc.paragraphs
    debut = None
    for i, p in enumerate(paras):
        if p.text.strip().upper().startswith('TABLE DES MATI'):
            suite = ' '.join(x.text for x in paras[i + 1:i + 8])
            if '\t' in suite and debut is None:
                debut = i
    if debut is None:
        return 0

    # Fin de la table : premiere rubrique de meme niveau rencontree apres
    fin = None
    for i in range(debut + 1, len(paras)):
        t = paras[i].text.strip().upper()
        if t.startswith(('RESUME', 'R\u00c9SUM\u00c9', 'ABSTRACT', 'ANNEXES')):
            fin = i
            break
    if fin is None:
        return 0

    # Suppression des entrees saisies a la main
    supprimes = 0
    for i in range(fin - 1, debut, -1):
        el = paras[i]._element
        el.getparent().remove(el)
        supprimes += 1

    # Insertion du champ TOC a la place
    ancre = paras[debut]._element
    p_champ = OxmlElement('w:p')
    ancre.addnext(p_champ)
    r = OxmlElement('w:r')
    p_champ.append(r)
    debut_champ = OxmlElement('w:fldChar')
    debut_champ.set(qn('w:fldCharType'), 'begin')
    r.append(debut_champ)

    r2 = OxmlElement('w:r')
    p_champ.append(r2)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    # \o "1-3" : niveaux 1 a 3 ; \h : liens ; \z : masque les numeros en mode
    # web ; \u : utilise les niveaux de plan.
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r2.append(instr)

    r3 = OxmlElement('w:r')
    p_champ.append(r3)
    sep = OxmlElement('w:fldChar')
    sep.set(qn('w:fldCharType'), 'separate')
    r3.append(sep)

    r4 = OxmlElement('w:r')
    p_champ.append(r4)
    t = OxmlElement('w:t')
    t.text = "Table des matieres a mettre a jour (F9 dans Word)"
    r4.append(t)

    r5 = OxmlElement('w:r')
    p_champ.append(r5)
    fin_champ = OxmlElement('w:fldChar')
    fin_champ.set(qn('w:fldCharType'), 'end')
    r5.append(fin_champ)

    return supprimes


def main():
    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)
    debut, fin = bornes_corps(doc)

    n_fig = reduire_figures(doc, debut, fin)
    print("  figures ramenees a %d %% de l'original : %d" % (CIBLE * 100, n_fig))

    n_tab = compacter_tableaux(doc)
    print("  tableaux compactes (9 pt) : %d" % n_tab)

    n_tdm = remplacer_tdm_par_champ(doc)
    print("  entrees de table des matieres remplacees par un champ : %d" % n_tdm)

    doc.save(SORTIE)
    print("\nEnregistre : %s" % SORTIE)


if __name__ == "__main__":
    main()
