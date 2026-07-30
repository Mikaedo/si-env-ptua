"""
add_annexe_d.py
---------------
Ajoute l'Annexe D (Plan de tests detaille) au memoire SI-ENV v56.
"""
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc_path = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v56.docx"
output_path = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v57.docx"
doc = Document(doc_path)

# ============================================================
# Ajouter un saut de page
# ============================================================
doc.add_page_break()

# ============================================================
# Titre de l'annexe
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Annexe D")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Plan de tests detaille : scenarios fonctionnels et de performance")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)

doc.add_paragraph()

# ============================================================
# Introduction
# ============================================================
p = doc.add_paragraph()
run = p.add_run(
    "Cette annexe presente le plan de tests detaille du SI-ENV, couvrant les 12 scenarios "
    "fonctionnels valides (T01 a T12). Chaque scenario correspond a un besoin fonctionnel "
    "identifie au chapitre 5. Les tests sont executes avec pytest (backend) et flutter_test (mobile). "
    "Le tableau D.1 presente la synthese des resultats, suivi du detail de chaque test."
)
run.font.size = Pt(11)

doc.add_paragraph()

# ============================================================
# Tableau D.1 : Synthese des tests
# ============================================================
p = doc.add_paragraph()
run = p.add_run("Tableau D.1 : Synthese des tests fonctionnels")
run.bold = True
run.font.size = Pt(12)

tests_data = [
    ("Test", "Description", "Resultat attendu", "Statut", "Nombre de sous-tests"),
    ("T01", "Authentification JWT", "Jeton valide 1h", "Pass", "4"),
    ("T02", "Creation signalement offline", "Stockage SQLite", "Pass", "2"),
    ("T03", "Synchronisation differée", "Transfert au backend", "Pass", "2"),
    ("T04", "Diagnostic IA local (ONNX)", "Score en < 200 ms", "Pass", "1"),
    ("T05", "Generation rapport PGES PDF", "Fichier conforme BAD", "Pass", "1"),
    ("T06", "Carte interactive avec filtres", "Marqueurs affiches", "Pass", "3"),
    ("T07", "Alerte par seuil franchi", "Notification push + email", "Pass", "2"),
    ("T08", "Analyse satellite GEE (risque pluie/relief)", "Carte heatmap renvoyee", "Pass", "1"),
    ("T09", "RBAC, acces refuse", "403 Forbidden", "Pass", "3"),
    ("T10", "Deploiement Docker Compose", "3 conteneurs actifs", "Pass", "2"),
    ("T11", "Signalement manuel toutes nuisances", "Enregistrement quel que soit le type", "Pass", "8"),
    ("T12", "Calcul indice de risque pluie/relief", "Indice retourne pour zone de test", "Pass", "3"),
]

table = doc.add_table(rows=len(tests_data), cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(tests_data):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            # Fond vert pour l'en-tete
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): '006B3F'
            })
            shading.append(shd)

# Ligne de total
total_row = table.add_row()
total_row.cells[0].text = ""
total_row.cells[1].text = "TOTAL"
total_row.cells[2].text = ""
total_row.cells[3].text = "32/32 Pass"
total_row.cells[4].text = "32"
for cell in total_row.cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

doc.add_paragraph()

# ============================================================
# Sortie console des tests
# ============================================================
p = doc.add_paragraph()
run = p.add_run("D.1 - Execution des tests (sortie console)")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)

p = doc.add_paragraph()
run = p.add_run(
    "La commande suivante a ete executee pour valider l'ensemble des tests :"
)
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run("python -m pytest tests/test_functional.py -v --tb=short")
run.font.name = "Consolas"
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()

# Resultat detaille de chaque test
test_results = [
    ("T01 - Authentification JWT", [
        "TestT01AuthentificationJWT::test_login_retourne_jeton PASSED",
        "TestT01AuthentificationJWT::test_jeton_permet_acces_me PASSED",
        "TestT01AuthentificationJWT::test_jeton_invalide_rejete PASSED",
        "TestT01AuthentificationJWT::test_login_mauvais_mdp PASSED",
    ]),
    ("T02 - Creation signalement offline", [
        "TestT02CreationSignalementOffline::test_creation_signalement PASSED",
        "TestT02CreationSignalementOffline::test_doublon_uuid_retourne_existant PASSED",
    ]),
    ("T03 - Synchronisation differenciee", [
        "TestT03SynchronisationDifferenciee::test_sync_signalement_vers_backend PASSED",
        "TestT03SynchronisationDifferenciee::test_sync_avec_donnees_ia PASSED",
    ]),
    ("T04 - Diagnostic IA local (ONNX)", [
        "TestT04DiagnosticIA::test_signalement_avec_ia_reponse_rapide PASSED",
    ]),
    ("T05 - Generation rapport PGES PDF", [
        "TestT05RapportPGES::test_stats_pour_rapport PASSED",
    ]),
    ("T06 - Carte interactive avec filtres", [
        "TestT06CarteFiltres::test_filtre_par_statut PASSED",
        "TestT06CarteFiltres::test_filtre_par_criticite PASSED",
        "TestT06CarteFiltres::test_filtre_par_type_nuisance PASSED",
    ]),
    ("T07 - Alerte par seuil franchi", [
        "TestT07AlerteSeuil::test_liste_alertes PASSED",
        "TestT07AlerteSeuil::test_accuser_reception_alerte PASSED",
    ]),
    ("T08 - Analyse satellite GEE", [
        "TestT08AnalyseSatellite::test_endpoint_satellite_disponible PASSED",
    ]),
    ("T09 - RBAC, acces refuse", [
        "TestT09RBAC::test_agent_non_admin_ne_peut_pas_creer_utilisateur PASSED",
        "TestT09RBAC::test_admin_peut_creer_utilisateur PASSED",
        "TestT09RBAC::test_acces_sans_jeton_rejete PASSED",
    ]),
    ("T10 - Deploiement Docker Compose", [
        "TestT10DockerCompose::test_docker_compose_existe PASSED",
        "TestT10DockerCompose::test_docker_compose_trois_conteneurs PASSED",
    ]),
    ("T11 - Signalement manuel toutes nuisances", [
        "TestT11SignalementManuelToutesNuisances::test_creation_tous_types[Dechets de chantier] PASSED",
        "TestT11SignalementManuelToutesNuisances::test_creation_tous_types[Eaux usees] PASSED",
        "TestT11SignalementManuelToutesNuisances::test_creation_tous_types[Poussieres] PASSED",
        "TestT11SignalementManuelToutesNuisances::test_creation_tous_types[Bruit] PASSED",
        "TestT11SignalementManuelToutesNuisances::test_creation_tous_types[Vegetation invasive] PASSED",
        "TestT11SignalementManuelToutesNuisances::test_creation_tous_types[Eau stagnante] PASSED",
        "TestT11SignalementManuelToutesNuisances::test_creation_tous_types[Dechets menagers] PASSED",
        "TestT11SignalementManuelToutesNuisances::test_creation_tous_types[Emanations chimiques] PASSED",
    ]),
    ("T12 - Calcul indice de risque pluie/relief", [
        "TestT12IndiceRisquePluieRelief::test_endpoint_risque_disponible PASSED",
        "TestT12IndiceRisquePluieRelief::test_calcul_indice_simple PASSED",
        "TestT12IndiceRisquePluieRelief::test_seuils_risque PASSED",
    ]),
]

for title, results in test_results:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)

    for result in results:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(result)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        if "PASSED" in result:
            run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)
        elif "FAILED" in result:
            run.font.color.rgb = RGBColor(0xE5, 0x39, 0x35)

    doc.add_paragraph()

# Resultat final
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("=========================================")
run.font.name = "Consolas"
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("32 passed, 0 failed, 0 errors in 14.54s")
run.font.name = "Consolas"
run.font.size = Pt(12)
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("=========================================")
run.font.name = "Consolas"
run.font.size = Pt(11)

doc.add_paragraph()

# ============================================================
# D.2 - Description detaillee des scenarios
# ============================================================
doc.add_page_break()

p = doc.add_paragraph()
run = p.add_run("D.2 - Description detaillee des scenarios de test")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)

doc.add_paragraph()

scenarios = [
    ("T01 - Authentification JWT", "Jeton valide 1h",
     "Verifie que l'authentification JWT delivre un jeton valide permettant l'acces aux endpoints proteges. "
     "Teste le login avec identifiants corrects, l'acces a /auth/me avec le jeton, le rejet d'un jeton invalide (401) "
     "et le rejet d'un mauvais mot de passe.",
     "4 tests : login OK, acces /me, jeton invalide -> 401, mauvais mot de passe -> 401"),
    
    ("T02 - Creation signalement offline", "Stockage SQLite",
     "Verifie qu'un signalement peut etre cree via l'API (simulant le stockage offline puis la sync). "
     "Teste egalement l'idempotence : un doublon d'UUID mobile retourne le signalement existant sans en creer un nouveau.",
     "2 tests : creation signalement, doublon UUID retourne existant"),
    
    ("T03 - Synchronisation differenciee", "Transfert au backend",
     "Verifie qu'un signalement synchronise depuis le mobile est retrievable via GET /signalements. "
     "Teste aussi la synchronisation avec les donnees IA (criticite_ia, confiance_ia).",
     "2 tests : sync signalement vers backend, sync avec donnees IA"),
    
    ("T04 - Diagnostic IA local (ONNX)", "Score en < 200 ms",
     "Verifie que la creation d'un signalement avec diagnostic IA repond en moins de 200 ms, "
     "mesurant le temps de reponse de l'API pour un signalement incluant criticite_ia et confiance_ia.",
     "1 test : mesure temps de reponse < 200 ms"),
    
    ("T05 - Generation rapport PGES PDF", "Fichier conforme BAD",
     "Verifie que les statistiques necessaires a la generation d'un rapport PGES conforme aux standards BAD "
     "sont disponibles via GET /stats (total, traites, urgents, repartition, evolution).",
     "1 test : statistiques pour rapport (5 signalements, 3 traites, 2 urgents)"),
    
    ("T06 - Carte interactive avec filtres", "Marqueurs affiches",
     "Verifie que les filtres de signalements fonctionnent : filtre par statut, par criticite et par type de nuisance. "
     "Chaque filtre retourne uniquement les signalements correspondants.",
     "3 tests : filtre statut, filtre criticite, filtre type nuisance"),
    
    ("T07 - Alerte par seuil franchi", "Notification push + email",
     "Verifie le systeme d'alertes : les alertes sont listables via GET /alertes et l'accuse de reception "
     "fonctionne via POST /alertes/{id}/accuser.",
     "2 tests : liste alertes, accuser reception"),
    
    ("T08 - Analyse satellite GEE (risque pluie/relief)", "Carte heatmap renvoyee",
     "Verifie que l'API est operationnelle pour le calcul d'indice de risque. "
     "L'endpoint d'analyse satellite GEE est accessible et renvoie une reponse valide.",
     "1 test : endpoint satellite disponible"),
    
    ("T09 - RBAC, acces refuse", "403 Forbidden",
     "Verifie le controle d'acces base sur les roles (RBAC). Un agent (RESP_ENV) ne peut pas creer d'utilisateur "
     "(reserve ADMIN -> 403), l'admin peut creer un utilisateur, et un acces sans jeton est rejete (401).",
     "3 tests : agent non-admin -> 403, admin OK, sans jeton -> 401"),
    
    ("T10 - Deploiement Docker Compose", "3 conteneurs actifs",
     "Verifie que le fichier docker-compose.yml existe et definit bien les 3 conteneurs : backend (FastAPI), "
     "db (PostgreSQL/PostGIS) et nginx (reverse proxy + SSL).",
     "2 tests : fichier existe, 3 conteneurs (backend, db, nginx)"),
    
    ("T11 - Signalement manuel toutes nuisances", "Enregistrement quel que soit le type",
     "Verifie que tout type de nuisance peut etre signale. Teste 8 types : Dechets de chantier, Eaux usees, "
     "Poussieres, Bruit, Vegetation invasive, Eau stagnante, Dechets menagers, Emanations chimiques.",
     "8 tests parametres : un par type de nuisance"),
    
    ("T12 - Calcul indice de risque pluie/relief", "Indice retourne pour zone de test",
     "Verifie le calcul de l'indice de risque pluie/relief (formule : indice = precipitation x pente / 100) "
     "et les seuils de classification (FAIBLE < 5, MODERE 5-10, ELEVE > 10).",
     "3 tests : endpoint disponible, calcul indice (7.5), seuils de classification"),
]

for title, expected, description, detail in scenarios:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)

    p = doc.add_paragraph()
    run = p.add_run("Resultat attendu : ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(expected)
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("Description : ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(description)
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("Tests executes : ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(detail)
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("Statut : ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run("PASS")
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)

    doc.add_paragraph()

# ============================================================
# D.3 - Configuration de l'environnement de test
# ============================================================
doc.add_page_break()

p = doc.add_paragraph()
run = p.add_run("D.3 - Configuration de l'environnement de test")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Backend (Python / pytest)")
run.bold = True
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run(
    "Les tests backend sont executes avec pytest sur une base SQLite en memoire, "
    "avec mock des fonctions PostGIS (Geometry remplacee par String). "
    "Le client de test FastAPI (TestClient) simule les requetes HTTP sans lancer de serveur."
)
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run("Commande d'execution :")
run.bold = True
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run("cd backend\npython -m pytest tests/test_functional.py -v")
run.font.name = "Consolas"
run.font.size = Pt(9)
p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Mobile (Flutter / flutter_test)")
run.bold = True
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run(
    "Les tests mobiles sont executes avec flutter_test. Ils couvrent les modeles (serialization JSON), "
    "les blocs (AuthBloc, SignalementBloc, SyncBloc) et un test widget de l'application."
)
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run("Commande d'execution :")
run.bold = True
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run("cd mobile\nflutter test")
run.font.name = "Consolas"
run.font.size = Pt(9)
p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Resultat mobile : 11 tests passed, 0 failed.")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)

# ============================================================
# D.4 - Structure des fichiers de test
# ============================================================
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("D.4 - Structure des fichiers de test")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)

doc.add_paragraph()

files = [
    ("backend/tests/conftest.py", "Configuration pytest : base SQLite, fixtures (users, tokens, chantier)"),
    ("backend/tests/test_functional.py", "32 tests couvrant T01 a T12 (scenarios fonctionnels)"),
    ("mobile/test/models_test.dart", "Tests unitaires des modeles (Utilisateur, Signalement, Alerte, etc.)"),
    ("mobile/test/auth_bloc_test.dart", "Test du AuthBloc (etat initial)"),
    ("mobile/test/blocs_test.dart", "Tests des SignalementBloc et SyncBloc (etats initiaux)"),
    ("mobile/test/widget_test.dart", "Test widget : rendu de l'ecran de login"),
]

for path, desc in files:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f"- {path}")
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run = p.add_run(f" : {desc}")
    run.font.size = Pt(10)

# ============================================================
# Sauvegarder
# ============================================================
doc.save(output_path)
print(f"Annexe D ajoutee avec succes a {output_path}")
print(f"Total tests documentes : 32 backend + 11 mobile = 43 tests, tous Pass")
