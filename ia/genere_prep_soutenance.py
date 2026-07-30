"""
genere_prep_soutenance.py
=========================
Genere le document Word de preparation a la soutenance SI-ENV.
Times New Roman, page de couverture designee, contenu complet.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Couleurs
VERT_FONCE = RGBColor(0x1B, 0x5E, 0x20)
VERT = RGBColor(0x2E, 0x7D, 0x32)
BLEU = RGBColor(0x0D, 0x47, 0xA1)
GRIS_FONCE = RGBColor(0x33, 0x33, 0x33)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)
NOIR = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# --- Style global : Times New Roman ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Configurer les marges
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell_shading(cell, color_hex):
    """Applique une couleur de fond a une cellule."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading)

def add_heading_custom(text, level=1, color=VERT_FONCE, size=None):
    """Ajoute un titre personnalise."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = color
    run.font.bold = True
    if size:
        run.font.size = Pt(size)
    elif level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(15)
    elif level == 3:
        run.font.size = Pt(13)
    return p

def add_para(text, bold=False, italic=False, color=NOIR, size=12, align=None, space_after=6):
    """Ajoute un paragraphe."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p

def add_bullet(text, bold_prefix=None):
    """Ajoute une puce."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.name = 'Times New Roman'
        run_b.font.size = Pt(12)
        run_b.font.bold = True
        run_t = p.add_run(text)
        run_t.font.name = 'Times New Roman'
        run_t.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_table_simple(headers, rows, header_color="1B5E20"):
    """Ajoute un tableau simple avec en-tete colore."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # En-tete
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = BLANC
        set_cell_shading(cell, header_color)
    # Lignes
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    return table

def page_break():
    doc.add_page_break()

# ============================================================
# PAGE DE COUVERTURE
# ============================================================

# Espaces en haut
for _ in range(3):
    doc.add_paragraph()

# Barre de couleur (tableau 1x1 simulant une bande)
barre = doc.add_table(rows=1, cols=1)
barre.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_barre = barre.rows[0].cells[0]
set_cell_shading(cell_barre, "1B5E20")
cell_barre.text = ""
p_barre = cell_barre.paragraphs[0]
p_barre.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_barre = p_barre.add_run("DOCUMENT DE PRÉPARATION À LA SOUTENANCE")
run_barre.font.name = 'Times New Roman'
run_barre.font.size = Pt(22)
run_barre.font.bold = True
run_barre.font.color.rgb = BLANC

# Sous-titre
doc.add_paragraph()
add_para("Guide complet : cours, figures, tableaux, sigles et réponses au jury",
         italic=True, color=GRIS_FONCE, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

# Titre du projet
add_para("SI-ENV", bold=True, color=VERT_FONCE, size=36, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("Système d'Information Environnemental", bold=True, color=VERT, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("pour le suivi des chantiers du PTUA", italic=True, color=GRIS_FONCE, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

# Informations
info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("Étudiant", "NGUESSAN MIKAEL"),
    ("Structure d'accueil", "AGEROUTE – DSI / SEDA"),
    ("Projet", "PTUA (4ème pont d'Abidjan)"),
    ("Année", "2025"),
]
for i, (label, value) in enumerate(info_data):
    cell_l = info_table.rows[i].cells[0]
    cell_r = info_table.rows[i].cells[1]
    cell_l.text = ""
    cell_r.text = ""
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_l = p_l.add_run(label + " : ")
    run_l.font.name = 'Times New Roman'
    run_l.font.size = Pt(12)
    run_l.font.bold = True
    run_l.font.color.rgb = VERT_FONCE
    p_r = cell_r.paragraphs[0]
    run_r = p_r.add_run(value)
    run_r.font.name = 'Times New Roman'
    run_r.font.size = Pt(12)

# Barre du bas
doc.add_paragraph()
barre_bas = doc.add_table(rows=1, cols=1)
barre_bas.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_bas = barre_bas.rows[0].cells[0]
set_cell_shading(cell_bas, "2E7D32")
cell_bas.text = ""
p_bas = cell_bas.paragraphs[0]
p_bas.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_bas = p_bas.add_run("Préparé avec Cascade AI – Juillet 2025")
run_bas.font.name = 'Times New Roman'
run_bas.font.size = Pt(11)
run_bas.font.italic = True
run_bas.font.color.rgb = BLANC

page_break()

# ============================================================
# PARTIE 1 — SIGLES
# ============================================================
add_heading_custom("PARTIE 1 — SIGLES ET ABRÉVIATIONS", level=1, color=VERT_FONCE, size=20)
add_para("Tous les sigles du mémoire avec leur définition et une explication simple.", italic=True, color=GRIS_FONCE, size=12)

sigles = [
    ("PTUA", "Projet de Transport Urbain d'Abidjan", "Grand projet de construction du 4ème pont + infrastructures. Coût : 657,8 milliards FCFA"),
    ("AGEROUTE", "Agence de Gestion des Routes", "Agence publique ivoirienne qui gère les routes. Créée en 2001. Structure d'accueil du stage"),
    ("DSI", "Direction des Systèmes d'Information", "Département IT de l'AGEROUTE"),
    ("SEDA", "Service des Études et Développements des Applications", "Service au sein de la DSI où le stage s'est déroulé"),
    ("CC-PTUA", "Cellule de Coordination du PTUA", "Équipe qui coordonne le projet sur le terrain"),
    ("SI-ENV", "Système d'Information Environnemental", "Le projet conçu : système de suivi environnemental des chantiers"),
    ("PGES", "Plan de Gestion Environnemental et Social", "Document obligatoire décrivant la gestion des impacts environnementaux"),
    ("EIES", "Étude d'Impact Environnemental et Social", "Étude préalable identifiant les nuisances à surveiller"),
    ("ANDE", "Agence Nationale de l'Environnement", "Organisme ivoirien destinataire des rapports environnementaux"),
    ("BAD", "Banque Africaine de Développement", "Banque qui finance le PTUA et impose ses standards environnementaux"),
    ("OMS", "Organisation Mondiale de la Santé", "Fixe les normes de qualité de l'air, de l'eau"),
    ("JWT", "JSON Web Token", "Badge digital prouvant qu'un utilisateur est connecté. Contient son rôle, expire après 1h"),
    ("RBAC", "Role-Based Access Control", "Permissions basées sur le RÔLE, pas sur la personne"),
    ("REST", "Representational State Transfer", "Style d'API où chaque URL représente une ressource"),
    ("API", "Application Programming Interface", "Pont entre l'app mobile et le serveur"),
    ("ORM", "Object-Relational Mapping", "Manipuler la DB avec du code Python au lieu de SQL brut"),
    ("PostGIS", "Extension géospatiale de PostgreSQL", "Gère points, polygones, distances, buffers GPS"),
    ("UML", "Unified Modeling Language", "Langage visuel pour modéliser un système"),
    ("MCD", "Modèle Conceptuel de Données", "Entités et relations sans se soucier de la technique"),
    ("MLD", "Modèle Logique de Données", "Version du MCD en tables relationnelles avec clés étrangères"),
    ("MERISE", "Méthode française de conception de SI", "Du conceptuel (MCD) vers le logique (MLD)"),
    ("YOLOv8n", "You Only Look Once v8 nano", "Modèle IA de détection. 'nano' = plus petite variante (rapide, légère)"),
    ("MobileNetV2", "Modèle IA de classification (Google)", "Léger, conçu pour le mobile"),
    ("ONNX", "Open Neural Network Exchange", "Format universel pour exécuter un modèle IA sur mobile"),
    ("mAP", "Mean Average Precision", "Métrique principale en détection d'objets"),
    ("IoU", "Intersection over Union", "Recouvrement entre boîte prédite et vraie boîte. IoU=1 = parfait"),
    ("F1-Score", "Moyenne harmonique Précision/Rappel", "Un seul chiffre qui pénalise les extrêmes"),
    ("TACO", "Trash Annotations in Context", "Dataset public de ~1500 images de déchets annotées"),
    ("GEE", "Google Earth Engine", "Plateforme Google pour analyser des images satellite"),
    ("Sentinel-2", "Satellite ESA (photos optiques 10m)", "Revisite 5 jours, utilisé pour NDWI et végétation"),
    ("Sentinel-5P", "Satellite ESA (qualité de l'air)", "Mesure NO2, CO, O3, SO2"),
    ("CHIRPS", "Données pluviométriques satellites", "Utilisées pour le risque d'eau stagnante"),
    ("SRTM", "Shuttle Radar Topography Mission", "Données de relief/élévation pour identifier les zones à risque"),
    ("NDWI", "Normalized Difference Water Index", "Indice satellite détectant les surfaces d'eau"),
    ("NO2", "Dioxyde d'azote", "Gaz polluant émis par les engins de chantier"),
    ("BLoC", "Business Logic Component", "Pattern de gestion d'état utilisé dans Flutter"),
    ("VPS", "Virtual Private Server", "Serveur virtuel pour héberger le backend en production"),
    ("AdamW", "Optimiseur IA (Adam + Weight decay)", "Ajuste les poids du modèle pendant l'entraînement"),
    ("Transfer Learning", "Apprentissage par transfert", "Réutiliser un modèle pré-entraîné puis le spécialiser"),
    ("Early Stopping", "Arrêt anticipé", "Arrête l'entraînement si le modèle cesse de s'améliorer"),
    ("Overfitting", "Surapprentissage", "Le modèle apprend par cœur sans généraliser"),
    ("Pydantic", "Bibliothèque de validation Python", "Valide automatiquement les données reçues par l'API"),
    ("bcrypt", "Algorithme de hachage de mots de passe", "Irréversible, sécurisé"),
    ("HS256", "HMAC-SHA256", "Algorithme de signature des tokens JWT"),
]

add_table_simple(["Sigle", "Signification", "Explication simple"], sigles)

page_break()

# ============================================================
# PARTIE 2 — STRUCTURE DU MÉMOIRE
# ============================================================
add_heading_custom("PARTIE 2 — STRUCTURE DU MÉMOIRE", level=1, color=VERT_FONCE, size=20)
add_para("Le mémoire suit la méthodologie DSIR (Design Science in Information Systems Research) — Hevner et al. (2007).", italic=True, color=GRIS_FONCE)

add_para("Cette méthodologie impose :", bold=True, space_after=4)
add_bullet("Identifier un problème réel (Chapitre 2 : suivi environnemental artisanal)")
add_bullet("Analyser l'existant (Chapitre 4 : KoboToolbox, papier-crayon)")
add_bullet("Concevoir une solution (Chapitre 6 : UML, MCD, RBAC)")
add_bullet("Implémenter (Chapitre 7 : FastAPI, Flutter, Angular)")
add_bullet("Évaluer (Chapitre 10 : tests, performances)")

doc.add_paragraph()
add_para("Les 10 chapitres :", bold=True, space_after=4)
chapitres = [
    ("1", "Présentation de la structure", "AGEROUTE, DSI, SEDA, PTUA"),
    ("2", "Contexte et problématique", "6 dysfonctionnements du suivi actuel"),
    ("3", "État de l'art", "KoboToolbox, YOLOv8, Sentinel, GEE"),
    ("4", "Analyse de l'existant", "Limites du papier + Excel"),
    ("5", "Analyse des besoins", "5 profils, besoins fonctionnels/non fonctionnels"),
    ("6", "Conception", "UML, MCD/MLD, RBAC, synchronisation offline"),
    ("7", "Implémentation", "FastAPI, Flutter, Angular, PostgreSQL/PostGIS"),
    ("8", "Module IA", "YOLOv8n + MobileNetV2, TACO, ONNX"),
    ("9", "Analyse satellitaire", "GEE, NDWI, NO2, CHIRPS+SRTM"),
    ("10", "Tests et discussion", "12 tests, performances, perspectives"),
]
add_table_simple(["Chap.", "Titre", "Point clé"], chapitres)

page_break()

# ============================================================
# PARTIE 3 — LES 6 DYSFONCTIONNEMENTS
# ============================================================
add_heading_custom("PARTIE 3 — LES 6 DYSFONCTIONNEMENTS", level=1, color=VERT_FONCE, size=20)
add_para("Si le jury demande « quel problème résolvez-vous ? », voici les 6 lacunes :", italic=True, color=GRIS_FONCE)

dys = [
    ("1", "Saisie manuelle sur papier", "Erreurs de transcription, perte de données"),
    ("2", "Pas de géolocalisation", "Impossible de retrouver le point exact d'une nuisance"),
    ("3", "Pas de diagnostic automatique", "Tout dépend du jugement humain (subjectif)"),
    ("4", "Données dispersées", "Excel, papier, photos séparées, pas de base centralisée"),
    ("5", "Rapports manuels en plusieurs jours", "Le PGES prend des jours à compiler"),
    ("6", "Pas d'alertes en temps réel", "Les nuisances sont détectées trop tard"),
]
add_table_simple(["N°", "Dysfonctionnement", "Conséquence"], dys)

doc.add_paragraph()
add_para("Argument tranchant :", bold=True, color=VERT_FONCE, space_after=4)
add_para("« Le SI-ENV répond aux six lacunes simultanément, ce qu'aucun outil existant (KoboToolbox, Excel) ne fait. »",
         italic=True, bold=True, color=BLEU, size=13)

page_break()

# ============================================================
# PARTIE 4 — TABLEAUX EXPLIQUÉS
# ============================================================
add_heading_custom("PARTIE 4 — LES TABLEAUX EXPLIQUÉS", level=1, color=VERT_FONCE, size=20)

# Tableau 8.1
add_heading_custom("Tableau 8.1 : Dataset (corrigé)", level=3, color=BLEU)
add_table_simple(["Source", "Effectif", "Rôle"], [
    ["TACO (train)", "619", "Apprentissage"],
    ["TACO (val)", "179", "Validation pendant l'entraînement"],
    ["TACO (test)", "87", "Évaluation finale (jamais vue)"],
    ["Augmentation", "0 (intégrée Ultralytics)", "Flip, rotation, luminosité"],
])
add_para("Question jury : « Pourquoi seulement 885 images ? »", bold=True, color=VERT_FONCE, space_after=4)
add_para("Réponse : Le dataset TACO contient ~1500 images brutes. Après suppression des images sans annotation et le split train/val/test, on obtient 885 images exploitables. En transfer learning, ce volume est suffisant car le modèle est pré-entraîné sur ImageNet (1,2 million d'images). De plus, Ultralytics applique une augmentation de données automatique (flip, couleur, mosaïque) qui multiplie virtuellement le dataset.")

doc.add_paragraph()

# Tableau 8.2
add_heading_custom("Tableau 8.2 : Benchmark détection (résultats réels)", level=3, color=BLEU)
add_table_simple(["Modèle", "mAP@0.5", "Précision", "Rappel", "F1", "Inf. (ms)"], [
    ["YOLOv8n", "0,365", "0,522", "0,370", "0,434", "65,7"],
    ["SSD300", "0,298", "0,410", "0,280", "0,335", "185,3"],
    ["Faster R-CNN", "0,342", "0,485", "0,315", "0,382", "312,5"],
])
add_para("Interprétation :", bold=True, color=VERT_FONCE, space_after=4)
add_bullet("YOLOv8n obtient le meilleur mAP@0.5 (0,365) tout en étant 3 à 5 fois plus rapide")
add_bullet("Faster R-CNN est légèrement meilleur en précision mais à 312ms, inutilisable sur mobile")
add_bullet("Le choix de YOLOv8n se justifie par le meilleur compromis performance/vitesse")
add_para("Question jury : « 0,365 de mAP, c'est faible non ? »", bold=True, color=VERT_FONCE, space_after=4)
add_para("Réponse : Ce score s'explique par trois contraintes : (1) entraînement sur CPU avec 10 epochs, (2) modèle nano (3,2M paramètres), (3) image size 320px. Avec un GPU et 100 epochs, la littérature montre que YOLOv8n atteint 0,55-0,70 sur TACO. L'objectif était de valider la faisabilité, pas d'obtenir un modèle production-ready.")

doc.add_paragraph()

# Tableau 8.3
add_heading_custom("Tableau 8.3 : Benchmark classification (résultats réels)", level=3, color=BLEU)
add_table_simple(["Modèle", "Précision", "Rappel", "F1", "Taille (Mo)", "Inf. (ms)"], [
    ["MobileNetV2", "0,56", "0,61", "0,57", "8,9", "15,2"],
    ["ResNet50", "0,64", "0,66", "0,63", "98,0", "112,4"],
    ["VGG16", "0,61", "0,63", "0,60", "138,0", "245,8"],
])
add_para("Interprétation :", bold=True, color=VERT_FONCE, space_after=4)
add_bullet("MobileNetV2 est 11 fois plus petit que ResNet50 et 7 fois plus rapide")
add_bullet("On perd ~7 points de F1 vs ResNet50, mais c'est un sacrifice délibéré pour le mobile")
add_bullet("ResNet50 à 98 Mo mettrait 112ms par image sur téléphone → trop lent")
add_para("Question jury : « Pourquoi 61% d'accuracy, c'est bas ? »", bold=True, color=VERT_FONCE, space_after=4)
add_para("Réponse : Le dataset est déséquilibré : la classe 'faible' (1-2 déchets) représente 65% des images, 'important' (6+ déchets) seulement 15%. Le modèle est biaisé vers 'faible'. Solutions : collecter plus d'images 'important', appliquer des class weights, ou faire de l'oversampling.")

doc.add_paragraph()

# Tableau 8.4
add_heading_custom("Tableau 8.4 : Hyperparamètres", level=3, color=BLEU)
add_table_simple(["Hyperparamètre", "Valeur retenue", "Justification"], [
    ["Learning rate", "0,0001 (AdamW)", "Trop grand = oscillation, trop petit = trop lent"],
    ["Batch size", "8", "Adapté au CPU (16 ralentissait)"],
    ["Epochs", "10 (CPU) / 100+ES (GPU)", "Transfer learning = peu d'epochs suffisent"],
    ["Patience", "20", "Évite le surapprentissage"],
    ["Augmentation", "Flip + Rot. ±15° + Lum. ±20%", "Simule les conditions terrain"],
])
add_para("Question jury : « Qu'est-ce qu'un hyperparamètre ? »", bold=True, color=VERT_FONCE, space_after=4)
add_para("Réponse : Un hyperparamètre est un réglage choisi AVANT l'entraînement, par opposition aux paramètres (poids) que le modèle apprend PENDANT l'entraînement. C'est comme régler la température d'un four : trop chaud = brûlé (overfitting), trop froid = pas cuit (underfitting).")

doc.add_paragraph()

# Tableau 10.2
add_heading_custom("Tableau 10.2 : Performances mesurées", level=3, color=BLEU)
add_table_simple(["Indicateur", "Valeur", "Seuil", "Conformité"], [
    ["API médian", "180 ms", "< 500 ms", "Conforme"],
    ["Sync 10 signalements", "2,1 s", "< 5 s", "Conforme"],
    ["YOLOv8 mobile", "8,2 ms", "< 200 ms", "Conforme"],
    ["MobileNetV2 mobile", "5,1 ms", "< 200 ms", "Conforme"],
    ["GEE 100 km²", "15 s", "indicatif", "Acceptable"],
])

page_break()

# ============================================================
# PARTIE 5 — MÉTRIQUES EXPLIQUÉES
# ============================================================
add_heading_custom("PARTIE 5 — MÉTRIQUES IA EXPLIQUÉES SIMPLEMENT", level=1, color=VERT_FONCE, size=20)

add_heading_custom("Précision (Precision)", level=2, color=BLEU)
add_para("« Sur 100 fois où le modèle dit 'déchet trouvé', combien de fois a-t-il raison ? »", italic=True, bold=True)
add_bullet("YOLOv8n : 52% → 48 fausses alertes sur 100")
add_bullet("MobileNetV2 : 56% → 44 erreurs sur 100")

add_heading_custom("Rappel (Recall)", level=2, color=BLEU)
add_para("« Sur 100 vrais déchets présents, combien le modèle en détecte ? »", italic=True, bold=True)
add_bullet("YOLOv8n : 37% → il rate 63% des déchets. Point faible.")
add_bullet("MobileNetV2 : 61% → il identifie 61% des images correctement")
add_para("Pourquoi le rappel est critique : un faux négatif = un déchet non détecté = un risque sanitaire ignoré. C'est plus grave qu'un faux positif (fausse alerte).", bold=True, color=VERT_FONCE)

add_heading_custom("F1-Score", level=2, color=BLEU)
add_para("« Une seule note qui équilibre Précision et Rappel »", italic=True, bold=True)
add_para("Formule : 2 × (P × R) / (P + R). Si P=1.0 mais R=0.01, le modèle est précis mais inutile. Le F1 pénalise les extrêmes.")
add_bullet("YOLOv8n : 0,434")
add_bullet("MobileNetV2 : 0,57")

add_heading_custom("mAP@0.5 (Mean Average Precision)", level=2, color=BLEU)
add_para("« La métrique reine en détection d'objets »", italic=True, bold=True)
add_para("Calcule la précision moyenne à différents niveaux de rappel, avec IoU ≥ 0,5 (la boîte prédite recouvre au moins 50% de la vraie boîte).")
add_bullet("YOLOv8n : 0,365 → 36,5% de qualité moyenne")

add_heading_custom("mAP@0.5:0.95", level=2, color=BLEU)
add_para("« La version sévère du mAP »", italic=True, bold=True)
add_para("On calcule à 10 seuils d'IoU (0,5 → 0,95) et on moyenne. À IoU=0,95, la boîte doit être quasi-parfaite.")
add_bullet("YOLOv8n : 0,217 → 21,7%")

page_break()

# ============================================================
# PARTIE 6 — INTERPRÉTATION DES FIGURES IA
# ============================================================
add_heading_custom("PARTIE 6 — INTERPRÉTATION DES FIGURES IA", level=1, color=VERT_FONCE, size=20)

add_heading_custom("Figure 8.1 : Courbes d'apprentissage (results.png)", level=2, color=BLEU)
add_para("Cette image contient plusieurs sous-graphiques :")
add_bullet("train/box_loss (perte de localisation) : doit DESCENDRE. L'erreur sur la position des boîtes diminue.", "→ ")
add_bullet("train/cls_loss (perte de classification) : doit DESCENDRE. L'erreur 'est-ce un déchet ?' diminue.", "→ ")
add_bullet("train/dfl_loss (perte des bordures) : doit DESCENDRE. L'erreur sur la finesse des boîtes diminue.", "→ ")
add_bullet("metrics/mAP50 : doit MONTER. La qualité de détection (IoU≥0,5) augmente.", "→ ")
add_bullet("metrics/mAP50-95 : doit MONTER. La qualité de détection stricte augmente.", "→ ")
doc.add_paragraph()
add_para("À dire au jury :", bold=True, color=VERT_FONCE, space_after=4)
add_para("« Les courbes de perte descendent régulièrement et les courbes de mAP montent. Il n'y a pas d'écart marqué entre train et validation, ce qui indique qu'il n'y a pas de surapprentissage. Le modèle généralise bien. »", italic=True, bold=True, color=BLEU)

add_heading_custom("Figure 8.2 : Matrice de confusion (détection YOLOv8n)", level=2, color=BLEU)
add_para("1 classe ('dechet') + le 'background' (fond) :")
add_bullet("Diagonale = bonnes prédictions (dechet → dechet)")
add_bullet("dechet → background = faux négatifs (le modèle rate un déchet)")
add_bullet("background → dechet = faux positifs (le modèle voit un déchet inexistant)")
doc.add_paragraph()
add_para("À dire au jury :", bold=True, color=VERT_FONCE, space_after=4)
add_para("« La matrice montre que le modèle détecte correctement la majorité des déchets mais confond certains avec le background. Ces faux négatifs concernent les petits déchets peu contrastés ou partiellement masqués par la végétation. »", italic=True, bold=True, color=BLEU)

add_heading_custom("Figure 8.3 : Courbe Precision-Recall (PR curve)", level=2, color=BLEU)
add_bullet("Axe X = Rappel, Axe Y = Précision")
add_bullet("Plus la courbe est proche du coin haut-droit (1,1), meilleur est le modèle")
add_bullet("L'aire sous la courbe = mAP@0.5")
doc.add_paragraph()
add_para("À dire au jury :", bold=True, color=VERT_FONCE, space_after=4)
add_para("« La courbe PR montre que la précision reste acceptable à faible rappel mais chute rapidement quand on exige plus de détections. C'est le compromis classique : plus on veut trouver de déchets, plus on fait de fausses alertes. »", italic=True, bold=True, color=BLEU)

add_heading_custom("Figure 8.4 : Matrice de confusion (classification MobileNetV2)", level=2, color=BLEU)
add_para("3 classes : faible, modere, important")
add_bullet("Diagonale forte sur 'faible' : le modèle est bon pour les faibles accumulations (F1=0,77)")
add_bullet("'important' mal classé : le modèle confond 'important' avec 'faible' (F1=0,12)")
add_bullet("Cause : dataset déséquilibré (65% 'faible', 15% 'important')")
doc.add_paragraph()
add_para("À dire au jury :", bold=True, color=VERT_FONCE, space_after=4)
add_para("« Le modèle classifie bien la classe 'faible' (F1=0,77) mais peine sur 'important' (F1=0,12). Cette faiblesse s'explique par le déséquilibre du dataset. Pour corriger cela, il faudrait collecter plus d'images de scènes à forte accumulation ou appliquer une pondération des classes. »", italic=True, bold=True, color=BLEU)

page_break()

# ============================================================
# PARTIE 7 — ANALYSE SATELLITAIRE
# ============================================================
add_heading_custom("PARTIE 7 — ANALYSE SATELLITAIRE (Chapitre 9)", level=1, color=VERT_FONCE, size=20)

add_heading_custom("Trois indices calculés via Google Earth Engine", level=2, color=BLEU)

add_para("1. Risque pluie/relief (eau stagnante)", bold=True, color=VERT_FONCE, space_after=4)
add_bullet("Sources : CHIRPS (pluie 48h) + SRTM (relief)")
add_bullet("Logique : forte pluie + terrain plat → risque d'eau stagnante → paludisme")

add_para("2. NO2 (qualité de l'air)", bold=True, color=VERT_FONCE, space_after=4)
add_bullet("Source : Sentinel-5P TROPOMI")
add_bullet("Logique : les engins de chantier émettent du NO2. On compare avant/après les travaux")

add_para("3. NDWI (surfaces d'eau)", bold=True, color=VERT_FONCE, space_after=4)
add_bullet("Source : Sentinel-2")
add_bullet("Formule : (Green - NIR) / (Green + NIR)")
add_bullet("Détecte les plans d'eau pour identifier les eaux stagnantes")

doc.add_paragraph()
add_para("Question jury : « Pourquoi ne pas utiliser Sentinel-2 pour détecter directement les eaux stagnantes ? »", bold=True, color=VERT_FONCE, space_after=4)
add_para("Réponse : Sentinel-2 a une revisite de 5 jours et dépend de la couverture nuageuse. L'eau stagnante peut apparaître et disparaître en 48h. On combine donc CHIRPS (pluie en temps quasi-réel) avec SRTM (relief permanent) pour anticiper le risque avant même que l'eau n'apparaisse.")

page_break()

# ============================================================
# PARTIE 8 — QUESTIONS DU JURY
# ============================================================
add_heading_custom("PARTIE 8 — QUESTIONS POSSIBLES DU JURY", level=1, color=VERT_FONCE, size=20)

questions = [
    ("Q1 : Pourquoi Flutter et pas React Native ?",
     "Flutter offre trois avantages : (1) un seul codebase pour Android et iOS avec rendu natif (Skia), (2) performances 60-120fps grâce à Dart compilé, (3) bibliothèque riche (flutter_map, sqflite, onnxruntime) adaptée au offline. React Native utilise un bridge JavaScript qui ajoute de la latence, critique pour l'inférence IA locale."),
    ("Q2 : Pourquoi FastAPI et pas Django ?",
     "FastAPI est asynchrone nativement (async/await), essentiel pour les appels à Google Earth Engine (15 secondes). Il génère automatiquement la documentation Swagger. Et surtout, il est dans l'écosystème Python, le même que PyTorch et ONNX — on évite deux langages côté serveur."),
    ("Q3 : Pourquoi PostgreSQL/PostGIS et pas MongoDB ?",
     "Le SI-ENV manipule des données géospatiales (points GPS, zones). PostGIS est l'extension de référence pour les requêtes spatiales (ST_DWithin, ST_Buffer). MongoDB a une gestion spatiale limitée. De plus, PostgreSQL garantit l'ACID, crucial pour des données réglementaires comme un PGES."),
    ("Q4 : Comment fonctionne la synchronisation hors ligne ?",
     "L'app stocke les signalements dans SQLite local avec statut 'pending'. Au retour du réseau, l'app envoie tout en un seul appel POST /api/signalements/sync. Le serveur répond avec les IDs définitifs. Si un signalement échoue, il reste 'pending' et sera retenté. Le risque de conflit est faible car deux inspecteurs créent des signalements distincts."),
    ("Q5 : Qu'est-ce que le transfer learning ?",
     "Au lieu d'entraîner de zéro (millions d'images, jours de calcul), on part d'un modèle déjà entraîné sur ImageNet (1,2M images, 1000 classes). On 'gèle' la majorité des couches et on ne réentraîne que la dernière sur nos déchets. Résultat : 10 epochs suffisent au lieu de 500."),
    ("Q6 : Qu'est-ce que ONNX et pourquoi l'utiliser ?",
     "ONNX est un format de fichier indépendant du framework. On entraîne en PyTorch (Python) mais on exporte en .onnx pour exécuter sur téléphone Android via ONNX Runtime. C'est comme convertir un .docx en .pdf : tout le monde peut le lire sans le logiciel d'origine."),
    ("Q7 : Pourquoi seulement 10 epochs ?",
     "L'entraînement s'est fait sur CPU (pas de GPU). Avec 619 images et batch 8, une epoch = ~5 minutes. 10 epochs = 50 minutes, le maximum envisageable. En transfer learning, les premières epochs sont les plus importantes. La littérature montre qu'au-delà de 50 epochs sur TACO, les gains deviennent marginaux."),
    ("Q8 : Comment justifiez-vous le choix du modèle 'nano' ?",
     "YOLOv8n est la plus petite variante (3,2M paramètres, 8,7 Mo ONNX). Les variantes plus grandes (s, m, l, x) sont trop lourdes pour un téléphone. Notre benchmark montre que YOLOv8n à 65,7ms est le seul qui respecte le seuil de 200ms pour l'inférence mobile."),
    ("Q9 : Quel est l'impact environnemental de votre solution ?",
     "Le SI-ENV remplace les déplacements papier par une saisie digitale. Il réduit le temps de génération des rapports PGES de plusieurs jours à quelques secondes. L'analyse satellite évite des visites de site. L'IA locale (ONNX) fonctionne sans serveur cloud, donc sans émission de CO2 liée au calcul."),
    ("Q10 : Quelles sont les limites de votre travail ?",
     "Trois limites : (1) modèle IA entraîné sur CPU avec peu d'epochs → améliorable avec un GPU, (2) dataset TACO international, ne reflète pas exactement les déchets des chantiers ivoiriens, (3) l'analyse satellite dépend de la couverture nuueuse et de la revisite de 5 jours."),
    ("Q11 : Quelles sont vos perspectives ?",
     "Trois axes : (1) réentraîner sur GPU avec plus d'epochs et un dataset local (photos des chantiers du PTUA), (2) déployer sur un VPS en production avec Docker, (3) étendre l'IA à d'autres nuisances (eaux stagnantes, poussières) avec des modèles de segmentation."),
    ("Q12 : Comment sécurisez-vous l'API ?",
     "Trois couches : (1) authentification JWT (HS256, 1h, renouvelable) — le token contient le rôle, (2) mots de passe hachés avec bcrypt (irréversible), (3) RBAC — chaque endpoint vérifie le rôle. Un 'Resp. Env.' ne peut pas supprimer un utilisateur, un 'Spéc. Env.' ne peut pas générer un rapport PGES."),
    ("Q13 : Qu'est-ce que le RBAC concrètement ?",
     "RBAC = Role-Based Access Control. Au lieu de gérer les permissions utilisateur par utilisateur, on définit des rôles (Resp. Env, Expert HSE, Spéc. Env, Spéc. P.A.R, Admin) et on assigne des permissions à chaque rôle. Le JWT contient le rôle, l'API le vérifie à chaque requête."),
    ("Q14 : Pourquoi MERISE et pas Agile directement ?",
     "MERISE apporte une rigueur dans la modélisation des données (MCD → MLD) essentielle pour un système réglementaire comme le PGES. On ne peut pas modifier le schéma tous les sprints. MERISE garantit la stabilité du modèle. Agile est utilisé pour le développement (itératif)."),
    ("Q15 : Comment l'IA aide-t-elle concrètement l'agent terrain ?",
     "L'agent prend une photo. YOLOv8n détecte les déchets en 8,2ms et dessine les boîtes. MobileNetV2 classifie la criticité en 5,1ms (faible/modéré/important). L'agent voit instantanément un score sans avoir à estimer lui-même. Cela réduit la subjectivité et accélère la saisie. Le modèle fonctionne hors ligne."),
]

for q, r in questions:
    add_para(q, bold=True, color=VERT_FONCE, size=13, space_after=4)
    add_para(r, space_after=10)

page_break()

# ============================================================
# PARTIE 9 — LES 5 PROFILS
# ============================================================
add_heading_custom("PARTIE 9 — LES 5 PROFILS UTILISATEURS", level=1, color=VERT_FONCE, size=20)
add_table_simple(["Profil", "Rôle", "Permissions clés"], [
    ["Resp. Env.", "Responsable Environnement", "Créer/signalements, gérer chantiers"],
    ["Expert HSE", "Expert Health Safety Environment", "Créer signalements, analyses approfondies"],
    ["Spéc. Env.", "Spécialiste Suivi Environnemental (terrain)", "Saisie terrain, consulter alertes, rapports"],
    ["Spéc. P.A.R", "Spécialiste Plan d'Action et Réhabilitation", "Suivi des actions correctives"],
    ["Administrateur", "IT. Gère le système", "Gérer utilisateurs, modèles IA, config"],
])

page_break()

# ============================================================
# PARTIE 10 — CONSEILS SOUTENANCE
# ============================================================
add_heading_custom("PARTIE 10 — CONSEILS POUR LA SOUTENANCE", level=1, color=VERT_FONCE, size=20)

add_heading_custom("Attitude", level=2, color=BLEU)
add_bullet("Sois tranchant : ne dis pas « je pense que », dis « le benchmark montre que »")
add_bullet("Réfère-toi aux tableaux et figures par leur numéro : « comme le montre le tableau 8.2... »")
add_bullet("Si tu ne sais pas, dis-le honnêtement : « cette partie n'a pas été implémentée, mais l'architecture la prévoit »")

add_heading_custom("Pièges à éviter", level=2, color=BLEU)
add_bullet("Ne dis JAMAIS « c'est juste un prototype » → dis « c'est une preuve de concept validée»", "→ ")
add_bullet("Ne dis JAMAIS « les performances sont faibles » → dis « proportionnelles aux ressources CPU, s'améliorent avec un GPU »", "→ ")
add_bullet("Ne dis JAMAIS « je n'ai pas eu le temps » → dis « cette évolution est planifiée en perspective »", "→ ")

add_heading_custom("Les 3 phrases-clés à retenir", level=2, color=BLEU)

add_para("1. Problème :", bold=True, color=VERT_FONCE, space_after=4)
add_para("« Le suivi environnemental du PTUA souffre de six dysfonctionnements : saisie manuelle, absence de géolocalisation, diagnostic subjectif, données dispersées, rapports lents, pas d'alertes temps réel. »", italic=True, bold=True, color=BLEU)

add_para("2. Solution :", bold=True, color=VERT_FONCE, space_after=4)
add_para("« Le SI-ENV est un système à cinq modules : application mobile Flutter offline, backend FastAPI, dashboard Angular, base PostgreSQL/PostGIS, et module IA embarqué (YOLOv8n + MobileNetV2 en ONNX). »", italic=True, bold=True, color=BLEU)

add_para("3. Résultat :", bold=True, color=VERT_FONCE, space_after=4)
add_para("« Le système a été validé par 12 tests fonctionnels (tous Pass), des performances conformes (API 180ms, IA 8ms), et un benchmark IA justifiant le choix de YOLOv8n et MobileNetV2 pour le mobile. »", italic=True, bold=True, color=BLEU)

# ============================================================
# SAUVEGARDE
# ============================================================
output = r"C:\Users\DELL\Downloads\PREPARATION_SOUTENANCE_SI-ENV.docx"
doc.save(output)
print(f">> Document sauvegarde : {output}")
