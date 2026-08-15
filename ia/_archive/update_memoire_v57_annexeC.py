# -*- coding: utf-8 -*-
"""
Ajoute le contenu de l'Annexe C (deja presente comme titre vide en v57) :
deux captures reelles de l'onglet "Evolution temporelle" du dashboard
(NO2 et risque pluie/relief, moyennes avant/pendant/apres travaux calculees
a partir de donnees Google Earth Engine reelles, capturees le 31/07/2026).
Suit la convention deja utilisee pour les Annexes A et B (pas de legende
numerotee "Figure X.X" separee, juste une ligne de description en gras
avant chaque image, coherent avec le reste du document).
"""
from docx import Document
from docx.shared import Inches

SRC = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v57.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v58.docx"
SCRATCH = r"C:\Users\DELL\AppData\Local\Temp\claude\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25\scratchpad"

doc = Document(SRC)

# Retrouver le placeholder de fin de document (dernier paragraphe)
placeholder = doc.paragraphs[-1]
assert "coller" in placeholder.text.lower(), f"Placeholder inattendu : {placeholder.text!r}"

def insert_label_and_image(label_text, image_path, width_inches=5.5):
    label_para = placeholder.insert_paragraph_before()
    r = label_para.add_run(label_text)
    r.bold = True
    img_para = placeholder.insert_paragraph_before()
    img_para.alignment = 1  # centre
    run = img_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    spacer = placeholder.insert_paragraph_before()
    print(f"[OK] Ajoute : {label_text}")

insert_label_and_image(
    "NO2 - moyennes avant/pendant/apres travaux (Sentinel-5P/TROPOMI, chantier Rocade Y4, "
    "janvier 2022 - juillet 2026).",
    f"{SCRATCH}\\annexe_c1_no2_serie.png"
)
insert_label_and_image(
    "Risque pluie/erosion - moyennes avant/pendant/apres travaux (CHIRPS + SRTM, chantier "
    "Rocade Y4, janvier 2022 - juillet 2026).",
    f"{SCRATCH}\\annexe_c2_risque_serie.png"
)

doc.save(DST)
print(f"\n=== MEMOIRE MIS A JOUR (Annexe C) : {DST} ===")
