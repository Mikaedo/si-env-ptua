# -*- coding: utf-8 -*-
"""
Complete la LISTE DES FIGURES du memoire avec les figures des annexes.

Les figures des annexes ont recu une numerotation (A.1 a D.5) pour respecter la
regle de forme « numerotation obligatoire ». Elles doivent donc apparaitre dans
la liste des figures : une figure numerotee et citee dans le texte, mais absente
de la liste, serait introuvable pour le lecteur.

Elles sont regroupees sous un intitule distinct : les annexes se situent hors du
corps du memoire, on les signale sans les melanger aux figures des chapitres.

Sortie : MEMOIRE_N'GUESSAN_v11.docx
"""
import copy
import os
import re
import shutil

from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph

DOSSIER = r"C:\Users\DELL\Downloads\MEMOIRE"
SOURCE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v10.docx")
SORTIE = os.path.join(DOSSIER, "MEMOIRE_N'GUESSAN_v11.docx")

#: Entrees a ajouter, dans l'ordre. Libelles alignes sur le style concis des
#: entrees existantes (titre de la figure, sans la phrase explicative complete).
ENTREES = [
    "Figure A.1 : Extrait de code - hyperparametres et parametres d'entrainement",
    "Figure B.1 : Extrait de code - calcul des metriques de detection",
    "Figure B.2 : Extrait de code - calcul des metriques de classification",
    "Figure C.1 : Concentrations moyennes de NO2 avant, pendant et apres travaux",
    "Figure C.2 : Indice de risque pluie/erosion avant, pendant et apres travaux",
    "Figure D.1 : Test T01 - Authentification JWT",
    "Figure D.2 : Test T06 - Carte interactive avec filtres",
    "Figure D.3 : Test T07 - Alertes par franchissement de seuil",
    "Figure D.4 : Test T08 - Analyse satellitaire Google Earth Engine",
    "Figure D.5 : Test T05 - Generation du rapport PGES",
]

SOUS_TITRE = "Figures en annexe"


def main():
    shutil.copy2(SOURCE, SORTIE)
    doc = Document(SORTIE)

    # Reperer la liste des figures et sa derniere entree
    debut = None
    for i, p in enumerate(doc.paragraphs[:140]):
        if p.text.strip().upper().startswith('LISTE DES FIGURES'):
            debut = i
            break
    if debut is None:
        raise SystemExit("Section LISTE DES FIGURES introuvable")

    derniere = None
    for i in range(debut + 1, min(debut + 60, len(doc.paragraphs))):
        t = doc.paragraphs[i].text.strip()
        if re.match(r'^Figure\s+\d+\.\d+', t):
            derniere = i
        elif t.upper().startswith('LISTE DES TABLEAUX'):
            break
    if derniere is None:
        raise SystemExit("Aucune entree de figure trouvee dans la liste")

    modele = doc.paragraphs[derniere]
    print("  liste des figures : p%d | derniere entree : p%d" % (debut, derniere))
    print("  modele : %s" % modele.text.strip()[:56])

    # On insere dans l'ordre inverse juste apres la derniere entree, de sorte
    # que l'ordre final soit celui de la liste ENTREES.
    def inserer(texte, gras=False):
        element = copy.deepcopy(modele._element)
        for enfant in list(element):
            if not enfant.tag.endswith('}pPr'):
                element.remove(enfant)
        modele._element.addnext(element)
        par = Paragraph(element, modele._parent)
        run = par.add_run(texte)
        run.bold = gras
        run.font.size = Pt(12)
        return par

    for texte in reversed(ENTREES):
        inserer(texte)
    # Le sous-titre est insere en dernier : il se retrouve ainsi en tete du bloc
    inserer(SOUS_TITRE, gras=True)
    inserer("")  # ligne de separation avec les figures du corps

    doc.save(SORTIE)

    # Verification
    verif = Document(SORTIE)
    txt = "\n".join(p.text for p in verif.paragraphs[:150])
    presentes = [e.split(' :')[0] for e in ENTREES
                 if e.split(' :')[0] in txt]
    print("\n  entrees d'annexe presentes dans la liste : %d / %d"
          % (len(presentes), len(ENTREES)))
    print("  sous-titre present : %s" % (SOUS_TITRE in txt))
    print("\nEnregistre : %s" % SORTIE)
    print("Paragraphes : %d | Tableaux : %d"
          % (len(verif.paragraphs), len(verif.tables)))


if __name__ == "__main__":
    main()
