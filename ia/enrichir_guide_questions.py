# -*- coding: utf-8 -*-
"""
Ajoute au guide les questions de fondamentaux qui manquaient.

Le guide couvrait le fonctionnement du systeme mais pas les notions sur
lesquelles il repose. Un jury de licence interroge volontiers le socle : quels
types d'apprentissage existent, ce qu'est le surapprentissage, d'ou sortent des
seuils. Ne pas savoir repondre a une question de cours apres avoir expose un
travail avance laisse une impression facheuse, celle d'un resultat mal
maitrise.

Les reponses sont ancrees sur les parametres reellement employes, lus dans
1_entrainer_detection.py et 2_entrainer_classification.py : transfert depuis
YOLOv8n et MobileNetV2 pre-entraines, dix epoques, lots de huit et seize, taux
d'apprentissage de 0,0001, decoupage a quatre-vingts et vingt pour cent, images
reduites a 320 pixels, aucune augmentation de donnees. Une reponse generique se
repere immediatement ; une reponse chiffree sur ses propres choix, non.

Deux formulations sont par ailleurs corrigees. Le guide decrivait le NDWI comme
un indice d'humidite des sols et le quatrieme comme un risque d'erosion, la ou
le memoire dit teneur en eau du couvert vegetal et risque d'eau stagnante. Un
jury qui compare les deux documents verrait l'ecart.
"""
import copy
import shutil
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\GUIDE_PREPARATION_SOUTENANCE_SI-ENV.docx")
SAUVEGARDE = SOURCE.with_name("GUIDE_avant_fondamentaux.docx")

CORRECTIONS = [
    ("le NDWI pour l'humidité des sols",
     "le NDWI pour la teneur en eau du couvert végétal"),
    ("un indice de risque d'érosion combinant",
     "un indice de risque d'eau stagnante combinant"),
    ("C'est un indice d'humidité de surface.",
     "C'est un indice de teneur en eau du couvert végétal."),
]

# Questions ajoutees en fin de la section consacree a l'intelligence
# artificielle, avant le titre de la section suivante.
IA = [
    ("Q. Quels types d'apprentissage existent, et lequel avez-vous utilisé ?",
     "On en distingue classiquement quatre. L'apprentissage supervisé, où chaque "
     "exemple porte la réponse attendue. Le non supervisé, où l'on cherche des "
     "regroupements sans étiquette. Le semi-supervisé, qui mêle une petite part "
     "de données étiquetées à beaucoup de non étiquetées. Et l'apprentissage par "
     "renforcement, où un agent apprend par essais et récompenses. "
     "Mes deux modèles relèvent du supervisé : chaque image porte soit des boîtes "
     "annotées avec leur classe de déchet, soit une étiquette de criticité. C'est "
     "le seul cadre qui permette de mesurer une précision, puisqu'il faut une "
     "vérité de référence à laquelle se comparer."),

    ("Q. Qu'est-ce que l'apprentissage par transfert, et pourquoi l'avoir choisi ?",
     "Je ne suis pas parti de zéro. YOLOv8n arrive déjà entraîné sur COCO, un jeu "
     "généraliste, et MobileNetV2 sur ImageNet. Ces modèles savent déjà "
     "reconnaître des contours, des textures, des formes : je n'ai eu qu'à les "
     "spécialiser sur mes classes, en remplaçant la dernière couche du "
     "classifieur par une sortie à trois niveaux de criticité. "
     "Deux raisons à ce choix. Le volume d'abord : mon corpus se compte en "
     "milliers d'images, très loin de ce qu'exigerait un entraînement complet. "
     "Les moyens ensuite : j'ai entraîné sur processeur, sans carte graphique. "
     "Partir de zéro aurait demandé des semaines pour un résultat inférieur."),

    ("Q. Quelle différence faites-vous entre détection et classification ?",
     "La classification répond à la question « qu'y a-t-il sur cette image », la "
     "détection à « qu'y a-t-il, et où ». La détection produit des boîtes "
     "englobantes avec leurs coordonnées, la classification une seule étiquette "
     "pour l'image entière. "
     "Chez moi les deux s'enchaînent : YOLOv8n détecte et localise les déchets, "
     "puis MobileNetV2 classe la scène en trois niveaux de criticité. C'est aussi "
     "pourquoi les métriques diffèrent, le mAP n'ayant de sens qu'en détection."),

    ("Q. Pourquoi un réseau convolutif plutôt qu'un réseau classique ?",
     "Parce qu'un pixel n'a de sens qu'avec ses voisins. La convolution fait "
     "glisser un petit filtre sur l'image et repère des motifs locaux, des bords "
     "d'abord, puis des formes, puis des objets à mesure qu'on empile les "
     "couches. "
     "L'autre raison tient au partage des poids : le même filtre sert partout "
     "dans l'image. Un réseau classique demanderait un poids par pixel et par "
     "neurone, ce qui serait ingérable, et il n'apprendrait pas qu'un déchet "
     "reste un déchet quel que soit l'endroit où il se trouve."),

    ("Q. Comment avez-vous séparé vos données, et pourquoi est-ce important ?",
     "Quatre-vingts pour cent pour l'entraînement, vingt pour cent pour la "
     "validation. La règle absolue est qu'aucune image de validation n'ait servi "
     "à l'entraînement : sans cela la mesure ne vaut rien, le modèle récitant au "
     "lieu de généraliser. "
     "Un point d'honnêteté : avec un jeu de validation de cette taille, les "
     "intervalles de confiance restent larges. Je les ai estimés par bootstrap "
     "plutôt que d'annoncer un chiffre unique comme s'il était exact."),

    ("Q. Qu'est-ce que le surapprentissage, et en souffrez-vous ?",
     "C'est le moment où le modèle apprend les exemples par cœur au lieu d'en "
     "tirer une règle. On le repère à un écart qui se creuse : l'erreur continue "
     "de baisser sur l'entraînement pendant qu'elle remonte sur la validation. "
     "Mes dix époques limitent mécaniquement ce risque, et le transfert aussi, "
     "l'essentiel des poids venant d'un entraînement sur des millions d'images. "
     "À vrai dire, avec dix époques seulement, mon risque est plutôt l'inverse, "
     "un sous-apprentissage. C'est d'ailleurs un de mes leviers d'amélioration."),

    ("Q. Que valent une époque, un lot et un taux d'apprentissage chez vous ?",
     "Une époque est un passage complet sur le jeu d'entraînement : j'en ai fait "
     "dix. Un lot est le nombre d'images traitées ensemble avant de corriger le "
     "modèle : huit en détection, seize en classification. Le taux "
     "d'apprentissage règle l'ampleur de chaque correction : j'ai retenu 0,0001, "
     "une valeur prudente qui évite d'effacer ce que le modèle pré-entraîné "
     "savait déjà. "
     "Ces valeurs sont commandées par l'entraînement sur processeur. Les images "
     "sont d'ailleurs réduites à 320 pixels pour la détection, pour la même "
     "raison."),

    ("Q. Qu'est-ce que le mAP@0.5, et pourquoi ne pas donner une exactitude ?",
     "C'est la précision moyenne sur l'ensemble des classes, à condition qu'une "
     "boîte prédite recouvre au moins la moitié de la boîte réelle. Le 0.5 "
     "désigne ce seuil de recouvrement. "
     "L'exactitude n'a pas de sens en détection : il faudrait rapporter les "
     "bonnes réponses à un total, or une image contient un nombre variable "
     "d'objets, et le modèle peut aussi bien en inventer que en manquer. Le mAP "
     "intègre ces deux erreurs à la fois."),

    ("Q. Entre précision et rappel, lequel compte le plus dans votre cas ?",
     "La précision dit quelle part de mes détections est juste, le rappel quelle "
     "part des déchets réellement présents a été trouvée. Chez moi, 0,797 de "
     "précision pour 0,717 de rappel : je manque plus d'objets que je n'en "
     "invente. "
     "Dans mon usage, le rappel importe davantage. Manquer une nuisance est plus "
     "grave que d'en signaler une à tort, puisqu'un agent valide de toute façon "
     "le diagnostic. C'est aussi pourquoi j'ai abaissé le seuil de confiance à "
     "0,25 plutôt que de le laisser à 0,5."),

    ("Q. Avez-vous fait de l'augmentation de données ?",
     "Non, et c'est un manque que j'assume. L'augmentation consiste à fabriquer "
     "des variantes des images d'entraînement, rotations, changements de "
     "luminosité, recadrages, afin que le modèle apprenne à reconnaître un objet "
     "indépendamment de ces variations. "
     "C'est précisément le levier qui répondrait à ma limite principale : un "
     "corpus issu de sources ouvertes, qui ne reflète ni la poussière ni la "
     "lumière crue d'un chantier d'Abidjan."),

    ("Q. Qu'est-ce qu'ONNX, et pourquoi convertir vos modèles ?",
     "ONNX est un format d'échange de modèles entraînés. J'entraîne avec "
     "Ultralytics et PyTorch, mais je ne peux pas embarquer ces bibliothèques "
     "dans une application Flutter. La conversion produit un fichier que lit "
     "ONNX Runtime, qui existe pour Android. "
     "C'est ce qui rend l'exécution locale possible, et donc le fonctionnement "
     "hors connexion, qui était ma contrainte de départ."),
]

# Questions ajoutees en fin de la section securite, donnees et architecture.
SATELLITE = [
    ("Q. D'où viennent vos seuils environnementaux ?",
     "De trois sources qu'il faut distinguer. Le PGES fixe les nuisances à "
     "surveiller et l'arrêté MINEEF les seuils sonores : c'est le niveau "
     "réglementaire. Le NDVI et le NDWI relèvent de la convention en "
     "télédétection, mes valeurs reprenant les paliers admis, du sol nu sous 0,2 "
     "au couvert établi au-dessus de 0,4, et le changement de signe pour le "
     "stress hydrique. Le NO2 et l'indice pluie-relief, enfin, sont calés "
     "empiriquement sur la ligne de base observée. "
     "Ce sont donc des seuils de vigilance et non de conformité : ils servent à "
     "décider quel chantier visiter en premier, pas à déclarer une infraction."),

    ("Q. Pourquoi ne pas reprendre la valeur limite de l'OMS pour le NO2 ?",
     "Parce que ce ne sont pas les mêmes grandeurs. Sentinel-5P mesure depuis "
     "l'espace la totalité du dioxyde d'azote contenu dans la colonne d'air "
     "au-dessus du site, en micromoles par mètre carré de sol. L'OMS fixe une "
     "concentration à hauteur de respiration, en microgrammes par mètre cube "
     "d'air. "
     "Passer de l'une à l'autre supposerait de connaître la répartition verticale "
     "du polluant, que le satellite ne fournit pas. Mes seuils sont donc "
     "relatifs, et c'est aussi pourquoi le mémoire précise que l'analyse "
     "satellitaire oriente les priorités sans remplacer la mesure instrumentée."),

    ("Q. Comment un indice comme le NDVI fonctionne-t-il concrètement ?",
     "Il compare deux couleurs. La chlorophylle absorbe le rouge pour la "
     "photosynthèse, tandis que la structure interne de la feuille renvoie "
     "fortement le proche infrarouge. Une végétation saine donne donc beaucoup "
     "d'infrarouge et peu de rouge, soit un grand écart entre les deux bandes. Du "
     "sol nu les renvoie à peu près pareil, et l'écart s'annule. "
     "Le NDWI suit le même principe avec l'infrarouge moyen, celui qu'absorbe "
     "l'eau contenue dans la feuille : moins la plante a d'eau, plus l'indice "
     "baisse, jusqu'à devenir négatif."),

    ("Q. Pourquoi parle-t-on d'indice « normalisé » ?",
     "Parce qu'on divise l'écart entre les deux bandes par leur somme. Sans "
     "cette division, une même parcelle paraîtrait différente selon la hauteur "
     "du soleil ou le passage d'un nuage. "
     "La division annule cet effet de luminosité et ne conserve que le rapport "
     "entre les couleurs. C'est ce qui rend deux dates comparables, et donc ce "
     "qui autorise la comparaison avant et après travaux."),
]


def inserer_apres(ancre, modele_q, modele_r, questions):
    """Pose les couples question et reponse a la suite de l'ancre."""
    precedent = ancre
    for question, reponse in questions:
        for modele, texte in ((modele_q, question), (modele_r, reponse)):
            element = copy.deepcopy(modele._element)
            precedent._element.addnext(element)
            p = Paragraph(element, modele._parent)
            for fragment in list(p.runs)[1:]:
                fragment._element.getparent().remove(fragment._element)
            if p.runs:
                p.runs[0].text = texte
            else:
                p.add_run(texte)
            precedent = p
    return len(questions)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    # 1. Les deux formulations qui divergeaient du memoire.
    for avant, apres in CORRECTIONS:
        for p in doc.paragraphs:
            if avant in p.text:
                entier = p.text.replace(avant, apres)
                for fragment in list(p.runs)[1:]:
                    fragment._element.getparent().remove(fragment._element)
                p.runs[0].text = entier
                print(f"  corrige : {apres[:56]}")
                break

    # 2. Modeles de mise en forme, pris sur des paragraphes existants.
    paras = doc.paragraphs
    modele_q = next(p for p in paras if p.style.name == "Question")
    indice_q = paras.index(modele_q)
    modele_r = next(p for p in paras[indice_q + 1:]
                    if p.style.name == "Normal" and len(p.text.strip()) > 80)

    # 3. Insertion en fin de section, juste avant le titre de la suivante.
    def dernier_avant(titre):
        rang = next(i for i, p in enumerate(paras)
                    if p.style.name == "Heading 1" and p.text.strip().startswith(titre))
        for p in reversed(paras[:rang]):
            if p.text.strip():
                return p
        raise SystemExit(f"rien avant {titre}")

    n = inserer_apres(dernier_avant("5. Questions de niveau 4"),
                      modele_q, modele_r, IA)
    print(f"  {n} questions ajoutees a la section intelligence artificielle")

    paras = doc.paragraphs
    n = inserer_apres(dernier_avant("6. Questions pièges"),
                      modele_q, modele_r, SATELLITE)
    print(f"  {n} questions ajoutees a la section donnees et architecture")

    doc.save(SOURCE)
    total = len([p for p in Document(SOURCE).paragraphs if p.style.name == "Question"])
    print(f"\n{total} questions au total dans le guide")
    print(f"Sauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
