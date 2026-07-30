"""
inserer_captures_output.py
==========================
Insere les captures d'output terminal dans les Annexes A et B du memoire.
Supprime les anciennes captures de code.
"""
from docx import Document
from docx.shared import Inches, Pt
import os

CHEMIN_ENTREE = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v48.docx"
CHEMIN_SORTIE = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v53.docx"
DOSSIER_CAPTURES = r"D:\etude_soutenance\SI-ENV\ia\captures_annexes"

doc = Document(CHEMIN_ENTREE)

# Trouver les paragraphes des annexes
annexe_a_idx = None
annexe_b_idx = None
placeholder_idx = None

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if "Annexe A" in t and "Extraits" in t:
        annexe_a_idx = i
    elif "Annexe B" in t:
        annexe_b_idx = i
    elif "contenu des annexes" in p.text:
        placeholder_idx = i

print(f"Annexe A: {annexe_a_idx} | Annexe B: {annexe_b_idx} | Placeholder: {placeholder_idx}")

# Captures d'output a inserer (sans legende)
captures = [
    (annexe_a_idx, "annexe_a1_output_yolo.png"),
    (annexe_a_idx, "annexe_a2_output_mobilenet.png"),
    (annexe_b_idx, "annexe_b1_output_metriques_yolo.png"),
    (annexe_b_idx, "annexe_b2_output_metriques_mobilenet.png"),
]

# Inserer en ordre inverse apres le paragraphe de reference
for ref_idx, img_file in reversed(captures):
    ref_para = doc.paragraphs[ref_idx]
    img_path = os.path.join(DOSSIER_CAPTURES, img_file)
    if not os.path.exists(img_path):
        print(f"ATTENTION: {img_path} non trouve")
        continue
    
    # Image
    img_p = doc.add_paragraph()
    img_p.alignment = 1
    run = img_p.add_run()
    run.add_picture(img_path, width=Inches(5.8))
    ref_para._element.addnext(img_p._element)
    
    # Espacement
    empty_p = doc.add_paragraph()
    ref_para._element.addnext(empty_p._element)
    
    print(f"  Insere: {img_file}")

# Vider le placeholder
if placeholder_idx is not None:
    placeholder = doc.paragraphs[placeholder_idx]
    for run in placeholder.runs:
        run.text = ""

doc.save(CHEMIN_SORTIE)
print(f"\n>> Document sauvegarde : {CHEMIN_SORTIE}")
