# -*- coding: utf-8 -*-
"""
Etape 15 : corriger les deux dernieres formes elidees non accentuees.

L'etape 14 tokenisait les mots avec l'apostrophe incluse : « l'etude » formait
un seul jeton, absent du dictionnaire, donc laisse tel quel. La correction se
fait ici sur la partie qui suit l'apostrophe, au niveau du paragraphe entier
pour franchir les coupures de runs.

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os
import re
import sys

from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

APRES_APOSTROPHE = {
    'etude': 'étude', 'etudes': 'études', 'etat': 'état',
    'element': 'élément', 'elements': 'éléments',
    'echelle': 'échelle', 'etape': 'étape', 'etapes': 'étapes',
    'evaluation': 'évaluation', 'evolution': 'évolution',
    'energie': 'énergie', 'entite': 'entité', 'identite': 'identité',
    'integration': 'intégration', 'integrite': 'intégrité',
    'interet': 'intérêt', 'operation': 'opération',
    'implementation': 'implémentation', 'interpretation': 'interprétation',
    'entrainement': 'entraînement', 'evenement': 'événement',
    'experience': 'expérience', 'ecosysteme': 'écosystème',
    'economie': 'économie', 'ecran': 'écran', 'echange': 'échange',
    'ecart': 'écart', 'enquete': 'enquête',
}

MOTIF = re.compile(
    r"([’'])(" + '|'.join(sorted(APRES_APOSTROPHE, key=len, reverse=True))
    + r")\b", re.IGNORECASE)


def corriger(texte):
    return MOTIF.sub(
        lambda m: m.group(1) + APRES_APOSTROPHE[m.group(2).lower()], texte)


def ecrire(paragraphe, contenu):
    noeuds = list(paragraphe._element.iter(W + 't'))
    if not noeuds:
        return
    noeuds[0].text = contenu
    for n in noeuds[1:]:
        n.text = ''


def main():
    doc = Document(CIBLE)
    n = 0
    cibles = []

    def traiter(p):
        nonlocal n
        avant = p.text
        apres = corriger(avant)
        if apres == avant:
            return
        # On tente d'abord au niveau du noeud de texte, pour ne pas perdre la
        # mise en forme interne du paragraphe.
        touche = False
        for t in p._element.iter(W + 't'):
            if not t.text:
                continue
            nv = corriger(t.text)
            if nv != t.text:
                t.text = nv
                touche = True
        if not touche:
            # La forme est coupee entre deux runs : on reecrit le paragraphe.
            ecrire(p, apres)
            cibles.append(avant[:70])
        n += 1

    for p in doc.paragraphs:
        traiter(p)
    for table in doc.tables:
        for ligne in table.rows:
            for cellule in ligne.cells:
                for p in cellule.paragraphs:
                    traiter(p)

    print("  paragraphes corriges : %d" % n)
    for c in cibles:
        print("    reecrit : %s" % c)

    doc.save(CIBLE)

    # Controle
    verif = Document(CIBLE)
    reste = set()
    for p in verif.paragraphs:
        for m in MOTIF.finditer(p.text):
            reste.add(m.group(0))
    print("  restant : %s" % (sorted(reste) if reste else 'aucun'))


if __name__ == "__main__":
    main()
