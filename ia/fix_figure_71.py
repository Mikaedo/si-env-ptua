"""
fix_figure_71.py
================
Ajoute la legende "Figure 7.1" sous l'image Swagger UI deja presente dans le memoire.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

MEMOIRE = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v43.docx"

doc = Document(MEMOIRE)

# Trouver le paragraphe avec l'image (index 451) et ajouter la legende apres
for i, p in enumerate(doc.paragraphs):
    has_img = False
    for run in p.runs:
        if run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
            has_img = True
    # L'image Swagger est apres le paragraphe sur FastAPI (index 450)
    if has_img and i > 448 and i < 455:
        print(f"Image trouvee a l'index {i}")

        # Verifier s'il y a deja une legende apres
        next_p = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
        if next_p and "figure 7.1" in next_p.text.lower():
            print(f"  Legende deja presente: {next_p.text}")
        else:
            # Creer un paragraphe legende
            legende_para = doc.add_paragraph()
            legende_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_leg = legende_para.add_run("Figure 7.1 : Documentation interactive des services web (Swagger UI / FastAPI)")
            run_leg.font.size = Pt(10)
            run_leg.font.italic = True
            p._element.addnext(legende_para._element)
            print(f"  [OK] Legende ajoutee apres l'image")
        break

output = MEMOIRE.replace("v43", "v44")
doc.save(output)
print(f"\n>> Memoire sauvegarde : {output}")
