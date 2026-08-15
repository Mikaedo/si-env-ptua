# -*- coding: utf-8 -*-
"""
Reprend les tableaux et passages du guide restes a cinq profils.

La section 12 ajoutee precedemment couvre les nouveautes, mais les sections
d'origine continuaient d'annoncer cinq roles et une seule application mobile.
Un guide qui se contredit d'une page a l'autre est pire qu'un guide incomplet :
le candidat ne sait plus lequel des deux passages fait foi, et c'est en general
pendant la soutenance qu'il s'en apercoit.

Deux tableaux sont donc repris a la source, celui des profils en tete de guide
et celui des comptes de la fiche de secours, plutot que d'etre laisses en
contradiction avec la section finale.
"""
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\GUIDE_PREPARATION_SOUTENANCE_SI-ENV.docx")
SAUVEGARDE = SOURCE.with_name("GUIDE_PREPARATION_SOUTENANCE_SI-ENV_avant_tableaux.docx")

# Tableau 0 : profils et outils, en tete de guide.
PROFILS = [
    ("Responsable Environnement", "Mobile agent",
     "Saisit les signalements sur le terrain, y compris hors connexion."),
    ("Expert HSE", "Mobile agent",
     "Contrôle externe : instruit les signalements et valide les actions correctives."),
    ("Spécialiste Suivi Environnemental", "Web",
     "Consolide les signalements, pilote l'analyse satellitaire, produit et transmet les rapports."),
    ("Spécialiste Suivi du P.A.R", "Web",
     "Instruit les doléances des riverains, quelle que soit leur provenance."),
    ("Administrateur", "Web",
     "Gère les comptes, le modèle embarqué et les journaux."),
    ("ANDE", "Web (consultation)",
     "Autorité de tutelle. Vérifie la conformité, reçoit les rapports. Aucune écriture."),
    ("BAD", "Web (consultation)",
     "Bailleur. Contrôle les sauvegardes, volet social compris. Aucune écriture."),
    ("Riverain", "Mobile citoyen",
     "Dépose des doléances depuis son téléphone, après vérification de sa position."),
]

# Tableau 6 : comptes de la fiche de secours.
COMPTES = [
    ("Administrateur", "admin@sienv.ci", "admin123"),
    ("Responsable Environnement (mobile agent)", "resp.env@ageroute.ci", "env123"),
    ("Expert HSE (mobile agent)", "expert.hse@ageroute.ci", "expert123"),
    ("Spécialiste Suivi Environnemental", "spec.env@ageroute.ci", "spec123"),
    ("Spécialiste Suivi du P.A.R", "spec.par@ageroute.ci", "spec123"),
    ("ANDE (consultation)", "controle@ande.ci", "ande123"),
    ("BAD (consultation)", "mission@afdb.org", "bad123"),
    ("Riverain (mobile citoyen)", "riverain@yopougon.ci", "riverain123"),
]


def ecrire(cellule, texte, gras=False, taille=9):
    cellule.text = ""
    r = cellule.paragraphs[0].add_run(texte)
    r.bold = gras
    r.font.size = Pt(taille)


def remplir(tableau, entetes, lignes):
    """Redimensionne un tableau existant et le remplit."""
    while len(tableau.rows) > 1:
        tableau._tbl.remove(tableau.rows[-1]._tr)
    for i, t in enumerate(entetes):
        if i < len(tableau.columns):
            ecrire(tableau.rows[0].cells[i], t, gras=True)
    for ligne in lignes:
        cells = tableau.add_row().cells
        for i, v in enumerate(ligne):
            if i < len(cells):
                ecrire(cells[i], v)


def reecrire_paragraphe(p, texte):
    for seg in list(p.runs)[1:]:
        seg._element.getparent().remove(seg._element)
    if p.runs:
        p.runs[0].text = texte
    else:
        p.add_run(texte)


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)

    remplir(doc.tables[0], ("Profil", "Outil", "Ce qu'il fait"), PROFILS)
    print("  tableau des profils : 8 lignes")

    remplir(doc.tables[6], ("Profil", "Adresse", "Mot de passe"), COMPTES)
    print("  tableau des comptes : 8 lignes")

    # La phrase sur l'exclusion mutuelle des deux applications reste vraie mais
    # devient incomplete : il en existe desormais trois surfaces.
    for p in doc.paragraphs:
        if p.text.strip().startswith("Précision qui fait bonne impression"):
            reecrire_paragraphe(p,
                "Précision qui fait bonne impression : les trois surfaces sont "
                "mutuellement exclusives. Un agent de terrain est refusé sur le "
                "web, un profil bureau est refusé sur le mobile des agents, et un "
                "riverain n'entre que dans l'application citoyenne. Le contrôle "
                "porte sur le rôle porté par le compte, vérifié à la connexion, "
                "et non sur l'application par laquelle on se présente.")
            print("  phrase sur l'exclusion mutuelle actualisee")
            break

    for p in doc.paragraphs:
        if p.text.strip().startswith("Rappel : un compte web est refusé"):
            reecrire_paragraphe(p,
                "Rappel : chaque compte n'ouvre qu'une surface. Si une connexion "
                "échoue en démonstration, vérifiez d'abord que vous utilisez le "
                "compte correspondant à l'application ouverte. Une même adresse "
                "ne peut porter qu'un seul rôle, la contrainte d'unicité "
                "s'appliquant à l'ensemble des comptes.")
            print("  rappel de la fiche de secours actualise")
            break

    doc.save(SOURCE)
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
