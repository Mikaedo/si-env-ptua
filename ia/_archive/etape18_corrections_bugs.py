# -*- coding: utf-8 -*-
"""
Etape 18 : corriger trois defauts introduits par les etapes precedentes.

  a) Table des matieres. Trois entrees de l'ancienne table saisie a la main ont
     survecu a son remplacement par un champ Word : RESUME, ABSTRACT et ANNEXES
     apparaissent deux fois, avec des numeros de page differents. Les entrees
     situees APRES la fin du champ sont des orphelines et sont supprimees. Elles
     se reconnaissent a leur position : le champ se termine par un caractere de
     champ « end », tout paragraphe de style « toc N » qui suit lui est etranger.

  b) Accentuation de l'anglais. L'etape 14 a accentue des mots anglais qui
     s'ecrivent comme le francais sans accent : « air emissions » est devenu
     « air émissions », « late detection » « late détection », et dans la
     bibliographie « Object Detection with Region Proposal Networks » a pris
     deux accents. Les zones anglaises, resume anglais et titres d'articles,
     sont desaccentuees.

  c) Accents oublies. Deux paragraphes des annexes E et F portaient des formes
     que le dictionnaire de l'etape 14 ne couvrait pas.

Sortie : mise a jour de MEMOIRE_FINAL.docx
"""
import os
import re
import sys

from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

CIBLE = os.path.join('C:', os.sep, 'Users', 'DELL', 'Downloads', 'MEMOIRE',
                     'MEMOIRE_FINAL.docx')
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

#: Mots dont la forme francaise accentuee est aussi un mot anglais sans accent.
#: Appliques uniquement aux passages rediges en anglais.
DESACCENTUER = {
    'détection': 'detection', 'différence': 'difference',
    'région': 'region', 'émission': 'emission', 'émissions': 'emissions',
    'référence': 'reference', 'références': 'references',
    'général': 'general', 'générale': 'general',
    'modèle': 'model', 'système': 'system', 'précision': 'precision',
    'séquence': 'sequence', 'intégration': 'integration',
    'opération': 'operation', 'évaluation': 'evaluation',
    'élément': 'element', 'éléments': 'elements',
    'série': 'series', 'scénario': 'scenario', 'métrique': 'metric',
    'prédiction': 'prediction', 'résolution': 'resolution',
    'végétation': 'vegetation', 'température': 'temperature',
}

#: Mots anglais courants servant a reconnaitre un passage en anglais.
INDICES_ANGLAIS = ('the', 'of', 'and', 'for', 'with', 'this', 'which',
                   'based', 'using', 'towards', 'shortcomings', 'keywords',
                   'proposal', 'networks', 'index', 'detector', 'learning',
                   'convolutional', 'image', 'handbook', 'dataset', 'report')

#: Formes oubliees par le dictionnaire de l'etape 14.
OUBLIS = [
    ("Les tableaux E.1 a E.6 detaillent",
     "Les tableaux E.1 à E.6 détaillent"),
    ("Les tableaux F.1 a F.5 rassemblent",
     "Les tableaux F.1 à F.5 rassemblent"),
]


def est_anglais(texte):
    mots = set(re.findall(r"[a-zA-Z']+", texte.lower()))
    return sum(1 for m in INDICES_ANGLAIS if m in mots) >= 2


def casse_de(source, cible):
    if source[:1].isupper():
        return cible[:1].upper() + cible[1:]
    return cible


def desaccentuer(texte):
    def remplace(m):
        mot = m.group(0)
        bon = DESACCENTUER.get(mot.lower())
        return casse_de(mot, bon) if bon else mot
    return re.sub(r"[A-Za-zÀ-ÿ]+", remplace, texte)


def ecrire(paragraphe, contenu):
    noeuds = list(paragraphe._element.iter(W + 't'))
    if not noeuds:
        return
    noeuds[0].text = contenu
    for n in noeuds[1:]:
        n.text = ''


def main():
    doc = Document(CIBLE)
    body = doc.element.body

    # ── a) Entrees de table des matieres orphelines ──────────────────────────
    # Position du dernier caractere de fin de champ du document.
    elements = list(body)
    fin_champ = -1
    for i, el in enumerate(elements):
        for fc in el.iter(W + 'fldChar'):
            if fc.get(W + 'fldCharType') == 'end':
                fin_champ = i
    print("  fin du champ de table des matieres : element %d" % fin_champ)

    orphelines = []
    for i in range(fin_champ + 1, len(elements)):
        el = elements[i]
        if not el.tag.endswith('}p'):
            continue
        ppr = el.find(W + 'pPr')
        style = ppr.find(W + 'pStyle') if ppr is not None else None
        nom = style.get(W + 'val') if style is not None else ''
        if nom and nom.lower().startswith('toc'):
            orphelines.append(el)

    for el in orphelines:
        body.remove(el)
    print("  entrees de table des matieres orphelines supprimees : %d"
          % len(orphelines))

    # ── b) Desaccentuation des passages anglais ──────────────────────────────
    n_ang = 0
    for p in doc.paragraphs:
        t = p.text
        if len(t) < 25 or not est_anglais(t):
            continue
        touche = False
        for noeud in p._element.iter(W + 't'):
            if not noeud.text:
                continue
            nv = desaccentuer(noeud.text)
            if nv != noeud.text:
                noeud.text = nv
                touche = True
        if touche:
            n_ang += 1
    print("  paragraphes anglais desaccentues : %d" % n_ang)

    # ── c) Accents oublies ──────────────────────────────────────────────────
    n_ou = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        for avant, apres in OUBLIS:
            if t.startswith(avant):
                ecrire(p, t.replace(avant, apres, 1))
                n_ou += 1
                break
    print("  formes oubliees corrigees : %d / %d" % (n_ou, len(OUBLIS)))

    doc.save(CIBLE)
    print("\nEnregistre : %s" % CIBLE)


if __name__ == "__main__":
    main()
