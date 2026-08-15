# -*- coding: utf-8 -*-
"""
Ajoute au sommaire la pagination de chacune de ses entrees.

Les numeros ne sont pas inventes : ils sont repris de deux sources deja
fiables. La table des matieres, champ Word mis a jour a chaque export, fournit
la page de tout titre indexe. Les trois pages de partie n'etant pas des titres,
elles n'y figurent pas ; leur numero est lu dans le PDF, sur la page elle-meme,
en reprenant le folio imprime plutot que le rang physique de la page. La
distinction compte : les pages liminaires portent des chiffres romains, si bien
que la conclusion generale est la cinquante-deuxieme page numerotee mais la
soixante-quatrieme feuille du document.

La mise en forme reprend celle d'une table des matieres : une tabulation
alignee a droite sur la marge, garnie de points de conduite. Sans elle les
numeros se colleraient au libelle et le sommaire se lirait mal.
"""
import re
import shutil
import unicodedata
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
PDF = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_N'GUESSAN_DIBY_FINAL.pdf")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_pagination.docx")

LARGEUR_UTILE = Cm(15.5)          # marges de 3 cm et 2,5 cm sur 21 cm
PARTIES = ["PREMIÈRE PARTIE", "DEUXIÈME PARTIE", "TROISIÈME PARTIE"]


def cle(texte):
    """Normalise un libelle pour le comparer sans risque d'accent ou de casse."""
    sans = unicodedata.normalize("NFD", texte)
    sans = "".join(c for c in sans if unicodedata.category(c) != "Mn")
    return re.sub(r"\s+", " ", sans).strip().upper()


def pages_des_parties():
    """Folio imprime de chaque page de partie, lu dans le PDF."""
    trouve = {}
    with pdfplumber.open(PDF) as pdf:
        for page in pdf.pages:
            lignes = [l.strip() for l in (page.extract_text() or "").split("\n")
                      if l.strip()]
            for partie in PARTIES:
                if partie in lignes and partie not in trouve:
                    folios = [l for l in lignes if re.fullmatch(r"\d+", l)]
                    if folios:
                        trouve[partie] = folios[-1]
    return trouve


def pages_de_la_table(paras):
    """Couples (libelle, page) extraits de la table des matieres."""
    depart = next(i for i, p in enumerate(paras)
                  if p.text.strip() == "TABLE DES MATIÈRES")
    simples, partielles = {}, []
    for p in paras[depart + 1:]:
        if not p.style.name.startswith("toc"):
            continue
        morceaux = p.text.strip().split("\t")
        if len(morceaux) < 2 or not morceaux[-1].strip():
            continue
        libelle, page = morceaux[0].strip(), morceaux[-1].strip()
        if cle(libelle) == "CONCLUSION PARTIELLE":
            partielles.append(page)
        else:
            simples.setdefault(cle(libelle), page)
    return simples, partielles


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    paras = doc.paragraphs

    simples, partielles = pages_de_la_table(paras)
    parties = pages_des_parties()
    print(f"{len(simples)} titres indexes, {len(partielles)} conclusions "
          f"partielles, {len(parties)} pages de partie")

    debut = next(i for i, p in enumerate(paras) if p.text.strip() == "SOMMAIRE")
    fin = next(i for i, p in enumerate(paras)
               if i > debut and p.text.strip() == "LISTE DES FIGURES")

    rang_partielle, rang_partie = 0, 0
    for p in paras[debut + 1:fin]:
        libelle = p.text.strip().split("\t")[0].strip()
        if not libelle:
            continue

        k = cle(libelle)
        page = None
        if k == "CONCLUSION PARTIELLE":
            if rang_partielle < len(partielles):
                page = partielles[rang_partielle]
            rang_partielle += 1
        elif "PARTIE" in k and k.split()[0] in {"PREMIERE", "DEUXIEME", "TROISIEME"}:
            if rang_partie < len(PARTIES):
                page = parties.get(PARTIES[rang_partie])
            rang_partie += 1
        elif k in simples:
            page = simples[k]
        else:
            # « Chapitre 3 : Analyse des besoins » est indexe sous « Chapitre 3 ».
            court = re.match(r"(CHAPITRE\s+\d+)", k)
            if court:
                page = simples.get(court.group(1))

        if page is None:
            print(f"  page introuvable : {libelle[:52]}")
            continue

        for seg in list(p.runs)[1:]:
            seg._element.getparent().remove(seg._element)
        if p.runs:
            p.runs[0].text = f"{libelle}\t{page}"
        else:
            p.add_run(f"{libelle}\t{page}")

        # Tabulation a droite, garnie de points, comme une table des matieres.
        taquets = p.paragraph_format.tab_stops
        for _ in range(len(taquets)):
            taquets._pPr.tabs.remove(taquets._pPr.tabs[0])
        taquets.add_tab_stop(LARGEUR_UTILE, WD_TAB_ALIGNMENT.RIGHT,
                             WD_TAB_LEADER.DOTS)
        print(f"  {libelle[:52]:<54}{page}")

    doc.save(SOURCE)
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
