# -*- coding: utf-8 -*-
"""
Met a jour le memoire v56 -> v57 : chapitre 10 (tests et performances) avec
des resultats reellement mesures le 31 juillet 2026 (suite pytest, deploiement
Docker reel, PDF PGES genere, alerte+email cables et verifies, satellite GEE
operationnel), au lieu des valeurs pre-remplies non verifiees de la v56.

Ne touche PAS au chapitre 8 (IA) : les figures/tableaux de la v56 proviennent
d'un entrainement GPU reel (Google Colab, dataset Recycle Trash) et ont ete
verifies coherents (moyenne des mAP par classe ~0.797 = valeur du tableau).
"""
from docx import Document
import copy

SRC = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v56.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v57.docx"

doc = Document(SRC)

def replace_paragraph_text(idx, new_text):
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)
    print(f"[OK] Paragraphe {idx} remplace")

# ============================================================
# 1. Intro section 10.2 (para 537) : mentionner la verification reelle
# ============================================================
replace_paragraph_text(537,
    "Les tests unitaires (pytest, PostGIS, Flutter) verifient les composants isoles. "
    "Les tests d'integration valident la chaine mobile -> backend -> dashboard. "
    "La suite pytest (32 tests automatises, fichier tests/test_functional.py) a ete executee "
    "le 31 juillet 2026 sur la stack Docker reelle (base PostGIS, backend FastAPI, nginx) : "
    "32/32 tests passes. Le tableau 10.2 synthetise les douze scenarios fonctionnels couverts."
)

# ============================================================
# 2. Tableau 10.2 (doc.tables[29]) : corriger T04 et T07
# ============================================================
t2 = doc.tables[29]
# Colonnes : 0=Test, 1=Description, 2=Resultat attendu, 3=Statut
# T04 : diagnostic IA local -> preciser le contexte de mesure reel (CPU desktop)
t2.rows[4].cells[2].text = "Score en < 200 ms (mesure CPU desktop, modele embarque actuel)"
t2.rows[4].cells[3].text = "Pass (13-27 ms)"
# T07 : la notification "push" n'existe pas dans le code ; seul l'email est implemente
t2.rows[7].cells[2].text = "Alerte creee automatiquement (seuil GEE depasse) + email SMTP"
t2.rows[7].cells[3].text = "Pass"
print("[OK] Tableau 10.2 (T04, T07) corrige")

# ============================================================
# 3. Tableau 10.3 (doc.tables[30]) : vraies valeurs mesurees le 31/07/2026
# ============================================================
t3 = doc.tables[30]
data_103 = [
    ["Indicateur", "Valeur mesuree", "Seuil ou reference", "Conformite"],
    ["Temps de reponse API median (via nginx, local)", "31 ms", "< 500 ms", "Conforme"],
    ["Synchronisation d'un lot de 10 signalements", "0,56 s", "< 5 s", "Conforme"],
    ["Inference YOLOv8 (ONNX Runtime, CPU desktop)", "27,0 ms", "< 200 ms", "Conforme (a confirmer sur mobile reel)"],
    ["Inference MobileNetV2 (ONNX Runtime, CPU desktop)", "13,1 ms", "< 200 ms", "Conforme (a confirmer sur mobile reel)"],
    ["Recuperation des indices satellitaires GEE (6 chantiers PTUA)", "19,2 s", "a titre indicatif", "Mesure"],
]
for i, row_data in enumerate(data_103):
    for j, val in enumerate(row_data):
        if i < len(t3.rows) and j < len(t3.rows[i].cells):
            t3.rows[i].cells[j].text = val
print("[OK] Tableau 10.3 (performances) remplace par des valeurs reellement mesurees")

# ============================================================
# 4. Paragraphe 542 (texte avant tableau 10.3) : lever la contradiction
#    "a mesurer lors du deploiement" vs tableau presentant des "valeurs mesurees"
# ============================================================
replace_paragraph_text(542,
    "Les indicateurs ci-dessous ont ete mesures le 31 juillet 2026 sur un deploiement Docker "
    "local (3 conteneurs : PostGIS, FastAPI, nginx), et non sur le VPS cible de production "
    "(2 vCPU / 4 Go RAM) : les temps reseau/API pourront differer legerement en production. "
    "Les temps d'inference ONNX ont ete mesures sur processeur d'ordinateur (CPU desktop) et "
    "restent a confirmer sur un telephone reel. Le modele de detection embarque dans "
    "l'application mobile est actuellement une version d'entrainement rapide mono-classe ; "
    "le modele Recycle Trash a six classes valide au chapitre 8 (GPU, mAP@0.5 = 0,798) n'a pas "
    "encore ete reexporte et integre a l'application (cf. discussion 10.5)."
)

# ============================================================
# 5. Paragraphe 545 (discussion) : adoucir "temps reel" + ajouter la limite
#    du modele mobile non encore mis a jour
# ============================================================
replace_paragraph_text(545,
    "Le SI-ENV repond aux six lacunes du chapitre 4 : signalement instantane avec diagnostic IA, "
    "subjectivite encadree, geolocalisation automatique, donnees centralisees, rapports en "
    "secondes, alertes quasi temps reel (rafraichissement automatique du tableau de bord toutes "
    "les 10 a 15 secondes ; la synchronisation mobile reste declenchee manuellement par "
    "l'utilisateur). Les limites : tests fonctionnels executes en environnement de developpement "
    "local (pas encore sur le VPS de production), dataset Recycle Trash non encore valide sur des "
    "photographies reelles des chantiers du PTUA, modele de detection six classes valide au "
    "chapitre 8 pas encore reexporte vers l'application mobile (version embarquee : modele "
    "mono-classe d'entrainement rapide), effet de propagation d'erreur inherent au pipeline en "
    "cascade (une detection manquee ou imprecise par YOLOv8 degrade la classification de "
    "criticite par MobileNetV2 en aval), dependance Internet pour le satellite, resolution "
    "Sentinel-5P trop coarse pour le niveau chantier."
)

# ============================================================
# 6. Insertion des captures d'ecran reelles (mobile, dashboard, satellite)
#    a la place des placeholders "[A coller : ... : Figure X.X]".
#    Insertion en ordre INVERSE (paragraphe le plus loin d'abord) pour ne
#    pas decaler les index des placeholders traites ensuite.
# ============================================================
from docx.shared import Inches

SCRATCH = r"C:\Users\DELL\AppData\Local\Temp\claude\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25\scratchpad"

def insert_figure_before(placeholder_idx, image_path, caption_text, width_inches=5.5):
    placeholder = doc.paragraphs[placeholder_idx]
    img_para = placeholder.insert_paragraph_before()
    img_para.alignment = placeholder.alignment
    run = img_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    for run in placeholder.runs:
        run.text = ""
    if placeholder.runs:
        placeholder.runs[0].text = caption_text
    else:
        r = placeholder.add_run(caption_text)
        r.italic = True
        r.bold = True
    print(f"[OK] Figure inseree avant paragraphe {placeholder_idx} : {caption_text[:60]}")

# Ajouter un renvoi explicite dans le texte avant chaque figure (guide UPB)
replace_paragraph_text(523,
    "Le Spec. Env. selectionne une zone d'interet sur le dashboard. FastAPI transmet la requete "
    "a GEE (source, dates, indicateur). Pour le NO2, GEE calcule la concentration pixel par pixel "
    "et applique un masque de nuages avant de renvoyer une heatmap ; pour le risque pluie/relief, "
    "GEE croise directement le cumul CHIRPS et la pente SRTM, sans masque de nuages necessaire. "
    "Le resultat est stocke dans PostGIS. Le dashboard permet de comparer deux periodes (avant/apres "
    "travaux) pour alimenter les rapports PGES. La figure 9.1 illustre ces indices, calcules a "
    "partir de donnees Google Earth Engine reelles, pour les six chantiers PTUA."
)
replace_paragraph_text(463,
    "Le module de rapports produit un PGES PDF en un clic : le serveur agrege les donnees par "
    "chantier, applique le modele BAD et renvoie le fichier. Cette operation, autrefois manuelle "
    "et sur plusieurs jours, prend quelques secondes. La figure 7.3 presente le tableau de bord, "
    "avec la carte des tracés PTUA et la repartition des signalements."
)
replace_paragraph_text(458,
    "L'application fonctionne hors connexion d'abord : SQLite stocke les signalements en attente, "
    "synchronises au retour du reseau. L'interface est optimisee terrain (boutons larges, contraste "
    "eleve, saisie < 1 minute). L'architecture suit le pattern BLoC ; onnxruntime execute YOLOv8 "
    "localement (13-27 ms sur CPU desktop selon le modele, cf. tableau 10.3). Le tableau 7.3 liste "
    "les packages principaux. La figure 7.2 presente la liste des signalements, la carte des "
    "chantiers et le detail d'un signalement."
)

insert_figure_before(524, f"{SCRATCH}\\fig_9_1_satellite.png",
    "Figure 9.1 : Indices environnementaux calcules via Google Earth Engine (NO2, NDVI, NDWI, "
    "risque pluie/relief) pour les six chantiers PTUA — donnees reelles Sentinel-5P/Sentinel-2, "
    "capture du 31 juillet 2026.")
insert_figure_before(464, f"{SCRATCH}\\fig_7_3_dashboard.png",
    "Figure 7.3 : Captures d'ecran du tableau de bord (vue d'ensemble, statistiques et carte des "
    "traces PTUA).")
insert_figure_before(459, f"{SCRATCH}\\fig_7_2_mobile.png",
    "Figure 7.2 : Captures d'ecran de l'application mobile (liste des signalements, carte des "
    "chantiers PTUA, detail d'un signalement).")

doc.save(DST)
print(f"\n=== MEMOIRE SAUVEGARDE : {DST} ===")
