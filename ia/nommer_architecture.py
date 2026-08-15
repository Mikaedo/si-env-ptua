# -*- coding: utf-8 -*-
"""
Nomme l'architecture en tete du paragraphe 4.2.

Le paragraphe enumerait cinq modules sans jamais dire de quel style
d'architecture il s'agissait. C'est une lacune que le jury releve : le mot
attendu, trois-tiers, n'y figurait pas, et un lecteur presse pouvait conclure
a une simple juxtaposition de composants.

La phrase ajoutee dit aussi pourquoi il s'agit de trois-tiers et non d'un
client-serveur a deux niveaux : les clients n'atteignent jamais la base
directement. C'est le seul critere qui separe reellement les deux modeles, et
c'est celui qu'une question de jury viendra chercher.

L'exception du module d'intelligence artificielle est mentionnee dans la
foulee. Mieux vaut la poser soi-meme que la laisser decouvrir : une partie du
traitement vit dans le niveau presentation, ce qui s'ecarte du modele et
demande une justification, donnee au paragraphe 5.13.

Le texte est lu depuis un fichier en UTF-8, les accents ne survivant pas a une
composition en ligne de commande.
"""
import copy
import shutil
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\Users\DELL\Downloads\MEMOIRE\MEMOIRE_FINAL.docx")
SAUVEGARDE = SOURCE.with_name("MEMOIRE_FINAL_avant_architecture.docx")
TEXTE = Path(r"C:\Users\DELL\AppData\Local\Temp\claude"
             r"\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25"
             r"\scratchpad\archi.txt")


def main():
    shutil.copy2(SOURCE, SAUVEGARDE)
    doc = Document(SOURCE)
    paras = doc.paragraphs

    titre = next(i for i, p in enumerate(paras)
                 if p.text.strip().startswith("4.2") and "Architecture" in p.text)
    modele = next(p for p in paras[titre + 1:]
                  if p.style.name == "Normal" and len(p.text.strip()) > 200)
    print(f"titre au rang {titre}, modele de mise en forme : "
          f"{modele.text.strip()[:46]}")

    # Le nouveau paragraphe est un clone du suivant : il en herite justification,
    # interligne et espacement, sans qu'aucune valeur ne soit ecrite en dur.
    element = copy.deepcopy(modele._element)
    paras[titre]._element.addnext(element)
    nouveau = Paragraph(element, modele._parent)
    for seg in list(nouveau.runs)[1:]:
        seg._element.getparent().remove(seg._element)
    nouveau.runs[0].text = TEXTE.read_text(encoding="utf-8").strip()

    doc.save(SOURCE)

    controle = next(p for p in Document(SOURCE).paragraphs
                    if p.text.strip().startswith("Le SI-ENV repose"))
    accents = sum(controle.text.count(c) for c in "éèêàçôûîœ")
    print(f"insere : {len(controle.text)} caracteres, {accents} accents")
    print(f"\nSauvegarde : {SAUVEGARDE.name}")


if __name__ == "__main__":
    main()
