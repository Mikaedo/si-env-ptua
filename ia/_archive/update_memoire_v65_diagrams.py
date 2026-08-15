# -*- coding: utf-8 -*-
"""v65 -> v66 : remplace les figures 6.3 (classes), 6.8 (MCD), 6.9 (MLD) par
des versions reconstruites avec cardinalites/multiplicites propres et des
traits qui touchent precisement les boites (calcul geometrique)."""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

SRC = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v65.docx"
DST = r"C:\Users\DELL\Downloads\MEMOIRE_SI-ENV_v66.docx"
SCRATCH = r"C:\Users\DELL\AppData\Local\Temp\claude\d--etude-soutenance-SI-ENV\3233866b-194c-446a-b8f6-65c65b911c25\scratchpad"

doc = Document(SRC)

def replace_image(pidx, image_path, width_inches=6.1):
    p = doc.paragraphs[pidx]
    for r in p.runs:
        drawings = r.element.findall(qn('w:drawing'))
        if drawings:
            for d in drawings:
                d.getparent().remove(d)
            r.add_picture(image_path, width=Inches(width_inches))
            print(f"[OK] Figure remplacee au paragraphe {pidx}")
            return
    raise RuntimeError(f"Aucune image trouvee au paragraphe {pidx}")

# Indices retrouves par legende dans v65 (chapitre 6 non modifie depuis v64)
replace_image(389, SCRATCH + "/class_render.png")
replace_image(416, SCRATCH + "/mcd_render.png")
replace_image(420, SCRATCH + "/mld_render.png")

doc.save(DST)
print(f"\n=== MEMOIRE SAUVEGARDE : {DST} ===")
